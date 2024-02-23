import datetime
# from datetime import datetime
from odoo import fields, models, api, _, exceptions
from odoo.exceptions import UserError
import math
import base64
from io import BytesIO
import decimal
try:
    from odoo.tools.misc import xlsxwriter
except ImportError:
    from odoo.addons.helper import xlsxwriter
from num2words import num2words


class CashIncentiveHead(models.Model):
    _name = "cash.incentive.head"
    _description = "Cash Incentive"
    _order = "name desc"
    _rec_name = "name"
    _inherit = ['mail.thread', 'mail.activity.mixin']

    file_data = fields.Binary(' Report')
    code = fields.Char(string='ERP Code',  copy=False, readonly=True,
                       states={'draft': [('readonly', False)]}, index=True, default=lambda self: _('New'), tracking=1)
    name = fields.Char(string='Reference', required=True, tracking=2)
    swift_customer_name = fields.Char(string='SWIFT Customer', tracking=True)
    date = fields.Date(string='Preparation Date', required=True, default=fields.Date.context_today, tracking=3)
    swift_date = fields.Date(string='SWIFT Date', tracking=4)
    application_deadline = fields.Date(string='Application Deadline', tracking=5, help='Application Deadline Will be 179 Days More than SWIFT Date', compute='_compute_application_deadline',store=True)
    remaining_days = fields.Integer(string='Remaining days', compute='_compute_remaining_day', tracking=6, search='_value_search')
    basis_fee_amt = fields.Float(string='BASIS Fee (BDT)', tracking=True)

    od_sight_rate = fields.Float(string='OD Sight Rate', digits=(16, 4), tracking=7)

    # not used -----------------
    incentive_rate_fc = fields.Float(string='Incentive Rate (FC)(%)', digits=(16, 4), default=10, tracking=8)
    encashment_rate_bdt = fields.Float(string='Encashment Rate (BDT)', digits=(16, 4), tracking=9)

    date_credited_beneficiaries = fields.Date(string='Date Credited Beneficiaries', tracking=10)

    reporting_st_to_bb = fields.Char(string='(g) Reporting statement/schedule to BB with Month', tracking=11)
    ref_online_to_bb = fields.Char(string='(h) Reference of Online reporting to BB', tracking=12)
    #not used -------------------

    partner_id = fields.Many2one('res.partner', string='Customer', domain="[('type', '=', 'contact'), ('active', '=', True), ('customer_rank', '>', 0)]", tracking=13)
    customer_address = fields.Text(tracking=14)
    institution_address = fields.Text(tracking=15)
    customer_country_id = fields.Many2one('res.country', string="Customer Country", tracking=16)

    bank_id = fields.Many2one('res.bank', string='Bank', tracking=17, domain="[('is_cash_incentive_bank', '=', True)]")

    invoice_line_ids = fields.One2many('cash.incentive.invoice', 'head_id', string='Invoices', tracking=18)
    swift_file_line_ids = fields.One2many('incentive.swift.message.download.line', 'head_id', string='SWIFT Files', tracking=18)
    swift_ids = fields.Many2many('swift.message', string='SWIFT Messages', tracking=True)
    #prc_line_ids = fields.One2many('cash.incentive.prc', 'head_id', string='PRC')

    state = fields.Selection([
        ('draft', 'Draft'),
        ('confirm', 'Confirmed'),
        ('approve', 'Approved'),
        ('reject', 'Rejected'),
        ('done', 'Completed'),
        ('cancel', 'Cancelled'),
    ], string='Status', copy=False, default='draft', tracking=19)

    fc_currency_id = fields.Many2one("res.currency", string="Currency (FC)", tracking=20)

    contract_price = fields.Float(string='Contract Price', tracking=21)
    contract_price_str = fields.Char(string='Contract Price', tracking=21)
    contract_no = fields.Float(string='Contract Number', tracking=True)
    contract_id = fields.Many2one("client.contract", string="Contract No.", tracking=20)
    contract_ids = fields.Many2many("client.contract", string="Contract No.", tracking=20)
    contract_number = fields.Char(string='Contract No.', tracking=True)
    contract_date = fields.Date(string='Contract Date', tracking=True)
    contract_date_str = fields.Char(string='Contract Date', tracking=True)

    #------------ SWIFT
    swift_file_file = fields.Binary(string='SWIFT Message File', attachment=True, tracking=22)
    swift_file_name = fields.Char("File Name", tracking=23)
    remiter_bank_name = fields.Char("Remiter Bank Name", tracking=24)
    remiter_bank_address = fields.Text("Remiter Bank Address", tracking=25)

    #--------- Invoice
    invoice_file_file = fields.Binary(string='Invoice File', attachment=True, tracking=26)
    invoice_file_name = fields.Char("File Name", tracking=27)

    # ------------- PRC
    prc_ref_code = fields.Char(string='PRC Reference', default="", tracking=28)
    prc_date = fields.Date(string='PRC Date', tracking=True)
    prc_file_file = fields.Binary(string='PRC File', attachment=True, tracking=29)
    prc_file_file_name = fields.Char("File Name", tracking=30)
    prc_letter_description = fields.Html(string="Description")

    # ------------- Forwarding Letter (BASIS)
    flbs_ref_code = fields.Char(string='Forwarding (BASIS) Reference', tracking=32)
    flbs_file_file = fields.Binary(string='Forwarding File', attachment=True, tracking=33)
    flbs_file_file_name = fields.Char("File Name", tracking=34)
    flbs_letter_description = fields.Html(string="Description")

    # ------------- Forwarding Letter (BANK)
    flbk_ref_code = fields.Char(string='Forwarding (Bank) Reference', tracking=36)
    flbk_file_file = fields.Binary(string='Forwarding File', attachment=True, tracking=37)
    flbk_file_file_name = fields.Char("File Name", tracking=38)
    flbk_letter_description = fields.Html(string="Description")

    #-------------------- Otehrs
    form_ka_basis_description = fields.Html(string="Letter Template")

    # ------------- Forwarding Letter (BANK)
    form_kha_ref_code = fields.Char(string='Form Kha Reference', tracking=41)
    form_kha_file_file = fields.Binary(string='Form Kha File', attachment=True, tracking=42)
    form_kha_file_file_name = fields.Char("File Name", tracking=43)
    form_kha_basis_description = fields.Html(string="Letter Template")

    form_c_description = fields.Html(string="Letter Template")
    form_ga_description = fields.Html(string="Letter Template")
    form_gha_description = fields.Html(string="Letter Template")

    inv_count = fields.Integer(compute="_compute_inv_count", string='Invoice Count', search='_value_search_inv', tracking=True)
    inv_names = fields.Char(string="Invoice Ref.", compute="_compute_inv_count", store=True, tracking=True)
    swift_date_str = fields.Char(string="SWIFT Date", compute="_compute_inv_count", store=True , tracking=True)
    swift_count = fields.Integer(string="SWIFT Count", compute="_compute_inv_count", search='_value_search_swift', tracking=True )
    swift_amount_fc = fields.Float(string="SWIFT Amount (FC)", compute="_compute_inv_count", tracking=True)
    incentive_amount_fc = fields.Float(string="Incentive Amount (FC)", compute="_compute_inv_count", tracking=True)
    incentive_amount_bdt = fields.Float(string="Incentive amount BDT", compute="_compute_inv_count", tracking=True)

    @api.depends("invoice_line_ids", "date")
    def _compute_inv_count(self):
        for rec in self:
            all_inv = self.env['cash.incentive.invoice'].search([('head_id', '=', rec.id)])
            incentive_amount_fc = sum(all_inv.mapped('incentive_amt_fc'))
            incentive_amt_bdt = sum(all_inv.mapped('incentive_amt_bdt'))
            swift_amount_fc = sum(all_inv.mapped('swift_amt'))
            total_count = len(all_inv)
            rec.inv_count = total_count
            inv_names = ''
            swift_date_str = ''
            swift_count = 0
            swift_ids = []
            for x in all_inv:
                if x.swift_message_id:
                    if x.swift_message_id.id not in swift_ids:
                        swift_count += 1
                        swift_ids.append(x.swift_message_id.id)
                if x.invoice_id.ref:
                    inv_names += x.invoice_id.ref if not inv_names else ', ' + x.invoice_id.ref
                if x.swift_message_id:
                    inv_date = datetime.datetime.strptime(str(x.swift_message_id.date), '%Y-%m-%d').strftime('%d/%m/%y')
                    swift_date_str += str(inv_date) if not swift_date_str else ', ' + str(inv_date)
            rec.inv_names = inv_names
            rec.swift_date_str = swift_date_str
            rec.swift_count = swift_count
            rec.swift_amount_fc = swift_amount_fc
            rec.incentive_amount_fc = incentive_amount_fc
            rec.incentive_amount_bdt = incentive_amt_bdt
    
    # @api.onchange('followup_line_ids','followup_line_ids.followup_date')
    # def onchange_followup_line_ids(self):
    #     a = []
    #     for rec in self.followup_line_ids:
    #         a.append(rec.followup_date)
    #     if a:
    #         b = max(a)
    #         self.latest_followup_date = b
    def _value_search_inv(self, operator, value):
        recs = self.search([]).filtered(lambda x: len(x.invoice_line_ids) <= 0)
        if recs:
            return [('id', 'in', [x.id for x in recs] if recs else False)]
        else:
            return [('id', '!=', 0)]

    def _value_search_swift(self, operator, value):
        recs = self.search([]).filtered(lambda x: x.swift_count <= 0)
        if recs:
            return [('id', 'in', [x.id for x in recs] if recs else False)]
        else:
            return [('id', '!=', 0)]

    def _value_search(self, operator, value):
        recs = self.search([]).filtered(lambda x: x.remaining_days <= value)
        if recs:
            return [('id', 'in', [x.id for x in recs] if recs else False)]
        else:
            return [('id', '!=', 0)]

    # @api.depends('date')
    # def _compute_application_deadline(self):
    #     for rec in self:
    #         if rec.date:
    #             rec.application_deadline = rec.date + datetime.timedelta(days=179)
    #         else:
    #             rec.application_deadline = None
    
    @api.model
    def create(self, vals):
        res = super(CashIncentiveHead, self).create(vals)
        if res:
            for rec in res.invoice_line_ids:
                rec.invoice_id.swift_id = rec.swift_message_id.id
                rec.invoice_id.cash_incentive_id = res.id
        return res

    def write(self, vals):
        super(CashIncentiveHead, self).write(vals)
        for rec in self:
            for x in rec.invoice_line_ids:
                x.invoice_id.swift_id = x.swift_message_id.id
                x.invoice_id.cash_incentive_id = rec.id

    @api.onchange('partner_id')
    def onchange_partner_id(self):
        if self.partner_id:
            self.swift_customer_name = self.partner_id.name.upper()

    @api.onchange('swift_customer_name', 'customer_address', 'institution_address')
    def onchange_uppercase_name(self):
        if self.swift_customer_name:
            self.swift_customer_name = self.swift_customer_name.upper()
        if self.customer_address:
            self.customer_address = self.customer_address.upper()
        if self.institution_address:
            self.institution_address = self.institution_address.upper()

    @api.onchange('invoice_line_ids.application_deadline')
    def _compute_application_deadline(self):
        for x in self:
            a = []
            if x.invoice_line_ids:
                for rec in x.invoice_line_ids:
                    if rec.application_deadline:
                        a.append(rec.application_deadline)
                if a:
                    b = min(a)
                    x.application_deadline = b
                else:
                    x.application_deadline = None
            else:
                x.application_deadline = None

    @api.depends("invoice_line_ids.od_sight_rate")
    def _compute_od_sight_rate(self):
        for x in self:
            a = []
            if x.invoice_line_ids:
                for rec in x.invoice_line_ids:
                    if rec.od_sight_rate:
                        a.append(rec.od_sight_rate)
                if a:
                    b = min(a)
                    x.od_sight_rate = b
                else:
                    x.od_sight_rate = 0
            else:
                x.od_sight_rate = 0

    @api.onchange('swift_ids')
    def onchange_swift_ids(self):
        if self.swift_ids:
            swift_ids = [x._origin.id for x in self.swift_ids]
            all_inv_ids = self.env['cash.incentive.invoice'].search(
                [('swift_message_id', 'in', swift_ids)])
            line_ids = []
            self.invoice_line_ids = all_inv_ids

    @api.onchange('contract_ids')
    def on_change_contract_ids(self):
        for rec in self.invoice_line_ids:
            if not rec.contract_ids:
                contract_ids = [x._origin.id for x in self.contract_ids]
                rec.contract_ids = contract_ids
                contract_number = ''
                date_str = ''
                contract_price_str = ''
                for x in self.contract_ids:
                    contract_number += x.reference if not contract_number else ', ' + x.reference
                    if x.date:
                        dates = datetime.datetime.strptime(str(x.date), '%Y-%m-%d').strftime('%d/%m/%y')
                        date_str += dates if not date_str else ', ' + dates
                    if x.range:
                        contract_price_str += x.range if not contract_price_str else ', ' + x.range
                rec.contract_number = contract_number
                rec.contract_date_str = date_str
                #print('mmm1',self.contract_id.range)
                rec.contract_price_str = contract_price_str
                #print('mmm2',rec.contract_price_str)
                # contract_ids = [x.id for x in rec.contract_id]
                rec.invoice_id.contract_ids = contract_ids
    
    # used ------------
    def get_contract_data(self):
        for rec in self.invoice_line_ids:
            # for x, y in zip(rec.contract_ids, self.contract_ids):
                # if x.id == y.id:
            previous_con_ids = [x.id for x in rec.contract_ids]
            contract_ids = [x.id for x in self.contract_ids]
            contract_ids += previous_con_ids
            rec.contract_ids = contract_ids
            contract_number = ''
            date_str = ''
            contract_price_str = ''
            for x in self.contract_ids:
                contract_number += x.reference if not contract_number else ', ' + x.reference
                if x.date:
                    dates = datetime.datetime.strptime(str(x.date), '%Y-%m-%d').strftime('%d/%m/%y')
                    date_str += dates if not date_str else ', ' + dates
                if x.range:
                    contract_price_str += x.range if not contract_price_str else ', ' + x.range
            rec.contract_number = contract_number
            rec.contract_date_str = date_str
            rec.contract_price_str = contract_price_str
            rec.invoice_id.contract_ids = contract_ids
                
        self.on_change_invoice_line_ids()

    # unused - -----------
    def contract_id_change(self, contract_id):
        for rec in self.invoice_line_ids:
            if rec.contract_id.id == contract_id:
                rec.contract_id = rec.contract_id.id
                rec.contract_number = rec.contract_id.code
                date_str = datetime.datetime.strptime(str(self.contract_id.date), '%Y-%m-%d').strftime('%d/%m/%y')
                rec.contract_price_str = rec.contract_id.range
                rec.contract_date_str = date_str
                contract_ids = [x.id for x in rec.contract_id]
                rec.invoice_id.contract_ids = contract_ids
                rec.save_data()

    @api.onchange('bank_id')
    def on_change_bank_ref(self):
        if self.bank_id:
            m_code_id = self.env['cash.incentive.head'].search([('bank_id', '=', self.bank_id.id)], limit=1, order="id DESC")
            m_code = m_code_id.name if m_code_id else ''
            if not m_code_id:
                if self.bank_id.code_prefix:
                    m_code = str(self.bank_id.code_prefix).strip()
                if self.bank_id.code_suffix:
                    if m_code:
                        m_code = m_code + str(self.bank_id.code_suffix)
                    else:
                        m_code = str(self.bank_id.code_suffix).strip()
            self.name = m_code

            # --------------PRC
            prc_ref = m_code_id.prc_ref_code if m_code_id else ''
            if not m_code_id:
                if self.bank_id.prc_ref_prefix:
                    prc_ref = str(self.bank_id.prc_ref_prefix).strip()
                if self.bank_id.prc_ref_suffix:
                    if prc_ref:
                        prc_ref = prc_ref + str(self.bank_id.prc_ref_suffix)
                    else:
                        prc_ref = str(self.bank_id.prc_ref_suffix).strip()
            self.prc_ref_code = prc_ref
            self.prc_letter_description = ''  # self.bank_id.prc_letter_description

            # --------------Forwading BASIS
            flbs_ref = ''
            if self.bank_id.flbs_ref_prefix:
                flbs_ref = str(self.bank_id.flbs_ref_prefix).strip()
            if self.bank_id.flbs_ref_suffix:
                if flbs_ref:
                    flbs_ref = flbs_ref + str(self.bank_id.flbs_ref_suffix)
                else:
                    flbs_ref = str(self.bank_id.flbs_ref_suffix).strip()
            self.flbs_ref_code = flbs_ref
            self.flbs_letter_description = ''  # self.bank_id.flbs_letter_description

            # --------------Forwading BANK
            flbk_ref = ''
            if self.bank_id.flbk_ref_prefix:
                flbk_ref = str(self.bank_id.flbk_ref_prefix).strip()
            if self.bank_id.flbk_ref_suffix:
                if flbk_ref:
                    flbk_ref = flbk_ref + str(self.bank_id.flbk_ref_suffix)
                else:
                    flbk_ref = str(self.bank_id.flbk_ref_suffix).strip()
            self.flbk_ref_code = flbk_ref
            self.flbk_letter_description = ''  # self.bank_id.flbk_letter_description

            # --------------Form Kha
            form_kha_ref = ''
            if self.bank_id.form_kha_ref_prefix:
                form_kha_ref = str(self.bank_id.form_kha_ref_prefix).strip()
            if self.bank_id.form_kha_ref_suffix:
                if form_kha_ref:
                    form_kha_ref = form_kha_ref + str(self.bank_id.form_kha_ref_suffix)
                else:
                    form_kha_ref = str(self.bank_id.form_kha_ref_suffix).strip()
            self.form_kha_ref_code = form_kha_ref

            # ---------
            self.form_ka_basis_description = ''
            self.form_kha_basis_description = ''
            self.form_c_description = ''
            self.form_ga_description = ''
            self.form_gha_description = ''
        else:
            self.prc_letter_description = ''
            self.flbs_letter_description = ''
            self.flbk_letter_description = ''

            self.form_ka_basis_description = ''
            self.form_kha_basis_description = ''
            self.form_c_description = ''
            self.form_ga_description = ''
            self.form_gha_description = ''

    @api.onchange('partner_id')
    def on_change_partner_id(self):
        if self.partner_id:
            street = self.partner_id.street
            street2 = self.partner_id.street2
            city = self.partner_id.city
            state = self.partner_id.state_id.name if self.partner_id.state_id else ''
            zip = self.partner_id.zip
            country = self.partner_id.country_id.name if self.partner_id.country_id else ''
            country_id = self.partner_id.country_id.id if self.partner_id.country_id else None

            address = ''
            if street:
                address = street
            if street2:
                if address:
                    address += ', ' + street2
                else:
                    address = street2
            if city:
                if address:
                    address += ', ' + city
                else:
                    address = city
            if state:
                if address:
                    address += ', ' + state
                else:
                    address = state
            if zip:
                if address:
                    address += ', ' + zip
                else:
                    address = zip
            if country:
                if address:
                    address += ', ' + country
                else:
                    address = country

            self.customer_address = address
            self.institution_address = address
            self.customer_country_id = country_id

        else:
            self.customer_address = ''
            self.institution_address = ''
            self.customer_country_id = None

    @api.onchange('invoice_line_ids')
    def on_change_invoice_line_ids(self):
        if self.invoice_line_ids:
            contract_number = ''
            contract_number_po = ''
            contract_number_prefix_po = ''
            contract_number_cntr = ''
            contract_number_prefix_cntr = ''
            contract_number_wo = ''
            contract_number_prefix_wo = ''

            contract_date_str = ''
            contract_price_str = ''
            basis_fee_amt = 0
            contract_id = []
            
            a = []
            for rec in self.invoice_line_ids:
                if rec.od_sight_rate:
                    a.append(rec.od_sight_rate)

                if rec.swift_customer_name:
                    self.swift_customer_name = rec.swift_customer_name.upper()
                basis_fee_amt += rec.invoice_amt
                l_contract_number = ''
                l_contract_date_str = ''
                l_contract_price_str = ''
                for l in rec.invoice_id.contract_ids:
                    date_str = datetime.datetime.strptime(str(l.date), '%Y-%m-%d').strftime('%d/%m/%Y')
                    l_contract_date_str += str(date_str) if not l_contract_date_str else ', ' + str(date_str)
                    l_contract_number += str(l.reference) if not l_contract_number else ', ' + str(l.reference)
                    if l.range is not False:
                        l_contract_price_str += str(l.range) if not l_contract_price_str else ', ' + str(l.range)

                    if l.id not in contract_id:
                        contract_id.append(l.id)
                        if l.type == '1':
                            contract_number_po += str(l.reference) if not contract_number_po else ', ' + str(l.reference)
                            contract_number_prefix_po = 'Purchase Order No. '
                        if l.type == '0':
                            contract_number_cntr += str(l.reference) if not contract_number_cntr else ', ' + str(l.reference)
                            contract_number_prefix_cntr = 'Contract No. '
                        if l.type == '2':
                            contract_number_wo += str(l.reference) if not contract_number_wo else ', ' + str(l.reference)
                            contract_number_prefix_wo = 'Work Order No. '

                        date_str_l = datetime.datetime.strptime(str(l.date), '%Y-%m-%d').strftime('%d/%m/%Y')
                        contract_date_str += str(date_str_l) if not contract_date_str else ', ' + str(date_str_l)
                        # contract_number += str(l.reference) if not contract_number else ', ' + str(l.reference)
                        if l.range:
                            contract_price_str += str(l.range) if not contract_price_str else ', ' + str(l.range)

                rec.contract_number = l_contract_number
                rec.contract_date_str = l_contract_date_str
                rec.contract_price_str = l_contract_price_str
                rec.partner_id = self.partner_id.id

            if a:
                b = min(a)
                self.od_sight_rate = b
            else:
                self.od_sight_rate = 0

            if contract_number_po:
                contract_number = contract_number_prefix_po + contract_number_po
            if contract_number_cntr:
                if contract_number:
                    contract_number += ' '+contract_number_prefix_cntr + contract_number_cntr
                else:
                    contract_number = contract_number_prefix_cntr + contract_number_cntr

            if contract_number_wo:
                if contract_number:
                    contract_number += ' '+contract_number_prefix_wo + contract_number_wo
                else:
                    contract_number = contract_number_prefix_wo + contract_number_wo

            self.contract_number = contract_number
            self.contract_date_str = contract_date_str
            self.contract_price_str = contract_price_str

            # Basis fee calculation
            config_rec = self.env['basis.fee.configuration'].sudo().search([('active', '=', True)], limit=1)
            if config_rec:
                config_rec_line = self.env['basis.fee.configuration.line'].sudo().search([
                    ('head_id', '<=', config_rec.id),
                    ('from_amount', '<=', basis_fee_amt), # Check if basis_fee_amt is greater than or equal to from_amount
                    ('to_amount', '>=', basis_fee_amt)  # Check if basis_fee_amt is less than or equal to to_amount
                ], limit=1)
                if config_rec_line:
                    self.basis_fee_amt = config_rec_line.fee_amount
                else:
                    if basis_fee_amt > config_rec.max_amount:
                        config_last_rec_line = self.env['basis.fee.configuration.line'].sudo().search(
                            [('head_id', '<=', config_rec.id)
                             ], order='fee_amount DESC', limit=1)
                        highest_fee_amount = config_last_rec_line.fee_amount
                        a = basis_fee_amt - config_rec.max_amount
                        if a >= config_rec.extend_amount_range:
                            b = a / config_rec.extend_amount_range
                            trunc_b = math.trunc(b)
                            c = trunc_b * config_rec.extend_amount
                            self.basis_fee_amt = highest_fee_amount + c
                        else:
                            self.basis_fee_amt = highest_fee_amount
            else:
                raise UserError(_('Set a Basis Fee Configuration.'))
                # --------------- previous calculation ---------
                # if basis_fee_amt < 5001:
                #     self.basis_fee_amt = 850
                # elif basis_fee_amt > 5000 and basis_fee_amt < 10001:
                #     self.basis_fee_amt = 1600
                # elif basis_fee_amt > 10000 and basis_fee_amt < 30001:
                #     self.basis_fee_amt = 3100
                # elif basis_fee_amt > 30000 and basis_fee_amt < 50001:
                #     self.basis_fee_amt = 6200
                # elif basis_fee_amt > 50000 and basis_fee_amt < 80001:
                #     self.basis_fee_amt = 12100
                # elif basis_fee_amt > 80000 and basis_fee_amt < 120001:
                #     self.basis_fee_amt = 24100
                # elif basis_fee_amt > 120000:
                #     a = basis_fee_amt - 120000
                #     if a > 50000:
                #         b = a / 50000
                #         truncA = math.trunc(b)
                #         c = truncA * 3000
                #         self.basis_fee_amt = 24100 + c
                #     else:
                #         self.basis_fee_amt = 24100

        else:
            self.contract_number = ''
            self.contract_date_str = ''
            self.contract_price_str = ''
            self.basis_fee_amt = 0

    # @api.onchange('swift_date')
    # def on_change_swift_date(self):
    #     print('swift_date',self.swift_date)
    #     if self.swift_date:
    #         self.application_deadline = self.swift_date + datetime.timedelta(days=179)

    @api.depends("application_deadline")
    def _compute_remaining_day(self):
        for rec in self:
            today = fields.Date.today()
            application_deadline = rec.application_deadline
            remaining_days=0
            if today and application_deadline:
                try:
                    remaining_days = (application_deadline - today).days
                except:
                    remaining_days = 0
            rec.remaining_days = remaining_days

    # prc ------------------
    def action_refresh_prc(self):
        html_thead = """
            <thead>
                <th style='vertical-align: middle; text-align:center; border: 1px solid; width:80px; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                    Invoice No.
                </th>
                <th style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                    (c) Amount in FC (%s)
                </th>
                <th style='vertical-align: middle; text-align:center; border: 1px solid; width:80px; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                    (d) Date of Credit in bank's nostro account
                </th>
            </thead>
        """

        # Create a dictionary to group rows by invoice number
        invoice_groups = {}
        for rec in self.invoice_line_ids:
            if rec.invoice_id.ref in invoice_groups:
                invoice_groups[rec.invoice_id.ref].append(rec)
            else:
                invoice_groups[rec.invoice_id.ref] = [rec]
        # Generate HTML for merged rows
        html = ""
        for invoice_num, rows in invoice_groups.items():
            html += "<tr>"
            html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{invoice_num}</td>"
            for i, row in enumerate(rows):
                if i != 0:
                    html += "<tr>"
                html += f"<td style='vertical-align: middle; text-align:right; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{str('{:,}'.format(round(row.swift_amt, 2)))}</td>"
                html += f"<td style='vertical-align: middle; text-align:right; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{row.swift_date}</td>"
                # if i == 0:
                #     html += f"<td rowspan='{len(rows) * 2}' style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{row.swift_date}</td>"
                html += "</tr>"

        # Generate the final HTML with the table
        expected_html = """
            <table class='table table-bordered'>
                {} <!-- table header -->
                <tbody>
                    {} <!-- table body -->
                </tbody>
            </table>
        """.format(html_thead, html)

        prc_text = self.bank_id.prc_letter_description

        final_text = ''
        if prc_text:
            final_text = prc_text.replace('$invoice_lines', expected_html)
        self.prc_letter_description = final_text

    def action_refresh_prc3(self):
        #self.prc_letter_description = self.bank_id.prc_letter_description
        if self.prc_ref_code =='' or self.prc_ref_code == False:
            prc_ref = ''
            if self.bank_id.prc_ref_prefix:
                prc_ref = str(self.bank_id.prc_ref_prefix).strip()
            if self.bank_id.prc_ref_suffix:
                if prc_ref:
                    prc_ref = prc_ref + str(self.bank_id.prc_ref_suffix)
                else:
                    prc_ref = str(self.bank_id.prc_ref_suffix).strip()
            self.prc_ref_code = prc_ref
        #---------------

        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name
        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

        html = ''
        html_thead = """
         <thead>

            <th style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                (a) Remitter
            </th>
            <th  style='vertical-align: middle; text-align:center; line-height:22px; border: 1px solid; width:300px; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                (b) Address
            </th>
            <th  style='vertical-align: middle; text-align:center; border: 1px solid; width:80px; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                Invoice No.
            </th>
            <th style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                (c) Amount in FC (%s)
            </th>
            <th  style='vertical-align: middle; text-align:center; border: 1px solid; width:80px; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                (d) Date of Credit in bank's nostro account
            </th>
            <th style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                (e) Equivalent Taka
            </th>
            <th  style='vertical-align: middle; text-align:center; border: 1px solid; width:80px; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                (f) Credited to beneficiary a/c
            </th>
            <th  style='vertical-align: middle; text-align:center; border: 1px solid'; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;>
                (g) Reporting statement/schedule to BB with Month
            </th>
            <th  style='vertical-align: middle; text-align:center; border: 1px solid; width:150px; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>
                (h) Reference of Online reporting to BB
            </th>

        </thead>
            """%(currency)
        inv_number = ''
        inv_date = ''
        inv_amount = 0
        encashment_amt_bdt = 0
        swift_amt_total = 0
        equivalent_taka_total = 0
        cus_name = ''
        if self.swift_customer_name:
            cus_name = self.swift_customer_name
        html += """
        <tr>
         <td style='vertical-align: middle; text-align:center; border: 1px solid; font-weight:bold;' rowspan=" """+str(len(self.invoice_line_ids)+1)+"""">
                """ + cus_name + """
          </td>
         <td style='vertical-align: middle; text-align:center; border: 1px solid; font-weight:bold;' rowspan=" """+str(len(self.invoice_line_ids)+1)+"""">
                """+str(self.customer_address)+"""
          </td>
        """
        # Create a dictionary to group rows by invoice number
        invoice_groups = {}
        invoice_ids = []
        for rec in self.invoice_line_ids:
            if rec.invoice_id.ref in invoice_groups:
                invoice_groups[rec.invoice_id.ref].append(rec)
            else:
                invoice_groups[rec.invoice_id.ref] = [rec]
        # Generate HTML for merged rows
        for invoice_num, rows in invoice_groups.items():
            inv_obj = self.env['cash.incentive.invoice'].search([('invoice_id.ref', '=', invoice_num),('head_id', '=', self.id)], limit=1)
            html += "<tr>"
            html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{invoice_num}</td>"
            inv_ids = []
            for i, rec in enumerate(rows):
                invoice_date = ''
                if rec.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
                if rec.invoice_id not in inv_ids:
                    inv_ids.append(rec.invoice_id)
                    inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
                    inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
                if i == 0:
                    inv_amount += rec.invoice_amt
                encashment_amt_bdt += rec.encashment_amt_bdt
                date_credited_beneficiaries = ''
                if rec.date_credited_beneficiaries:
                    date_credited_beneficiaries = datetime.datetime.strptime(str(rec.date_credited_beneficiaries), '%Y-%m-%d').strftime('%d-%b-%y')
                swift_date = ''
                if rec.swift_message_id.date:
                    swift_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime(
                        '%d-%b-%y')
                ref_online_to_bb = rec.ref_online_to_bb if rec.ref_online_to_bb else ''
                reporting_st_to_bb = rec.reporting_st_to_bb if rec.reporting_st_to_bb else ''
                equivalent_taka = round(rec.swift_amt * rec.encashment_rate_bdt, 2)

                equivalent_taka_decimal = "{:,.2f}".format(equivalent_taka)

                # if equivalent_taka == int(equivalent_taka):
                #     equivalent_taka_decimal = "{:,.0f}".format(equivalent_taka)
                # else:
                #     # if the number has decimal value
                #     if equivalent_taka - int(equivalent_taka) >= 0.5:
                #         equivalent_taka_decimal = "{:,.0f}".format(round(equivalent_taka))
                #     else:
                #         equivalent_taka_decimal = "{:,.0f}".format(equivalent_taka)

                equivalent_taka_total += equivalent_taka
                swift_amt_total += rec.swift_amt
                if i != 0:
                    html += "<tr>"
                html += f"<td style='vertical-align: middle; text-align:right; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{str('{:,}'.format(rec.swift_amt))}</td>"
                html += """

                     <td style="vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">
                        """+str(swift_date)+"""
                    </td>
                    <td style="vertical-align: middle; text-align:right; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">
                        """+str(equivalent_taka_decimal)+"""
                    </td>
                    <td style="vertical-align: middle; text-align:center; border: 1px solid; width:200px !important'; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">
                        """+str(date_credited_beneficiaries)+"""
                    </td>
                    <td style="vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">
                        """+str(reporting_st_to_bb)+"""
                    </td>
                    <td style="vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">
                        """+str(ref_online_to_bb)+"""
                    </td >
                """
                # if i == 0:
                #     html += f"<td rowspan='{len(rows) * 2}' style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{row.swift_date}</td>"
                html += "</tr>"
        equivalent_taka_total_decimal = "{:,.0f}".format(round(equivalent_taka_total))
        html += """
                </tr>
                 <tr style="font-weight:bold;">
                    <td style="border: 1px solid"></td>
                    <td style="border: 1px solid"></td>
                    <td style="text-align:right; border: 1px solid"> Total  </td>
                    <td style="text-align:right; border: 1px solid"> """+str("{:,.2f}".format(swift_amt_total))+""" </td>
                    <td style="border: 1px solid; width:200px !important"> </td>
                    <td style="text-align:right; border: 1px solid"> """+str(equivalent_taka_total_decimal)+""" </td>
                    <td style="border: 1px solid"> </td>
                    <td style="border: 1px solid"> </td>
                    <td style="border: 1px solid"> </td>
                </tr>
                    """

        expected_html = """
        <table class="table table-bordered">
            """+html_thead+"""
            <tbody>
                """+html+"""
            </tbody>
        </table>
        """
        prc_text = self.bank_id.prc_letter_description

        final_text = ''
        if prc_text:
            final_text = prc_text.replace('$ref_name', self.name).replace('$date', str(self_date)).replace('$contract_number', '' if not self.contract_number else str(self.contract_number))\
                .replace('$contract_date', '' if not self.contract_date_str else '& date: ' + str(self.contract_date_str)).replace('$invoice_number', inv_number)\
                .replace('$invoice_date', '& date: '+inv_date).replace('$currency', currency).replace('$total_amount',str("{:,}".format(round(inv_amount, 2)))).replace('$invoice_lines', expected_html)
        self.prc_letter_description = final_text

    # BASIS ------------------------------
    def action_refresh_flbs(self):
        # self.flbs_letter_description = self.bank_id.flbs_letter_description
        # --------------Forwading BASIS
        if self.flbs_ref_code == '' or self.flbs_ref_code == False:
            flbs_ref = ''
            if self.bank_id.flbs_ref_prefix:
                flbs_ref = str(self.bank_id.flbs_ref_prefix).strip()
            if self.bank_id.flbs_ref_suffix:
                if flbs_ref:
                    flbs_ref = flbs_ref + str(self.bank_id.flbs_ref_suffix)
                else:
                    flbs_ref = str(self.bank_id.flbs_ref_suffix).strip()
            self.flbs_ref_code = flbs_ref

        inv_number = ''
        # inv_date = ''
        inv_ids = []
        for rec in self.invoice_line_ids:
            if rec.invoice_id not in inv_ids:
                inv_ids.append(rec.invoice_id)
                inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
            # inv_date += str(rec.invoice_date) if not inv_date else ', ' + str(rec.invoice_date)
        current_date = datetime.datetime.now().date()
        date_str = ''
        if current_date:
            date_str = datetime.datetime.strptime(str(current_date),'%Y-%m-%d').strftime('%d/%m/%y')
        flbs_text = self.bank_id.flbs_letter_description
        final_text = ''
        if flbs_text:
            final_text = flbs_text.replace('$ref_name', self.name).replace('$date', str(date_str)).replace(
                '$contract_number', '' if not self.contract_number else 'and ' + str(self.contract_number)).replace('$invoice_number', str(inv_number))
        self.flbs_letter_description = final_text

    # .replace('$invoice_number', str(inv_number)).replace('<p><br></p>', '')
    # ka -------------------
    def action_refresh_form_ka(self):
        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name
        flbs_text = self.bank_id.form_ka_basis_description
        if flbs_text:
            flbs_text = self.bank_id.form_ka_basis_description.replace('<tr ', '<tr style="border: 1px solid"').replace('<td ', '<td style="border: 1px solid"').replace('<th ', '<th style="border: 1px solid"') \
                .replace('<tr>', '<tr style="border: 1px solid">').replace('<td>', '<td style="border: 1px solid">').replace('<th>', '<th style="border: 1px solid">')
        html = ''
        html_thead = """
                         <thead style='font-size:24px'>
                         <tr>
                            <th style="text-align:center; border: 1px solid">
                                পণ্যের বর্ণনা
                            </th>
                            <th style="text-align:center; border: 1px solid; width:300px;">
                                পরিমাণ
                            </th>
                            <th style="text-align:center; border: 1px solid">
                                আমদানিকারকের দেশের নাম
                            </th>
                            <th style="text-align:right; border: 1px solid; width:180px;">
                                ইনভয়েস মূল্য (বৈদেশিক মুদ্রায়)
                            </th>
                            <th style="border: 1px solid">
                                জাহাজীকরণ/রপ্তানির তারিখ
                            </th>
                            <th style="text-align:center; border: 1px solid">
                                ইএ·পি নম্বর*
                            </th>
                            <th colspan='2' style="text-align:center; border: 1px solid">
                                বৈদেশিক মুদ্রায় প্রত্যাবাসিত রপ্তানিমূল্য ও প্রত্যাবাসনের তারিখ
                            </th>                 
                        </tr>
                        <tr>
                            <th style="text-align:center; border: 1px solid">
                                ১
                            </th>
                            <th style="text-align:center; border: 1px solid">
                                ২
                            </th>
                            <th style="text-align:center; border: 1px solid">
                                ৩
                            </th>
                            <th style="text-align:center; border: 1px solid">
                                ৪
                            </th >
                            <th style="text-align:center; border: 1px solid">
                                ৫
                            </th>
                            <th style="text-align:center; border: 1px solid">
                                ৬
                            </th>
                            <th colspan='2' style="text-align:center; border: 1px solid">
                                ৭
                            </th>                  
                        </tr>
                        </thead>
                            """

        inv_number = ''
        # inv_date = ''
        inv_amount = 0
        incentive_rate_fc = 0
        swift_amt = 0
        encashment_amt_bdt = 0
        total_incentive_rate_fc = 0
        sl = 0
        html += """
                 <td style="text-align:center; border: 1px solid" rowspan=" """ + str(len(self.invoice_line_ids) + 1) + """">
                        Software Development
                  </td>
                """
        customer = False
        # processed_invoices = set(self.invoice_line_ids.invoice_id)
        # dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in self.invoice_line_ids.invoice_id.invoice_line_ids)
        inv_ids = []
        dev_total1 = 0
        hour_total1 = 0

        for line in self.invoice_line_ids:
            if line.invoice_id.id not in inv_ids:
                dev_total1 += sum(r.quantity if r.quantity_type == '1' else 0 for r in line.invoice_id.invoice_line_ids)
                hour_total1 += sum(r.quantity if r.quantity_type == '0' else 0 for r in line.invoice_id.invoice_line_ids)
            inv_ids.append(line.invoice_id.id)

        # hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in self.invoice_line_ids.invoice_id.invoice_line_ids)
        hour_total = float(format(hour_total1, '.2f'))
        dev_total = float(format(dev_total1, '.2f'))

        all_invoice_qty_str = ''
        if dev_total > 0:
            all_invoice_qty_str += str('ITES/') + str(dev_total) if not all_invoice_qty_str else ' & ' + str('ITES/') + str(
                dev_total)
        if hour_total > 0:
            all_invoice_qty_str += str('ITES/') + str(hour_total) + str(' hrs') if not all_invoice_qty_str else ' & ' + str(
                hour_total) + str(' hrs')
        currency_symbol = ''
        if self.fc_currency_id.symbol:
            currency_symbol = self.fc_currency_id.symbol
        country_name = ''
        if self.partner_id.country_id:
            country_name = self.partner_id.country_id.name
        # Create a dictionary to group rows by invoice number
        invoice_groups = {}
        a = []
        c = []
        for rec in self.invoice_line_ids:
            a.append(rec.od_sight_rate)
            if rec.encashment_rate_bdt:
                c.append(rec.encashment_rate_bdt)
            if rec.invoice_id.ref in invoice_groups:
                invoice_groups[rec.invoice_id.ref].append(rec)
            else:
                invoice_groups[rec.invoice_id.ref] = [rec]
        if a:
            b = min(a)
            od_s_r = b
        else:
            od_s_r = 0

        if c:
            d = min(c)
            encashment_rate_bdt = d
        else:
            encashment_rate_bdt = 0

        # Generate HTML for merged rows
        for invoice_num, rows in invoice_groups.items():
            html += "<tr>"
            for i, rec in enumerate(rows):
                # if rec.swift_message_id.encashment_rate_bdt:
                #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
                # usd_price = sum(rec.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
                usd_price = rec.invoice_id.invoice_total_fc
                incentive_rate_fc = usd_price / self.incentive_rate_fc if self.incentive_rate_fc else 0
                swift_amt += rec.swift_amt
                if i == 0:
                    inv_amount += usd_price
                encashment_amt_bdt += rec.encashment_amt_bdt
                total_incentive_rate_fc += incentive_rate_fc
                inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
                invoice_date = ''
                swift_message_date = ''
                if rec.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d-%b-%y')
                if rec.swift_message_id.date:
                    swift_message_date = datetime.datetime.strptime(str(rec.swift_message_id.date),
                                                                    '%Y-%m-%d').strftime('%d-%b-%y')

                # getting hours and developers ===========
                l_dev_total = sum(r.quantity if r.quantity_type == '1' else 0 for r in rec.invoice_id.invoice_line_ids)
                l_hour_total = sum(r.quantity if r.quantity_type == '0' else 0 for r in rec.invoice_id.invoice_line_ids)
                l_all_invoice_qty_str = ''
                if l_dev_total > 0:
                    l_all_invoice_qty_str += str('ITES/') + str(
                        l_dev_total) if not l_all_invoice_qty_str else ' & ' + str(
                        'ITES/') + str(
                        l_dev_total)
                if l_hour_total > 0:
                    l_all_invoice_qty_str += str('ITES/') + str(l_hour_total) + str(
                        ' hrs') if not l_all_invoice_qty_str else ' & ' + str(
                        l_hour_total) + str(' hrs')

                sl += 1
                if i == 0:
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; border: 1px solid; text-align: center'>{l_all_invoice_qty_str}</td>"

                if not customer:
                    html += """
                            <td style="text-align:center; border: 1px solid" rowspan=" """ + str(
                        len(self.invoice_line_ids)) + """">
                                """ + str(country_name) + """
                            </td>                         
                            """
                str_invoice_amt = str(currency_symbol)+ ' ' + str("{:,}".format(round(rec.invoice_amt, 2)))
                if i == 0:
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; border: 1px solid; text-align: center'>{str_invoice_amt}</td>"
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; border: 1px solid; text-align: center'>{invoice_date}</td>"

                if not customer:
                    customer = True
                    html += """<td style="text-align:center; border: 1px solid" rowspan=" """ + str(
                                    len(self.invoice_line_ids)) + """">'N/A' 
                               </td>                         
                               """
                html += """
                           <td style="text-align:right; border: 1px solid">
                           """ + str(currency_symbol) + ' ' + str("{:,}".format(round(rec.swift_amt, 2))) + """
                           </td>
                           <td style="text-align:center; border: 1px solid">
                           """ + str(swift_message_date) + """
                           </td>                          
                       </tr>
                       """

        html += """
                 <tr>
                    <td style="text-align:center; font-weight:bold; border: 1px solid" > Total  </td>
                    <td style="text-align:center; font-weight:bold; border: 1px solid"> """ + str(all_invoice_qty_str) + """ </td>
                    <td style="border: 1px solid"> </td>
                    <td style="text-align:right; font-weight:bold; border: 1px solid"> """ + str(currency_symbol)+ ' ' + str("{:,}".format(round(inv_amount, 2))) + """ </td>
                    <td style="border: 1px solid"> </td>
                    <td style="border: 1px solid"> </td>
                    <td style="text-align:right; font-weight:bold; border: 1px solid"> """ + str(currency_symbol)+ ' ' + str("{:,}".format(round(swift_amt, 2))) + """ </td> 
                    <td style="border: 1px solid"> </td>
                </tr>
                    """

        expected_html = """
                    <table class="table table-bordered">
                        """ + html_thead + """
                        <tbody style='font-size:22px'>
                            """ + html + """
                        </tbody>
                    </table>
                    """

        # if len(self.invoice_line_ids) < 2:
        #     flbs_text = flbs_text.replace('$break_point', """<div style='margin-top:50px'></div>""")
        # else:
        #
        flbs_text = flbs_text.replace('$break_point', '<p style="page-break-after: always;"/>')

        if self.od_sight_rate:
            od_sight_rate = self.od_sight_rate
        else:
            if od_s_r:
                od_sight_rate = od_s_r
            else:
                od_sight_rate = encashment_rate_bdt

        rate_fc = swift_amt / self.incentive_rate_fc
        # encashment_final_amt = encashment_amt_bdt / self.incentive_rate_fc
        encashment_final_amt = round(rate_fc, 2) * od_sight_rate
        encashment_final_amt_a = self.custom_round(encashment_final_amt)
        # bdt_amount = rate_fc * self.encashment_rate_bdt
        usd_bdt_amount = currency + ' ' + str(round(rate_fc, 2)) + '\n' + '@' + str(
            round(od_sight_rate, 2)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a

        final_text = ''
        vortuki_swift_amt = 0
        if self.incentive_rate_fc > 0:
            vortuki_swift_amt = swift_amt / self.incentive_rate_fc

        # amount in word
        amount_in_word_str = ''
        if encashment_final_amt:
            # amount_in_word = num2words(round(encashment_final_amt))
            amount_in_word = self.num2words_fun(round(encashment_final_amt))
            # upper case function call
            amount_in_word_str = self.upper_case(amount_in_word)

        # dynamic variable replace
        if flbs_text:
            final_text = flbs_text.replace('$contract_price', str(self.contract_price_str) if self.contract_price_str else '').replace('$swift_amnt', str(currency_symbol)+ ' '+ str("{:,}".format(round(swift_amt, 2))))\
                .replace('$vortuki_amnt', str(currency_symbol)+ ' ' + str(round(vortuki_swift_amt, 2))).replace('$usd_bdt_amount', str(usd_bdt_amount)).replace(
                '$contract_number', str(self.contract_number) if self.contract_number else '').replace('$contract_date',str(self.contract_date_str) if self.contract_date_str else '')\
                .replace('$invoice_lines', expected_html).replace('$amount_in_word', amount_in_word_str + ' Only')
        self.form_ka_basis_description = final_text

    # upper case function
    def upper_case(self, ff):
        a = ff.split(' ')
        for i in range(len(a)):
            a[i] = a[i].capitalize()
        return ' '.join(a)

    # kha -----------------------
    def action_refresh_basis_kha(self):
        # --------------Form Kha
        if self.form_kha_ref_code == '' or self.form_kha_ref_code == False:
            form_kha_ref = ''
            if self.bank_id.form_kha_ref_prefix:
                form_kha_ref = str(self.bank_id.form_kha_ref_prefix).strip()
            if self.bank_id.form_kha_ref_suffix:
                if form_kha_ref:
                    form_kha_ref = form_kha_ref + str(self.bank_id.form_kha_ref_suffix)
                else:
                    form_kha_ref = str(self.bank_id.form_kha_ref_suffix).strip()
            self.form_kha_ref_code = form_kha_ref

        flbs_text = self.bank_id.form_kha_basis_description
        if flbs_text:
            flbs_text = self.bank_id.form_kha_basis_description.replace('<tr ', '<tr style="border: 1px solid"').replace(
                '<td ', '<td style="border: 1px solid"'). replace('<th ', '<th style="border: 1px solid"') \
                .replace('<tr>', '<tr style="border: 1px solid">').replace('<td>', '<td style="border: 1px solid">').replace('<th>', '<th style="border: 1px solid">')
        currency = ''
        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date),'%Y-%m-%d').strftime('%d/%m/%y')
        if self.fc_currency_id:
            currency = self.fc_currency_id.name

        html = ''
        html_thead = """
         <thead>
            <th style="border: 1px solid; text-align:center;margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">
                ক) ইনভয়েস নম্বর
            </th>
            <th style="border: 1px solid; text-align: center;margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">
                তারিখ
            </th>
            <th style="border: 1px solid;margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">
                খ) ইনভয়েসে উল্লেখিত সেবা/পণ্যের পরিমাণ (Qty & Hours)
            </th>
            <th style="text-align:right; border: 1px solid;margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">
                মূল্য (%s)
            </th>
        </thead>
            """%(currency)
        inv_number = ''
        inv_date = ''
        inv_amount = 0
        encashment_amt_bdt = 0
        total_qty = 0

        dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in self.invoice_line_ids.invoice_id.invoice_line_ids)
        hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in self.invoice_line_ids.invoice_id.invoice_line_ids)
        hour_total = float(format(hour_total1, '.2f'))
        dev_total = float(format(dev_total1, '.2f'))

        all_invoice_qty_str = ''
        qty_d_total=''
        qty_h_total=''
        if dev_total > 0:
            if dev_total>1:
                qty_d_total = str(dev_total) + str(' Developers')
            else:
                qty_d_total = str(dev_total) + str(' Developer')
            #all_invoice_qty_str += str(dev_total) + str(' Developers') if not all_invoice_qty_str else ', ' + str(dev_total) + str(' Developers')

        if hour_total > 0:
            if hour_total > 1:
                qty_h_total = str(hour_total) + str(' Hours')
            else:
                qty_h_total = str(hour_total) + str(' Hour')

            #all_invoice_qty_str += str(hour_total) + str(' Hours') if not all_invoice_qty_str else ', ' + str(hour_total) + str(' Hours')

        if qty_d_total:
            all_invoice_qty_str = qty_d_total
        if qty_h_total:
            if all_invoice_qty_str:
                all_invoice_qty_str = all_invoice_qty_str + ' & ' + qty_h_total
            else:
                all_invoice_qty_str = qty_h_total
        #--------------

        swift_id = None
        swift_amt = 0
        # Create a dictionary to group rows by invoice number
        invoice_groups = {}
        for rec in self.invoice_line_ids:
            if rec.invoice_id.ref in invoice_groups:
                invoice_groups[rec.invoice_id.ref].append(rec)
            else:
                invoice_groups[rec.invoice_id.ref] = [rec]

        # Generate HTML for merged rows
        for invoice_num, rows in invoice_groups.items():
            inv_obj = self.env['cash.incentive.invoice'].search(
                [('invoice_id.ref', '=', invoice_num), ('head_id', '=', self.id)], limit=1)
            html += "<tr>"
            # html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{invoice_num}</td>"
            inv_ids = []
            for i, rec in enumerate(rows):
                str_invoice_amt = str("{:,}".format(round(rec.invoice_amt, 2)))
                swift_amt += rec.swift_amt
                invoice_date = ''
                if rec.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
                if i == 0:
                    inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
                if i == 0:
                    inv_amount += rec.invoice_amt
                encashment_amt_bdt += rec.encashment_amt_bdt
                if not swift_id:
                    swift_id = rec.swift_message_id

                qty_h = ''
                qty_d = ''
                h_q = 0
                d_q = 0
                for l in rec.invoice_id.invoice_line_ids:
                    if l.quantity_type == '0':
                        h_q += l.quantity
                        total_qty += l.quantity
                    else:
                        d_q += l.quantity
                        total_qty += l.quantity

                if d_q:
                    if d_q > 1:
                        qty_d = str(d_q) + ' Developers'
                    else:
                        qty_d = str(d_q) + ' Developer'

                if h_q:
                    if h_q > 1:
                        qty_h = str(h_q) + ' Hours'
                    else:
                        qty_h = str(h_q) + ' Hour'

                invoice_qty_str = ''
                if qty_d:
                    invoice_qty_str = qty_d
                if qty_h:
                    if invoice_qty_str:
                        invoice_qty_str = invoice_qty_str + ' & ' + qty_h
                    else:
                        invoice_qty_str = qty_h
                if i != 0:
                    html += "<tr>"
                if i == 0:
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{invoice_num}</td>"
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{str(invoice_date)}</td>"
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; text-align:left; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{'Software Development / ' + invoice_qty_str}</td>"
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; text-align:right; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{currency + ' ' + str_invoice_amt}</td>"
                    html += "</tr>"

        # for rec in self.invoice_line_ids:
        #     swift_amt += rec.swift_amt
        #     invoice_date = ''
        #     if rec.invoice_date:
        #         invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
        #     inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
        #     inv_amount += rec.invoice_amt
        #     encashment_amt_bdt += rec.encashment_amt_bdt
        #     if not swift_id:
        #         swift_id = rec.swift_message_id
        #
        #     qty_h = ''
        #     qty_d = ''
        #     h_q = 0
        #     d_q = 0
        #     for l in rec.invoice_id.invoice_line_ids:
        #         if l.quantity_type == '0':
        #             h_q += l.quantity
        #             total_qty += l.quantity
        #         else:
        #             d_q += l.quantity
        #             total_qty += l.quantity
        #
        #     if d_q:
        #         if d_q > 1:
        #             qty_d = str(d_q) + ' Developers'
        #         else:
        #             qty_d = str(d_q) + ' Developer'
        #
        #     if h_q:
        #         if h_q > 1:
        #             qty_h = str(h_q) + ' Hours'
        #         else:
        #             qty_h = str(h_q) + ' Hour'
        #
        #     invoice_qty_str=''
        #     if qty_d:
        #         invoice_qty_str = qty_d
        #     if qty_h:
        #         if invoice_qty_str:
        #             invoice_qty_str = invoice_qty_str +' & '+ qty_h
        #         else:
        #             invoice_qty_str = qty_h
        #
        #     html += """
        #     <tr>
        #         <td style="border: 1px solid">
        #         """+str(rec.invoice_id.ref)+"""
        #         </td>
        #         <td style="border: 1px solid; text-align: center">
        #         """+str(invoice_date)+"""
        #         </td>
        #         <td style="border: 1px solid">Software Development /
        #         """+invoice_qty_str+"""
        #         </td>
        #         <td style="text-align:right; border: 1px solid">
        #         """+currency + ' ' + str("{:,}".format(round(rec.invoice_amt, 2)))+"""
        #         </td>
        #     </tr>
        #     """
        
        html += """
                </tr>
                 <tr style="font-weight:bold;">
                    <td style="text-align:left; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;" colspan = '2'> Total  </td>
                    <td style="text-align:left; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;">Software Development /  """+str(all_invoice_qty_str)+""" </td>
                    <td style="text-align:right; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;"> """+currency + ' ' + str("{:,.2f}".format(inv_amount))+""" </td> 
                </tr>
                    """
        expected_html = """
        <table class="table table-bordered">
            """+html_thead+"""
            <tbody>
                """+html+"""
            </tbody>
        </table>
        """

        final_text = ''
        customer = '' if not self.swift_customer_name else str(self.swift_customer_name) + ', '+ str(self.customer_address)
        country = ''
        if self.partner_id.country_id:
            country = self.partner_id.country_id.name

        remiter_address = ''
        if swift_id:
            if swift_id.remiter_bank_name:
                remiter_address += swift_id.remiter_bank_name
            if swift_id.remiter_bank_address:
                remiter_address += swift_id.remiter_bank_address if not remiter_address else ', ' + swift_id.remiter_bank_address
        prc_date = ''
        if self.prc_date:
            prc_date = datetime.datetime.strptime(str(self.prc_date), '%Y-%m-%d').strftime('%d/%m/%y')
        prc_ref_code = ''
        if self.prc_ref_code:
            prc_ref_code = self.prc_ref_code
        if flbs_text:
            final_text = flbs_text.replace('$contract_number', '' if not self.contract_number else str(self.contract_number))\
                .replace('$contract_date', '' if not self.contract_date_str else str(self.contract_date_str))\
                .replace('$contract_price', str(self.contract_price_str) if self.contract_price_str else '').replace('$customer_address', str(customer))\
                .replace('$ref_name', str(self.form_kha_ref_code) if self.form_kha_ref_code else '')\
                .replace('$date', str(self_date))\
                .replace('$quantity', str(all_invoice_qty_str)).replace('$invoice_lines', expected_html)\
                .replace('$inv_amount', currency+ ' ' + str("{:,}".format(round(inv_amount, 2)))).replace('$remiter_address',  str(remiter_address))\
                .replace('$swift_amount', currency+ ' ' + str("{:,}".format(round(swift_amt, 2))))\
                .replace('$inv_date', str(inv_date)).replace('$prc_num', str(prc_ref_code)).replace('$prc_date', str(prc_date))\
                .replace('$today_date', str(self_date)).replace('$country', str(country))
        self.form_kha_basis_description = final_text

    # c -------------------
    def action_refresh_form_c(self):
        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')
        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name
        client_name = ''
        customer_address = ''
        if self.partner_id:
            client_name = '' if not self.swift_customer_name else self.swift_customer_name

        customer_address = self.customer_address

        inv_number = ''
        encashment_amt_bdt_total = 0
        inv_amount = 0
        # #inv_date = ''
        swift_id = None
        for rec in self.invoice_line_ids:
            inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
            inv_amount += rec.invoice_amt
            encashment_amt_bdt_total += rec.encashment_amt_bdt
            swift_id = rec.swift_message_id

            #inv_date += str(rec.invoice_date) if not inv_date else ', ' + str(rec.invoice_date)
        remiter_address = ''
        if swift_id:
            if swift_id.remiter_bank_name:
                remiter_address += swift_id.remiter_bank_name
            if swift_id.remiter_bank_address:
                remiter_address += swift_id.remiter_bank_address if not remiter_address else ', ' + swift_id.remiter_bank_address
        description_text = self.bank_id.form_c_description
        final_text = ''
        if description_text:
            final_text = description_text.replace('$currency', str(currency)).replace('$total_amount', str(self.contract_price))\
                .replace('$client_name', str(client_name)).replace('$client_address', str(customer_address))\
                .replace('$invoice_number', str(inv_number)).replace('$bdt_amount', str("{:,}".format(round(inv_amount, 2))))\
                .replace('$date', str(self_date)).replace('$remiter_address',  str(remiter_address))\
                .replace('<tr ', '<tr style="border: 1px solid"').replace('<td ', '<td style="border: 1px solid"')\
                .replace('<th ', '<th style="border: 1px solid"') \
                .replace('<tr>', '<tr style="border: 1px solid">').replace('<td>', '<td style="border: 1px solid">')\
                .replace('<th>', '<th style="border: 1px solid">')
        self.form_c_description = final_text

    # forwarding -------------------
    def action_refresh_flbk(self):
        #self.flbs_letter_description = self.bank_id.flbs_letter_description
        if self.flbk_ref_code =='' or self.flbk_ref_code==False:
            flbk_ref = ''
            if self.bank_id.flbk_ref_prefix:
                flbk_ref = str(self.bank_id.flbk_ref_prefix).strip()
            if self.bank_id.flbk_ref_suffix:
                if flbk_ref:
                    flbk_ref = flbk_ref + str(self.bank_id.flbk_ref_suffix)
                else:
                    flbk_ref = str(self.bank_id.flbk_ref_suffix).strip()
            self.flbk_ref_code = flbk_ref

        flbs_text = self.bank_id.flbk_letter_description
        if flbs_text:
            flbs_text = self.bank_id.flbk_letter_description.replace('<tr ', '<tr style="border: 1px solid"').replace('<td ', '<td style="border: 1px solid"').\
                replace('<th ', '<th style="border: 1px solid"') \
                .replace('<tr>', '<tr style="border: 1px solid">').replace('<td>', '<td style="border: 1px solid">').replace('<th>', '<th style="border: 1px solid">')

        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name
        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

        html = ''
        html_thead = """
                 <thead>
                 <tr>
                    <th style='width: 67px; text-align:center; border: 1px solid'>
                        SL#
                    </th>
                    <th style="border: 1px solid;text-align:center">
                        Client Name
                    </th>
                    <th style="border: 1px solid; text-align: center">
                        Invoice No
                    </th>
                    <th style="text-align:right; border: 1px solid">
                        Invoice amount (%s)
                    </th>
                    <th style="border: 1px solid;text-align:center">
                        Invoice Date
                    </th>
                    <th style="border: 1px solid;text-align:center">
                        Swift/nostro date
                    </th>
                    <th style="text-align:right; border: 1px solid">
                        Realize Amount (%s)
                    </th>
                    <th style="text-align:right; border: 1px solid">
                        Incentive claim (%s)
                    </th>                    
                </tr>
                </thead>
                    """%(currency, currency, currency)

        inv_number = ''
        #inv_date = ''
        inv_amount = 0
        encashment_amt_bdt = 0
        total_incentive_rate_fc = 0
        sl=0
        swift_amt = 0
        customer = False
        cus_name = ''
        if self.swift_customer_name:
            cus_name = self.swift_customer_name

        # Create a dictionary to group rows by invoice number
        invoice_groups = {}
        for rec in self.invoice_line_ids:
            if rec.invoice_id.ref in invoice_groups:
                invoice_groups[rec.invoice_id.ref].append(rec)
            else:
                invoice_groups[rec.invoice_id.ref] = [rec]

        # Generate HTML for merged rows
        for invoice_num, rows in invoice_groups.items():
            inv_obj = self.env['cash.incentive.invoice'].search([('invoice_id.ref', '=', invoice_num), ('head_id', '=', self.id)], limit=1)
            # usd_price = sum(inv_obj.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
            usd_price = inv_obj.invoice_id.invoice_total_fc
            inv_amount += usd_price
            str_usd_price = str("{:,}".format(round(usd_price, 2)))
            html += "<tr>"
            # html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; text-align:center; border: 1px solid; margin-top:0px; margin-bottom:0px; padding-top:0px; padding-bottom:0px;'>{invoice_num}</td>"
            for i, rec in enumerate(rows):
                swift_amt += rec.swift_amt
                # incentive_rate_fc = usd_price / self.incentive_rate_fc if self.incentive_rate_fc else 0
                incentive_rate_fc = (rec.swift_amt * self.incentive_rate_fc) / 100 if self.incentive_rate_fc else 0
                total_incentive_rate_fc += incentive_rate_fc
                inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
                # inv_date += str(rec.invoice_date) if not inv_date else ', ' + str(rec.invoice_date)
                swift_date = self.swift_date if self.swift_date else ''
                swift_message_date = ''
                if rec.swift_message_id.date:
                    swift_message_date = datetime.datetime.strptime(str(rec.swift_message_id.date),
                                                                    '%Y-%m-%d').strftime('%d-%b-%y')
                invoice_date = ''
                if rec.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d-%b-%y')

                sl += 1
                html += """
                           <td style='vertical-align: middle;text-align:center; border: 1px solid'>
                               """ + str(sl) + """
                           </td>                         
                       """
                if not customer:
                    customer = True
                    html += """
                              <td style="vertical-align: middle;text-align:center;border: 1px solid" rowspan=" """ + str(len(self.invoice_line_ids)) + """"> """ + cus_name + """</td>                         
                               """
                if i == 0:
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; border: 1px solid; text-align: center'>{invoice_num}</td>"
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; border: 1px solid; text-align: right'>{str_usd_price}</td>"
                    html += f"<td rowspan='{len(rows)}' style='vertical-align: middle; border: 1px solid; text-align: center'>{invoice_date}</td>"

                # if i != 0:
                #     html += "<tr>"

                html += """
                        <td style="vertical-align: middle;text-align:center;border: 1px solid">
                        """ + str(swift_message_date) + """
                        </td>
                        <td style="vertical-align: middle;text-align:right; border: 1px solid">
                        """ + str("{:,.2f}".format(rec.swift_amt)) + """
                        </td>
                        <td style="vertical-align: middle;text-align:right; border: 1px solid">
                        """ + str("{:,}".format(float(format(incentive_rate_fc, '.2f')))) + """
                        </td>                            
                    </tr>
                    """
                html += "</tr>"
        html += """
                 <tr>
                    <td style=" border: 1px solid">  </td>
                    <td style=" border: 1px solid">  </td>
                    <td style="text-align:right; font-weight:bold; border: 1px solid"> Total  </td>
                    <td style="text-align:right; font-weight:bold; border: 1px solid"> """ + str("{:,.2f}".format(inv_amount)) + """ </td>
                    <td style="border: 1px solid"> </td> 
                    <td style="border: 1px solid"> </td>
                    <td style="text-align:right; font-weight:bold; border: 1px solid"> """ + str("{:,.2f}".format(swift_amt)) + """ </td>
                    <td style="text-align:right; font-weight:bold; border: 1px solid"> """ + str("{:,.2f}".format(total_incentive_rate_fc) ) + """ </td>
                </tr> 
                    """

        expected_html = """
            <table class="table table-bordered">
                """ + html_thead + """
                <tbody>
                    """ + html + """
                </tbody>
            </table>
            """

        final_text=''
        if flbs_text:
            final_text = flbs_text.replace('$ref_name', self.name).replace('$date', str(self_date)).replace('$invoice_lines', expected_html).replace('<ol>', "<ol style='margin-left:40px'>")
        self.flbk_letter_description = final_text

    # ga ---------------------------
    def action_refresh_form_ga(self):
        a = []
        c = []
        for rec in self.invoice_line_ids:
            a.append(rec.od_sight_rate)
            if rec.encashment_rate_bdt:
                c.append(rec.encashment_rate_bdt)
            # if rec.swift_message_id.encashment_rate_bdt:
            #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
        if a:
            b = min(a)
            od_s_r = b
        else:
            od_s_r = 0
        if c:
            d = min(c)
            encashment_rate_bdt = d
        else:
            encashment_rate_bdt = 0


        inv_amount = sum(self.invoice_line_ids.mapped('swift_amt'))
        encashment_amt_bdt = sum(self.invoice_line_ids.mapped('encashment_amt_bdt'))
        description_text = self.bank_id.form_ga_description
        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name

        if self.od_sight_rate:
            od_sight_rate = self.od_sight_rate
        else:
            if od_s_r:
                od_sight_rate = od_s_r
            else:
                od_sight_rate = encashment_rate_bdt

        rate_fc = inv_amount / self.incentive_rate_fc
        # encashment_final_amt = encashment_amt_bdt / self.incentive_rate_fc
        encashment_final_amt = round(rate_fc, 2) * od_sight_rate
        encashment_final_amt_a = self.custom_round(encashment_final_amt)
        # bdt_amount = rate_fc * self.encashment_rate_bdt
        usd_bdt_amount = currency + ' ' + str("{:,}".format(round(rate_fc, 2))) + '\n' + '@' + str("{:,}".format(round(od_sight_rate, 2))) + '=' + '\n' + 'BDT ' + encashment_final_amt_a
        # usd_bdt_amount = currency + str(sum(self.invoice_line_ids.mapped('invoice_amt'))) + '\n' + '@' + str(
        #     self.encashment_rate_bdt) + '=' + '\n' + 'BDT' + str(bdt_amount)
        final_text = ''
        if description_text:
            final_text = description_text.replace('$usd_amount', currency + ' ' + str("{:,}".format(round(inv_amount, 2)))).replace('$usd_bdt_amount',str(usd_bdt_amount)) \
                .replace('<tr ', '<tr style="border: 1px solid;vertical-align: middle"').replace('<td ', '<td style="border: 1px solid;vertical-align: middle"').replace('<th ', '<th style="border: 1px solid;vertical-align: middle"') \
                .replace('<tr>', '<tr style="border: 1px solid;vertical-align: middle">').replace('<td>', '<td style="border: 1px solid;vertical-align: middle">').replace('<th>', '<th style="border: 1px solid;vertical-align: middle">')
        self.form_ga_description = final_text

    # gha -------------------------
    def action_refresh_form_gha(self):
        a = []
        c = []
        for rec in self.invoice_line_ids:
            a.append(rec.od_sight_rate)
            if rec.encashment_rate_bdt:
                c.append(rec.encashment_rate_bdt)
            # if rec.swift_message_id.encashment_rate_bdt:
            #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
        if a:
            b = min(a)
            od_s_r = b
        else:
            od_s_r = 0

        if c:
            d = min(c)
            encashment_rate_bdt = d
        else:
            encashment_rate_bdt = 0

        inv_amount = sum(self.invoice_line_ids.mapped('swift_amt'))
        encashment_amt_bdt = sum(self.invoice_line_ids.mapped('encashment_amt_bdt'))
        description_text = self.bank_id.form_gha_description
        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name

        if self.od_sight_rate:
            od_sight_rate = self.od_sight_rate
        else:
            if od_s_r:
                od_sight_rate = od_s_r
            else:
                od_sight_rate = encashment_rate_bdt

        rate_fc = inv_amount / self.incentive_rate_fc
        # encashment_final_amt = encashment_amt_bdt / self.incentive_rate_fc
        encashment_final_amt = round(rate_fc, 2) * od_sight_rate
        encashment_final_amt_a = self.custom_round(encashment_final_amt)
        usd_bdt_amount = currency + ' ' + str("{:,}".format(round(rate_fc, 2))) + '\n' + '@' + str("{:,}".format(round(od_sight_rate, 2))) + '=' + '\n' + 'BDT ' + encashment_final_amt_a
        final_text = ''
        if description_text:
            final_text = description_text.replace('$usd_amount', currency + ' ' + str("{:,}".format(round(inv_amount, 2)))).replace('$usd_bdt_amount', str(usd_bdt_amount)) \
                .replace('<tr ', '<tr style="text-align: center;border: 1px solid;vertical-align: middle"').replace('<td ', '<td style="text-align: center;border: 1px solid;vertical-align: middle"')\
                .replace('<th ', '<th style="text-align: center;border: 1px solid;vertical-align: middle"') \
                .replace('<tr>', '<tr style="text-align: center;border: 1px solid;vertical-align: middle">')\
                .replace('<td>', '<td style="text-align: center;border: 1px solid;vertical-align: middle">')\
                .replace('<th>', '<th style="text-align: center;border: 1px solid;vertical-align: middle">')
        self.form_gha_description = final_text

    def action_print_with_head(self):
        type = self.env.context.get('type')
        with_head = self.env.context.get('with_head')
        data = {}
        result = []
        rpt_name = ''
        # if type == 'PRC':
        #     # result.append({'details': self.prc_letter_description})
        #     data['incentive_id'] = self.id
        #     rpt_name = 'PRC'
        # elif type == 'BASIS_KA':
        #     # result.append({'details': self.form_ka_basis_description})
        #     rpt_name = 'KA'
        # elif type == 'BASIS_KHA':
        #     # result.append({'details': self.form_kha_basis_description})
        #     rpt_name = 'KHA'
        # elif type == 'FLBS':
        #     # result.append({'details': self.flbs_letter_description})
        #     rpt_name = 'BASIS'
        # elif type == 'FLBK':
        #     # result.append({'details': self.flbk_letter_description})
        #     rpt_name = 'BANK'
        # elif type == 'FORM_C':
        #     # result.append({'details': self.form_c_description})
        #     rpt_name = 'ICT'
        # elif type == 'FORM_GA':
        #     # result.append({'details': self.form_ga_description})
        #     rpt_name = 'GA'
        # elif type == 'FORM_GHA':
        #     # result.append({'details': self.form_gha_description})
        #     rpt_name = 'GHA'
        data['ids'] = result
        data['with_head'] = with_head
        data['rpt_name'] = type
        data['incentive_id'] = self.id
        # print(self.env.user.company_id.id)
        data['company_idjh'] = self.env.user.company_id.id

        if type == 'FORM_GA':
            if with_head:
                return self.env.ref('cash_incentive.report_incentive_letter_w_head_'+type).with_context(
                    landscape=True).report_action(self, data=data)
            else:
                return self.env.ref('cash_incentive.report_incentive_letter_wo_head_'+type).with_context(
                    landscape=True).report_action(self, data=data)

        # if type in ('PRC', 'BASIS_KHA'):

        if with_head:
            return self.env.ref('cash_incentive.report_incentive_letter_w_head_'+type).with_context(
                landscape=False).report_action(self, data=data)
        else:
            return self.env.ref('cash_incentive.report_incentive_letter_wo_head_'+type).with_context(
                landscape=False).report_action(self, data=data)

        # else:
        #     if with_head:
        #         return self.env.ref('cash_incentive.report_incentive_letter_w_head').with_context(landscape=False).report_action(self, data=data)
        #     else:
        #         return self.env.ref('cash_incentive.report_incentive_letter_wo_head').with_context(landscape=False).report_action(self, data=data)

    def action_draft(self):
        self.state = 'draft'

    def action_confirm(self):
        # if len(self.invoice_line_ids) > 5:
        #     raise UserError(_('Only Five Invoices Allowed.'))
        for rec in self:
            rec.state = 'confirm'
            rec.code = self.env['ir.sequence'].get('cash_incentive_code')
            m_code = rec.bank_id.code_suffix
            if m_code>=0:
                rec.bank_id.code_suffix = m_code+1
            #------- PRC
            prc_code = rec.bank_id.prc_ref_suffix
            if prc_code >= 0:
                rec.bank_id.prc_ref_suffix = prc_code + 1
            # ------- Forwarding BASIS
            flbs_code = rec.bank_id.flbs_ref_suffix
            if flbs_code >= 0:
                rec.bank_id.flbs_ref_suffix = flbs_code + 1

            # ------- Forwarding BANK
            flbk_code = rec.bank_id.flbk_ref_suffix
            if flbk_code >= 0:
                rec.bank_id.flbk_ref_suffix = flbk_code + 1

            # ------- Forwarding form_kha
            form_kha_code = rec.bank_id.form_kha_ref_suffix
            if form_kha_code >= 0:
                rec.bank_id.form_kha_ref_suffix = form_kha_code + 1

    def action_approve(self):
        self.state = 'approve'

    def action_done(self):
        self.state = 'done'

    def action_cancel(self):
        for record in self:
            if record.state !='draft':
                raise UserError(_('Only Draft record can be cancelled!.'))
            else:
                self.state = 'cancel'

    @api.constrains('name')
    def _check_unique_name(self):
        for rec in self:
            envobj = self.env['cash.incentive.head']
            msg = 'Reference "%s"' % rec.name
            records = envobj.sudo().search([('id', '!=', rec.id), ('name', '=', rec.name)], limit=1)
            if records:
                raise exceptions.ValidationError("'" + msg + "' already exists!")

    @api.constrains('prc_ref_code')
    def _check_prc_ref_code(self):
        for rec in self:
            envobj = self.env['cash.incentive.head']
            msg = 'PRC Ref "%s"' % rec.prc_ref_code
            records = envobj.sudo().search([('id', '!=', rec.id), ('prc_ref_code', '=', rec.prc_ref_code)], limit=1)
            if records:
                raise exceptions.ValidationError("'" + msg + "' already exists!")

    def action_get_swift_files(self):
        self.swift_file_line_ids = None
        invoice_line1 = []
        swift_ids = []
        for inv_obj in self.invoice_line_ids:
            if inv_obj.swift_message_id.id not in swift_ids:
                moveLineData4 = {
                    'head_id': inv_obj.head_id.id,
                    'swift_id': inv_obj.swift_message_id.id,
                }
                invoice_line1.append((0, 0, moveLineData4))
                swift_ids.append(inv_obj.swift_message_id.id)
        self.swift_file_line_ids = invoice_line1

    def action_download_docx_file3(self):
        type = self.env.context.get('type')
        from docx import Document
        from htmldocx import HtmlToDocx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.text import WD_BREAK
        from docx.shared import Pt
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.section import WD_ORIENT

        from docx.shared import Inches
        # create a new document
        document = Document()
        # Set the default section properties to landscape orientation
        # section = document.sections[0]
        # new_width, new_height = section.page_height, section.page_width
        # section.orientation = WD_ORIENT.LANDSCAPE
        # section.page_width = new_width
        # section.page_height = new_height
        # get the section object and adjust the left and right margins
        from docx.shared import Cm
        section = document.sections[0]
        section.left_margin = Cm(1.5)  # 1.5 cm
        section.right_margin = Cm(1.5)  # 1.5 cm
        country = ''
        if self.partner_id.country_id:
            country = self.partner_id.country_id.name

        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name
        inv_number = ''
        inv_date = ''
        inv_amount = 0
        swift_amt = 0
        swift_id = None
        for rec in self.invoice_line_ids:
            swift_amt += rec.swift_amt
            invoice_date = ''
            if rec.invoice_date:
                invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
            inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
            inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
            inv_amount += rec.invoice_amt
            if not swift_id:
                swift_id = rec.swift_message_id

        remiter_address = ''
        if swift_id:
            if swift_id.remiter_bank_name:
                remiter_address += swift_id.remiter_bank_name
            if swift_id.remiter_bank_address:
                remiter_address += swift_id.remiter_bank_address if not remiter_address else ', ' + swift_id.remiter_bank_address

        file_name = ''
        if type == 'BASIS_KHA':
            self_date = ''
            if self.date:
                self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

            paragraph = document.add_paragraph()
            paragraph.add_run("(অনুচ্ছেদ ০৬ (ক), একই সার্কুলার নং-০৩/২০১৮ দ্রষ্টব্য).")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("ফরম-'খ'")

            paragraph = document.add_paragraph()
            paragraph.add_run("Ref: " + str(self.form_kha_ref_code))
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("Date: " + self_date)

            custom_style13 = document.styles.add_style('BanglaStyle', WD_STYLE_TYPE.PARAGRAPH)
            custom_style13.font.name = 'KacstPen'
            custom_style13.font.size = Pt(10)

            # Add some Bangla text using the custom style
            paragraph = document.add_paragraph('বাংলা টেক্সট', style='BanglaStyle')

            # document.styles['Normal'].font.size = Pt(10)
            #
            # custom_styler = document.styles.add_style('BanglaStyle', WD_STYLE_TYPE.PARAGRAPH)
            # custom_styler.font.name = 'Jamrul'
            # custom_styler.font.size = Pt(10)
            #
            # # Add some Bangla text using the custom style
            # document.add_paragraph('বাংলা টেক্সট', style='BanglaStyle')
            #
            # document.styles['Normal'].font.size = Pt(10)
            #
            # # add a new paragraph with Bangla text
            # p = document.add_paragraph('বাংলা টেক্সট')
            #
            # # set the font size for Bangla text
            # for run in p.runs:
            #     run.font.name = 'SolaimanLipi'
            #     run.font.size = Pt(20)

            # Add a custom font style


            # Add some text with the custom font

            # document.add_paragraph('Ref :' + self.name + '\nDate: '+ self_date + '\n' )
            # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

            # Define a custom style

            # custom_style2.font.size = Pt(13)
            document.styles['Normal'].font.size = Pt(9)
            # Add a paragraph and apply the custom style
            document.add_paragraph(
                'বেসিস প্রদেয় প্রত্যয়ন সনদপত্র \nবাংলাদেশ হতে সফটওয়্যার, আইটিইএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকি প্রাপ্তির প্রত্যয়ন সনদপত্র')

            customer = '' if not self.swift_customer_name else str(self.swift_customer_name) + ', ' + str(
                self.customer_address)
            document.add_paragraph(
                "১। আবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road,"
                " Mohakhali C/A, Dhaka-1212.       ইআরসি নম্বরঃ র ০১৪৩০৪৯ \n২। রপ্তানি ঋণপত্র/চুক্তপিত্র নম্বরঃ " + self.contract_number + '\t তারিখঃ' +
                self.contract_date_str + '\t মূল্যঃ ' + self.contract_price_str + '\n৩। বিদেশি ক্রেতার নাম ও ঠিকানাঃ" ' + customer +
                "\n৪। বিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ" + remiter_address + '\n' + '৫।')

            # paragraph3 = document.add_paragraph()
            # paragraph3.add_run("২। রপ্তানি ঋণপত্র/চুক্তপিত্র নম্বরঃ " + self.contract_number)
            # paragraph3.add_run("\t")
            # paragraph3.add_run("তারিখঃ " +self.contract_date_str)
            # paragraph3.add_run("\t")
            # paragraph3.add_run("মূল্যঃ " +self.contract_price_str)
            # paragraph3.add_run("\t")
            # paragraph3.style = custom_style2

            # customer = '' if not self.swift_customer_name else str(self.swift_customer_name) + ', ' + str(
            #     self.customer_address)
            # paragraph2 = document.add_paragraph("৩। বিদেশি ক্রেতার নাম ও ঠিকানাঃ" + customer)
            # paragraph2.style = custom_style2

            # remiter_address = ''
            # if swift_id:
            #     if swift_id.remiter_bank_name:
            #         remiter_address += swift_id.remiter_bank_name
            #     if swift_id.remiter_bank_address:
            #         remiter_address += swift_id.remiter_bank_address if not remiter_address else ', ' + swift_id.remiter_bank_address
            #
            # paragraph2 = document.add_paragraph("৪। বিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ" + remiter_address + '\n' + '৫।')
            # paragraph2.style = custom_style2

            # Set the paragraph alignment to center
            table = document.add_table(rows=1, cols=4)
            table.style = 'TableGrid'

            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('ক) ইনভয়েস নম্বর')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(0, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('তারিখ')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(0, 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('খ) ইনভয়েসে উল্লেখিত সেবা/পণ্যের পরিমাণ (Qty & Hours)')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(0, 3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('মূল্য (USD)')
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # add a dynamic row to the table
            cus_name = ''
            if self.swift_customer_name:
                cus_name = self.swift_customer_name

            # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
            total_qty = 0
            dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in
                             self.invoice_line_ids.invoice_id.invoice_line_ids)
            hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in
                              self.invoice_line_ids.invoice_id.invoice_line_ids)
            hour_total = float(format(hour_total1, '.2f'))
            dev_total = float(format(dev_total1, '.2f'))

            all_invoice_qty_str = ''
            if hour_total > 0:
                all_invoice_qty_str += str(hour_total) + str(' Hours') if not all_invoice_qty_str else ', ' + str(
                    hour_total) + str(' Hours')
            if dev_total > 0:
                all_invoice_qty_str += str(dev_total) + str(' Developers') if not all_invoice_qty_str else ', ' + str(
                    dev_total) + str(' Developers')
            row_len = len(self.invoice_line_ids)
            swift_amt = 0
            for rec in self.invoice_line_ids:
                swift_amt += rec.swift_amt
                invoice_date = ''
                if rec.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d-%b-%y')

                qty = ''
                h_q = 0
                d_q = 0
                for l in rec.invoice_id.invoice_line_ids:
                    if l.quantity_type == '0':
                        h_q += l.quantity
                        total_qty += l.quantity
                    else:
                        d_q += l.quantity
                        total_qty += l.quantity
                if h_q:
                    if not qty:
                        qty += str(h_q) + ' Hours'
                    else:
                        qty += ', ' + str(h_q) + ' Hours'
                if d_q:
                    if not qty:
                        qty += str(d_q) + ' Developers'
                    else:
                        qty += ', ' + str(d_q) + ' Developers'
                invoice_qty_str = qty

                # add a new row to the table
                new_row = table.add_row()

                # add data to the cells in the new row
                new_row.cells[0].text = rec.invoice_id.ref
                new_row.cells[1].text = invoice_date
                new_row.cells[2].text = invoice_qty_str
                new_row.cells[3].text = currency + ' ' + str("{:,}".format(round(rec.invoice_amt, 2)))

                first_cell = new_row.cells[0]
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_cell = new_row.cells[1]
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_cell = new_row.cells[2]
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_cell = new_row.cells[3]
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # total
            total_row = table.add_row()
            total_row.cells[1].text = 'Total'
            total_row.cells[2].text = all_invoice_qty_str
            total_row.cells[3].text = currency + ' ' + str("{:,}".format(round(inv_amount, 2)))

            first_cell = total_row.cells[1]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_cell = total_row.cells[2]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            first_cell = total_row.cells[3]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            self_date = ''
            if self.date:
                self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

            # footer -------------------
            document.add_paragraph(
                "৬। রপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত স্থানীয় সেবা/উপকরণাদরি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): 100% in house production of Brain Station 23 Limited \nপরিমাণঃ N/A  \t \t মূল্যঃ N/A"
                "\n৭। রপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত আমদানিকৃত আনুসাঙ্গিক সেবা/উপকরণাদরি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): N/A  \n পরিমাণঃ N/A   \t\t  মূল্যঃ  N/A"
                "\n৮। রপ্তানি সেবা/পণ্যের বিবরণঃ Software Development Services  \t\t পরিমাণঃ " + all_invoice_qty_str + " \t\t  মূল্যঃ " + currency + ' ' + str(
                    "{:,}".format(round(inv_amount, 2))) +
                "\n৯। জাহাজীকরণের তারখিঃ " + inv_date + " \t\t  গন্তব্য বন্দরঃ " + country + "\n১০। ইএক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানরি ক্ষেত্রে): N/A    মূল্যঃ N/A       তারিখঃ N/A"
                                                                                             "\n১১। োট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str(
                    "{:,}".format(
                        round(swift_amt, 2))) + " নীট এফওবি রপ্তানি মূল্য ( বৈদেশিক মুদ্রায়): " + currency + ' ' + str(
                    "{:,}".format(round(swift_amt, 2))) +
                "\n১২। প্রত্যাবসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ " + self.prc_ref_code + ' ' + str(
                    "{:,}".format(round(swift_amt, 2))) + " নতারিখঃ " + self_date)

            document.add_paragraph("রপ্তানিকারকের স্বাক্ষর ও তারিখ")

            paragraph2 = document.add_paragraph("রপ্তানিকারকের স্বাক্ষর ও তারিখ")

            paragraph2 = document.add_paragraph(
                "\nএতদ্বারা প্রত্যয়ন করা যাচ্ছে যে, আমাদের নিজস্ব কারখানায়/প্রতিষ্ঠানে তৈরীকৃত/উৎপাদিত সফটওয়্যার/আইটিইএস/হার্ডওয়্যার উপরোক্ত ৬ ও ৭ নং ক্রমিক বর্ণিত সূত্র হতে সেবা/উপকরাণাদি সংগ্রহের মাধ্যমে রপ্তানির বিপরীতে ভর্তুকির জন্য উপরোক্ত অনুচ্ছেদগুলোতে উল্লিখিত বক্তব্য সঠিক ও নির্ভুল। বিদেশী ক্রেতা/ আমদানিকারকের ক্রয়াদেশের যথার্থতা/বিশ্বাসযোগ্যতা সম্পর্কেও নিশ্চিত করা হলো। ")

            paragraph2 = document.add_paragraph("রপ্তানিকারকের স্বাক্ষর ও তারিখ")

            paragraph2 = document.add_paragraph(
                "\nরপ্তানিকারকের উপরোক্ত ঘোষণার যথার্থতা যাচাইয়ান্তে সঠিক পাওয়া গিয়েছে। ৮নং ক্রমিকে উল্লিখিত ঘোষিত রপ্তানিমূল্য যৌক্তিক ও বিদ্যমান আন্তর্জাতিক বাজার মুল্যের সঙ্গে সংগতিপূর্ণ পাওয়া গিয়েছে এবং বিদেশী ক্রেতার যথার্থতা/বিশ্বাসযোগ্যতা সর্ম্পকেও নিশ্চিত হওয়া গিয়েছে। প্রত্যাবাসিত রপ্তানি মূল্যের (নীট এফওবি মূল্য) রপ্তানি ভর্তুকি পরিশোধের সুপারিশ করা হলো।")

            paragraph2 = document.add_paragraph("এসোসিয়েশন এর দুইজন উপযুক্ত কর্মকর্তার স্বাক্ষর, তারিখ ও সীল")

            paragraph2 = document.add_paragraph(
                "[কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।]")


            file_name = '%s_kha' % (datetime.datetime.now())

        # -------------------
        import os
        dir_path = os.path.dirname(os.path.abspath(__file__))
        base_path = str(dir_path).replace('/models', '')
        docxfile = base_path + '/static/docx/' + file_name + '.docx'
        document.save(docxfile)
        # docx.save(docxfile)

        # return document.save('/home/jobaer/Downloads/jh3.docx')

        return {
            'type': 'ir.actions.act_url',
            'url': 'cash_incentive/static/docx/' + file_name + '.docx',
            'target': 'self',
        }

    # word -----------------------------------
    def action_download_docx_file1(self):
        type = self.env.context.get('type')
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        from docx.oxml.shared import OxmlElement, qn

        from docx.shared import Inches
        # create a new document
        document = Document()
        # get the section object and adjust the left and right margins
        from docx.shared import Cm
        section = document.sections[0]
        section.left_margin = Cm(1.5)  # 1.5 cm
        section.right_margin = Cm(1.5)  # 1.5 cm
        # jh = Document()
        # new_parser = HtmlToDocx()

        country = ''
        if self.partner_id.country_id:
            country = self.partner_id.country_id.name

        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name
        inv_number = ''
        inv_date = ''
        inv_amount = 0
        swift_amt = 0
        swift_id = None
        inv_ids = []
        for rec in self.invoice_line_ids:
            swift_amt += rec.swift_amt
            invoice_date = ''
            if rec.invoice_date:
                invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
            if rec.invoice_id not in inv_ids:
                inv_ids.append(rec.invoice_id)
                inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
                inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
                inv_amount += rec.invoice_amt

            if not swift_id:
                swift_id = rec.swift_message_id
        # print(inv_amount)
        client_name = ''
        if self.partner_id:
            client_name = '' if not self.swift_customer_name else self.swift_customer_name

        customer_address = self.customer_address

        remiter_address = ''
        if swift_id:
            if swift_id.remiter_bank_name:
                remiter_address += swift_id.remiter_bank_name
            if swift_id.remiter_bank_address:
                remiter_address += swift_id.remiter_bank_address if not remiter_address else ', ' + swift_id.remiter_bank_address

        ref_name_val = self.name
        split_values = ref_name_val.split("/")
        desired_value = split_values[-1]

        file_name = ''
        if type == 'PRC':
            self_date = ''
            if self.date:
                self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

            document.add_paragraph('\n\nRef: ' + self.name + '\t\t\t\t\t\t\t\t\t Format-A' + '\nDate: '+ self_date + '\n' )
            # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

            # Define a custom style
            custom_style = document.styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
            custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            custom_style2 = document.styles.add_style('MyStyle2', WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(11)
            custom_style2.font.name = 'Arial Narrow'
            custom_style2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            custom_style3 = document.styles.add_style('MyStyle3', WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.bold = True
            custom_style3.font.size = Pt(10)
            custom_style3.font.name = 'Arial Narrow'

            custom_style4 = document.styles.add_style('MyStyle4', WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.size = Pt(11)
            custom_style4.font.name = 'Arial Narrow'

            # Add a paragraph and apply the custom style
            paragraph1 = document.add_paragraph('CERTIFICATE OF AUTHORIZED DEALER')
            paragraph1.style = custom_style
            contract_number = self.contract_number
            if not self.contract_number:
                contract_number = ''
            contract_date_str = self.contract_date_str
            if not self.contract_date_str:
                contract_date_str = ''
            paragraph2 = document.add_paragraph("This is to certify that we have received following inward remittance in "
                                                "the name of Brain Station 23 Limited, Plot 02  (8th Floor), Bir Uttam A. K."
                                                " Khandakar Road, Mohakhali C/A, Dhaka-1212 against "
                                                + contract_number + ' & date: ' + str(contract_date_str) + " invoice no. "
                                                + inv_number + " & date:  " + inv_date +" amount in " +currency + ' ' +  str("{:,}".format(round(inv_amount, 2)))
                                                + " for rendering of software development service. Summary of the transaction is as follows: ")
            paragraph2.style = custom_style2
            # Set the paragraph alignment to center

            table = document.add_table(rows=2, cols=9)
            table.style = 'TableGrid'

            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(a) Remitter')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
            # cell.merge(table.cell(0, 1))

            cell = table.cell(0, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(b) Address')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            cell = table.cell(0, 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Invoice No.')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            cell = table.cell(0, 3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name = '(c) Amount in FC (' + self.fc_currency_id.name + ')'
            p3.add_run(cur_name)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            cell = table.cell(0, 4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(d) Date of Credit in banks nostro account')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            cell = table.cell(0, 5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(e) Equivalent Taka')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            cell = table.cell(0, 6)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(f) Credited to beneficiary a/c')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            cell = table.cell(0, 7)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(g) Reporting statement/schedule to BB with Month')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            cell = table.cell(0, 8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(h) Reference of Online reporting to BB')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            # add a dynamic row to the table
            cus_name = ''
            if self.swift_customer_name:
                cus_name = self.swift_customer_name

            # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
            row_len = len(self.invoice_line_ids)
            inv_amount = 0
            encashment_amt_bdt = 0
            equivalent_taka_total = 0

            from itertools import groupby

            # group the invoice_line_ids by invoice_id.ref
            groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                             key=lambda x: x.invoice_id.ref)
            # iterate over each group and add a new row for each unique invoice_id.ref
            row = 1
            swift_amt_total = 0
            for invoice_ref, group in groups:
                gr = 0
                for rec in group:
                    gr += 1
                    new_row = table.add_row()
                    inv_amount += rec.invoice_amt
                    swift_amt_total += rec.swift_amt
                    encashment_amt_bdt += rec.encashment_amt_bdt

                    swift_date = ''
                    if rec.swift_message_id.date:
                        swift_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime(
                            '%d-%b-%y')

                    date_credited_beneficiaries = ''
                    if rec.date_credited_beneficiaries:
                        date_credited_beneficiaries = datetime.datetime.strptime(str(rec.date_credited_beneficiaries),
                                                                                 '%Y-%m-%d').strftime('%d-%b-%y')
                    reporting_st_to_bb = rec.reporting_st_to_bb if rec.reporting_st_to_bb else ''
                    ref_online_to_bb = rec.ref_online_to_bb if rec.ref_online_to_bb else ''
                    equivalent_taka = round(rec.swift_amt * rec.encashment_rate_bdt, 2)
                    equivalent_taka_total += equivalent_taka

                    equivalent_taka_decimal = "{:,.2f}".format(equivalent_taka)

                    # add data to the cells in the new row
                    # new_row.cells[2].text = rec.invoice_id.ref
                    new_row.cells[3].text = str("{:,.2f}".format(rec.swift_amt))
                    new_row.cells[4].text = str(swift_date)
                    new_row.cells[5].text = str(equivalent_taka_decimal)
                    new_row.cells[6].text = str(date_credited_beneficiaries)
                    new_row.cells[7].text = str(reporting_st_to_bb)
                    new_row.cells[8].text = str(ref_online_to_bb)

                    first_cell = new_row.cells[3]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style4

                    first_cell1 = new_row.cells[5]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = custom_style4

                    first_cell1 = new_row.cells[2]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # paragraph22.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph22.style = custom_style4

                    first_cell1 = new_row.cells[0]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[1]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[4]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[6]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[7]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[8]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style4

                cell_1_0 = table.cell(row, 2)
                p3 = cell_1_0.paragraphs[0]
                p3.add_run(invoice_ref)
                cell_2_0 = table.cell((row + gr) - 1, 2)
                cell_1_0.merge(cell_2_0)
                row = gr + row

            # for rec in self.invoice_line_ids:
            #     inv_amount += rec.invoice_amt
            #     encashment_amt_bdt += rec.encashment_amt_bdt
            #
            #     swift_date = ''
            #     if rec.swift_message_id.date:
            #         swift_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime(
            #             '%d-%b-%y')
            #
            #     date_credited_beneficiaries = ''
            #     if rec.date_credited_beneficiaries:
            #         date_credited_beneficiaries = datetime.datetime.strptime(str(rec.date_credited_beneficiaries),
            #                                                                  '%Y-%m-%d').strftime('%d-%b-%y')
            #     reporting_st_to_bb = rec.reporting_st_to_bb if rec.reporting_st_to_bb else ''
            #     ref_online_to_bb = rec.ref_online_to_bb if rec.ref_online_to_bb else ''
            #     equivalent_taka = round(rec.swift_amt * rec.encashment_rate_bdt, 2)
            #     equivalent_taka_total += equivalent_taka
            #     # add a new row to the table
            #     new_row = table.add_row()
            #
            #     # add data to the cells in the new row
            #     new_row.cells[2].text = rec.invoice_id.ref
            #     new_row.cells[3].text = str("{:,}".format(round(rec.invoice_amt, 2)))
            #     new_row.cells[4].text = str(swift_date)
            #     new_row.cells[5].text = str("{:,}".format(round(equivalent_taka, 2)))
            #     new_row.cells[6].text = str(date_credited_beneficiaries)
            #     new_row.cells[7].text = str(reporting_st_to_bb)
            #     new_row.cells[8].text = str(ref_online_to_bb)
            #
            #     first_cell = new_row.cells[3]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph21 = first_cell.paragraphs[0]
            #     paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #     paragraph21.style = custom_style4
            #
            #     first_cell1 = new_row.cells[5]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #     paragraph22.style = custom_style4
            #
            #     first_cell1 = new_row.cells[2]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     # paragraph22.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            #     paragraph22.style = custom_style4
            #
            #     first_cell1 = new_row.cells[0]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.style = custom_style4
            #     first_cell1 = new_row.cells[1]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.style = custom_style4
            #     first_cell1 = new_row.cells[4]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style4
            #     first_cell1 = new_row.cells[6]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style4
            #     first_cell1 = new_row.cells[7]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style4
            #     first_cell1 = new_row.cells[8]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style4
            # from itertools import groupby
            #
            # # group the invoice_line_ids by invoice_id.ref
            # groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
            #                  key=lambda x: x.invoice_id.ref)
            # print(groups)
            # # iterate over each group and add a new row for each unique invoice_id.ref
            # row = 1
            # for invoice_ref, group in groups:
            #     new_row = table
            #     print(invoice_ref)
            #     print(group)
            #     # add data to the cells in the new row
            #     gr = 0
            #     for rec in group:
            #         gr += 1
            #         new_row = table.add_row()
            #         # add data to the cells in the new row for each row in the group
            #         new_row.cells[0].text = str("{:,}".format(round(rec.invoice_amt, 2)))
            #         new_row.cells[1].text = str("{:,}".format(round(rec.invoice_amt, 2)))
            #         new_row.cells[3].text = str("{:,}".format(round(rec.invoice_amt, 2)))
            #         new_row.cells[4].text = str(12)
            #         new_row.cells[5].text = str("{:,}".format(round(100, 2)))
            #         new_row.cells[6].text = str(200)
            #         new_row.cells[7].text = str(290)
            #         new_row.cells[8].text = str(12)
            #         # row += 1
            #     # new_row.cells[2].text = invoice_ref
            #     cell_1_0 = table.cell(row, 2)
            #     p3 = cell_1_0.paragraphs[0]
            #     p3.add_run(invoice_ref)
            #     cell_2_0 = table.cell((row + gr) - 1, 2)
            #     cell_1_0.merge(cell_2_0)
            #     row = gr + row

            # cus name merge
            cell_1_0 = table.cell(1, 0)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run(cus_name)
            cell_2_0 = table.cell(row_len, 0)
            cell_1_0.merge(cell_2_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            # address merge
            cell_3_0 = table.cell(1, 1)
            p3 = cell_3_0.paragraphs[0]
            p3.add_run(self.customer_address)
            cell_4_0 = table.cell(row_len, 1)
            cell_3_0.merge(cell_4_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            # total
            equivalent_taka_total_decimal = "{:,.0f}".format(equivalent_taka_total)
            total_row = table.add_row()
            total_row.cells[2].text = 'Total'
            total_row.cells[3].text = "{:,.2f}".format(swift_amt_total)
            total_row.cells[5].text = str(equivalent_taka_total_decimal)

            first_cell = total_row.cells[2]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = custom_style3

            first_cell = total_row.cells[3]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = custom_style3

            first_cell1 = total_row.cells[5]
            paragraph22 = first_cell1.paragraphs[0]
            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph22.style = custom_style3

            paragraph = document.add_paragraph()
            paragraph.add_run("\n\n\n\n\n Signature of Head of the branch.")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("Signature of the Issuing Officer.")

            file_name = 'PRC_%s_%s' % (desired_value, datetime.datetime.now())
            # file_name = '%s_prc' % (datetime.datetime.now())


            # html = self.get_word_prc()
            # html_text1 = html.replace('<br>', '<p style="display:block"></p>')
            # html_text1 = "<p></p><p></p><p></p><p></p> <div style='display:inline-flex'><p style='font-size:9px; margin-right:10px'>Signature of Head of the Branch</p> &nbsp;&nbsp;&nbsp; &nbsp;&nbsp;&nbsp;  <p>Signature of the Issuing Officer</p></div>"
            # new_parser.add_html_to_document(html_text1, document)

        if type == 'FLBS':
            section.page_width = Cm(22)  # set the page width to 21 centimeters
            section.page_height = Cm(29.9)  # set the page width to 21 centimeters
            section.left_margin = Cm(3)  # 1.5 cm
            section.right_margin = Cm(2)
            # Define a custom style
            custom_style = document.styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
            custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            custom_style.font.name = 'Arial Narrow'

            custom_style2 = document.styles.add_style('MyStyle2', WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.name = 'Arial Narrow'
            custom_style2.font.size = Pt(11)

            custom_style3 = document.styles.add_style('MyStyle3', WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(11)
            custom_style3.font.name = 'Arial Narrow'
            custom_style3.font.bold = True

            # Add a paragraph and apply the custom style
            # paragraph1 = document.add_paragraph('CERTIFICATE OF AUTHORIZED DEALER \n ')
            # paragraph1.style = custom_style

            # paragraph2 = document.add_paragraph("Ref No.: " + self.name + " \nDate: "+ self_date + '\n\n")
            current_date = datetime.datetime.now().date()

            self_date = ''
            if current_date:
                self_date = datetime.datetime.strptime(str(current_date), '%Y-%m-%d').strftime('%d/%m/%y')
            paragraph2 = document.add_paragraph('\n\n\n\nRef: ' + self.name + '\nDate: ' + str(self_date) + '\n')
            paragraph2.style = custom_style2

            paragraph4 = document.add_paragraph("The President \nBASIS \nBDBL Bhabon (5th Floor-West), \n12 Kawran Bazar, Dhaka-1215 \n")
            paragraph4.style = custom_style2

            paragraph5 = document.add_paragraph("Subject: Request to Issue BASIS Certificate for Cash Assistance.\n")
            paragraph5.style = custom_style3

            paragraph6 = document.add_paragraph("Dear Sir,")
            paragraph6.style = custom_style2
            paragraph6.paragraph_format.space_after = Pt(0)

            contract_number = ''
            if self.contract_number:
                contract_number = " and "+ self.contract_number

            paragraph6 = document.add_paragraph("With reference to the above-mentioned subject, we would like to draw your kind attention to the fact that we are going to draw cash subsidy against our following Invoice No: " + inv_number + contract_number + ".\n")
            paragraph6.style = custom_style2
            paragraph6.style = custom_style2

            paragraph6 = document.add_paragraph("There is a checklist as following:  ")
            paragraph6.style = custom_style2

            # Create a list of items
            items = ["Forwarding Letter Addressing BASIS president", "Export Agreement", "Commercial Invoice Related to Export", "Format -A Certificate of Authorized Dealer Issued by Bank",
                     "ICT form-c", "Company's Undertaking", "Form-Ka (Bangladesh Bank)", "Form-Kha (Bangladesh Bank)", "Pay order"]

            # Create a paragraph object for each list item and set the paragraph style to "List Bullet"
            # for item in items:
            #     paragraph7 = document.add_paragraph(style="List Bullet")
            #     paragraph7.add_run(item)

            left_indent = Inches(.8)  # Adjust the left indentation as needed

            for item in items:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.left_indent = left_indent

                run = paragraph.add_run(item)
                run.font.name = 'Arial Narrow'
                paragraph.paragraph_format.line_spacing = Pt(12)

            paragraph8 = document.add_paragraph("\nYour kind co-operation in this regard will be highly appreciated.\n\nThanking You.\n\nYours Truly\n\n\n\n\nMasudur Rahman \nSenior Manager, Finance, Accounts & Admin \nBrain Station 23 Ltd. \nPlot-2, (5th Floor), Bir Uttam Ak Khandaker Road \nMohakhali C/A, Dhaka-1212")
            paragraph8.style = custom_style2
            paragraph8.paragraph_format.space_before = Pt(0)
            paragraph8.paragraph_format.space_after = Pt(0)

            file_name = 'Forwarding_Letter_BASIS_%s_%s' % (desired_value, datetime.datetime.now())
            # file_name = '%s_flbs' % (datetime.datetime.now())

            # html_text1 = self.flbs_letter_description.replace('<br>', '<p style="display:block"></p>')

        if type == 'BASIS_KA':
            # section.page_height = Cm(30)  # set the page width to 21 centimeters
            # section.left_margin = Cm(2)  # 1.5 cm
            # section.right_margin = Cm(1)

            section.page_width = Cm(22)  # set the page width to 21 centimeters
            section.page_height = Cm(29.9)  # set the page width to 21 centimeters
            section.left_margin = Cm(2)  # 1.5 cm
            section.right_margin = Cm(2)

            self_date = ''
            if self.date:
                self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')
            # document.add_paragraph('Ref :' + self.name + '\nDate: '+ self_date + '\n' )
            # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)
            # bangla font maker =========================================== left align
            my_stylel = document.styles['Normal']
            # define your desired fonts
            user_cs_font_size = 7
            user_cs_font_name = 'Shonar Bangla'
            rpr = my_stylel.element.get_or_add_rPr()
            if rpr.rFonts is None:
                rpr._add_rFonts()
            if rpr.sz is None:
                rpr._add_sz()
            w_nsmap = '{' + rpr.nsmap['w'] + '}'
            szCs = None
            lang = None

            for element in rpr:
                if element.tag == w_nsmap + 'szCs':
                    szCs = element
                elif element.tag == w_nsmap + 'lang':
                    lang = element

            if szCs is None:
                szCs = rpr.makeelement(w_nsmap + 'szCs', nsmap=rpr.nsmap)
            if lang is None:
                lang = rpr.makeelement(w_nsmap + 'lang', nsmap=rpr.nsmap)

            rpr.append(szCs)
            rpr.append(lang)

            szCs_attrib = szCs.attrib
            lang_attrib = lang.attrib
            rFonts_atr = rpr.rFonts.attrib
            szCs_attrib[w_nsmap + 'val'] = str(int(user_cs_font_size * 2))
            rFonts_atr[w_nsmap + 'cs'] = user_cs_font_name
            lang_attrib[w_nsmap + 'bidi'] = user_cs_font_name  # For Persian

            # bangla font maker =========================================== underline, bold and center align
            my_styleu = document.styles.add_style('MyStyles21', WD_STYLE_TYPE.PARAGRAPH)
            my_styleu.base_style = my_stylel
            # define your desired fonts
            user_cs_font_size = 9
            user_cs_font_name = 'SutonnyOMJ'
            rpr = my_styleu.element.get_or_add_rPr()
            if rpr.rFonts is None:
                rpr._add_rFonts()
            if rpr.sz is None:
                rpr._add_sz()
            w_nsmap = '{' + rpr.nsmap['w'] + '}'
            szCs = None
            lang = None

            for element in rpr:
                if element.tag == w_nsmap + 'szCs':
                    szCs = element
                elif element.tag == w_nsmap + 'lang':
                    lang = element

            if szCs is None:
                szCs = rpr.makeelement(w_nsmap + 'szCs', nsmap=rpr.nsmap)
            if lang is None:
                lang = rpr.makeelement(w_nsmap + 'lang', nsmap=rpr.nsmap)

            rpr.append(szCs)
            rpr.append(lang)

            szCs_attrib = szCs.attrib
            lang_attrib = lang.attrib
            rFonts_atr = rpr.rFonts.attrib
            szCs_attrib[w_nsmap + 'val'] = str(int(user_cs_font_size * 2))
            rFonts_atr[w_nsmap + 'cs'] = user_cs_font_name
            # lang_attrib[w_nsmap + 'bidi'] = 'fa-IR'  # For Persian
            bCs = OxmlElement('w:bCs')
            rpr.append(bCs)
            bCs.set(qn('w:val'), "True")
            my_styleu.font.underline = True

            # bangla font maker =========================================== bold and center align
            my_styleub = document.styles.add_style('MyStyles211', WD_STYLE_TYPE.PARAGRAPH)
            my_styleub.base_style = my_stylel
            # define your desired fonts
            user_cs_font_size = 9
            user_cs_font_name = 'SolaimanLipi'
            rpr = my_styleub.element.get_or_add_rPr()
            if rpr.rFonts is None:
                rpr._add_rFonts()
            if rpr.sz is None:
                rpr._add_sz()
            w_nsmap = '{' + rpr.nsmap['w'] + '}'
            szCs = None
            lang = None

            for element in rpr:
                if element.tag == w_nsmap + 'szCs':
                    szCs = element
                elif element.tag == w_nsmap + 'lang':
                    lang = element

            if szCs is None:
                szCs = rpr.makeelement(w_nsmap + 'szCs', nsmap=rpr.nsmap)
            if lang is None:
                lang = rpr.makeelement(w_nsmap + 'lang', nsmap=rpr.nsmap)

            rpr.append(szCs)
            rpr.append(lang)

            szCs_attrib = szCs.attrib
            lang_attrib = lang.attrib
            rFonts_atr = rpr.rFonts.attrib
            szCs_attrib[w_nsmap + 'val'] = str(int(user_cs_font_size * 2))
            rFonts_atr[w_nsmap + 'cs'] = user_cs_font_name
            # lang_attrib[w_nsmap + 'bidi'] = 'fa-IR'  # For Persian
            bCs = OxmlElement('w:bCs')
            rpr.append(bCs)
            bCs.set(qn('w:val'), "True")

            # end bangla font maker ====================================
            # Define a custom style
            custom_style = document.styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
            # custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            custom_style2 = document.styles.add_style('MyStyle2', WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(11)

            custom_style3 = document.styles.add_style('MyStyle3', WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(11)
            custom_style3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            custom_style4 = document.styles.add_style('MyStyle4', WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.bold = True
            custom_style4.font.size = Pt(9)
            custom_style4.font.name = 'Arial Narrow'

            tab_total = document.styles.add_style('tab_total', WD_STYLE_TYPE.PARAGRAPH)
            tab_total.font.bold = True
            tab_total.font.size = Pt(9)
            tab_total.font.name = 'Arial Narrow'

            custom_styleextra = document.styles.add_style('custom_styleextra', WD_STYLE_TYPE.PARAGRAPH)
            custom_styleextra.font.size = Pt(9)
            custom_styleextra.font.name = 'Arial Narrow'

            custom_style5 = document.styles.add_style('MyStyle5', WD_STYLE_TYPE.PARAGRAPH)
            custom_style5.font.size = Pt(9)
            custom_style5.font.name = 'Arial Narrow'

            custom_stylebl = document.styles.add_style('custom_stylebl', WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebl.font.size = Pt(9)
            custom_stylebl.font.name = 'SutonnyOMJ'

            custom_na = document.styles.add_style('custom_na', WD_STYLE_TYPE.PARAGRAPH)
            custom_na.font.size = Pt(9)
            custom_na.font.name = 'Arial Narrow'

            custom_stylebold = document.styles.add_style('custom_stylebold', WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebold.font.size = Pt(11)
            custom_stylebold.font.name = 'SutonnyOMJ'
            custom_stylebold.font.bold = True

            custom_styleubl = document.styles.add_style('custom_styleubl', WD_STYLE_TYPE.PARAGRAPH)
            custom_styleubl.font.size = Pt(11)
            custom_styleubl.font.name = 'SutonnyOMJ'
            custom_styleubl.font.underline = True
            custom_styleubl.font.bold = True

            document.styles['Normal'].font.size = Pt(8)
            paragraph = document.add_paragraph()
            paragraph.add_run("অনুচ্ছেদ ০৫(খ) এফই সার্কুলার নং-০৩/২০১৮ দ্রষ্টব্য)")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run(' ' * 7 + "ফরম-‘ক’")
            paragraph.style = custom_stylebl

            # Add a paragraph and apply the custom style
            paragraph1 = document.add_paragraph('বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও')
            paragraph1.style = custom_stylebold
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph1.paragraph_format.space_before = Pt(0)
            paragraph1.paragraph_format.space_after = Pt(0)
            paragraph1 = document.add_paragraph('হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকির জন্য আবেদনপত্র ')
            paragraph1.style = custom_styleubl
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph1.paragraph_format.space_before = Pt(0)

            paragraph2 = document.add_paragraph("(ক) আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানাঃ ব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী বাণিজ্যিক এলাকা, ঢাকা - ১২১২। রপ্তানি নিবন্ধন সনদপত্র (ইআরসি) নম্বরঃ ২৬০৩২৬২১০৬৬৬৪২০")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)

            con_no = ''
            if self.contract_number:
                con_no = self.contract_number.replace('Contract No.', '').replace('Purchase Order No.', '')
            paragraph2 = document.add_paragraph("(খ) রপ্তানি চুক্তিপত্র নম্বরঃ"+ con_no + " \nতারিখঃ " + str(self.contract_date_str) + "\nমূল্যঃ " +self.contract_price_str + "\n(পাঠ্যযোগ সত্যায়িত কপি দাখিল করতে হবে)" )
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)
            paragraph2 = document.add_paragraph("(গ) রপ্তানিকৃত সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের স্থানীয় সংগ্রহসূত্র, পরিমাণ ও মূল্যঃ" )
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)
            # Set the paragraph alignment to center

            # table1 ------------------------------
            table5 = document.add_table(rows=3, cols=3)
            table5.alignment = 2
            table5.left_indent = Inches(100)
            table5.style = 'TableGrid'
            table5.autofit = False  # Disable automatic column width adjustment
              # Set the width of the table to 6 inches (adjust the value according to your desired width)

            cell = table5.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('সরবরাহকারীর নাম ও ঠিকানা')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_stylebl

            cell = table5.cell(0, 1)
            p1 = cell.paragraphs[0]
            p1.add_run('পরিমাণ')
            p1.style = custom_stylebl
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table5.cell(0, 2)
            p2 = cell.paragraphs[0]
            p2.add_run('মূল্য')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table5.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table5.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table5.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table5.cell(2, 0)
            p7 = cell.paragraphs[0]
            p7.add_run('N/A')
            p7.style = custom_na
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table5.cell(2, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table5.cell(2, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('N/A')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # table 2 ----------------------------
            paragraph2 = document.add_paragraph("(রপ্তানিকৃত সেবা/পণ্যের বর্ণনা, মূল্য ও সংগ্রহসূত্রের বিষয়ে সেবা/পণ্য সংশ্লিষ্ট এসোসিয়েশন এর প্রত্যয়নপত্র দাখিল করতে হবে) \n(ঘ) রপ্তানিকৃত সেবা/পণ্য উৎপাদনে ব্যবহৃত আমাদানিকৃত সেবা/উপকরণাদিঃ" )
            paragraph2.style = custom_stylebl
            # paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(12)

            # Set the paragraph alignment to center
            table1 = document.add_table(rows=3, cols=4)
            table1.autofit = False
            table1.width = Inches(6)
            table1.alignment = 2
            table1.style = 'TableGrid'

            cell = table1.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('সরবরাহকারীর নাম ও ঠিকানা')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table1.columns[0].width = Inches(.6)
            p3.style = custom_stylebl

            cell = table1.cell(0, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('সেবা/পণ্যের নাম ও পরিমাণ ')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table1.columns[1].width = Inches(.6)
            p3.style = custom_stylebl

            cell = table1.cell(0, 2)
            p2 = cell.paragraphs[0]
            p2.add_run('ঋণপত্র/ব্যাংক টু ব্যাংক ঋণপত্র/ডকুমেন্টরী কালেকশন/টিটি রেমিটেন্স নম্বর, তারিখ')
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table1.columns[2].width = Inches(5.2)
            p2.style = custom_stylebl

            cell = table1.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('মূল্য')
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table1.columns[3].width = Inches(1)
            p2.style = custom_stylebl

            cell = table1.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table1.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table1.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table1.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('৪')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table1.cell(2, 0)
            p7 = cell.paragraphs[0]
            p7.add_run('N/A')
            p7.style = custom_na
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table1.cell(2, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table1.cell(2, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('N/A')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table1.cell(2, 3)
            p9 = cell.paragraphs[0]
            p9.add_run('N/A')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # table 3 ----------------------
            paragraph2 = document.add_paragraph("(৩) নং কলামের ঋণপত্রে পাঠযোগ্য সত্যায়িত কপি দাখিল করতে হবে। সেবা আমাদানির ক্ষেত্রে যথাযথ পদ্ধতি অনুসরণ করা হয়েছে মর্মে অনুমোদিত ডিলার শাখাকে নিশ্চিত হতে হবে। উৎপাদন প্রক্রিয়ায় ব্যবহৃত উপকরণাদির জন্য শুল্ক বন্ড সুবিধা ভোগ করা হয়নি/ডিউটি ড্র-ব্যাংক সুবিধা গ্রহণ করা হয়নি ও ভবিষ্যতে আবেদনও করা হবে না মর্মে রপ্তানিকরাকের ঘোষণাপত্র দাখিল করতে হবে।) \n(ঙ) রপ্তানি চালানের বিবরণঃ" )
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(12)

            table = document.add_table(rows=3, cols=8)
            table.alignment = 2  # Left alignment
            # table.left_indent = Inches(5)
            table.style = 'TableGrid'
            # table.autofit = False
            # table.width = Inches(6) it is worked...............
            # table.autofit = False
            # table.width = Inches(4)
            # table.alignment = 1

            column_width = Inches(.2)
            table.columns[0].width = column_width
            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('   পণ্যের বর্ণনা   ')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl
            # table.columns[0].width = Inches(.3)
            # cell.merge(table.cell(0, 1))

            column_width = Inches(1.8)
            table.columns[1].width = column_width
            cell = table.cell(0, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('পরিমাণ')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl
            # table.columns[1].width = Inches(1.69)

            column_width = Inches(.1)
            table.columns[2].width = column_width
            cell = table.cell(0, 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('আমদানিকারকের দেশের নাম')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl
            # table.columns[2].width = Inches(.8)

            table.columns[3].width = Inches(.6)
            cell = table.cell(0, 3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('ইনভয়েস মূল্য (বৈদেশিক মুদ্রায়)')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl

            table.columns[4].width = Inches(.6)
            cell = table.cell(0, 4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('জাহাজীকরণ/রপ্তানির তারিখ')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl

            column_width = Inches(.2)
            table.columns[5].width = column_width
            cell = table.cell(0, 5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('   ইএক্সপি নম্বর*   ')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl

            table.columns[6].width = Inches(.8)
            table.columns[7].width = Inches(.3)
            cell_1_0 = table.cell(0, 6)
            cell_1_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell_1_0.paragraphs[0]
            p10.add_run('বৈদেশিক মুদ্রায় প্রত্যাবাসিত রপ্তানিমূল্য ও প্রত্যাবাসনের তারিখ')
            cell_2_0 = table.cell(0, 7)
            cell_1_0.merge(cell_2_0)
            p10.style = custom_stylebl
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER

            left_indent = Inches(2)  # Adjust the indent as needed

            for row in table.rows:
                for cell in row.cells:
                    cell.left_indent = left_indent

            # cus name merge
            # cell_1_0 = table.cell(0, 6)
            # p10 = cell_1_0.paragraphs[0]
            # p10.add_run('ৈদেশিক মুদ্রায় প্রত্যাবাসিত রপ্তানিমূল্য ও প্রত্যাবাসনের তারিখ')
            # cell_2_0 = table.cell(0, 7)
            # cell_1_0.merge(cell_2_0)

            # row 2 ----------------------------
            cell = table.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('৪')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 4)
            p6 = cell.paragraphs[0]
            p6.add_run('৫')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 5)
            p6 = cell.paragraphs[0]
            p6.add_run('৬')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 6)
            p10 = cell_1_1.paragraphs[0]
            p10.add_run('৭')
            p6.style = custom_stylebl
            cell_2_1 = table.cell(1, 7)
            cell_1_1.merge(cell_2_1)
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # cell_1_1 = table.cell(1, 6)
            # p10 = cell_1_1.paragraphs[0]
            # p10.add_run('৭')
            # cell_2_1 = table.cell(1, 7)
            # cell_1_1.merge(cell_2_1)

            # # add a dynamic row to the table
            # inv_date = ''
            inv_amount = 0
            incentive_rate_fc = 0
            swift_amt = 0
            incentive_amt_fc = 0
            encashment_amt_bdt = 0
            total_incentive_rate_fc = 0
            sl = 0

            inv_ids = []
            dev_total1 = 0
            hour_total1 = 0
            a = []
            c = []
            for line in self.invoice_line_ids:
                a.append(line.od_sight_rate)
                c.append(line.encashment_rate_bdt)
                if line.invoice_id.id not in inv_ids:
                    dev_total1 += sum(
                        r.quantity if r.quantity_type == '1' else 0 for r in line.invoice_id.invoice_line_ids)
                    hour_total1 += sum(
                        r.quantity if r.quantity_type == '0' else 0 for r in line.invoice_id.invoice_line_ids)
                inv_ids.append(line.invoice_id.id)
            
            if a:
                b = min(a)
                od_s_r = b
            else:
                od_s_r = 0
            
            if c:
                d = min(c)
                encashment_rate_bdt = d
            else:
                encashment_rate_bdt = 0
                
            # dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in
            #                  self.invoice_line_ids.invoice_id.invoice_line_ids)
            # hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in
            #                   self.invoice_line_ids.invoice_id.invoice_line_ids)
            hour_total = float(format(hour_total1, '.2f'))
            dev_total = float(format(dev_total1, '.2f'))

            if hour_total == int(hour_total):
                hour_total_num = "{:.0f}".format(hour_total)
            else:
                hour_total_num = "{:.2f}".format(hour_total)

            if dev_total == int(dev_total):
                dev_total_num = "{:.0f}".format(dev_total)
            else:
                dev_total_num = "{:.2f}".format(dev_total)

            all_invoice_qty_str = ''
            if dev_total > 0:
                all_invoice_qty_str += str('ITES/ ') + str(dev_total_num) if not all_invoice_qty_str else ' & ' + str(
                    'ITES/ ') + str(
                    dev_total_num)
            if hour_total > 0:
                all_invoice_qty_str += str('ITES/ ') + str(hour_total_num) + str(' HRS') if not all_invoice_qty_str else ' & ' + str(
                    hour_total_num) + str(' HRS')

            currency_symbol = ''
            if self.fc_currency_id.symbol:
                currency_symbol = self.fc_currency_id.symbol
            country_name = ''
            if self.partner_id.country_id:
                country_name = self.partner_id.country_id.name
            row_len = len(self.invoice_line_ids)

            from itertools import groupby
            # group the invoice_line_ids by invoice_id.ref
            groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                             key=lambda x: x.invoice_id.ref)
            # iterate over each group and add a new row for each unique invoice_id.ref
            row = 2

            for invoice_ref, group in groups:
                inv_obj = self.env['cash.incentive.invoice'].search(
                    [('invoice_id.ref', '=', invoice_ref), ('head_id', '=', self.id)], limit=1)
                # usd_price = sum(inv_obj.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
                usd_price = inv_obj.invoice_id.invoice_total_fc
                incentive_rate_fc = usd_price / self.incentive_rate_fc if self.incentive_rate_fc else 0
                total_incentive_rate_fc += incentive_rate_fc
                inv_amount += inv_obj.invoice_amt
                encashment_amt_bdt += inv_obj.encashment_amt_bdt
                invoice_date = ''
                swift_message_date = ''
                if inv_obj.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime(
                        '%d-%b-%y')
                if inv_obj.swift_message_id.date:
                    swift_message_date = datetime.datetime.strptime(str(inv_obj.swift_message_id.date),
                                                                    '%Y-%m-%d').strftime('%d-%b-%y')

                # getting hours and developers ===========
                l_dev_total = sum(
                    r.quantity if r.quantity_type == '1' else 0 for r in inv_obj.invoice_id.invoice_line_ids)
                l_hour_total = sum(
                    r.quantity if r.quantity_type == '0' else 0 for r in inv_obj.invoice_id.invoice_line_ids)
                l_all_invoice_qty_str = ''
                if l_dev_total > 0:
                    if l_dev_total == int(l_dev_total):
                        l_dev_total = "{:.0f}".format(l_dev_total)
                    else:
                        l_dev_total = "{:.2f}".format(l_dev_total)
                    l_all_invoice_qty_str += str('ITES/ ') + str(
                        l_dev_total) if not l_all_invoice_qty_str else ' & ' + str(
                        'ITES/ ') + str(l_dev_total)
                if l_hour_total > 0:
                    if l_hour_total == int(l_hour_total):
                        l_hour_total = "{:.0f}".format(l_hour_total)
                    else:
                        l_hour_total = "{:.2f}".format(l_hour_total)
                    l_all_invoice_qty_str += str('ITES/ ') + str(l_hour_total) + str(
                        ' HRS') if not l_all_invoice_qty_str else ' & ' + str(
                        l_hour_total) + str(' HRS')
                sl += 1
                gr = 0
                for rec in group:
                    swift_amt += rec.swift_amt
                    incentive_amt_fc += rec.incentive_amt_fc
                    # if rec.swift_message_id.encashment_rate_bdt:
                    #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
                    gr += 1
                    new_row = table.add_row()
                    new_row.cells[6].text = str(currency_symbol) + ' ' + str("{:,.2f}".format(rec.swift_amt))
                    new_row.cells[7].text = str(swift_message_date)

                    first_cell = new_row.cells[3]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style5

                    first_cell1 = new_row.cells[6]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = custom_style5

                    first_cell1 = new_row.cells[0]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[1]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[2]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[4]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[5]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[7]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5

                cell1_1_0 = table.cell(row, 1)
                p3 = cell1_1_0.paragraphs[0]
                p3.add_run(l_all_invoice_qty_str)
                cell1_2_0 = table.cell((row + gr) - 1, 1)
                cell1_1_0.merge(cell1_2_0)

                cell_4_0 = table.cell(row, 3)
                p3 = cell_4_0.paragraphs[0]
                p3.add_run(str(currency_symbol)+ ' ' + str("{:,.2f}".format(inv_obj.invoice_amt)))
                cell_3_0 = table.cell((row + gr) - 1, 3)
                cell_4_0.merge(cell_3_0)

                cell_5_0 = table.cell(row, 4)
                p3 = cell_5_0.paragraphs[0]
                p3.add_run(invoice_date)
                cell_6_0 = table.cell((row + gr) - 1, 4)
                cell_5_0.merge(cell_6_0)
                row = gr + row

            # #column merge -----------------------
            column_width = Inches(.3)
            # table.columns[0].width = column_width
            cell_1_2 = table.cell(2, 0)
            cell_1_2.width = column_width
            p3 = cell_1_2.paragraphs[0]
            p3.add_run('Software Development')
            cell_2_2 = table.cell(1+row_len, 0)
            cell_1_2.merge(cell_2_2)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_styleextra

            cell_1_2 = table.cell(2, 2)
            p3 = cell_1_2.paragraphs[0]
            p3.add_run(country_name)
            cell_2_2 = table.cell(1+row_len, 2)
            cell_1_2.merge(cell_2_2)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_styleextra

            cell_1_2 = table.cell(2, 5)
            p3 = cell_1_2.paragraphs[0]
            p3.add_run('N/A')
            cell_2_2 = table.cell(1+row_len, 5)
            cell_1_2.merge(cell_2_2)
            p3.style = custom_styleextra

            # total ---------------------------
            total_row = table.add_row()
            total_row.cells[0].text = 'Total'
            total_row.cells[1].text = all_invoice_qty_str

            # inv_amount_a = "{:.2f}".format(inv_amount).rstrip('0').rstrip('.') + ('0' if inv_amount % 1 else '')
            inv_amount_a = str("{:,.2f}".format(inv_amount))
            total_row.cells[3].text = str(currency_symbol)+ ' ' + inv_amount_a

            # total_row.cells[3].text = str(currency_symbol)+ ' ' + str("{:,}".format(round(inv_amount, 2)))
            # formatted_number = "{:.2f}".format(inv_amount).rstrip('0').rstrip('.') + ('0' if inv_amount % 1 else '')

            # swift_amt_a = "{:.2f}".format(swift_amt).rstrip('0').rstrip('.') + ('0' if swift_amt % 1 else '')
            swift_amt_a = str("{:,.2f}".format(swift_amt))
            total_row.cells[6].text = str(currency_symbol)+ ' ' + swift_amt_a

            first_cell = total_row.cells[0]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph21.style = tab_total

            first_cell = total_row.cells[1]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph21.style = tab_total

            first_cell = total_row.cells[3]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = tab_total

            first_cell1 = total_row.cells[6]
            paragraph22 = first_cell1.paragraphs[0]
            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph22.style = tab_total

            # table 4 ----------------------------------
            paragraph2 = document.add_paragraph("*দৃশ্যমান আকারে পণ্য রপ্তানির ক্ষেত্রে প্রযোজ্য \n (কমার্শিয়াল ইনভয়েস, প্যাকিং লিষ্ট এবং জাহাজীকরণের প্রমাণ স্বরূপ পরিবহন কর্তৃপক্ষ ইস্যুকৃত এবং প্রত্যয়নকৃত বিল অব লোডিং/এয়ারওয়ে বিল, বিল অব এক্সপোর্ট (শুল্ক কর্তৃপক্ষ কর্তৃক ইস্যুকৃত ও পীক্ষিত এবং on-hand হওয়ার স্বপক্ষে পরিবহন কর্তৃপক্ষ প্রত্যয়নকৃত) এর পূর্ণাঙ্গ সেট ইত্যাদির সত্যায়িত পাঠযোগ্য কপি এবং রপ্তানিমূল্য প্রত্যাবাসন সনদপত্র দাখিল করতে হবে। তবে অদৃশ্যকারে সেবা রপ্তানির ক্ষেত্রে জাহাজীকরণের দলিল ও বিল অব এক্সপোর্ট  আবশ্যকতা থাকবে না।) \n(চ) ভর্তুকির আবেদনকৃত অংকঃ")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(12)

            # Set the paragraph alignment to center
            table = document.add_table(rows=3, cols=4)
            table.style = 'TableGrid'
            table.autofit = False
            table.width = Inches(6)
            table.alignment = 2

            cell = table.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('প্রত্যাবাসিত রপ্তানিমূল্য')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[0].width = Inches(1.2)

            cell = table.cell(0, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার পরিমাণ')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[1].width = Inches(1.7)

            cell = table.cell(0, 2)
            p2 = cell.paragraphs[0]
            p2.add_run('বৈদেশিক মুদ্রায় পরিশোধ্য কমিশন ইন্সুরেন্স ইত্যাদি (যদি থাকে)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[2].width = Inches(2.58)

            cell = table.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('নীট এফওবি রপ্তানিমূল্য ১-(২+৩)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('৪')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 0)
            p7 = cell.paragraphs[0]
            p7.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p7.style = custom_style4

            cell = table.cell(2, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('N/A')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 3)
            p9 = cell.paragraphs[0]
            p9.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p9.style = custom_style4

            # table 5 ----------------------------------
            paragraph2 = document.add_paragraph("(প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার উল্লেখ সম্বলিত ফ্রেইট সার্টিফিকেটের সত্যায়িত কপি দাখিল করতে হবে)")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(12)

            vortuki_swift_amt = 0
            if self.incentive_rate_fc > 0:
                vortuki_swift_amt = swift_amt / self.incentive_rate_fc

            # Set the paragraph alignment to center
            table = document.add_table(rows=4, cols=4)
            table.style = 'TableGrid'
            table.autofit = False
            table.width = Inches(6)
            table.alignment = 2

            cell = table.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('রপ্তানি সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের মূল্য')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(0, 1)
            cell.merge(cell2)

            cell = table.cell(0, 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p2 = cell.paragraphs[0]
            p2.add_run('স্থানীয় মূল্য সংযোজনের হার')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(1, 2)
            cell.merge(cell2)

            cell = table.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('প্রাপ্য ভর্তুকি* ৪x১০%')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('দেশীয় পণ্য/সেবা')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('আমদানিকৃত সেবা/পণ্য')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # cell = table.cell(1, 2)
            # p6 = cell.paragraphs[0]
            # p6.add_run('')
            # p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('')
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('৫')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('৬')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৭')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 3)
            p6 = cell.paragraphs[0]
            p6.add_run(' ৮')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 0)
            p7 = cell.paragraphs[0]
            p7.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p7.style = custom_style4

            cell = table.cell(3, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('100%')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

            incentive_amt_fc_a = str("{:,.2f}".format(incentive_amt_fc))
            cell = table.cell(3, 3)
            p9 = cell.paragraphs[0]
            p9.add_run(str(currency_symbol)+ ' ' + incentive_amt_fc_a)
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p9.style = custom_style4

            # signature ------------------------------------------
            paragraph2 = document.add_paragraph("(*৭ নম্বর কলামের হার আলোচ্য সার্কুলারের ৪ নম্বর অনুচ্ছেদের সাথে সামঞ্জস্যতার ক্ষেত্রে ভর্তুকি প্রাপ্য হবে।) \nএ মর্মে অঙ্গীকার করা হচ্ছে যে, আমাদের নিজস্ব কারখানায়/প্রতিষ্ঠানে তৈরী/উৎপাদিত সফটওয়্যার/আইটিইএস/হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকির জন্য আবেদন করা হলো। এ আবেদনপত্রে প্রদত্ত সকল তথ্য/ঘোষণা সম্পূর্ণ ও সঠিক। যদি পরবর্তীতে কোন ভুল/অসত্য তথ্য/প্রতারণা/জালিয়াতি উদঘাটিত হয় তবে গৃহীত ভর্তুকির সমুদয় অর্থ বা এর অংশবিশেষ আমার/আমাদের নিকট হইতে এবং/অথবা আমার/আমাদের ব্যাংক হিসাব থেকে আদায়/ফেরত নেয়া যাবে। \n\nতারিখঃ..................................... ")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)

            paragraph2 = document.add_paragraph("........................................ \nআবেদনকারী প্রতিষ্ঠানের স্বত্তাধিকারী/\n ক্ষমতাপ্রাপ্ত  কর্মকর্তার স্বাক্ষর ও পদবী")
            paragraph2.style = custom_stylebl
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)

            # table 6 ----------------------------------
            document.add_page_break()
            paragraph2 = document.add_paragraph("(ছ) ভর্তুকি প্রদানকারী ব্যাংক শাখা কর্তৃক পূরণীয়ঃ \t\t\t\t\t\t\t\t\t" + ' ' * 11 + "(বৈদেশিক মুদ্রায়)")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)

            # Set the paragraph alignment to center
            table = document.add_table(rows=3, cols=4)
            table.style = 'TableGrid'
            table.autofit = False
            table.width = Inches(6)
            table.alignment = 2

            cell = table.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('প্রত্যাবাসিত রপ্তানিমূল্য')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[0].width = Inches(1.21)

            cell = table.cell(0, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার পরিমাণ')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[1].width = Inches(1.7)

            cell = table.cell(0, 2)
            p2 = cell.paragraphs[0]
            p2.add_run('বৈদেশিক মুদ্রায় পরিশোধ কমিশন, ইন্সুরেন্স ইত্যাদি (যদি থাকে)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[2].width = Inches(2.6)

            cell = table.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('নীট এফওবি রপ্তানিমূল্য ১-(২+৩)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[3].width = Inches(1.8)

            cell = table.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('৪')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 0)
            p4 = cell.paragraphs[0]
            p4.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.style = custom_style4

            cell = table.cell(2, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('N/A')
            p5.style = custom_na
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('N/A')
            p6.style = custom_na
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 3)
            p6 = cell.paragraphs[0]
            p6.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p6.style = custom_style4

            # table 7 ----------------------------------
            if self.od_sight_rate:
                od_sight_rate = self.od_sight_rate
            else:
                if od_s_r:
                    od_sight_rate = od_s_r 
                else:
                    od_sight_rate = encashment_rate_bdt

            rate_fc = swift_amt / self.incentive_rate_fc
            # encashment_final_amt = encashment_amt_bdt / self.incentive_rate_fc
            # encashment_final_amt = rate_fc / self.incentive_rate_fc
            rate_fc_a = "{:,.2f}".format(rate_fc)
            # od_sight_rate = "{:.2f}".format(self.od_sight_rate)
            encashment_final_amt_a = round(rate_fc, 2) * od_sight_rate
            # encashment_final_amt = "{:,.0f}".format(round(encashment_final_amt_a))
            encashment_final_amt = self.custom_round(encashment_final_amt_a)
            usd_bdt_amount = currency + ' ' + str(rate_fc_a) + '\n' + '@' + str("{:,}".format(round(od_sight_rate, 2))) + '=' + '\n' + 'BDT ' + encashment_final_amt
            # amount in word
            amount_in_word_str = ''
            if encashment_final_amt_a:
                # amount_in_word = num2words(round(encashment_final_amt_a))
                amount_in_word = self.num2words_fun(round(encashment_final_amt_a))
                # upper case function call
                amount_in_word_str = self.upper_case(amount_in_word)

            paragraph2 = document.add_paragraph("(প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়া উল্লেখ সম্বলিত ফ্রেইট সার্টিফিকেটের সত্যায়িত কপি দাখিল করতে হবে)")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.line_spacing = Pt(12)

            # Set the paragraph alignment to center
            table = document.add_table(rows=4, cols=4)
            table.style = 'TableGrid'
            table.autofit = False
            table.width = Inches(6)
            table.alignment = 2

            cell = table.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('রপ্তানি সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের মূল্য')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(0, 1)
            cell.merge(cell2)

            cell = table.cell(0, 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p2 = cell.paragraphs[0]
            p2.add_run('স্থানীয় মূল্য সংযোজনের হার [(৪-৬)/৪]x ১০০')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(1, 2)
            cell.merge(cell2)

            cell = table.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('পরিশোধ ভর্তুকির পরিমাণ (টাকায়)*(রপ্তানিমূল্য প্রত্যাবাসনের তারিখে সংশ্লিষ্ট ৪x ১০% বৈদেশিক মুদ্রার ওডি সাইট)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(1, 3)
            cell.merge(cell2)

            cell = table.cell(1, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p4 = cell.paragraphs[0]
            p4.add_run('দেশীয় পণ্য/সেবা')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(1, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p5 = cell.paragraphs[0]
            p5.add_run('আমদানিকৃত সেবা/পণ্য বৈদেশিক মুদ্রায়')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # cell = table.cell(1, 3)
            # p6 = cell.paragraphs[0]
            # p6.add_run('')
            # p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('৫')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('৬')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৭')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 3)
            p6 = cell.paragraphs[0]
            p6.add_run(' ৮')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 0)
            p7 = cell.paragraphs[0]
            p7.add_run(str(currency_symbol) + ' ' + swift_amt_a)
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p7.style = custom_style4

            cell = table.cell(3, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('100%')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 3)
            p9 = cell.paragraphs[0]
            p9.add_run(usd_bdt_amount)
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p9.style = custom_style4

            # conclusion ----------------
            paragraph2 = document.add_paragraph("(*৭ নম্বর কলামের হার আলোচ্য সার্কুলারের ৪ নম্বর অনুচ্ছেদের সাথে সামঞ্জস্যতার ক্ষেত্রে ভর্তুকি প্রাপ্য হবে) \nভর্তুকি পরিমাণ: " + amount_in_word_str + ' Only')
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.line_spacing = Pt(12)

            paragraph2 = document.add_paragraph("\n\n\n\n\n\n\nপরিশোধের তারিখঃ --------------------------")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.line_spacing = Pt(10)

            paragraph2 = document.add_paragraph("----------------------------------------- \nভর্তুকির অনুমোদনের ক্ষমতাপ্রাপ্ত ব্যাংক  \nকর্মকর্তার স্বাক্ষর, নাম ও পদবী")
            paragraph2.style = custom_stylebl
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.line_spacing = Pt(10)

            # paragraph = document.add_paragraph()
            # paragraph.add_run("\n\n\n\nপরিশোধের তারিখঃ --------------------------")
            # paragraph.add_run("\t")
            # paragraph.add_run("\t")
            # paragraph.add_run("........................................ \nভর্তুকির অনুমোদনের ক্ষমতাপ্রাপ্ত ব্যাংক  \nকর্মকর্তার স্বাক্ষর, নাম ও পদবী")

            file_name = 'Form_KA_%s_%s' % (desired_value, datetime.datetime.now())
            # file_name = '%s_ka' % (datetime.datetime.now())

            # html = self.get_word_prc()
            # html_text1 = html.replace('<br>', '<p style="display:block"></p>')
            # html_text1 = "<p></p><p></p><p></p><p></p> <div style='display:inline-flex'><p style='font-size:9px; margin-right:10px'>Signature of Head of the Branch</p> &nbsp;&nbsp;&nbsp; &nbsp;&nbsp;&nbsp;  <p>Signature of the Issuing Officer</p></div>"
            # new_parser.add_html_to_document(html_text1, document)

        if type == 'BASIS_KHA':
            # modify the page setup
            section.page_width = Cm(22)  # set the page width to 21 centimeters
            section.page_height = Cm(29.9)  # set the page width to 21 centimeters
            section.left_margin = Cm(2.5)  # 1.5 cm
            section.right_margin = Cm(2)  # 1.5 cm
            self_date = ''
            if self.date:
                self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

            # document.add_paragraph('কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।')
            # en_font = my_stylew.font
            # document.add_paragraph('some text')

            custom_style = document.styles.add_style('MyStyled', WD_STYLE_TYPE.PARAGRAPH)
            # custom_style.font.bold = True
            custom_style.font.size = Pt(8)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            custom_style.font.name = 'Arial Narrow'

            custom_style_table = document.styles.add_style('custom_style_table', WD_STYLE_TYPE.PARAGRAPH)
            custom_style_table.font.size = Pt(10)
            custom_style_table.font.name = 'Arial Narrow'

            custom_stylebsmallbold = document.styles.add_style('custom_stylebsmallbold', WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebsmallbold.font.size = Pt(9)
            custom_stylebsmallbold.font.name = 'SutonnyOMJ'
            custom_stylebsmallbold.font.bold = True

            custom_stylebsmall = document.styles.add_style('custom_stylebsmall', WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebsmall.font.size = Pt(9)
            custom_stylebsmall.font.name = 'SutonnyOMJ'

            custom_stylebl = document.styles.add_style('custom_stylebl', WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebl.font.size = Pt(9)
            custom_stylebl.font.name = 'SutonnyOMJ'

            custom_sonar_bangla = document.styles.add_style('custom_sonar_bangla', WD_STYLE_TYPE.PARAGRAPH)
            custom_sonar_bangla.font.size = Pt(9)
            custom_sonar_bangla.font.name = 'Shonar Bangla'

            custom_tableh = document.styles.add_style('custom_tableh', WD_STYLE_TYPE.PARAGRAPH)
            custom_tableh.font.size = Pt(9)
            custom_tableh.font.name = 'SutonnyOMJ'

            custom_table = document.styles.add_style('custom_table', WD_STYLE_TYPE.PARAGRAPH)
            custom_table.font.size = Pt(9)
            custom_table.font.name = 'SutonnyOMJ'

            custom_stylebold = document.styles.add_style('custom_stylebold', WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebold.font.size = Pt(11)
            custom_stylebold.font.name = 'SutonnyOMJ'
            custom_stylebold.font.bold = True

            custom_styleubl = document.styles.add_style('custom_styleubl', WD_STYLE_TYPE.PARAGRAPH)
            custom_styleubl.font.size = Pt(11)
            custom_styleubl.font.name = 'SutonnyOMJ'
            custom_styleubl.font.underline = True
            custom_styleubl.font.bold = True

            document.styles['Normal'].font.size = Pt(8)
            paragraph = document.add_paragraph()
            paragraph.add_run("\n(অনুচ্ছেদ ০৬ (ক), এফই সার্কুলার নং-০৩/২০১৮ দ্রষ্টব্য)")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.style = custom_stylebl
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.add_run("ফরম- 'খ'")
            paragraph = document.add_paragraph()
            paragraph.add_run("Ref: " + str(self.form_kha_ref_code))
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run(" Date: " + self_date)
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.style = custom_stylebl
            
            # paragraph.style = custom_style

            # Define a custom style
            custom_style2 = document.styles.add_style('MyStyle21', WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(8)

            custom_style3 = document.styles.add_style('MyStyle3', WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(8)
            custom_style3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            custom_style4 = document.styles.add_style('MyStyle4', WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.bold = True
            custom_style4.font.size = Pt(10)
            custom_style4.font.name = 'Arial Narrow'

            custom_total = document.styles.add_style('custom_total', WD_STYLE_TYPE.PARAGRAPH)
            custom_total.font.bold = True
            custom_total.font.size = Pt(9)
            custom_total.font.name = 'Arial Narrow'

            custom_style5 = document.styles.add_style('MyStyle5', WD_STYLE_TYPE.PARAGRAPH)
            custom_style5.font.size = Pt(8)
            custom_style5.font.name = 'Arial Narrow'

            # Add a paragraph and apply the custom style
            paragraph = document.add_paragraph('বেসিস প্রদেয় প্রত্যয়ন সনদপত্র \nবাংলাদেশ হতে সফটওয়্যার, আইটিইএস (Information Technology Enabled Services)')
            paragraph.style = custom_stylebold
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph = document.add_paragraph('ও হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকি প্রাপ্তির প্রত্যয়ন সনদপত্র।')
            paragraph.style = custom_styleubl
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)

            customer = '' if not self.swift_customer_name else str(self.swift_customer_name) + ', ' + str(
                self.customer_address)
            con_no = ''
            if self.contract_number:
                con_no = self.contract_number.replace('Contract No.', '').replace('Purchase Order No.', '')

            # new_parser = HtmlToDocx()
            # html_text = """
            #             <div style='position: relative'>
            #                 <p style='float:left'> ১। </p>
            #                 <p  style='margin-left: 18px; float:right'> আবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road, Mohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ RA- 260326210666420 </p>
            #             </div>
            #
            #             """
            # new_parser.add_html_to_document(html_text, document)

            # paragraph = document.add_paragraph()
            # paragraph.add_run('Word1')
            # paragraph.add_run(' ' * 3)
            # paragraph.add_run('Word2')
            import textwrap
            width = 140
            # long_text = "১।\t  আবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road, Mohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ 260326210666420"
            # wrapped_lines = textwrap.wrap(long_text, width=width)
            # formatted_lines = [line + '\t' for line in wrapped_lines]
            # formatted_text = '\n\t'.join(formatted_lines)
            # paragraph2 = document.add_paragraph(formatted_text)
            # paragraph2.style = custom_stylebl
            # paragraph2.paragraph_format.space_before = Pt(1)
            # paragraph2.paragraph_format.space_after = Pt(1)
            #
            # long_text = "২।\tরপ্তানি ঋণপত্র/চুক্তিপত্র নম্বরঃ" + con_no + '\t তারিখঃ ' + self.contract_date_str + '\t মূল্যঃ ' +self.contract_price_str
            # wrapped_lines = textwrap.wrap(long_text, width=width)
            # formatted_lines = [line + '\t' for line in wrapped_lines]
            # formatted_text = '\n\t'.join(formatted_lines)
            # paragraph2 = document.add_paragraph(formatted_text)
            # paragraph2.paragraph_format.space_before = Pt(1)
            # paragraph2.paragraph_format.space_after = Pt(1)
            #
            # long_text = "৩।\tবিদেশি ক্রেতার নাম ও ঠিকানাঃ "+ customer
            # wrapped_lines = textwrap.wrap(long_text, width=width)
            # formatted_lines = [line + '\t' for line in wrapped_lines]
            # formatted_text = '\n\t'.join(formatted_lines)
            # paragraph2 = document.add_paragraph(formatted_text)
            # paragraph2.paragraph_format.space_before = Pt(1)
            # paragraph2.paragraph_format.space_after = Pt(1)
            #
            # long_text = "৪।\tবিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ " + remiter_address + '\n' + '৫।'
            # wrapped_lines = textwrap.wrap(long_text, width=width)
            # formatted_lines = [line + '\t' for line in wrapped_lines]
            # formatted_text = '\n\t'.join(formatted_lines)
            # paragraph2 = document.add_paragraph(formatted_text)
            # paragraph2.paragraph_format.space_before = Pt(1)
            # paragraph2.paragraph_format.space_after = Pt(1)

            # paragraph2 = document.add_paragraph("১।\tআবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road,"
            #                                     " \tMohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ 260326210666420 \n২।\tরপ্তানি ঋণপত্র/চুক্তিপত্র নম্বরঃ" + con_no + '\t তারিখঃ ' +
            #                                     self.contract_date_str + '\t মূল্যঃ ' +self.contract_price_str + '\n৩।\tবিদেশি ক্রেতার নাম ও ঠিকানাঃ '+ customer +
            #                                     "\n৪।\tবিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ " + remiter_address + '\n' + '৫।' )

            paragraph2 = document.add_paragraph("১।\tআবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road,"
                                                " \tMohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ 260326210666420")
            paragraph2.style = custom_stylebl
            # paragraph2.add_run("\t")
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)

            long_text1 = "২।\t" + ' ' * 5 + "রপ্তানি ঋণপত্র/চুক্তিপত্র নম্বরঃ" + con_no + '\t তারিখঃ ' + self.contract_date_str + '\t মূল্যঃ ' +self.contract_price_str
            wrapped_lines1 = textwrap.wrap(long_text1, width=120)
            formatted_lines1 = [line + '\t' for line in wrapped_lines1]
            formatted_text1 = '\n\t'.join(formatted_lines1)
            paragraph2 = document.add_paragraph(formatted_text1)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            long_text2 = "৩।\t" + ' ' * 5 + "বিদেশি ক্রেতার নাম ও ঠিকানাঃ " + customer
            # long_text2 = "৩।\t" + ' ' * 5 + "বিদেশি ক্রেতার নাম ও ঠিকানাঃ ABBVIE INC AP, DEPT V312 AP34-2 1 N WAUKEGAN RD, NORTH CHICAGO"
            wrapped_lines2 = textwrap.wrap(long_text2, width=123)
            formatted_lines2 = [line + '\n\t' for line in wrapped_lines2]
            # formatted_text2 = ''.join(formatted_lines2)
            formatted_text2 = ''.join(formatted_lines2).rstrip('\n\t')
            paragraph2 = document.add_paragraph(formatted_text2)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            long_text1 = "৪।\t" + ' ' * 5 + "বিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ " + remiter_address
            wrapped_lines1 = textwrap.wrap(long_text1, width=120)
            formatted_lines1 = [line + '\t' for line in wrapped_lines1]
            formatted_text1 = '\n\t'.join(formatted_lines1)
            paragraph2 = document.add_paragraph(formatted_text1)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph('৫।')
            paragraph2.style = custom_stylebl

            # table = document.add_table(rows=3, cols=4)
            #
            # # Set the width of the first column to 1 inch
            # column_width = Inches(1)
            # table.columns[0].width = column_width
            #
            # # Set the width of the second column to 2 inches
            # column_width = Inches(1)
            # table.columns[1].width = column_width
            #
            # # Set the width of the third column to 0.5 inches
            # column_width = Inches(4)
            # table.columns[2].width = column_width
            #
            # # Set the width of the third column to 0.5 inches
            # column_width = Inches(1)
            # table.columns[3].width = column_width

            table = document.add_table(rows=1, cols=4)
            table.style = 'TableGrid'
            table.alignment = 1
            # Set the width of the first column to 1 inch
            column_width = Inches(1.2)
            table.columns[0].width = column_width
            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('ক) ইনভয়েস নম্বর')
            p3.style = custom_tableh
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Set the width of the second column to 2 inches
            column_width = Inches(1)
            table.columns[1].width = column_width
            cell = table.cell(0, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('তারিখ')
            p3.style = custom_tableh
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Set the width of the third column to 0.5 inches
            column_width = Inches(4.3)
            table.columns[2].width = column_width
            cell1 = table.cell(0, 2)
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell1.paragraphs[0]
            p3.add_run('খ) ইনভয়েসে উল্লেখিত সেবা/পণ্যের পরিমাণ (Qty & Hours)')
            p3.style = custom_tableh
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[2].width = Inches(3.8)

            # Set the width of the third column to 0.5 inches
            column_width = Inches(1)
            table.columns[3].width = column_width
            cell = table.cell(0, 3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name = 'মূল্য (' + self.fc_currency_id.name + ')'
            # p3.add_run('মূল্য (USD)')
            p3.add_run(cur_name)
            p3.style = custom_tableh
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Set the paragraph alignment to center
            # table = document.add_table(rows=1, cols=4)
            # table.style = 'TableGrid'
            # table.autofit = False
            # # table = document.tables[0]
            # # Adjust the left indentation of the table
            # table.alignment = 1  # Center alignment
            # # table.left_indent = Pt(144)
            #
            # cell = table.cell(0, 0)
            # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # p3 = cell.paragraphs[0]
            # p3.add_run('ক) ইনভয়েস নম্বর'...............)
            # p3.style = custom_tableh
            # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[0].width = Inches(1.2)
            #
            # cell = table.cell(0, 1)
            # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # p3 = cell.paragraphs[0]
            # p3.add_run('তারিখ')
            # p3.style = custom_tableh
            # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[1].width = Inches(1)
            #
            # cell1 = table.cell(0, 2)
            # cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # p3 = cell1.paragraphs[0]
            # p3.add_run('খ) ইনভয়েসে উল্লেখিত সেবা/পণ্যের পরিমাণ (Qty & Hours)')
            # p3.style = custom_tableh
            # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[2].width = Inches(3.8)
            #
            # cell = table.cell(0, 3)
            # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # p3 = cell.paragraphs[0]
            # p3.add_run('মূল্য (USD)')
            # p3.style = custom_tableh
            # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[3].width = Inches(1)
            #
            # # add a dynamic row to the table
            cus_name = ''
            if self.swift_customer_name:
                cus_name = self.swift_customer_name

            # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
            total_qty = 0
            dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in
                             self.invoice_line_ids.invoice_id.invoice_line_ids)
            hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in
                              self.invoice_line_ids.invoice_id.invoice_line_ids)
            dev_total = dev_total1
            hour_total = hour_total1

            if hour_total == int(hour_total):
                hour_total_num = "{:.0f}".format(hour_total)
            else:
                hour_total_num = "{:.2f}".format(hour_total)

            if dev_total == int(dev_total):
                dev_total_num = "{:.0f}".format(dev_total)
            else:
                dev_total_num = "{:.2f}".format(dev_total)

            qty_str = ''
            if dev_total > 0:
                qty_str += str(dev_total_num) + (str(' Developers') if dev_total > 1 else str(' Developer'))
            if hour_total > 0:
                qty_str += str(hour_total_num) + str(' HRS') if not qty_str else ' & ' + str(
                    hour_total_num) + str(' HRS')
            all_invoice_qty_str = 'Software Development /' + qty_str
            all_invoice_qty_str_a = qty_str
            row_len = len(self.invoice_line_ids)
            swift_amt = 0
            inv_amount = 0

            from itertools import groupby

            # group the invoice_line_ids by invoice_id.ref
            groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                             key=lambda x: x.invoice_id.ref)
            # iterate over each group and add a new row for each unique invoice_id.ref
            row = 1
            for invoice_ref, group in groups:
                inv_obj = self.env['cash.incentive.invoice'].search(
                    [('invoice_id.ref', '=', invoice_ref), ('head_id', '=', self.id)], limit=1)
                inv_amount += inv_obj.invoice_amt
                invoice_date = ''
                if inv_obj.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime(
                        '%d-%b-%y')

                qty = 'Software Development /'
                h_q = 0
                d_q = 0
                for l in inv_obj.invoice_id.invoice_line_ids:
                    if l.quantity_type == '0':
                        h_q += l.quantity
                    else:
                        d_q += l.quantity
                if d_q:
                    if d_q == int(d_q):
                        d_q_num = "{:.0f}".format(d_q)
                    else:
                        d_q_num = "{:.2f}".format(d_q)
                    qty += str(d_q_num) + (str(' Developers') if d_q > 1 else str(' Developer'))
                if h_q:
                    if h_q == int(h_q):
                        h_q_new = "{:.0f}".format(h_q)
                    else:
                        h_q_new = "{:.2f}".format(h_q)
                    if not d_q:
                        qty += str(h_q_new) + ' HRS'
                    else:
                        qty += ' & ' + str(h_q_new) + ' HRS'
                invoice_qty_str = qty

                gr = 0
                for rec in group:
                    gr += 1
                    swift_amt += rec.swift_amt
                    # add a new row to the table
                    new_row = table.add_row()
                    # add data to the cells in the new row
                    # new_row.cells[0].text = rec.invoice_id.ref
                    # new_row.cells[1].text = invoice_date
                    # new_row.cells[2].text = invoice_qty_str
                    # new_row.cells[3].text = currency + ' ' + str("{:,}".format(round(rec.invoice_amt, 2)))

                    first_cell = new_row.cells[0]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = custom_style5
                    first_cell = new_row.cells[1]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = custom_style5
                    first_cell = new_row.cells[2]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = custom_style5
                    first_cell = new_row.cells[3]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style5

                cell_1_0 = table.cell(row, 0)
                p3 = cell_1_0.paragraphs[0]
                p3.add_run(invoice_ref)
                p3.style = custom_style_table
                cell_2_0 = table.cell((row + gr) - 1, 0)
                cell_1_0.merge(cell_2_0)

                cell_4_0 = table.cell(row, 1)
                p3 = cell_4_0.paragraphs[0]
                p3.add_run(invoice_date)
                p3.style = custom_style_table
                cell_3_0 = table.cell((row + gr) - 1, 1)
                cell_4_0.merge(cell_3_0)

                cell_5_0 = table.cell(row, 2)
                p3 = cell_5_0.paragraphs[0]
                p3.add_run(invoice_qty_str)
                p3.style = custom_style_table
                cell_6_0 = table.cell((row + gr) - 1, 2)
                cell_5_0.merge(cell_6_0)

                cell_5_0 = table.cell(row, 3)
                p3 = cell_5_0.paragraphs[0]
                p3.style = custom_style_table
                # p3.add_run(str("{:,}".format(round(inv_obj.invoice_amt, 2))))
                p3.add_run(str("{:,.2f}".format(inv_obj.invoice_amt)))
                cell_6_0 = table.cell((row + gr) - 1, 3)
                cell_5_0.merge(cell_6_0)
                row = gr + row

            # total
            total_row = table.add_row()
            total_row.cells[0].text = 'Total'
            total_row.cells[2].text = all_invoice_qty_str
            total_row.cells[3].text = str("{:,.2f}".format(inv_amount))

            first_cell = total_row.cells[0]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph21.style = custom_total
            first_cell = total_row.cells[2]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph21.style = custom_total
            first_cell = total_row.cells[3]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = custom_total

            self_date = ''
            if self.date:
                self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

            # footer -------------------
            # paragraph2 = document.add_paragraph("৬।\tরপ্তানিকৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত স্থানীয় সেবা/উপকরণাদির সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): 100% in house production \tof Brain Station 23 Limited \t\t\t\t পরিমাণঃ N/A  \t\t মূল্যঃ N/A"
            #                                     "\n৭।\tরপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত আমদানিকৃত আনুসাংগিক সেবা/উপরকরণাদি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): N/A  \t\t\tপরিমাণঃ N/A   \t\t\t\t\t  মূল্যঃ  N/A"
            #                                     "\n৮।\tরপ্তানি সেবা/পণ্যের বিবরণঃ Software Development Services  \tপরিমাণঃ " + all_invoice_qty_str_a + " \t  মূল্যঃ " + currency+ ' ' + str("{:,.2f}".format(inv_amount)) +
            #                                     "\n৯।\tজাহাজীকরণের তারিখঃ " + inv_date + "\tগন্তব্য বন্দরঃ " + country + "\n১০।\tইএক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানরি ক্ষেত্রে): N/A  \t\t\t  মূল্যঃ N/A   \t\t তারিখঃ N/A"
            #                                     "\n১১।\tমোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency+ ' ' + str("{:,.2f}".format(swift_amt)) + "\t\t নীট এফওবি রপ্তানি মূল্য ( বৈদেশিক মুদ্রায়): "+ currency+ ' ' + str("{:,.2f}".format(swift_amt)) +
            #                                     "\n১২।\tপ্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ \t\t\t\t\t\t তারিখঃ")
            # paragraph2.style = custom_stylebl

            paragraph2 = document.add_paragraph("৬।\tরপ্তানিকৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত স্থানীয় সেবা/উপকরণাদির সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): 100% in house \tproduction of Brain Station 23 Limited \t\t পরিমাণঃ N/A  \t\t মূল্যঃ N/A"
                                                "\n৭।\tরপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত আমদানিকৃত আনুসাংগিক সেবা/উপরকরণাদি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): N/A  \t\tপরিমাণঃ N/A   \t\t\t\t\t  মূল্যঃ  N/A")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_after = Pt(1)

            # long_text = "৮। \t " + ' ' * 4 + "রপ্তানি সেবা/পণ্যের বিবরণঃ Software Development Services পরিমাণঃ " + all_invoice_qty_str_a + " " + "  মূল্যঃ " + currency+ ' ' + str("{:,.2f}".format(inv_amount))
            # long_text = "৯। \t " + ' ' * 4 + "জাহাজীকরণের তারিখঃ 31/10/22, 31/10/22, 30/11/22, 30/11/22, 30/11/22, 30/11/22, 31/10/22, 30/11/22, 30/11/22, 30/11/22, 30/11/22 \t\tগন্তব্য বন্দরঃ Norway"
            # wrapped_lines = textwrap.wrap(long_text, width=120)
            # formatted_lines = [line + '\t' for line in wrapped_lines]
            # formatted_text = '\n\t'.join(formatted_lines)
            # paragraph2 = document.add_paragraph(formatted_text)
            # paragraph2.style = custom_stylebl
            # paragraph2.paragraph_format.space_before = Pt(1)
            # paragraph2.paragraph_format.space_after = Pt(1)

            if len(self.invoice_line_ids) > 10:
                long_text = "৯। \t " + ' ' * 4 + "জাহাজীকরণের তারিখঃ " + inv_date + " \t\tগন্তব্য বন্দরঃ " + country
                wrapped_lines = textwrap.wrap(long_text, width=120)
                formatted_lines = [line + '\t' for line in wrapped_lines]
                formatted_text = '\n\t'.join(formatted_lines)
                paragraph2 = document.add_paragraph(formatted_text)
                paragraph2.style = custom_stylebl
                paragraph2.paragraph_format.space_before = Pt(1)
                paragraph2.paragraph_format.space_after = Pt(1)
            else:
                if len(self.invoice_line_ids) > 7:
                    paragraph2 = document.add_paragraph("৯।\tজাহাজীকরণের তারিখঃ " + inv_date + "\n\tগন্তব্য বন্দরঃ " + country )
                    paragraph2.style = custom_stylebl
                else:
                    paragraph2 = document.add_paragraph("৯।\tজাহাজীকরণের তারিখঃ " + inv_date + " \tগন্তব্য বন্দরঃ " + country )
                    paragraph2.style = custom_stylebl
                paragraph2.paragraph_format.space_before = Pt(1)
                paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("১০।\tইএক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানরি ক্ষেত্রে): N/A  \t\t\t  মূল্যঃ N/A   \t\t তারিখঃ N/A")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            # long_text1 = "১১। \t " + ' ' * 5 +  "মোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): USD 158,793.64          নীট এফওবি রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): USD 158,793.64"
            long_text1 = "১১। \t " + ' ' * 4 + "মোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str("{:,.2f}".format(swift_amt)) + "\t\tনীট এফওবি রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str("{:,.2f}".format(swift_amt))
            wrapped_lines1 = textwrap.wrap(long_text1, width=150)
            formatted_lines1 = [line + '\t' for line in wrapped_lines1]
            formatted_text1 = '\n\t'.join(formatted_lines1)
            paragraph2 = document.add_paragraph(formatted_text1)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            prc_date = ''
            if self.prc_date:
                prc_date = datetime.datetime.strptime(str(self.prc_date), '%Y-%m-%d').strftime('%d/%m/%y')
            prc_ref_code = ''
            if self.prc_ref_code:
                prc_ref_code = self.prc_ref_code

            long_text1 ="১২।\t" + ' ' * 4 + "প্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ " + prc_ref_code+ "\t\t\t তারিখঃ " + prc_date
            wrapped_lines1 = textwrap.wrap(long_text1, width=120)
            formatted_lines1 = [line + '\t' for line in wrapped_lines1]
            formatted_text1 = '\n\t'.join(formatted_lines1)
            paragraph2 = document.add_paragraph(formatted_text1)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            # paragraph2 = document.add_paragraph("১২।\tপ্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ " + prc_ref_code+ "\t\t\t\t\t\t তারিখঃ " + prc_date)
            # paragraph2.style = custom_stylebl

            paragraph2 = document.add_paragraph("\n\n\nরপ্তানিকারকের স্বাক্ষর ও তারিখ")
            # paragraph2.style = custom_style3
            # paragraph2.style = my_styler
            paragraph2.style = custom_stylebsmallbold
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph4 = document.add_paragraph("এতদ্বারা প্রত্যয়ন করা যাচ্ছে যে, আমাদের নিজস্ব কারখানায়/প্রতিষ্ঠানে তৈরীকৃত/উৎপাদিত সফটওয়্যার/আইটিইএস/হার্ডওয়্যার উপরোক্ত ৬ ও ৭ নং ক্রমিক বর্ণিত সূত্র হতে সেবা/উপকরাণাদি সংগ্রহের মাধ্যমে রপ্তানির বিপরীতে ভর্তুকির জন্য উপরোক্ত অনুচ্ছেদগুলোতে উল্লিখিত বক্তব্য সঠিক ও নির্ভুল। বিদেশী ক্রেতা/ আমদানিকারকের ক্রয়াদেশের যথার্থতা/বিশ্বাসযোগ্যতা সম্পর্কেও নিশ্চিত করা হলো। ")
            paragraph4.style = custom_stylebsmall
            paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # paragraph4 = document.paragraphs[1]
            # Modify the paragraph format to reduce spacing
            paragraph4.paragraph_format.space_before = Pt(1)
            paragraph4.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("\n\n\nরপ্তানিকারকের স্বাক্ষর ও তারিখ")
            # paragraph2.style = custom_style3
            # paragraph2.style = my_styler
            paragraph2.style = custom_stylebsmallbold
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("রপ্তানিকারকের উপরোক্ত ঘোষণার যথার্থতা যাচাইয়ান্তে সঠিক পাওয়া গিয়েছে। ৮নং ক্রমিকে উল্লিখিত ঘোষিত রপ্তানিমূল্য যৌক্তিক ও বিদ্যমান আন্তর্জাতিক বাজার মুল্যের সঙ্গে সংগতিপূর্ণ পাওয়া গিয়েছে এবং বিদেশী ক্রেতার যথার্থতা/বিশ্বাসযোগ্যতা সর্ম্পকেও নিশ্চিত হওয়া গিয়েছে। প্রত্যাবাসিত রপ্তানি মূল্যের (নীট এফওবি মূল্য) রপ্তানি ভর্তুকি পরিশোধের সুপারিশ করা হলো।")
            paragraph2.style = custom_stylebsmall
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("\n\n\n\n----------------------------------- এসোসিয়েশন এর দুইজন উপযুক্ত কর্মকর্তার স্বাক্ষর, তারিখ ও সীল")
            # paragraph2.style = my_styler
            paragraph2.style = custom_stylebsmall
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("[কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।]")
            # paragraph2.style = my_stylec
            paragraph2.style = custom_stylebsmall
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(0)

            file_name = 'Form_KHA_%s_%s' % (desired_value, datetime.datetime.now())
            # file_name = '%s_kha' % (datetime.datetime.now())

        if type == 'FORM_C':
            self_date = ''
            if self.date:
                self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

            # Define a custom style
            custom_style = document.styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
            custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            custom_style4 = document.styles.add_style('MyStyle4', WD_STYLE_TYPE.PARAGRAPH)
            # custom_style.font.bold = True
            custom_style4.font.size = Pt(11)
            custom_style4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            custom_style2 = document.styles.add_style('MyStyle2', WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(11)

            custom_style3 = document.styles.add_style('MyStyle3', WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(11)
            custom_style3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Add a paragraph and apply the custom style
            paragraph1 = document.add_paragraph('Form- C (ICT)')
            paragraph1.style = custom_style
            paragraph1 = document.add_paragraph('Declaration for inward remittance on account of ICT \n related services of amount exceeding USD 10,000 or equivlent')
            paragraph1.style = custom_style4

            paragraph2 = document.add_paragraph("\nI/We do hereby declare that I/we have received remittance of  7,000.0 (amount) which is a fair value against ICT elated services described below in respect of which this declaration is made out and that the particulars given below are true:")
            paragraph2.style = custom_style2

            paragraph2 = document.add_paragraph("\n     a) Remitter's name and address: " + client_name + ', ' + customer_address )
            paragraph2.style = custom_style2
            paragraph2 = document.add_paragraph("     b) Remitting bank and address: " + remiter_address)
            paragraph2.style = custom_style2
            paragraph2 = document.add_paragraph("     c) Reference No. (contract/invoice/electric communication etc.): Invoice " + inv_number)
            paragraph2.style = custom_style2
            paragraph2 = document.add_paragraph("     d) Purpose (please tick): \n")
            paragraph2.style = custom_style2

            # Create a list of items
            items = ["Information Technology Enabled Services (IT Enabled Services) and Business Process Outsourcing (BPO) services- code 2410",
                     "Computer and Information Technology consultancy and management services– code 2411",
                     "Export of computer software  including turn-key basis (customized)– code 2412 \ncomputer software including turn-key basis (non-customized) – code 2413",
                     "Installation services concerning hardware and software maintenance and repairs of computers and peripheral equipment services- code 2414",
                     "....................................... (please  specify)"]

            # Create a paragraph object for each list item and set the paragraph style to "List Bullet"
            for item in items:
                paragraph7 = document.add_paragraph()
                paragraph7.add_run('\t' + u'\u25A0 ')
                paragraph7.add_run(item)

            paragraph2 = document.add_paragraph("\n\n\n\nNote: Please see the “Code lists for Reporting of External Sector transactions by the Authorized dealers” for explanatory notes of above purposes. \n\n\nSignature with name of applicant: \nAddress: \nDate: \n\n\n")
            paragraph2.style = custom_style2

            # Set the paragraph alignment to center ----------------
            table = document.add_table(rows=2, cols=15)
            table.style = 'TableGrid'
            cell = table.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('Month')
            p3.alignment = WD_ALIGN_PARAGRAPH.LEFT

            cell_1_2 = table.cell(0, 1)
            cell_1_2.text = 'Country of ordering customer'
            cell_2_2 = table.cell(0, 4)
            cell_1_2.merge(cell_2_2)

            cell_1_3 = table.cell(0, 5)
            cell_1_3.text = 'Purpose'
            cell_2_4 = table.cell(0, 8)
            cell_1_3.merge(cell_2_4)

            cell_1_6 = table.cell(0, 9)
            cell_1_6.text = 'Amount in (state currency'
            cell_2_5 = table.cell(0, 14)
            cell_1_6.merge(cell_2_5)

            cell = table.cell(1, 0)
            p = cell.paragraphs[0]
            p.add_run('')

            cell = table.cell(1, 1)
            p = cell.paragraphs[0]
            p.add_run('')

            paragraph2 = document.add_paragraph("\n\n\nCoded by: \nChecked by: \nPurpose of remittance: \n\nSignature and stamp of Authorized Dealer: \nDate:")
            paragraph2.style = custom_style2

            file_name = 'Form_C_%s_%s' % (desired_value, datetime.datetime.now())

            # file_name = '%s_Form-C' % (datetime.datetime.now())

        if type == 'FLBK':
            section.page_width = Cm(22)  # set the page width to 21 centimeters
            section.page_height = Cm(29.9)  # set the page width to 21 centimeters
            section.left_margin = Cm(2.5)  # 1.5 cm
            section.right_margin = Cm(2)

            self_date = ''
            current_date = datetime.datetime.now().date()
            if self.date:
                self_date = datetime.datetime.strptime(str(current_date), '%Y-%m-%d').strftime('%d/%m/%y')

            # Define a custom style
            custom_style = document.styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
            custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            custom_style.font.name = 'Arial Narrow'

            custom_style2 = document.styles.add_style('MyStyle2', WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(10)
            custom_style2.font.name = 'Calibri'

            custom_style3 = document.styles.add_style('MyStyle3', WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(10)
            custom_style3.font.bold = True
            custom_style3.font.name = 'Calibri'

            custom_style4 = document.styles.add_style('MyStyle4', WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.bold = True
            custom_style4.font.size = Pt(9)
            custom_style4.font.name = 'Arial Narrow'

            table_total = document.styles.add_style('table_total', WD_STYLE_TYPE.PARAGRAPH)
            table_total.font.bold = True
            table_total.font.size = Pt(10)
            table_total.font.name = 'Arial Narrow'

            custom_style5 = document.styles.add_style('MyStyle5', WD_STYLE_TYPE.PARAGRAPH)
            custom_style5.font.size = Pt(9)
            custom_style5.font.name = 'Arial Narrow'

            paragraph2 = document.add_paragraph('\n\n\n\n\nRef: ' + self.name + '\n\nDate: ' + self_date) # + '\n'
            paragraph2.style = custom_style3
            paragraph2.paragraph_format.left_indent = Pt(-6)

            # paragraph2 = document.add_paragraph("The Branch Manager  \nGulshan Branch \n" + self.bank_id.name + "\nHolding No. 75, Gulshan Avenue \nGulshan, Dhaka \n\n\nDear Sir\n")
            paragraph2 = document.add_paragraph("The Manager \n" + self.bank_id.name + "\nIslamic Banking Window. \n4 Bir Uttam AK Khandakar Rd, Mohakhali C/A \nDhaka 1212 \n\n\nDear Sir\n")
            paragraph2.style = custom_style2
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(1)
            paragraph2.paragraph_format.left_indent = Pt(-6)

            paragraph2 = document.add_paragraph("For cash incentive claim.\n")
            paragraph2.style = custom_style3
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
            paragraph2.paragraph_format.left_indent = Pt(-6)

            paragraph2 = document.add_paragraph("We are submitting herewith necessary documents against following Invoices:")
            paragraph2.style = custom_style2
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.left_indent = Pt(-6)

            # table -----------------------

            table = document.add_table(rows=1, cols=8)
            table.style = 'TableGrid'
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # table.autofit = False  # Disable auto-fit behavior
            # table.left_indent = Inches(6)

            cell = table.cell(0, 0)
            # table.columns[0].width = Inches(.2)
            # table.columns[0].left_indent = Pt(100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('SL#')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4

            cell = table.cell(0, 1)
            table.columns[1].width = Inches(1.8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Client Name')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4

            cell = table.cell(0, 2)
            table.columns[2].width = Inches(.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Invoice No')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4

            cell = table.cell(0, 3)
            table.columns[3].width = Inches(.7)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name = 'Invoice amount (' + self.fc_currency_id.name + ')'
            p3.add_run(cur_name)
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p3.style = custom_style4

            cell = table.cell(0, 4)
            table.columns[4].width = Inches(.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Invoice Date')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4

            cell = table.cell(0, 5)
            table.columns[5].width = Inches(.6)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Swift/nostro date')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4

            cell = table.cell(0, 6)
            table.columns[6].width = Inches(.6)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name_r = 'Realize Amount (' + self.fc_currency_id.name + ')'
            p3.add_run(cur_name_r)
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p3.style = custom_style4

            cell = table.cell(0, 7)
            table.columns[7].width = Inches(.6)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name_i = 'Incentive claim (' + self.fc_currency_id.name + ')'
            p3.add_run(cur_name_i)
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p3.style = custom_style4

            # add a dynamic row to the table
            cus_name = ''
            if self.swift_customer_name:
                cus_name = self.swift_customer_name

            # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
            row_len = len(self.invoice_line_ids)
            inv_amount = 0
            encashment_amt_bdt = 0
            sl = 1
            total_incentive_rate_fc = 0

            from itertools import groupby

            # group the invoice_line_ids by invoice_id.ref
            groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                             key=lambda x: x.invoice_id.ref)
            # iterate over each group and add a new row for each unique invoice_id.ref
            row = 1
            for invoice_ref, group in groups:
                inv_obj = self.env['cash.incentive.invoice'].search([('invoice_id.ref', '=', invoice_ref), ('head_id', '=', self.id)], limit=1)
                # usd_price = sum(inv_obj.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
                usd_price = inv_obj.invoice_id.invoice_total_fc
                # incentive_rate_fc = usd_price / self.incentive_rate_fc if self.incentive_rate_fc else 0

                inv_amount += inv_obj.invoice_amt
                invoice_date = ''
                if inv_obj.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime('%d-%b-%y')
                gr = 0
                for rec in group:
                    gr += 1
                    encashment_amt_bdt += rec.encashment_amt_bdt
                    incentive_rate_fc = (rec.swift_amt * self.incentive_rate_fc) / 100 if self.incentive_rate_fc else 0
                    total_incentive_rate_fc += incentive_rate_fc
                    swift_message_date = ''
                    if rec.swift_message_id.date:
                        swift_message_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime('%d-%b-%y')

                    new_row = table.add_row()
                    # add data to the cells in the new row
                    new_row.cells[0].text = str(sl)
                    new_row.cells[5].text = swift_message_date
                    new_row.cells[6].text = str("{:,}".format(round(rec.swift_amt, 2)))
                    new_row.cells[7].text = str("{:,}".format(float(format(incentive_rate_fc, '.2f'))))

                    first_cell = new_row.cells[0]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph24 = first_cell.paragraphs[0]
                    paragraph24.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph24.style = custom_style5

                    first_cell = new_row.cells[2]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style5

                    first_cell = new_row.cells[3]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style5

                    first_cell1 = new_row.cells[6]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = custom_style5

                    first_cell1 = new_row.cells[7]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = custom_style5

                    # first_cell1 = new_row.cells[2]
                    # first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # paragraph22 = first_cell1.paragraphs[0]
                    # paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # paragraph22.style = custom_style5

                    first_cell1 = new_row.cells[1]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.style = custom_style5

                    first_cell1 = new_row.cells[4]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5

                    first_cell1 = new_row.cells[5]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5

                    sl += 1

                cell_1_0 = table.cell(row, 2)
                p3 = cell_1_0.paragraphs[0]
                p3.add_run(invoice_ref)
                cell_2_0 = table.cell((row + gr) - 1, 2)
                cell_1_0.merge(cell_2_0)

                cell_4_0 = table.cell(row, 3)
                p3 = cell_4_0.paragraphs[0]
                p3.add_run(str("{:,}".format(round(usd_price, 2))))
                cell_3_0 = table.cell((row + gr) - 1, 3)
                cell_4_0.merge(cell_3_0)

                cell_5_0 = table.cell(row, 4)
                p3 = cell_5_0.paragraphs[0]
                p3.add_run(invoice_date)
                cell_6_0 = table.cell((row + gr) - 1, 4)
                cell_5_0.merge(cell_6_0)
                row = gr + row

            # for rec in self.invoice_line_ids:
            #     usd_price = sum(rec.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
            #     incentive_rate_fc = usd_price / self.incentive_rate_fc if self.incentive_rate_fc else 0
            #     total_incentive_rate_fc += incentive_rate_fc
            #     inv_amount += rec.invoice_amt
            #     encashment_amt_bdt += rec.encashment_amt_bdt
            #
            #     invoice_date = ''
            #     if rec.invoice_date:
            #         invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d-%b-%y')
            #     swift_message_date = ''
            #     if rec.swift_message_id.date:
            #         swift_message_date = datetime.datetime.strptime(str(rec.swift_message_id.date),
            #                                                         '%Y-%m-%d').strftime('%d-%b-%y')
            #     # add a new row to the table
            #     new_row = table.add_row()
            #
            #     # add data to the cells in the new row
            #     new_row.cells[0].text = str(sl)
            #     new_row.cells[2].text = rec.invoice_id.ref
            #     new_row.cells[3].text = str("{:,}".format(round(usd_price, 2)))
            #     new_row.cells[4].text = invoice_date
            #     new_row.cells[5].text = swift_message_date
            #     new_row.cells[6].text = str("{:,}".format(round(rec.swift_amt, 2)))
            #     new_row.cells[7].text = str("{:,}".format(float(format(incentive_rate_fc, '.2f'))))
            #
            #     first_cell = new_row.cells[0]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph24 = first_cell.paragraphs[0]
            #     paragraph24.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph24.style = custom_style5
            #
            #     first_cell = new_row.cells[3]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph21 = first_cell.paragraphs[0]
            #     paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #     paragraph21.style = custom_style5
            #
            #     first_cell1 = new_row.cells[6]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[7]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[2]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[1]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[4]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[5]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style5
            #
            #     sl += 1

            # cus name merge
            cell_1_0 = table.cell(1, 1)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run(cus_name)
            cell_2_0 = table.cell(row_len, 1)
            cell_1_0.merge(cell_2_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4

            # total
            total_row = table.add_row()
            total_row.cells[2].text = 'Total'
            total_row.cells[3].text = str("{:,}".format(round(inv_amount, 2)))
            total_row.cells[6].text = str("{:,}".format(round(swift_amt, 2)))
            total_row.cells[7].text = str("{:,}".format(float(format(total_incentive_rate_fc, '.2f'))))

            first_cell = total_row.cells[3]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = table_total

            first_cell1 = total_row.cells[6]
            paragraph22 = first_cell1.paragraphs[0]
            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph22.style = table_total

            first_cell2 = total_row.cells[7]
            paragraph23 = first_cell2.paragraphs[0]
            paragraph23.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph23.style = table_total

            first_cell2 = total_row.cells[2]
            paragraph23 = first_cell2.paragraphs[0]
            paragraph23.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph23.style = table_total

            # after table
            paragraph2 = document.add_paragraph("\nThose documents are as follows:")
            paragraph2.style = custom_style2
            paragraph2.paragraph_format.left_indent = Pt(-6)

            # paragraph2 = document.add_paragraph("\t 1. Form Ka\n\t 2. Commercial Invoice\n\t 3. Form-C (ICT)\n\t 4. Agreement\n\t 5. Certificate of Authorized Dealer\n\t 6. Copy of Swift Message\n\t 7. নগদ ভর্তুকি প্রত্যয়নপত্র \n\t 8. Company's undertaking")
            # paragraph2.style = custom_style2

            # Create a table with 1 row and 2 columns
            table = document.add_table(rows=1, cols=2)
            table.border = None
            table.columns[0].width = Inches(3)  # Adjust the width as needed
            table.columns[1].width = Inches(3)  # Adjust the width as needed

            # Get the first row of the table
            row = table.rows[0]
            # Remove cell borders
            for cell in row.cells:
                cell.border = None

            left_column = table.cell(0, 0)
            left_column = left_column.paragraphs[0]
            left_column.add_run("1. Form Ka\n")
            left_column.add_run("2. Commercial Invoice\n")
            left_column.add_run("3. Form-C (ICT)\n")
            left_column.add_run("4. Agreement\n")
            left_column.paragraph_format.left_indent = Inches(.7)
            left_column.style = custom_style2
            # left_column.paragraph_format.left_indent = Pt(-6)

            # Add the last four items to the right column
            right_column = table.cell(0, 1)
            right_column = right_column.paragraphs[0]
            right_column.add_run("5. Certificate of Authorized Dealer\n")
            right_column.add_run("6. Copy of Swift Message\n")
            run2 = right_column.add_run("7. নগদ ভর্তুকি প্রত্যয়নপত্র\n")
            run2.font.size = Pt(11)
            right_column.add_run("8. Company's undertaking\n")
            right_column.style = custom_style2
            # right_column.paragraph_format.left_indent = Pt(-6)

            paragraph2 = document.add_paragraph("Please proceed at your earliest possible time. \n\nThanking You \n\n\nMasudur Rahman \nSenior Manager, Finance, Accounts & Admin \nBrain Station 23 Limited")
            paragraph2.style = custom_style2
            paragraph2.paragraph_format.left_indent = Pt(-6)

            # file_name = '%s_flfb' % (datetime.datetime.now())
            file_name = 'Forwarding_Letter_Bank_%s_%s' % (desired_value, datetime.datetime.now())
            # html = self.get_word_prc()
            # html_text1 = html.replace('<br>', '<p style="display:block"></p>')
            # html_text1 = "<p></p><p></p><p></p><p></p> <div style='display:inline-flex'><p style='font-size:9px; margin-right:10px'>Signature of Head of the Branch</p> &nbsp;&nbsp;&nbsp; &nbsp;&nbsp;&nbsp;  <p>Signature of the Issuing Officer</p></div>"
            # new_parser.add_html_to_document(html_text1, document)

        if type == 'FORM_GA':
            from docx.enum.section import WD_ORIENT
            section = document.sections[0]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            self_date = ''
            if self.date:
                self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

            custom_bangla_style = document.styles.add_style('custom_bangla_style', WD_STYLE_TYPE.PARAGRAPH)
            custom_bangla_style.font.size = Pt(11)
            custom_bangla_style.font.name = 'SutonnyOMJ'

            custom_bangla_style_ga = document.styles.add_style('custom_bangla_style_ga', WD_STYLE_TYPE.PARAGRAPH)
            custom_bangla_style_ga.font.size = Pt(11)
            custom_bangla_style_ga.font.name = 'SutonnyOMJ'

            custom_stylebold = document.styles.add_style('custom_stylebold', WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebold.font.size = Pt(14)
            custom_stylebold.font.name = 'SutonnyOMJ'
            custom_stylebold.font.bold = True
            custom_stylebold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            custom_styleubl = document.styles.add_style('custom_styleubl', WD_STYLE_TYPE.PARAGRAPH)
            custom_styleubl.font.size = Pt(11)
            custom_styleubl.font.name = 'SutonnyOMJ'
            custom_styleubl.font.underline = True
            custom_styleubl.font.bold = True

            # p1 = document.add_paragraph('(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য) \t\t\t \t \t \t \t  \t \t  \t \t \t \t   ফরম-‘গ’ \n')
            # p1.style = custom_bangla_style

            paragraph = document.add_paragraph()
            # paragraph.add_run("(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
            run1 = paragraph.add_run("\n(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
            run1.font.size = Pt(11)
            run1.font.name = 'SutonnyOMJ'
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t ")
            paragraph.add_run("\t ")
            paragraph.add_run("\t ")
            run2 = paragraph.add_run("ফরম-‘গ’")
            run2.font.size = Pt(14)
            run2.font.bold = True
            run2.font.name = 'SutonnyOMJ'
            # paragraph.add_run("ফরম-‘গ’")
            # run2.style = custom_bangla_style_ga

            # paragraph = document.add_paragraph()
            #
            # # Add runs with different font sizes to the paragraph
            # run1 = paragraph.add_run("This is text with font size 12.")
            # run1.font.size = Pt(12)
            #
            # run2 = paragraph.add_run(" This is text with font size 16.")
            # run2.font.size = Pt(16)

            # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

            # Define a custom style
            custom_style = document.styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
            # custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            custom_style2 = document.styles.add_style('MyStyle2', WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(11)
            custom_style2.font.name = 'Arial Narrow'

            custom_style3 = document.styles.add_style('MyStyle3', WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(11)
            custom_style3.font.bold = True

            paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক")
            paragraph2.style = custom_bangla_style_ga
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            # paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক \nবাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব \n-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
            paragraph2 = document.add_paragraph("বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব")
            paragraph2.style = custom_stylebold
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
            paragraph2.style = custom_bangla_style_ga
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)

            # table -----------------------
            currency = ''
            if self.fc_currency_id:
                currency = self.fc_currency_id.name

            a = []
            c = []
            for rec in self.invoice_line_ids:
                a.append(rec.od_sight_rate)
                if rec.encashment_rate_bdt:
                    c.append(rec.encashment_rate_bdt)
                # if rec.swift_message_id.encashment_rate_bdt:
                #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
            inv_amount = sum(self.invoice_line_ids.mapped('swift_amt'))
            if a:
                b = min(a)
                od_s_r = b
            else:
                od_s_r = 0
                
            if c:
                d = min(c)
                encashment_rate_bdt = d
            else:
                encashment_rate_bdt = 0

            if self.od_sight_rate:
                od_sight_rate = self.od_sight_rate
            else:
                if od_s_r:
                    od_sight_rate = od_s_r 
                else:
                    od_sight_rate = encashment_rate_bdt
            
            rate_fc = inv_amount / self.incentive_rate_fc
            encashment_final_amt = round(rate_fc, 2) * od_sight_rate
            # encashment_final_amt_a = "{:,.0f}".format(round(encashment_final_amt))
            encashment_final_amt_a = self.custom_round(encashment_final_amt)
            usd_bdt_amount = currency + ' ' + str("{:,.2f}".format(rate_fc)) + '\n' + '@' + \
                             str("{:,.2f}".format(od_sight_rate)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a

            table = document.add_table(rows=4, cols=9)
            table.style = 'TableGrid'
            cell_1_0 = table.cell(0, 0)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run('ভর্তুকি আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানা')
            p3.style = custom_bangla_style
            cell_2_0 = table.cell(1, 0)
            cell_1_0.merge(cell_2_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(0, 1)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('রপ্তানি সংশ্লিষ্ট তথ্য')
            p3.style = custom_bangla_style
            cell_2_1 = table.cell(0, 5)
            cell_1_1.merge(cell_2_1)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 1)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('সেবা/পণ্যের বিবরণ')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 2)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('জাহাজীকরণের তারিখ ও গন্তব্যস্থল')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 3)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('ইক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানির ক্ষেত্রে)')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 4)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('মোট প্রত্যাবাসিত রপ্তানি (মূল্য  বৈদেশিক মুদ্রা)')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 5)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('নীট এফওবি মূল্য (বৈদেশিক মুদ্রা)')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_3 = table.cell(0, 6)
            p3 = cell_1_3.paragraphs[0]
            p3.add_run('শাখা কর্তৃক আবেদনপত্র গ্রহণের তারিখ')
            p3.style = custom_bangla_style
            cell_2_4 = table.cell(1, 6)
            cell_1_3.merge(cell_2_4)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_5 = table.cell(0, 7)
            p3 = cell_1_5.paragraphs[0]
            p3.add_run('আবেদনকৃত অর্থের পরিমাণ')
            p3.style = custom_bangla_style
            cell_2_6 = table.cell(1, 7)
            cell_1_5.merge(cell_2_6)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_6 = table.cell(0, 8)
            p3 = cell_1_6.paragraphs[0]
            p3.add_run('পরিশোধ্য দাবির পরিমাণ')
            p3.style = custom_bangla_style
            cell_2_7 = table.cell(1, 8)
            cell_1_6.merge(cell_2_7)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('১')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('২')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 2)
            p3 = cell.paragraphs[0]
            p3.add_run('৩')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 3)
            p3 = cell.paragraphs[0]
            p3.add_run('৪')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 4)
            p3 = cell.paragraphs[0]
            p3.add_run('৫')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 5)
            p3 = cell.paragraphs[0]
            p3.add_run('৬')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 6)
            p3 = cell.paragraphs[0]
            p3.add_run('৭')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 7)
            p3 = cell.paragraphs[0]
            p3.add_run('৮')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 8)
            p3 = cell.paragraphs[0]
            p3.add_run('৯')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 0)

            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('\n\nব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী  বাণিজ্যিক এলাকা, ঢাকা- ১২১২।\n\n')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # p3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            cell = table.cell(3, 1)
            cell.height = Inches(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Software Development')
            p3.style = custom_style2
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 2)
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 3)
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            cell = table.cell(3, 5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3

            cell = table.cell(3, 6)
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 7)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run(usd_bdt_amount)
            p3.style = custom_style3
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 8)
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            column_index = 0
            column_cells = [row.cells[column_index] for row in table.rows]

            # Set the desired height for each cell in the column
            cell_height = Inches(12)  # Adjust the height value as needed
            for cell in column_cells:
                cell.height = cell_height

            p3 = document.add_paragraph('\n\n\n\n\n\t\t দাপ্তরিক সীল  \t\t\t\t\t\t \t\t\t\t  ক্ষমতাপ্রাপ্ত ব্যাংক কর্মকর্তার স্বাক্ষর, তারিখ, নাম ও পদবী')
            p3.style = custom_bangla_style
            file_name = 'Form_GA_%s_%s' % (desired_value, datetime.datetime.now())
            # file_name = '%s_ga' % (datetime.datetime.now())

            # html = self.get_word_prc()
            # html_text1 = html.replace('<br>', '<p style="display:block"></p>')
            # html_text1 = "<p></p><p></p><p></p><p></p> <div style='display:inline-flex'><p style='font-size:9px; margin-right:10px'>Signature of Head of the Branch</p> &nbsp;&nbsp;&nbsp; &nbsp;&nbsp;&nbsp;  <p>Signature of the Issuing Officer</p></div>"
            # new_parser.add_html_to_document(html_text1, document)

        if type == 'FORM_GHA':
            from docx.enum.section import WD_ORIENT
            section = document.sections[0]
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            self_date = ''
            if self.date:
                self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')
            # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

            # Define a custom style
            custom_style = document.styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
            # custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            custom_style2 = document.styles.add_style('MyStyle2', WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(11)
            custom_style2.font.name = 'Arial Narrow'

            custom_style3 = document.styles.add_style('MyStyle3', WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(11)
            custom_style3.font.bold = True

            custom_style4 = document.styles.add_style('MyStyle4', WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.bold = True
            custom_style4.font.size = Pt(14)
            custom_style4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            custom_style4.font.name = 'SutonnyOMJ'

            custom_bangla_style = document.styles.add_style('custom_bangla_style', WD_STYLE_TYPE.PARAGRAPH)
            custom_bangla_style.font.size = Pt(11)
            custom_bangla_style.font.name = 'SutonnyOMJ'

            custom_bangla_style_gha = document.styles.add_style('custom_bangla_style_gha', WD_STYLE_TYPE.PARAGRAPH)
            custom_bangla_style_gha.font.size = Pt(11)
            custom_bangla_style_gha.font.name = 'SutonnyOMJ'

            custom_stylebold = document.styles.add_style('custom_stylebold', WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebold.font.size = Pt(16)
            custom_stylebold.font.name = 'SutonnyOMJ'
            custom_stylebold.font.bold = True
            custom_stylebold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            custom_styleubl = document.styles.add_style('custom_styleubl', WD_STYLE_TYPE.PARAGRAPH)
            custom_styleubl.font.size = Pt(13)
            custom_styleubl.font.name = 'SutonnyOMJ'
            custom_styleubl.font.bold = True
            custom_styleubl.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph1 = document.add_paragraph('\nফরম-‘ঘ’')
            paragraph1.style = custom_style4

            p3 = document.add_paragraph('(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য) \n')
            p3.style = custom_bangla_style_gha

            # paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক \nবাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধিত হিসাবের .........সালের ...... মাসের বিবরণী \nঅনুমোদিত ডিলার ব্যাংক শাখার নামঃ...................................")

            paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক")
            paragraph2.style = custom_bangla_style_gha
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও")
            paragraph2.style = custom_stylebold
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("হার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধিত হিসাবের .........সালের ...... মাসের বিবরণী")
            paragraph2.style = custom_styleubl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("অনুমোদিত ডিলার ব্যাংক শাখার নামঃ...................................................................................")
            paragraph2.style = custom_bangla_style_gha
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)

            # table -----------------------
            currency = ''
            if self.fc_currency_id:
                currency = self.fc_currency_id.name

            a = []
            c = []
            for rec in self.invoice_line_ids:
                a.append(rec.od_sight_rate)
                if rec.encashment_rate_bdt:
                    c.append(rec.encashment_rate_bdt)
                # if rec.swift_message_id.encashment_rate_bdt:
                #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
            inv_amount = sum(self.invoice_line_ids.mapped('swift_amt'))

            if a:
                b = min(a)
                od_s_r = b
            else:
                od_s_r = 0
                
            if c:
                d = min(c)
                encashment_rate_bdt = d
            else:
                encashment_rate_bdt = 0

            if self.od_sight_rate:
                od_sight_rate = self.od_sight_rate
            else:
                if od_s_r:
                    od_sight_rate = od_s_r 
                else:
                    od_sight_rate = encashment_rate_bdt

            rate_fc = inv_amount / self.incentive_rate_fc
            encashment_final_amt = round(rate_fc, 2) * od_sight_rate
            # encashment_final_amt_a = "{:,.0f}".format(round(encashment_final_amt))
            encashment_final_amt_a = self.custom_round(encashment_final_amt)
            usd_bdt_amount = currency + ' ' + str("{:,.2f}".format(rate_fc)) + '\n' + '@' + \
                             str("{:,.2f}".format(od_sight_rate)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a
            
            table = document.add_table(rows=4, cols=10)
            table.style = 'TableGrid'

            cell_1_0 = table.cell(0, 0)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run('ভর্তুকি আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানা')
            p3.style = custom_bangla_style_gha
            cell_2_0 = table.cell(1, 0)
            cell_1_0.merge(cell_2_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(0, 1)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('রপ্তানি সংশ্লিষ্ট তথ্য')
            p3.style = custom_bangla_style_gha
            cell_2_1 = table.cell(0, 5)
            cell_1_1.merge(cell_2_1)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_2_1 = table.cell(0, 8)
            p3 = cell_2_1.paragraphs[0]
            p3.add_run('শাখায় প্রাপ্ত অনিস্পন্ন আবেদন')
            p3.style = custom_bangla_style_gha
            cell_3_1 = table.cell(0, 9)
            cell_2_1.merge(cell_3_1)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 1)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('সেবা/পণ্যের বিবরণ')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 2)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('জাহাজীকরণের তারিখ ও গন্তব্যস্থল')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 3)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('ইক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানির ক্ষেত্রে)')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 4)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('মোট প্রত্যাবাসিত রপ্তানি (মূল্য  বৈদেশিক মুদ্রা)')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_1 = table.cell(1, 5)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('নীট এফওবি মূল্য (বৈদেশিক মুদ্রা)')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_3 = table.cell(0, 6)
            table.columns[7].width = Inches(.6)
            p3 = cell_1_3.paragraphs[0]
            p3.add_run('পরিশোধিত ভর্তুকি (টাকা)')
            p3.style = custom_bangla_style_gha
            cell_2_4 = table.cell(1, 6)
            cell_1_3.merge(cell_2_4)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_5 = table.cell(0, 7)
            table.columns[7].width = Inches(.6)
            p3 = cell_1_5.paragraphs[0]
            p3.add_run('পরিশোধের তারিখ')
            p3.style = custom_bangla_style_gha
            cell_2_6 = table.cell(1, 7)
            cell_1_5.merge(cell_2_6)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_6 = table.cell(1, 8)
            table.columns[8].width = Inches(.5)
            p3 = cell_1_6.paragraphs[0]
            p3.add_run('সংখ্যা')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_1_6 = table.cell(1, 9)
            table.columns[9].width = Inches(1.7)
            p3 = cell_1_6.paragraphs[0]
            p3.add_run('দাবীকৃত অর্থ (টাকা)')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('১')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('২')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 2)
            p3 = cell.paragraphs[0]
            p3.add_run('৩')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 3)
            p3 = cell.paragraphs[0]
            p3.add_run('৪')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 4)
            p3 = cell.paragraphs[0]
            p3.add_run('৫')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 5)
            p3 = cell.paragraphs[0]
            p3.add_run('৬')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 6)
            p3 = cell.paragraphs[0]
            p3.add_run('৭')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 7)
            p3 = cell.paragraphs[0]
            p3.add_run('৮')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 8)
            p3 = cell.paragraphs[0]
            p3.add_run('৯')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(2, 9)
            p3 = cell.paragraphs[0]
            p3.add_run('১০')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run(
                '\nব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী  বাণিজ্যিক এলাকা, ঢাকা- ১২১২।\n\n')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_bangla_style_gha

            cell = table.cell(3, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Software Development')
            p3.style = custom_style2
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 2)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 3)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
            p3.style = custom_style3
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
            p3.style = custom_style3
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 6)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 7)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 8)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell = table.cell(3, 9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run(usd_bdt_amount)
            p3.style = custom_style3
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p3 = document.add_paragraph(
                '\n\n\n\n\t\t দাপ্তরিক সীল  \t\t\t\t\t\t\t\t\t  ক্ষমতাপ্রাপ্ত ব্যাংক কর্মকর্তার স্বাক্ষর, তারিখ, নাম ও পদবী')
            p3.style = custom_bangla_style_gha
            file_name = 'Form_GHA_%s_%s' % (desired_value, datetime.datetime.now())
            # file_name = '%s_gha' % (datetime.datetime.now())

            # html = self.get_word_prc()
            # html_text1 = html.replace('<br>', '<p style="display:block"></p>')
            # html_text1 = "<p></p><p></p><p></p><p></p> <div style='display:inline-flex'><p style='font-size:9px; margin-right:10px'>Signature of Head of the Branch</p> &nbsp;&nbsp;&nbsp; &nbsp;&nbsp;&nbsp;  <p>Signature of the Issuing Officer</p></div>"
            # new_parser.add_html_to_document(html_text1, document)
        # new_parser.add_html_to_document(html_text1, document)
        # do more stuff to document
        # print(datetime.datetime.now())
        # docx = new_parser.parse_html_string(html_text1)

        # -------------------
        import os
        dir_path = os.path.dirname(os.path.abspath(__file__))
        base_path = str(dir_path).replace('/models', '')
        docxfile = base_path + '/static/docx/' + file_name + '.docx'
        document.save(docxfile)
        # docx.save(docxfile)

        # return document.save('/home/jobaer/Downloads/jh3.docx')

        return {
            'type': 'ir.actions.act_url',
            'url': 'cash_incentive/static/docx/' + file_name + '.docx',
            'target': 'self',
        }

    # all_form_word -----------------------------------
    def action_download_all_docx_file(self):
        type = self.env.context.get('type')
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        from docx.oxml.shared import OxmlElement, qn

        from docx.shared import Inches
        # create a new document

        # get the section object and adjust the left and right margins
        from docx.shared import Cm
        # section = document.sections[0]
        # section.left_margin = Cm(1.5)  # 1.5 cm
        # section.right_margin = Cm(1.5)  # 1.5 cm
        # jh = Document()
        # new_parser = HtmlToDocx()

        country = ''
        if self.partner_id.country_id:
            country = self.partner_id.country_id.name

        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name
        inv_number = ''
        inv_date = ''
        inv_amount = 0
        swift_amt = 0
        swift_id = None
        inv_ids = []
        for rec in self.invoice_line_ids:
            swift_amt += rec.swift_amt
            invoice_date = ''
            if rec.invoice_date:
                invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
            if rec.invoice_id not in inv_ids:
                inv_ids.append(rec.invoice_id)
                inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
                inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
            inv_amount += rec.invoice_amt
            if not swift_id:
                swift_id = rec.swift_message_id
        client_name = ''
        if self.partner_id:
            client_name = '' if not self.swift_customer_name else self.swift_customer_name

        customer_address = self.customer_address

        remiter_address = ''
        if swift_id:
            if swift_id.remiter_bank_name:
                remiter_address += swift_id.remiter_bank_name
            if swift_id.remiter_bank_address:
                remiter_address += swift_id.remiter_bank_address if not remiter_address else ', ' + swift_id.remiter_bank_address
        document = Document()
        # document.add_section()
        section = document.sections[-1]
        # section.page_width = Pt(612)
        # section.page_height = Pt(792)

        file_name = ''
        # prc 11111111111-------------------------------------------------------------------------------------------------------
        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

        document.add_paragraph('\n\nRef: ' + self.name + '\t\t\t\t\t\t\t\t\t Format-A' + '\nDate: '+ self_date + '\n' )
        # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

        # Define a custom style
        custom_style = document.styles.add_style('MyStyle', WD_STYLE_TYPE.PARAGRAPH)
        custom_style.font.bold = True
        custom_style.font.size = Pt(14)
        custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        custom_style2 = document.styles.add_style('MyStyle2', WD_STYLE_TYPE.PARAGRAPH)
        custom_style2.font.size = Pt(11)
        custom_style2.font.name = 'Arial Narrow'
        custom_style2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        custom_style3 = document.styles.add_style('MyStyle3', WD_STYLE_TYPE.PARAGRAPH)
        custom_style3.font.bold = True
        custom_style3.font.size = Pt(10)
        custom_style3.font.name = 'Arial Narrow'

        custom_style4 = document.styles.add_style('MyStyle4', WD_STYLE_TYPE.PARAGRAPH)
        custom_style4.font.size = Pt(11)
        custom_style4.font.name = 'Arial Narrow'

        # Add a paragraph and apply the custom style
        paragraph1 = document.add_paragraph('CERTIFICATE OF AUTHORIZED DEALER')
        paragraph1.style = custom_style
        contract_number = self.contract_number
        if not self.contract_number:
            contract_number = ''
        contract_date_str = self.contract_date_str
        if not self.contract_date_str:
            contract_date_str = ''
        paragraph2 = document.add_paragraph("This is to certify that we have received following inward remittance in "
                                            "the name of Brain Station 23 Limited, Plot 02  (8th Floor), Bir Uttam A. K."
                                            " Khandakar Road, Mohakhali C/A, Dhaka-1212 against "
                                            + contract_number + ' & date: ' + str(contract_date_str) + " invoice no. "
                                            + inv_number + " & date:  " + inv_date +" amount in " +currency + ' ' +  str("{:,}".format(round(inv_amount, 2)))
                                            + " for rendering of software development service. Summary of the transaction is as follows: ")
        paragraph2.style = custom_style2
        # Set the paragraph alignment to center

        table = document.add_table(rows=2, cols=9)
        table.style = 'TableGrid'

        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('(a) Remitter')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3
        # cell.merge(table.cell(0, 1))

        cell = table.cell(0, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('(b) Address')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        cell = table.cell(0, 2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('Invoice No.')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        cell = table.cell(0, 3)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        cur_name = '(c) Amount in FC (' + str(self.fc_currency_id.name) + ')'
        p3.add_run(cur_name)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        cell = table.cell(0, 4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('(d) Date of Credit in banks nostro account')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        cell = table.cell(0, 5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('(e) Equivalent Taka')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        cell = table.cell(0, 6)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('(f) Credited to beneficiary a/c')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        cell = table.cell(0, 7)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('(g) Reporting statement/schedule to BB with Month')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        cell = table.cell(0, 8)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('(h) Reference of Online reporting to BB')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        # add a dynamic row to the table
        cus_name = ''
        if self.swift_customer_name:
            cus_name = self.swift_customer_name

        # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
        row_len = len(self.invoice_line_ids)
        inv_amount = 0
        encashment_amt_bdt = 0
        equivalent_taka_total = 0

        from itertools import groupby

        # group the invoice_line_ids by invoice_id.ref
        groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                         key=lambda x: x.invoice_id.ref)
        # iterate over each group and add a new row for each unique invoice_id.ref
        row = 1
        swift_amt_total = 0
        for invoice_ref, group in groups:
            gr = 0
            for rec in group:
                gr += 1
                new_row = table.add_row()
                inv_amount += rec.invoice_amt
                swift_amt_total += rec.swift_amt
                encashment_amt_bdt += rec.encashment_amt_bdt

                swift_date = ''
                if rec.swift_message_id.date:
                    swift_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime(
                        '%d-%b-%y')

                date_credited_beneficiaries = ''
                if rec.date_credited_beneficiaries:
                    date_credited_beneficiaries = datetime.datetime.strptime(str(rec.date_credited_beneficiaries),
                                                                             '%Y-%m-%d').strftime('%d-%b-%y')
                reporting_st_to_bb = rec.reporting_st_to_bb if rec.reporting_st_to_bb else ''
                ref_online_to_bb = rec.ref_online_to_bb if rec.ref_online_to_bb else ''
                equivalent_taka = round(rec.swift_amt * rec.encashment_rate_bdt, 2)
                equivalent_taka_total += equivalent_taka

                equivalent_taka_decimal = "{:,.2f}".format(equivalent_taka)

                # add data to the cells in the new row
                # new_row.cells[2].text = rec.invoice_id.ref
                new_row.cells[3].text = str("{:,.2f}".format(rec.swift_amt))
                new_row.cells[4].text = str(swift_date)
                new_row.cells[5].text = str(equivalent_taka_decimal)
                new_row.cells[6].text = str(date_credited_beneficiaries)
                new_row.cells[7].text = str(reporting_st_to_bb)
                new_row.cells[8].text = str(ref_online_to_bb)

                first_cell = new_row.cells[3]
                first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph21.style = custom_style4

                first_cell1 = new_row.cells[5]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph22.style = custom_style4

                first_cell1 = new_row.cells[2]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # paragraph22.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph22.style = custom_style4

                first_cell1 = new_row.cells[0]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.style = custom_style4
                first_cell1 = new_row.cells[1]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.style = custom_style4
                first_cell1 = new_row.cells[4]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style4
                first_cell1 = new_row.cells[6]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style4
                first_cell1 = new_row.cells[7]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style4
                first_cell1 = new_row.cells[8]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style4

            cell_1_0 = table.cell(row, 2)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run(invoice_ref)
            cell_2_0 = table.cell((row + gr) - 1, 2)
            cell_1_0.merge(cell_2_0)
            row = gr + row

        # for rec in self.invoice_line_ids:
        #     inv_amount += rec.invoice_amt
        #     encashment_amt_bdt += rec.encashment_amt_bdt
        #
        #     swift_date = ''
        #     if rec.swift_message_id.date:
        #         swift_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime(
        #             '%d-%b-%y')
        #
        #     date_credited_beneficiaries = ''
        #     if rec.date_credited_beneficiaries:
        #         date_credited_beneficiaries = datetime.datetime.strptime(str(rec.date_credited_beneficiaries),
        #                                                                  '%Y-%m-%d').strftime('%d-%b-%y')
        #     reporting_st_to_bb = rec.reporting_st_to_bb if rec.reporting_st_to_bb else ''
        #     ref_online_to_bb = rec.ref_online_to_bb if rec.ref_online_to_bb else ''
        #     equivalent_taka = round(rec.swift_amt * rec.encashment_rate_bdt, 2)
        #     equivalent_taka_total += equivalent_taka
        #     # add a new row to the table
        #     new_row = table.add_row()
        #
        #     # add data to the cells in the new row
        #     new_row.cells[2].text = rec.invoice_id.ref
        #     new_row.cells[3].text = str("{:,}".format(round(rec.invoice_amt, 2)))
        #     new_row.cells[4].text = str(swift_date)
        #     new_row.cells[5].text = str("{:,}".format(round(equivalent_taka, 2)))
        #     new_row.cells[6].text = str(date_credited_beneficiaries)
        #     new_row.cells[7].text = str(reporting_st_to_bb)
        #     new_row.cells[8].text = str(ref_online_to_bb)
        #
        #     first_cell = new_row.cells[3]
        #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph21 = first_cell.paragraphs[0]
        #     paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        #     paragraph21.style = custom_style4
        #
        #     first_cell1 = new_row.cells[5]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        #     paragraph22.style = custom_style4
        #
        #     first_cell1 = new_row.cells[2]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     # paragraph22.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        #     paragraph22.style = custom_style4
        #
        #     first_cell1 = new_row.cells[0]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.style = custom_style4
        #     first_cell1 = new_row.cells[1]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.style = custom_style4
        #     first_cell1 = new_row.cells[4]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     paragraph22.style = custom_style4
        #     first_cell1 = new_row.cells[6]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     paragraph22.style = custom_style4
        #     first_cell1 = new_row.cells[7]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     paragraph22.style = custom_style4
        #     first_cell1 = new_row.cells[8]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     paragraph22.style = custom_style4
        # from itertools import groupby
        #
        # # group the invoice_line_ids by invoice_id.ref
        # groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
        #                  key=lambda x: x.invoice_id.ref)
        # print(groups)
        # # iterate over each group and add a new row for each unique invoice_id.ref
        # row = 1
        # for invoice_ref, group in groups:
        #     new_row = table
        #     print(invoice_ref)
        #     print(group)
        #     # add data to the cells in the new row
        #     gr = 0
        #     for rec in group:
        #         gr += 1
        #         new_row = table.add_row()
        #         # add data to the cells in the new row for each row in the group
        #         new_row.cells[0].text = str("{:,}".format(round(rec.invoice_amt, 2)))
        #         new_row.cells[1].text = str("{:,}".format(round(rec.invoice_amt, 2)))
        #         new_row.cells[3].text = str("{:,}".format(round(rec.invoice_amt, 2)))
        #         new_row.cells[4].text = str(12)
        #         new_row.cells[5].text = str("{:,}".format(round(100, 2)))
        #         new_row.cells[6].text = str(200)
        #         new_row.cells[7].text = str(290)
        #         new_row.cells[8].text = str(12)
        #         # row += 1
        #     # new_row.cells[2].text = invoice_ref
        #     cell_1_0 = table.cell(row, 2)
        #     p3 = cell_1_0.paragraphs[0]
        #     p3.add_run(invoice_ref)
        #     cell_2_0 = table.cell((row + gr) - 1, 2)
        #     cell_1_0.merge(cell_2_0)
        #     row = gr + row

        # cus name merge
        cell_1_0 = table.cell(1, 0)
        p3 = cell_1_0.paragraphs[0]
        p3.add_run(cus_name)
        cell_2_0 = table.cell(row_len, 0)
        cell_1_0.merge(cell_2_0)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        # address merge
        cell_3_0 = table.cell(1, 1)
        p3 = cell_3_0.paragraphs[0]
        p3.add_run(self.customer_address)
        cell_4_0 = table.cell(row_len, 1)
        cell_3_0.merge(cell_4_0)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        # total
        equivalent_taka_total_decimal = "{:,.0f}".format(equivalent_taka_total)
        total_row = table.add_row()
        total_row.cells[2].text = 'Total'
        total_row.cells[3].text = "{:,.2f}".format(swift_amt_total)
        total_row.cells[5].text = str(equivalent_taka_total_decimal)

        first_cell = total_row.cells[2]
        paragraph21 = first_cell.paragraphs[0]
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph21.style = custom_style3

        first_cell = total_row.cells[3]
        paragraph21 = first_cell.paragraphs[0]
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph21.style = custom_style3

        first_cell1 = total_row.cells[5]
        paragraph22 = first_cell1.paragraphs[0]
        paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph22.style = custom_style3

        paragraph = document.add_paragraph()
        paragraph.add_run("\n\n\n\n\n Signature of Head of the branch.")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("Signature of the Issuing Officer.")

        document.add_page_break()
        # file_name = '%s_prc' % (datetime.datetime.now())
        #forwarding 2222222222-------------------------------------------------
        section.page_width = Cm(22)  # set the page width to 21 centimeters
        section.page_height = Cm(29.9)  # set the page width to 21 centimeters
        section.left_margin = Cm(3)  # 1.5 cm
        section.right_margin = Cm(2)
        # Define a custom style
        custom_style = document.styles.add_style('MyStyle5', WD_STYLE_TYPE.PARAGRAPH)
        custom_style.font.bold = True
        custom_style.font.size = Pt(14)
        custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        custom_style.font.name = 'Arial Narrow'

        custom_style2 = document.styles.add_style('MyStyle6', WD_STYLE_TYPE.PARAGRAPH)
        custom_style2.font.name = 'Arial Narrow'
        custom_style2.font.size = Pt(11)

        custom_style3 = document.styles.add_style('MyStyle7', WD_STYLE_TYPE.PARAGRAPH)
        custom_style3.font.size = Pt(11)
        custom_style3.font.name = 'Arial Narrow'
        custom_style3.font.bold = True

        # Add a paragraph and apply the custom style
        # paragraph1 = document.add_paragraph('CERTIFICATE OF AUTHORIZED DEALER \n ')
        # paragraph1.style = custom_style

        # paragraph2 = document.add_paragraph("Ref No.: " + self.name + " \nDate: "+ self_date + '\n\n")
        current_date = datetime.datetime.now().date()
        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')
        paragraph2 = document.add_paragraph('\n\n\n\nRef: ' + self.name + '\nDate: ' + str(current_date) + '\n')
        paragraph2.style = custom_style2

        paragraph4 = document.add_paragraph("The President \nBASIS \nBDBL Bhabon (5th Floor-West), \n12 Kawran Bazar, Dhaka-1215 \n")
        paragraph4.style = custom_style2

        paragraph5 = document.add_paragraph("Subject: Request to Issue BASIS Certificate for Cash Assistance.\n")
        paragraph5.style = custom_style3

        paragraph6 = document.add_paragraph("Dear Sir,")
        paragraph6.style = custom_style2
        paragraph6.paragraph_format.space_after = Pt(0)

        contract_number = ''
        if self.contract_number:
            contract_number = " and "+ self.contract_number

        paragraph6 = document.add_paragraph("With reference to the above-mentioned subject, we would like to draw your kind attention to the fact that we are going to draw cash subsidy against our following Invoice No: " + inv_number + contract_number + ".\n")
        paragraph6.style = custom_style2
        paragraph6.style = custom_style2

        paragraph6 = document.add_paragraph("There is a checklist as following:  ")
        paragraph6.style = custom_style2

        # Create a list of items
        items = ["Forwarding Letter Addressing BASIS president", "Export Agreement", "Commercial Invoice Related to Export", "Format -A Certificate of Authorized Dealer Issued by Bank",
                 "ICT form-c", "Company's Undertaking", "Form-Ka (Bangladesh Bank)", "Form-Kha (Bangladesh Bank)", "Pay order"]

        # Create a paragraph object for each list item and set the paragraph style to "List Bullet"
        # for item in items:
        #     paragraph7 = document.add_paragraph(style="List Bullet")
        #     paragraph7.add_run(item)

        left_indent = Inches(.8)  # Adjust the left indentation as needed

        for item in items:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = left_indent

            run = paragraph.add_run(item)
            run.font.name = 'Arial Narrow'
            paragraph.paragraph_format.line_spacing = Pt(12)

        paragraph8 = document.add_paragraph("\nYour kind co-operation in this regard will be highly appreciated.\n\nThanking You.\n\nYours Truly\n\n\n\n\nMasudur Rahman \nSenior Manager, Finance, Accounts & Admin \nBrain Station 23 Ltd. \nPlot-2, (5th Floor), Bir Uttam Ak Khandaker Road \nMohakhali C/A, Dhaka-1212")
        paragraph8.style = custom_style2
        paragraph8.paragraph_format.space_before = Pt(0)
        paragraph8.paragraph_format.space_after = Pt(0)
        document.add_page_break()
        # file_name = '%s_flbs' % (datetime.datetime.now())

        # ka 333333333333333 -------------------------------------------
        section.page_width = Cm(22)  # set the page width to 21 centimeters
        section.page_height = Cm(29.9)  # set the page width to 21 centimeters
        section.left_margin = Cm(2)  # 1.5 cm
        section.right_margin = Cm(2)

        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

        # end bangla font maker ====================================
        # Define a custom style
        custom_style = document.styles.add_style('MyStyle8', WD_STYLE_TYPE.PARAGRAPH)
        # custom_style.font.bold = True
        custom_style.font.size = Pt(14)
        custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        custom_style2 = document.styles.add_style('MyStyle9', WD_STYLE_TYPE.PARAGRAPH)
        custom_style2.font.size = Pt(11)

        custom_style3 = document.styles.add_style('MyStyle10', WD_STYLE_TYPE.PARAGRAPH)
        custom_style3.font.size = Pt(11)
        custom_style3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        custom_style4 = document.styles.add_style('MyStyle11', WD_STYLE_TYPE.PARAGRAPH)
        custom_style4.font.bold = True
        custom_style4.font.size = Pt(9)
        custom_style4.font.name = 'Arial Narrow'

        tab_total = document.styles.add_style('tab_total12', WD_STYLE_TYPE.PARAGRAPH)
        tab_total.font.bold = True
        tab_total.font.size = Pt(9)
        tab_total.font.name = 'Arial Narrow'

        custom_styleextra = document.styles.add_style('custom_styleextra13', WD_STYLE_TYPE.PARAGRAPH)
        custom_styleextra.font.size = Pt(9)
        custom_styleextra.font.name = 'Arial Narrow'

        custom_style5 = document.styles.add_style('MyStyle14', WD_STYLE_TYPE.PARAGRAPH)
        custom_style5.font.size = Pt(9)
        custom_style5.font.name = 'Arial Narrow'

        custom_stylebl = document.styles.add_style('custom_stylebl15', WD_STYLE_TYPE.PARAGRAPH)
        custom_stylebl.font.size = Pt(9)
        custom_stylebl.font.name = 'SutonnyOMJ'

        custom_na = document.styles.add_style('custom_na16', WD_STYLE_TYPE.PARAGRAPH)
        custom_na.font.size = Pt(9)
        custom_na.font.name = 'Arial Narrow'

        custom_stylebold = document.styles.add_style('custom_stylebold17', WD_STYLE_TYPE.PARAGRAPH)
        custom_stylebold.font.size = Pt(11)
        custom_stylebold.font.name = 'SutonnyOMJ'
        custom_stylebold.font.bold = True

        custom_styleubl = document.styles.add_style('custom_styleubl18', WD_STYLE_TYPE.PARAGRAPH)
        custom_styleubl.font.size = Pt(11)
        custom_styleubl.font.name = 'SutonnyOMJ'
        custom_styleubl.font.underline = True
        custom_styleubl.font.bold = True

        document.styles['Normal'].font.size = Pt(8)
        paragraph = document.add_paragraph()
        paragraph.add_run("অনুচ্ছেদ ০৫(খ) এফই সার্কুলার নং-০৩/২০১৮ দ্রষ্টব্য)")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run(' ' * 7 + "ফরম-‘ক’")
        paragraph.style = custom_stylebl

        # Add a paragraph and apply the custom style
        paragraph1 = document.add_paragraph('বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও')
        paragraph1.style = custom_stylebold
        paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph1.paragraph_format.space_before = Pt(0)
        paragraph1.paragraph_format.space_after = Pt(0)
        paragraph1 = document.add_paragraph('হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকির জন্য আবেদনপত্র ')
        paragraph1.style = custom_styleubl
        paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph1.paragraph_format.space_before = Pt(0)

        paragraph2 = document.add_paragraph("(ক) আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানাঃ ব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী বাণিজ্যিক এলাকা, ঢাকা - ১২১২। রপ্তানি নিবন্ধন সনদপত্র (ইআরসি) নম্বরঃ ২৬০৩২৬২১০৬৬৬৪২০")
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(10)

        con_no = ''
        if self.contract_number:
            con_no = self.contract_number.replace('Contract No.', '').replace('Purchase Order No.', '')
        paragraph2 = document.add_paragraph("(খ) রপ্তানি চুক্তিপত্র নম্বরঃ"+ con_no + " \nতারিখঃ " + str(self.contract_date_str) + "\nমূল্যঃ " +self.contract_price_str + "\n(পাঠ্যযোগ সত্যায়িত কপি দাখিল করতে হবে)" )
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(10)
        paragraph2 = document.add_paragraph("(গ) রপ্তানিকৃত সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের স্থানীয় সংগ্রহসূত্র, পরিমাণ ও মূল্যঃ" )
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(10)
        # Set the paragraph alignment to center

        # table1 ------------------------------
        table5 = document.add_table(rows=3, cols=3)
        table5.alignment = 2
        table5.left_indent = Inches(100)
        table5.style = 'TableGrid'
        table5.autofit = False  # Disable automatic column width adjustment
          # Set the width of the table to 6 inches (adjust the value according to your desired width)

        cell = table5.cell(0, 0)
        p3 = cell.paragraphs[0]
        p3.add_run('সরবরাহকারীর নাম ও ঠিকানা')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_stylebl

        cell = table5.cell(0, 1)
        p1 = cell.paragraphs[0]
        p1.add_run('পরিমাণ')
        p1.style = custom_stylebl
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table5.cell(0, 2)
        p2 = cell.paragraphs[0]
        p2.add_run('মূল্য')
        p2.style = custom_stylebl
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table5.cell(1, 0)
        p4 = cell.paragraphs[0]
        p4.add_run('১')
        p4.style = custom_stylebl
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table5.cell(1, 1)
        p5 = cell.paragraphs[0]
        p5.add_run('২')
        p5.style = custom_stylebl
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table5.cell(1, 2)
        p6 = cell.paragraphs[0]
        p6.add_run('৩')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table5.cell(2, 0)
        p7 = cell.paragraphs[0]
        p7.add_run('N/A')
        p7.style = custom_na
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table5.cell(2, 1)
        p8 = cell.paragraphs[0]
        p8.add_run('N/A')
        p8.style = custom_na
        p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table5.cell(2, 2)
        p9 = cell.paragraphs[0]
        p9.add_run('N/A')
        p9.style = custom_na
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # table 2 ----------------------------
        paragraph2 = document.add_paragraph("(রপ্তানিকৃত সেবা/পণ্যের বর্ণনা, মূল্য ও সংগ্রহসূত্রের বিষয়ে সেবা/পণ্য সংশ্লিষ্ট এসোসিয়েশন এর প্রত্যয়নপত্র দাখিল করতে হবে) \n(ঘ) রপ্তানিকৃত সেবা/পণ্য উৎপাদনে ব্যবহৃত আমাদানিকৃত সেবা/উপকরণাদিঃ" )
        paragraph2.style = custom_stylebl
        # paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(12)

        # Set the paragraph alignment to center
        table1 = document.add_table(rows=3, cols=4)
        table1.autofit = False
        table1.width = Inches(6)
        table1.alignment = 2
        table1.style = 'TableGrid'

        cell = table1.cell(0, 0)
        p3 = cell.paragraphs[0]
        p3.add_run('সরবরাহকারীর নাম ও ঠিকানা')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table1.columns[0].width = Inches(.6)
        p3.style = custom_stylebl

        cell = table1.cell(0, 1)
        p3 = cell.paragraphs[0]
        p3.add_run('সেবা/পণ্যের নাম ও পরিমাণ ')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table1.columns[1].width = Inches(.6)
        p3.style = custom_stylebl

        cell = table1.cell(0, 2)
        p2 = cell.paragraphs[0]
        p2.add_run('ঋণপত্র/ব্যাংক টু ব্যাংক ঋণপত্র/ডকুমেন্টরী কালেকশন/টিটি রেমিটেন্স নম্বর, তারিখ')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table1.columns[2].width = Inches(5.2)
        p2.style = custom_stylebl

        cell = table1.cell(0, 3)
        p2 = cell.paragraphs[0]
        p2.add_run('মূল্য')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table1.columns[3].width = Inches(1)
        p2.style = custom_stylebl

        cell = table1.cell(1, 0)
        p4 = cell.paragraphs[0]
        p4.add_run('১')
        p4.style = custom_stylebl
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table1.cell(1, 1)
        p5 = cell.paragraphs[0]
        p5.add_run('২')
        p5.style = custom_stylebl
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table1.cell(1, 2)
        p6 = cell.paragraphs[0]
        p6.add_run('৩')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table1.cell(1, 3)
        p6 = cell.paragraphs[0]
        p6.add_run('৪')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table1.cell(2, 0)
        p7 = cell.paragraphs[0]
        p7.add_run('N/A')
        p7.style = custom_na
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table1.cell(2, 1)
        p8 = cell.paragraphs[0]
        p8.add_run('N/A')
        p8.style = custom_na
        p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table1.cell(2, 2)
        p9 = cell.paragraphs[0]
        p9.add_run('N/A')
        p9.style = custom_na
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table1.cell(2, 3)
        p9 = cell.paragraphs[0]
        p9.add_run('N/A')
        p9.style = custom_na
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # table 3 ----------------------
        paragraph2 = document.add_paragraph("(৩) নং কলামের ঋণপত্রে পাঠযোগ্য সত্যায়িত কপি দাখিল করতে হবে। সেবা আমাদানির ক্ষেত্রে যথাযথ পদ্ধতি অনুসরণ করা হয়েছে মর্মে অনুমোদিত ডিলার শাখাকে নিশ্চিত হতে হবে। উৎপাদন প্রক্রিয়ায় ব্যবহৃত উপকরণাদির জন্য শুল্ক বন্ড সুবিধা ভোগ করা হয়নি/ডিউটি ড্র-ব্যাংক সুবিধা গ্রহণ করা হয়নি ও ভবিষ্যতে আবেদনও করা হবে না মর্মে রপ্তানিকরাকের ঘোষণাপত্র দাখিল করতে হবে।) \n(ঙ) রপ্তানি চালানের বিবরণঃ" )
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(12)

        table = document.add_table(rows=3, cols=8)
        table.alignment = 2  # Left alignment
        # table.left_indent = Inches(5)
        table.style = 'TableGrid'
        # table.autofit = False
        # table.width = Inches(6) it is worked...............
        # table.autofit = False
        # table.width = Inches(4)
        # table.alignment = 1

        column_width = Inches(.2)
        table.columns[0].width = column_width
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p10 = cell.paragraphs[0]
        p10.add_run('   পণ্যের বর্ণনা   ')
        p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p10.style = custom_stylebl
        # table.columns[0].width = Inches(.3)
        # cell.merge(table.cell(0, 1))

        column_width = Inches(1.8)
        table.columns[1].width = column_width
        cell = table.cell(0, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p10 = cell.paragraphs[0]
        p10.add_run('পরিমাণ')
        p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p10.style = custom_stylebl
        # table.columns[1].width = Inches(1.69)

        column_width = Inches(.1)
        table.columns[2].width = column_width
        cell = table.cell(0, 2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p10 = cell.paragraphs[0]
        p10.add_run('আমদানিকারকের দেশের নাম')
        p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p10.style = custom_stylebl
        # table.columns[2].width = Inches(.8)

        table.columns[3].width = Inches(.6)
        cell = table.cell(0, 3)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p10 = cell.paragraphs[0]
        p10.add_run('ইনভয়েস মূল্য (বৈদেশিক মুদ্রায়)')
        p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p10.style = custom_stylebl

        table.columns[4].width = Inches(.6)
        cell = table.cell(0, 4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p10 = cell.paragraphs[0]
        p10.add_run('জাহাজীকরণ/রপ্তানির তারিখ')
        p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p10.style = custom_stylebl

        column_width = Inches(.2)
        table.columns[5].width = column_width
        cell = table.cell(0, 5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p10 = cell.paragraphs[0]
        p10.add_run('   ইএক্সপি নম্বর*   ')
        p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p10.style = custom_stylebl

        table.columns[6].width = Inches(.8)
        table.columns[7].width = Inches(.3)
        cell_1_0 = table.cell(0, 6)
        cell_1_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p10 = cell_1_0.paragraphs[0]
        p10.add_run('বৈদেশিক মুদ্রায় প্রত্যাবাসিত রপ্তানিমূল্য ও প্রত্যাবাসনের তারিখ')
        cell_2_0 = table.cell(0, 7)
        cell_1_0.merge(cell_2_0)
        p10.style = custom_stylebl
        p10.alignment = WD_ALIGN_PARAGRAPH.CENTER

        left_indent = Inches(2)  # Adjust the indent as needed

        for row in table.rows:
            for cell in row.cells:
                cell.left_indent = left_indent

        # cus name merge
        # cell_1_0 = table.cell(0, 6)
        # p10 = cell_1_0.paragraphs[0]
        # p10.add_run('ৈদেশিক মুদ্রায় প্রত্যাবাসিত রপ্তানিমূল্য ও প্রত্যাবাসনের তারিখ')
        # cell_2_0 = table.cell(0, 7)
        # cell_1_0.merge(cell_2_0)

        # row 2 ----------------------------
        cell = table.cell(1, 0)
        p4 = cell.paragraphs[0]
        p4.add_run('১')
        p4.style = custom_stylebl
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 1)
        p5 = cell.paragraphs[0]
        p5.add_run('২')
        p5.style = custom_stylebl
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 2)
        p6 = cell.paragraphs[0]
        p6.add_run('৩')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 3)
        p6 = cell.paragraphs[0]
        p6.add_run('৪')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 4)
        p6 = cell.paragraphs[0]
        p6.add_run('৫')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 5)
        p6 = cell.paragraphs[0]
        p6.add_run('৬')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 6)
        p10 = cell_1_1.paragraphs[0]
        p10.add_run('৭')
        p6.style = custom_stylebl
        cell_2_1 = table.cell(1, 7)
        cell_1_1.merge(cell_2_1)
        p10.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # cell_1_1 = table.cell(1, 6)
        # p10 = cell_1_1.paragraphs[0]
        # p10.add_run('৭')
        # cell_2_1 = table.cell(1, 7)
        # cell_1_1.merge(cell_2_1)

        # # add a dynamic row to the table
        # inv_date = ''
        inv_amount = 0
        incentive_rate_fc = 0
        swift_amt = 0
        incentive_amt_fc = 0
        encashment_amt_bdt = 0
        total_incentive_rate_fc = 0
        sl = 0

        inv_ids = []
        dev_total1 = 0
        hour_total1 = 0
        a = []
        c = []
        for line in self.invoice_line_ids:
            a.append(line.od_sight_rate)
            c.append(line.encashment_rate_bdt)
            if line.invoice_id.id not in inv_ids:
                dev_total1 += sum(
                    r.quantity if r.quantity_type == '1' else 0 for r in line.invoice_id.invoice_line_ids)
                hour_total1 += sum(
                    r.quantity if r.quantity_type == '0' else 0 for r in line.invoice_id.invoice_line_ids)
            inv_ids.append(line.invoice_id.id)

        if a:
            b = min(a)
            od_s_r = b
        else:
            od_s_r = 0

        if c:
            d = min(c)
            encashment_rate_bdt = d
        else:
            encashment_rate_bdt = 0

        # dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in
        #                  self.invoice_line_ids.invoice_id.invoice_line_ids)
        # hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in
        #                   self.invoice_line_ids.invoice_id.invoice_line_ids)
        hour_total = float(format(hour_total1, '.2f'))
        dev_total = float(format(dev_total1, '.2f'))

        if hour_total == int(hour_total):
            hour_total_num = "{:.0f}".format(hour_total)
        else:
            hour_total_num = "{:.2f}".format(hour_total)

        if dev_total == int(dev_total):
            dev_total_num = "{:.0f}".format(dev_total)
        else:
            dev_total_num = "{:.2f}".format(dev_total)

        all_invoice_qty_str = ''
        if dev_total > 0:
            all_invoice_qty_str += str('ITES/ ') + str(dev_total_num) if not all_invoice_qty_str else ' & ' + str(
                'ITES/ ') + str(
                dev_total_num)
        if hour_total > 0:
            all_invoice_qty_str += str(hour_total_num) + str(' HRS') if not all_invoice_qty_str else ' & ' + str(
                hour_total_num) + str(' HRS')

        currency_symbol = ''
        if self.fc_currency_id.symbol:
            currency_symbol = self.fc_currency_id.symbol
        country_name = ''
        if self.partner_id.country_id:
            country_name = self.partner_id.country_id.name
        row_len = len(self.invoice_line_ids)

        from itertools import groupby
        # group the invoice_line_ids by invoice_id.ref
        groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                         key=lambda x: x.invoice_id.ref)
        # iterate over each group and add a new row for each unique invoice_id.ref
        row = 2

        for invoice_ref, group in groups:
            inv_obj = self.env['cash.incentive.invoice'].search(
                [('invoice_id.ref', '=', invoice_ref), ('head_id', '=', self.id)], limit=1)
            # usd_price = sum(inv_obj.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
            usd_price = inv_obj.invoice_id.invoice_total_fc
            incentive_rate_fc = usd_price / self.incentive_rate_fc if self.incentive_rate_fc else 0
            total_incentive_rate_fc += incentive_rate_fc
            inv_amount += inv_obj.invoice_amt
            encashment_amt_bdt += inv_obj.encashment_amt_bdt
            invoice_date = ''
            swift_message_date = ''
            if inv_obj.invoice_date:
                invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime(
                    '%d-%b-%y')
            if inv_obj.swift_message_id.date:
                swift_message_date = datetime.datetime.strptime(str(inv_obj.swift_message_id.date),
                                                                '%Y-%m-%d').strftime('%d-%b-%y')

            # getting hours and developers ===========
            l_dev_total = sum(
                r.quantity if r.quantity_type == '1' else 0 for r in inv_obj.invoice_id.invoice_line_ids)
            l_hour_total = sum(
                r.quantity if r.quantity_type == '0' else 0 for r in inv_obj.invoice_id.invoice_line_ids)
            l_all_invoice_qty_str = ''
            if l_dev_total > 0:
                if l_dev_total == int(l_dev_total):
                    l_dev_total = "{:.0f}".format(l_dev_total)
                else:
                    l_dev_total = "{:.2f}".format(l_dev_total)
                l_all_invoice_qty_str += str('ITES/ ') + str(
                    l_dev_total) if not l_all_invoice_qty_str else ' & ' + str(
                    'ITES/ ') + str(l_dev_total)
            if l_hour_total > 0:
                if l_hour_total == int(l_hour_total):
                    l_hour_total = "{:.0f}".format(l_hour_total)
                else:
                    l_hour_total = "{:.2f}".format(l_hour_total)
                l_all_invoice_qty_str += str(l_hour_total) + str(
                    ' HRS') if not l_all_invoice_qty_str else ' & ' + str(
                    l_hour_total) + str(' HRS')
            sl += 1
            gr = 0
            for rec in group:
                swift_amt += rec.swift_amt
                incentive_amt_fc += rec.incentive_amt_fc
                # if rec.swift_message_id.encashment_rate_bdt:
                #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
                gr += 1
                new_row = table.add_row()
                new_row.cells[6].text = str(currency_symbol) + ' ' + str("{:,.2f}".format(rec.swift_amt))
                new_row.cells[7].text = str(swift_message_date)

                first_cell = new_row.cells[3]
                first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph21.style = custom_style5

                first_cell1 = new_row.cells[6]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph22.style = custom_style5

                first_cell1 = new_row.cells[0]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.style = custom_style5
                first_cell1 = new_row.cells[1]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style5
                first_cell1 = new_row.cells[2]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style5
                first_cell1 = new_row.cells[4]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style5
                first_cell1 = new_row.cells[5]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style5
                first_cell1 = new_row.cells[7]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style5

            cell1_1_0 = table.cell(row, 1)
            p3 = cell1_1_0.paragraphs[0]
            p3.add_run(l_all_invoice_qty_str)
            cell1_2_0 = table.cell((row + gr) - 1, 1)
            cell1_1_0.merge(cell1_2_0)

            cell_4_0 = table.cell(row, 3)
            p3 = cell_4_0.paragraphs[0]
            p3.add_run(str(currency_symbol)+ ' ' + str("{:,.2f}".format(inv_obj.invoice_amt)))
            cell_3_0 = table.cell((row + gr) - 1, 3)
            cell_4_0.merge(cell_3_0)

            cell_5_0 = table.cell(row, 4)
            p3 = cell_5_0.paragraphs[0]
            p3.add_run(invoice_date)
            cell_6_0 = table.cell((row + gr) - 1, 4)
            cell_5_0.merge(cell_6_0)
            row = gr + row

        # #column merge -----------------------
        column_width = Inches(.3)
        # table.columns[0].width = column_width
        cell_1_2 = table.cell(2, 0)
        cell_1_2.width = column_width
        p3 = cell_1_2.paragraphs[0]
        p3.add_run('Software Development')
        cell_2_2 = table.cell(1+row_len, 0)
        cell_1_2.merge(cell_2_2)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_styleextra

        cell_1_2 = table.cell(2, 2)
        p3 = cell_1_2.paragraphs[0]
        p3.add_run(country_name)
        cell_2_2 = table.cell(1+row_len, 2)
        cell_1_2.merge(cell_2_2)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_styleextra

        cell_1_2 = table.cell(2, 5)
        p3 = cell_1_2.paragraphs[0]
        p3.add_run('N/A')
        cell_2_2 = table.cell(1+row_len, 5)
        cell_1_2.merge(cell_2_2)
        p3.style = custom_styleextra

        # total ---------------------------
        total_row = table.add_row()
        total_row.cells[0].text = 'Total'
        total_row.cells[1].text = all_invoice_qty_str

        # inv_amount_a = "{:.2f}".format(inv_amount).rstrip('0').rstrip('.') + ('0' if inv_amount % 1 else '')
        inv_amount_a = str("{:,.2f}".format(inv_amount))
        total_row.cells[3].text = str(currency_symbol)+ ' ' + inv_amount_a

        # total_row.cells[3].text = str(currency_symbol)+ ' ' + str("{:,}".format(round(inv_amount, 2)))
        # formatted_number = "{:.2f}".format(inv_amount).rstrip('0').rstrip('.') + ('0' if inv_amount % 1 else '')

        # swift_amt_a = "{:.2f}".format(swift_amt).rstrip('0').rstrip('.') + ('0' if swift_amt % 1 else '')
        swift_amt_a = str("{:,.2f}".format(swift_amt))
        total_row.cells[6].text = str(currency_symbol)+ ' ' + swift_amt_a

        first_cell = total_row.cells[0]
        paragraph21 = first_cell.paragraphs[0]
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph21.style = tab_total

        first_cell = total_row.cells[1]
        paragraph21 = first_cell.paragraphs[0]
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph21.style = tab_total

        first_cell = total_row.cells[3]
        paragraph21 = first_cell.paragraphs[0]
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph21.style = tab_total

        first_cell1 = total_row.cells[6]
        paragraph22 = first_cell1.paragraphs[0]
        paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph22.style = tab_total

        # table 4 ----------------------------------
        paragraph2 = document.add_paragraph("*দৃশ্যমান আকারে পণ্য রপ্তানির ক্ষেত্রে প্রযোজ্য \n (কমার্শিয়াল ইনভয়েস, প্যাকিং লিষ্ট এবং জাহাজীকরণের প্রমাণ স্বরূপ পরিবহন কর্তৃপক্ষ ইস্যুকৃত এবং প্রত্যয়নকৃত বিল অব লোডিং/এয়ারওয়ে বিল, বিল অব এক্সপোর্ট (শুল্ক কর্তৃপক্ষ কর্তৃক ইস্যুকৃত ও পীক্ষিত এবং on-hand হওয়ার স্বপক্ষে পরিবহন কর্তৃপক্ষ প্রত্যয়নকৃত) এর পূর্ণাঙ্গ সেট ইত্যাদির সত্যায়িত পাঠযোগ্য কপি এবং রপ্তানিমূল্য প্রত্যাবাসন সনদপত্র দাখিল করতে হবে। তবে অদৃশ্যকারে সেবা রপ্তানির ক্ষেত্রে জাহাজীকরণের দলিল ও বিল অব এক্সপোর্ট  আবশ্যকতা থাকবে না।) \n(চ) ভর্তুকির আবেদনকৃত অংকঃ")
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(12)

        # Set the paragraph alignment to center
        table = document.add_table(rows=3, cols=4)
        table.style = 'TableGrid'
        table.autofit = False
        table.width = Inches(6)
        table.alignment = 2

        cell = table.cell(0, 0)
        p3 = cell.paragraphs[0]
        p3.add_run('প্রত্যাবাসিত রপ্তানিমূল্য')
        p3.style = custom_stylebl
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Inches(1.2)

        cell = table.cell(0, 1)
        p3 = cell.paragraphs[0]
        p3.add_run('প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার পরিমাণ')
        p3.style = custom_stylebl
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[1].width = Inches(1.7)

        cell = table.cell(0, 2)
        p2 = cell.paragraphs[0]
        p2.add_run('বৈদেশিক মুদ্রায় পরিশোধ্য কমিশন ইন্সুরেন্স ইত্যাদি (যদি থাকে)')
        p2.style = custom_stylebl
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[2].width = Inches(2.58)

        cell = table.cell(0, 3)
        p2 = cell.paragraphs[0]
        p2.add_run('নীট এফওবি রপ্তানিমূল্য ১-(২+৩)')
        p2.style = custom_stylebl
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 0)
        p4 = cell.paragraphs[0]
        p4.add_run('১')
        p4.style = custom_stylebl
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 1)
        p5 = cell.paragraphs[0]
        p5.add_run('২')
        p5.style = custom_stylebl
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 2)
        p6 = cell.paragraphs[0]
        p6.add_run('৩')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 3)
        p6 = cell.paragraphs[0]
        p6.add_run('৪')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 0)
        p7 = cell.paragraphs[0]
        p7.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p7.style = custom_style4

        cell = table.cell(2, 1)
        p8 = cell.paragraphs[0]
        p8.add_run('N/A')
        p8.style = custom_na
        p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 2)
        p9 = cell.paragraphs[0]
        p9.add_run('N/A')
        p9.style = custom_na
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 3)
        p9 = cell.paragraphs[0]
        p9.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p9.style = custom_style4

        # table 5 ----------------------------------
        paragraph2 = document.add_paragraph("(প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার উল্লেখ সম্বলিত ফ্রেইট সার্টিফিকেটের সত্যায়িত কপি দাখিল করতে হবে)")
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(12)

        vortuki_swift_amt = 0
        if self.incentive_rate_fc > 0:
            vortuki_swift_amt = swift_amt / self.incentive_rate_fc

        # Set the paragraph alignment to center
        table = document.add_table(rows=4, cols=4)
        table.style = 'TableGrid'
        table.autofit = False
        table.width = Inches(6)
        table.alignment = 2

        cell = table.cell(0, 0)
        p3 = cell.paragraphs[0]
        p3.add_run('রপ্তানি সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের মূল্য')
        p3.style = custom_stylebl
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2 = table.cell(0, 1)
        cell.merge(cell2)

        cell = table.cell(0, 2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = cell.paragraphs[0]
        p2.add_run('স্থানীয় মূল্য সংযোজনের হার')
        p2.style = custom_stylebl
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2 = table.cell(1, 2)
        cell.merge(cell2)

        cell = table.cell(0, 3)
        p2 = cell.paragraphs[0]
        p2.add_run('প্রাপ্য ভর্তুকি* ৪x১০%')
        p2.style = custom_stylebl
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 0)
        p4 = cell.paragraphs[0]
        p4.add_run('দেশীয় পণ্য/সেবা')
        p4.style = custom_stylebl
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 1)
        p5 = cell.paragraphs[0]
        p5.add_run('আমদানিকৃত সেবা/পণ্য')
        p5.style = custom_stylebl
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # cell = table.cell(1, 2)
        # p6 = cell.paragraphs[0]
        # p6.add_run('')
        # p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 3)
        p6 = cell.paragraphs[0]
        p6.add_run('')
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 0)
        p4 = cell.paragraphs[0]
        p4.add_run('৫')
        p4.style = custom_stylebl
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 1)
        p5 = cell.paragraphs[0]
        p5.add_run('৬')
        p5.style = custom_stylebl
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 2)
        p6 = cell.paragraphs[0]
        p6.add_run('৭')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 3)
        p6 = cell.paragraphs[0]
        p6.add_run(' ৮')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 0)
        p7 = cell.paragraphs[0]
        p7.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p7.style = custom_style4

        cell = table.cell(3, 1)
        p8 = cell.paragraphs[0]
        p8.add_run('N/A')
        p8.style = custom_na
        p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 2)
        p9 = cell.paragraphs[0]
        p9.add_run('100%')
        p9.style = custom_na
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

        incentive_amt_fc_a = str("{:,.2f}".format(incentive_amt_fc))
        cell = table.cell(3, 3)
        p9 = cell.paragraphs[0]
        p9.add_run(str(currency_symbol)+ ' ' + incentive_amt_fc_a)
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p9.style = custom_style4

        # signature ------------------------------------------
        paragraph2 = document.add_paragraph("(*৭ নম্বর কলামের হার আলোচ্য সার্কুলারের ৪ নম্বর অনুচ্ছেদের সাথে সামঞ্জস্যতার ক্ষেত্রে ভর্তুকি প্রাপ্য হবে।) \nএ মর্মে অঙ্গীকার করা হচ্ছে যে, আমাদের নিজস্ব কারখানায়/প্রতিষ্ঠানে তৈরী/উৎপাদিত সফটওয়্যার/আইটিইএস/হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকির জন্য আবেদন করা হলো। এ আবেদনপত্রে প্রদত্ত সকল তথ্য/ঘোষণা সম্পূর্ণ ও সঠিক। যদি পরবর্তীতে কোন ভুল/অসত্য তথ্য/প্রতারণা/জালিয়াতি উদঘাটিত হয় তবে গৃহীত ভর্তুকির সমুদয় অর্থ বা এর অংশবিশেষ আমার/আমাদের নিকট হইতে এবং/অথবা আমার/আমাদের ব্যাংক হিসাব থেকে আদায়/ফেরত নেয়া যাবে। \n\nতারিখঃ..................................... ")
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(10)

        paragraph2 = document.add_paragraph("........................................ \nআবেদনকারী প্রতিষ্ঠানের স্বত্তাধিকারী/\n ক্ষমতাপ্রাপ্ত  কর্মকর্তার স্বাক্ষর ও পদবী")
        paragraph2.style = custom_stylebl
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(10)

        # table 6 ----------------------------------
        document.add_page_break()
        paragraph2 = document.add_paragraph("(ছ) ভর্তুকি প্রদানকারী ব্যাংক শাখা কর্তৃক পূরণীয়ঃ \t\t\t\t\t\t\t\t\t" + ' ' * 11 + "(বৈদেশিক মুদ্রায়)")
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.line_spacing = Pt(10)

        # Set the paragraph alignment to center
        table = document.add_table(rows=3, cols=4)
        table.style = 'TableGrid'
        table.autofit = False
        table.width = Inches(6)
        table.alignment = 2

        cell = table.cell(0, 0)
        p3 = cell.paragraphs[0]
        p3.add_run('প্রত্যাবাসিত রপ্তানিমূল্য')
        p3.style = custom_stylebl
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Inches(1.21)

        cell = table.cell(0, 1)
        p3 = cell.paragraphs[0]
        p3.add_run('প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার পরিমাণ')
        p3.style = custom_stylebl
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[1].width = Inches(1.7)

        cell = table.cell(0, 2)
        p2 = cell.paragraphs[0]
        p2.add_run('বৈদেশিক মুদ্রায় পরিশোধ কমিশন, ইন্সুরেন্স ইত্যাদি (যদি থাকে)')
        p2.style = custom_stylebl
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[2].width = Inches(2.6)

        cell = table.cell(0, 3)
        p2 = cell.paragraphs[0]
        p2.add_run('নীট এফওবি রপ্তানিমূল্য ১-(২+৩)')
        p2.style = custom_stylebl
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[3].width = Inches(1.8)

        cell = table.cell(1, 0)
        p4 = cell.paragraphs[0]
        p4.add_run('১')
        p4.style = custom_stylebl
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 1)
        p5 = cell.paragraphs[0]
        p5.add_run('২')
        p5.style = custom_stylebl
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 2)
        p6 = cell.paragraphs[0]
        p6.add_run('৩')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 3)
        p6 = cell.paragraphs[0]
        p6.add_run('৪')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 0)
        p4 = cell.paragraphs[0]
        p4.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.style = custom_style4

        cell = table.cell(2, 1)
        p5 = cell.paragraphs[0]
        p5.add_run('N/A')
        p5.style = custom_na
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 2)
        p6 = cell.paragraphs[0]
        p6.add_run('N/A')
        p6.style = custom_na
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 3)
        p6 = cell.paragraphs[0]
        p6.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p6.style = custom_style4

        # table 7 ----------------------------------
        if self.od_sight_rate:
            od_sight_rate = self.od_sight_rate
        else:
            if od_s_r:
                od_sight_rate = od_s_r
            else:
                od_sight_rate = encashment_rate_bdt

        rate_fc = swift_amt / self.incentive_rate_fc
        # encashment_final_amt = encashment_amt_bdt / self.incentive_rate_fc
        # encashment_final_amt = rate_fc / self.incentive_rate_fc
        rate_fc_a = "{:,.2f}".format(rate_fc)
        # od_sight_rate = "{:.2f}".format(self.od_sight_rate)
        encashment_final_amt_a = round(rate_fc, 2) * od_sight_rate
        # encashment_final_amt = "{:,.0f}".format(round(encashment_final_amt_a))
        encashment_final_amt = self.custom_round(encashment_final_amt_a)
        usd_bdt_amount = currency + ' ' + str(rate_fc_a) + '\n' + '@' + str("{:,}".format(round(od_sight_rate, 2))) + '=' + '\n' + 'BDT ' + encashment_final_amt
        # amount in word
        amount_in_word_str = ''
        if encashment_final_amt_a:
            # amount_in_word = num2words(round(encashment_final_amt_a))
            amount_in_word = self.num2words_fun(round(encashment_final_amt_a))
            # upper case function call
            amount_in_word_str = self.upper_case(amount_in_word)

        paragraph2 = document.add_paragraph("(প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়া উল্লেখ সম্বলিত ফ্রেইট সার্টিফিকেটের সত্যায়িত কপি দাখিল করতে হবে)")
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.line_spacing = Pt(12)

        # Set the paragraph alignment to center
        table = document.add_table(rows=4, cols=4)
        table.style = 'TableGrid'
        table.autofit = False
        table.width = Inches(6)
        table.alignment = 2

        cell = table.cell(0, 0)
        p3 = cell.paragraphs[0]
        p3.add_run('রপ্তানি সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের মূল্য')
        p3.style = custom_stylebl
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2 = table.cell(0, 1)
        cell.merge(cell2)

        cell = table.cell(0, 2)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = cell.paragraphs[0]
        p2.add_run('স্থানীয় মূল্য সংযোজনের হার [(৪-৬)/৪]x ১০০')
        p2.style = custom_stylebl
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2 = table.cell(1, 2)
        cell.merge(cell2)

        cell = table.cell(0, 3)
        p2 = cell.paragraphs[0]
        p2.add_run('পরিশোধ ভর্তুকির পরিমাণ (টাকায়)*(রপ্তানিমূল্য প্রত্যাবাসনের তারিখে সংশ্লিষ্ট ৪x ১০% বৈদেশিক মুদ্রার ওডি সাইট)')
        p2.style = custom_stylebl
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2 = table.cell(1, 3)
        cell.merge(cell2)

        cell = table.cell(1, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p4 = cell.paragraphs[0]
        p4.add_run('দেশীয় পণ্য/সেবা')
        p4.style = custom_stylebl
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(1, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p5 = cell.paragraphs[0]
        p5.add_run('আমদানিকৃত সেবা/পণ্য বৈদেশিক মুদ্রায়')
        p5.style = custom_stylebl
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # cell = table.cell(1, 3)
        # p6 = cell.paragraphs[0]
        # p6.add_run('')
        # p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 0)
        p4 = cell.paragraphs[0]
        p4.add_run('৫')
        p4.style = custom_stylebl
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 1)
        p5 = cell.paragraphs[0]
        p5.add_run('৬')
        p5.style = custom_stylebl
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 2)
        p6 = cell.paragraphs[0]
        p6.add_run('৭')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 3)
        p6 = cell.paragraphs[0]
        p6.add_run(' ৮')
        p6.style = custom_stylebl
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 0)
        p7 = cell.paragraphs[0]
        p7.add_run(str(currency_symbol) + ' ' + swift_amt_a)
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p7.style = custom_style4

        cell = table.cell(3, 1)
        p8 = cell.paragraphs[0]
        p8.add_run('N/A')
        p8.style = custom_na
        p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 2)
        p9 = cell.paragraphs[0]
        p9.add_run('100%')
        p9.style = custom_na
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 3)
        p9 = cell.paragraphs[0]
        p9.add_run(usd_bdt_amount)
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p9.style = custom_style4

        # conclusion ----------------
        paragraph2 = document.add_paragraph("(*৭ নম্বর কলামের হার আলোচ্য সার্কুলারের ৪ নম্বর অনুচ্ছেদের সাথে সামঞ্জস্যতার ক্ষেত্রে ভর্তুকি প্রাপ্য হবে) \nভর্তুকি পরিমাণ: " + amount_in_word_str + ' Only')
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.line_spacing = Pt(12)

        paragraph2 = document.add_paragraph("\n\n\n\n\n\n\nপরিশোধের তারিখঃ --------------------------")
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.line_spacing = Pt(10)

        paragraph2 = document.add_paragraph("----------------------------------------- \nভর্তুকির অনুমোদনের ক্ষমতাপ্রাপ্ত ব্যাংক  \nকর্মকর্তার স্বাক্ষর, নাম ও পদবী")
        paragraph2.style = custom_stylebl
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph2.paragraph_format.line_spacing = Pt(10)
        document.add_page_break()
        # file_name = '%s_ka' % (datetime.datetime.now())
        # KHA 444444444444444 -------------------------------------------
        # modify the page setup
        section.page_width = Cm(22)  # set the page width to 21 centimeters
        section.page_height = Cm(29.9)  # set the page width to 21 centimeters
        section.left_margin = Cm(2.5)  # 1.5 cm
        section.right_margin = Cm(2)  # 1.5 cm
        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

        # document.add_paragraph('কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।')
        # en_font = my_stylew.font
        # document.add_paragraph('some text')

        custom_style = document.styles.add_style('MyStyled19', WD_STYLE_TYPE.PARAGRAPH)
        # custom_style.font.bold = True
        custom_style.font.size = Pt(8)
        custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        custom_style.font.name = 'Arial Narrow'

        custom_style_table = document.styles.add_style('custom_style_table20', WD_STYLE_TYPE.PARAGRAPH)
        custom_style_table.font.size = Pt(10)
        custom_style_table.font.name = 'Arial Narrow'

        custom_stylebsmallbold = document.styles.add_style('custom_stylebsmallbold21', WD_STYLE_TYPE.PARAGRAPH)
        custom_stylebsmallbold.font.size = Pt(9)
        custom_stylebsmallbold.font.name = 'SutonnyOMJ'
        custom_stylebsmallbold.font.bold = True

        custom_stylebsmall = document.styles.add_style('custom_stylebsmall22', WD_STYLE_TYPE.PARAGRAPH)
        custom_stylebsmall.font.size = Pt(9)
        custom_stylebsmall.font.name = 'SutonnyOMJ'

        custom_stylebl = document.styles.add_style('custom_stylebl23', WD_STYLE_TYPE.PARAGRAPH)
        custom_stylebl.font.size = Pt(9)
        custom_stylebl.font.name = 'SutonnyOMJ'

        custom_sonar_bangla = document.styles.add_style('custom_sonar_bangla24', WD_STYLE_TYPE.PARAGRAPH)
        custom_sonar_bangla.font.size = Pt(9)
        custom_sonar_bangla.font.name = 'Shonar Bangla'

        custom_tableh = document.styles.add_style('custom_tableh25', WD_STYLE_TYPE.PARAGRAPH)
        custom_tableh.font.size = Pt(9)
        custom_tableh.font.name = 'SutonnyOMJ'

        custom_table = document.styles.add_style('custom_table26', WD_STYLE_TYPE.PARAGRAPH)
        custom_table.font.size = Pt(9)
        custom_table.font.name = 'SutonnyOMJ'

        custom_stylebold = document.styles.add_style('custom_stylebold27', WD_STYLE_TYPE.PARAGRAPH)
        custom_stylebold.font.size = Pt(11)
        custom_stylebold.font.name = 'SutonnyOMJ'
        custom_stylebold.font.bold = True

        custom_styleubl = document.styles.add_style('custom_styleubl28', WD_STYLE_TYPE.PARAGRAPH)
        custom_styleubl.font.size = Pt(11)
        custom_styleubl.font.name = 'SutonnyOMJ'
        custom_styleubl.font.underline = True
        custom_styleubl.font.bold = True

        document.styles['Normal'].font.size = Pt(8)
        paragraph = document.add_paragraph()
        paragraph.add_run("\n(অনুচ্ছেদ ০৬ (ক), এফই সার্কুলার নং-০৩/২০১৮ দ্রষ্টব্য)")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.style = custom_stylebl
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.add_run("ফরম- 'খ'")
        paragraph = document.add_paragraph()
        paragraph.add_run("Ref: " + str(self.form_kha_ref_code))
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run(" Date: " + self_date)
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.style = custom_stylebl

        # paragraph.style = custom_style

        # Define a custom style
        custom_style2 = document.styles.add_style('MyStyle56', WD_STYLE_TYPE.PARAGRAPH)
        custom_style2.font.size = Pt(8)

        custom_style3 = document.styles.add_style('MyStyle57', WD_STYLE_TYPE.PARAGRAPH)
        custom_style3.font.size = Pt(8)
        custom_style3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        custom_style4 = document.styles.add_style('MyStyle58', WD_STYLE_TYPE.PARAGRAPH)
        custom_style4.font.bold = True
        custom_style4.font.size = Pt(10)
        custom_style4.font.name = 'Arial Narrow'

        custom_total = document.styles.add_style('custom_total59', WD_STYLE_TYPE.PARAGRAPH)
        custom_total.font.bold = True
        custom_total.font.size = Pt(9)
        custom_total.font.name = 'Arial Narrow'

        custom_style5 = document.styles.add_style('MyStyle60', WD_STYLE_TYPE.PARAGRAPH)
        custom_style5.font.size = Pt(8)
        custom_style5.font.name = 'Arial Narrow'

        # Add a paragraph and apply the custom style
        paragraph = document.add_paragraph('বেসিস প্রদেয় প্রত্যয়ন সনদপত্র \nবাংলাদেশ হতে সফটওয়্যার, আইটিইএস (Information Technology Enabled Services)')
        paragraph.style = custom_stylebold
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph = document.add_paragraph('ও হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকি প্রাপ্তির প্রত্যয়ন সনদপত্র।')
        paragraph.style = custom_styleubl
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)

        customer = '' if not self.swift_customer_name else str(self.swift_customer_name) + ', ' + str(
            self.customer_address)
        con_no = ''
        if self.contract_number:
            con_no = self.contract_number.replace('Contract No.', '').replace('Purchase Order No.', '')

        # new_parser = HtmlToDocx()
        # html_text = """
        #             <div style='position: relative'>
        #                 <p style='float:left'> ১। </p>
        #                 <p  style='margin-left: 18px; float:right'> আবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road, Mohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ RA- 260326210666420 </p>
        #             </div>
        #
        #             """
        # new_parser.add_html_to_document(html_text, document)

        # paragraph = document.add_paragraph()
        # paragraph.add_run('Word1')
        # paragraph.add_run(' ' * 3)
        # paragraph.add_run('Word2')
        import textwrap
        width = 140
        # long_text = "১।\t  আবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road, Mohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ 260326210666420"
        # wrapped_lines = textwrap.wrap(long_text, width=width)
        # formatted_lines = [line + '\t' for line in wrapped_lines]
        # formatted_text = '\n\t'.join(formatted_lines)
        # paragraph2 = document.add_paragraph(formatted_text)
        # paragraph2.style = custom_stylebl
        # paragraph2.paragraph_format.space_before = Pt(1)
        # paragraph2.paragraph_format.space_after = Pt(1)
        #
        # long_text = "২।\tরপ্তানি ঋণপত্র/চুক্তিপত্র নম্বরঃ" + con_no + '\t তারিখঃ ' + self.contract_date_str + '\t মূল্যঃ ' +self.contract_price_str
        # wrapped_lines = textwrap.wrap(long_text, width=width)
        # formatted_lines = [line + '\t' for line in wrapped_lines]
        # formatted_text = '\n\t'.join(formatted_lines)
        # paragraph2 = document.add_paragraph(formatted_text)
        # paragraph2.paragraph_format.space_before = Pt(1)
        # paragraph2.paragraph_format.space_after = Pt(1)
        #
        # long_text = "৩।\tবিদেশি ক্রেতার নাম ও ঠিকানাঃ "+ customer
        # wrapped_lines = textwrap.wrap(long_text, width=width)
        # formatted_lines = [line + '\t' for line in wrapped_lines]
        # formatted_text = '\n\t'.join(formatted_lines)
        # paragraph2 = document.add_paragraph(formatted_text)
        # paragraph2.paragraph_format.space_before = Pt(1)
        # paragraph2.paragraph_format.space_after = Pt(1)
        #
        # long_text = "৪।\tবিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ " + remiter_address + '\n' + '৫।'
        # wrapped_lines = textwrap.wrap(long_text, width=width)
        # formatted_lines = [line + '\t' for line in wrapped_lines]
        # formatted_text = '\n\t'.join(formatted_lines)
        # paragraph2 = document.add_paragraph(formatted_text)
        # paragraph2.paragraph_format.space_before = Pt(1)
        # paragraph2.paragraph_format.space_after = Pt(1)

        # paragraph2 = document.add_paragraph("১।\tআবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road,"
        #                                     " \tMohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ 260326210666420 \n২।\tরপ্তানি ঋণপত্র/চুক্তিপত্র নম্বরঃ" + con_no + '\t তারিখঃ ' +
        #                                     self.contract_date_str + '\t মূল্যঃ ' +self.contract_price_str + '\n৩।\tবিদেশি ক্রেতার নাম ও ঠিকানাঃ '+ customer +
        #                                     "\n৪।\tবিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ " + remiter_address + '\n' + '৫।' )

        paragraph2 = document.add_paragraph("১।\tআবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road,"
                                            " \tMohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ 260326210666420")
        paragraph2.style = custom_stylebl
        # paragraph2.add_run("\t")
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(0)

        long_text1 = "২।\t" + ' ' * 5 + "রপ্তানি ঋণপত্র/চুক্তিপত্র নম্বরঃ" + con_no + '\t তারিখঃ ' + self.contract_date_str + '\t মূল্যঃ ' +self.contract_price_str
        wrapped_lines1 = textwrap.wrap(long_text1, width=120)
        formatted_lines1 = [line + '\t' for line in wrapped_lines1]
        formatted_text1 = '\n\t'.join(formatted_lines1)
        paragraph2 = document.add_paragraph(formatted_text1)
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        long_text2 = "৩।\t" + ' ' * 5 + "বিদেশি ক্রেতার নাম ও ঠিকানাঃ " + customer
        # long_text2 = "৩।\t" + ' ' * 5 + "বিদেশি ক্রেতার নাম ও ঠিকানাঃ ABBVIE INC AP, DEPT V312 AP34-2 1 N WAUKEGAN RD, NORTH CHICAGO"
        wrapped_lines2 = textwrap.wrap(long_text2, width=123)
        formatted_lines2 = [line + '\n\t' for line in wrapped_lines2]
        # formatted_text2 = ''.join(formatted_lines2)
        formatted_text2 = ''.join(formatted_lines2).rstrip('\n\t')
        paragraph2 = document.add_paragraph(formatted_text2)
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        long_text1 = "৪।\t" + ' ' * 5 + "বিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ " + remiter_address
        wrapped_lines1 = textwrap.wrap(long_text1, width=120)
        formatted_lines1 = [line + '\t' for line in wrapped_lines1]
        formatted_text1 = '\n\t'.join(formatted_lines1)
        paragraph2 = document.add_paragraph(formatted_text1)
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph('৫।')
        paragraph2.style = custom_stylebl

        # table = document.add_table(rows=3, cols=4)
        #
        # # Set the width of the first column to 1 inch
        # column_width = Inches(1)
        # table.columns[0].width = column_width
        #
        # # Set the width of the second column to 2 inches
        # column_width = Inches(1)
        # table.columns[1].width = column_width
        #
        # # Set the width of the third column to 0.5 inches
        # column_width = Inches(4)
        # table.columns[2].width = column_width
        #
        # # Set the width of the third column to 0.5 inches
        # column_width = Inches(1)
        # table.columns[3].width = column_width

        table = document.add_table(rows=1, cols=4)
        table.style = 'TableGrid'
        table.alignment = 1
        # Set the width of the first column to 1 inch
        column_width = Inches(1.2)
        table.columns[0].width = column_width
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('ক) ইনভয়েস নম্বর')
        p3.style = custom_tableh
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set the width of the second column to 2 inches
        column_width = Inches(1)
        table.columns[1].width = column_width
        cell = table.cell(0, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('তারিখ')
        p3.style = custom_tableh
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set the width of the third column to 0.5 inches
        column_width = Inches(4.3)
        table.columns[2].width = column_width
        cell1 = table.cell(0, 2)
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell1.paragraphs[0]
        p3.add_run('খ) ইনভয়েসে উল্লেখিত সেবা/পণ্যের পরিমাণ (Qty & Hours)')
        p3.style = custom_tableh
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[2].width = Inches(3.8)

        # Set the width of the third column to 0.5 inches
        column_width = Inches(1)
        table.columns[3].width = column_width
        cell = table.cell(0, 3)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        cur_name = 'মূল্য (' + self.fc_currency_id.name + ')'
        # p3.add_run('মূল্য (USD)')
        p3.add_run(cur_name)
        p3.style = custom_tableh
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set the paragraph alignment to center
        # table = document.add_table(rows=1, cols=4)
        # table.style = 'TableGrid'
        # table.autofit = False
        # # table = document.tables[0]
        # # Adjust the left indentation of the table
        # table.alignment = 1  # Center alignment
        # # table.left_indent = Pt(144)
        #
        # cell = table.cell(0, 0)
        # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # p3 = cell.paragraphs[0]
        # p3.add_run('ক) ইনভয়েস নম্বর'...............)
        # p3.style = custom_tableh
        # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[0].width = Inches(1.2)
        #
        # cell = table.cell(0, 1)
        # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # p3 = cell.paragraphs[0]
        # p3.add_run('তারিখ')
        # p3.style = custom_tableh
        # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[1].width = Inches(1)
        #
        # cell1 = table.cell(0, 2)
        # cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # p3 = cell1.paragraphs[0]
        # p3.add_run('খ) ইনভয়েসে উল্লেখিত সেবা/পণ্যের পরিমাণ (Qty & Hours)')
        # p3.style = custom_tableh
        # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[2].width = Inches(3.8)
        #
        # cell = table.cell(0, 3)
        # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # p3 = cell.paragraphs[0]
        # p3.add_run('মূল্য (USD)')
        # p3.style = custom_tableh
        # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # table.columns[3].width = Inches(1)
        #
        # # add a dynamic row to the table
        cus_name = ''
        if self.swift_customer_name:
            cus_name = self.swift_customer_name

        # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
        total_qty = 0
        dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in
                         self.invoice_line_ids.invoice_id.invoice_line_ids)
        hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in
                          self.invoice_line_ids.invoice_id.invoice_line_ids)
        dev_total = dev_total1
        hour_total = hour_total1

        if hour_total == int(hour_total):
            hour_total_num = "{:.0f}".format(hour_total)
        else:
            hour_total_num = "{:.2f}".format(hour_total)

        if dev_total == int(dev_total):
            dev_total_num = "{:.0f}".format(dev_total)
        else:
            dev_total_num = "{:.2f}".format(dev_total)

        qty_str = ''
        if dev_total > 0:
            qty_str += str(dev_total_num) + (str(' Developers') if dev_total > 1 else str(' Developer'))
        if hour_total > 0:
            qty_str += str(hour_total_num) + str(' HRS') if not qty_str else ' & ' + str(
                hour_total_num) + str(' HRS')
        all_invoice_qty_str = 'Software Development /' + qty_str
        all_invoice_qty_str_a = qty_str
        row_len = len(self.invoice_line_ids)
        swift_amt = 0
        inv_amount = 0

        from itertools import groupby

        # group the invoice_line_ids by invoice_id.ref
        groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                         key=lambda x: x.invoice_id.ref)
        # iterate over each group and add a new row for each unique invoice_id.ref
        row = 1
        for invoice_ref, group in groups:
            inv_obj = self.env['cash.incentive.invoice'].search(
                [('invoice_id.ref', '=', invoice_ref), ('head_id', '=', self.id)], limit=1)
            inv_amount += inv_obj.invoice_amt
            invoice_date = ''
            if inv_obj.invoice_date:
                invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime(
                    '%d-%b-%y')

            qty = 'Software Development /'
            h_q = 0
            d_q = 0
            for l in inv_obj.invoice_id.invoice_line_ids:
                if l.quantity_type == '0':
                    h_q += l.quantity
                else:
                    d_q += l.quantity
            if d_q:
                if d_q == int(d_q):
                    d_q_num = "{:.0f}".format(d_q)
                else:
                    d_q_num = "{:.2f}".format(d_q)
                qty += str(d_q_num) + (str(' Developers') if d_q > 1 else str(' Developer'))
            if h_q:
                if h_q == int(h_q):
                    h_q_new = "{:.0f}".format(h_q)
                else:
                    h_q_new = "{:.2f}".format(h_q)
                if not d_q:
                    qty += str(h_q_new) + ' HRS'
                else:
                    qty += ' & ' + str(h_q_new) + ' HRS'
            invoice_qty_str = qty

            gr = 0
            for rec in group:
                gr += 1
                swift_amt += rec.swift_amt
                # add a new row to the table
                new_row = table.add_row()
                # add data to the cells in the new row
                # new_row.cells[0].text = rec.invoice_id.ref
                # new_row.cells[1].text = invoice_date
                # new_row.cells[2].text = invoice_qty_str
                # new_row.cells[3].text = currency + ' ' + str("{:,}".format(round(rec.invoice_amt, 2)))

                first_cell = new_row.cells[0]
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph21.style = custom_style5
                first_cell = new_row.cells[1]
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph21.style = custom_style5
                first_cell = new_row.cells[2]
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph21.style = custom_style5
                first_cell = new_row.cells[3]
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph21.style = custom_style5

            cell_1_0 = table.cell(row, 0)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run(invoice_ref)
            p3.style = custom_style_table
            cell_2_0 = table.cell((row + gr) - 1, 0)
            cell_1_0.merge(cell_2_0)

            cell_4_0 = table.cell(row, 1)
            p3 = cell_4_0.paragraphs[0]
            p3.add_run(invoice_date)
            p3.style = custom_style_table
            cell_3_0 = table.cell((row + gr) - 1, 1)
            cell_4_0.merge(cell_3_0)

            cell_5_0 = table.cell(row, 2)
            p3 = cell_5_0.paragraphs[0]
            p3.add_run(invoice_qty_str)
            p3.style = custom_style_table
            cell_6_0 = table.cell((row + gr) - 1, 2)
            cell_5_0.merge(cell_6_0)

            cell_5_0 = table.cell(row, 3)
            p3 = cell_5_0.paragraphs[0]
            p3.style = custom_style_table
            # p3.add_run(str("{:,}".format(round(inv_obj.invoice_amt, 2))))
            p3.add_run(str("{:,.2f}".format(inv_obj.invoice_amt)))
            cell_6_0 = table.cell((row + gr) - 1, 3)
            cell_5_0.merge(cell_6_0)
            row = gr + row

        # total
        total_row = table.add_row()
        total_row.cells[0].text = 'Total'
        total_row.cells[2].text = all_invoice_qty_str
        total_row.cells[3].text = str("{:,.2f}".format(inv_amount))

        first_cell = total_row.cells[0]
        paragraph21 = first_cell.paragraphs[0]
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph21.style = custom_total
        first_cell = total_row.cells[2]
        paragraph21 = first_cell.paragraphs[0]
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph21.style = custom_total
        first_cell = total_row.cells[3]
        paragraph21 = first_cell.paragraphs[0]
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph21.style = custom_total

        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

        # footer -------------------
        # paragraph2 = document.add_paragraph("৬।\tরপ্তানিকৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত স্থানীয় সেবা/উপকরণাদির সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): 100% in house production \tof Brain Station 23 Limited \t\t\t\t পরিমাণঃ N/A  \t\t মূল্যঃ N/A"
        #                                     "\n৭।\tরপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত আমদানিকৃত আনুসাংগিক সেবা/উপরকরণাদি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): N/A  \t\t\tপরিমাণঃ N/A   \t\t\t\t\t  মূল্যঃ  N/A"
        #                                     "\n৮।\tরপ্তানি সেবা/পণ্যের বিবরণঃ Software Development Services  \tপরিমাণঃ " + all_invoice_qty_str_a + " \t  মূল্যঃ " + currency+ ' ' + str("{:,.2f}".format(inv_amount)) +
        #                                     "\n৯।\tজাহাজীকরণের তারিখঃ " + inv_date + "\tগন্তব্য বন্দরঃ " + country + "\n১০।\tইএক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানরি ক্ষেত্রে): N/A  \t\t\t  মূল্যঃ N/A   \t\t তারিখঃ N/A"
        #                                     "\n১১।\tমোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency+ ' ' + str("{:,.2f}".format(swift_amt)) + "\t\t নীট এফওবি রপ্তানি মূল্য ( বৈদেশিক মুদ্রায়): "+ currency+ ' ' + str("{:,.2f}".format(swift_amt)) +
        #                                     "\n১২।\tপ্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ \t\t\t\t\t\t তারিখঃ")
        # paragraph2.style = custom_stylebl

        paragraph2 = document.add_paragraph("৬।\tরপ্তানিকৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত স্থানীয় সেবা/উপকরণাদির সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): 100% in house \tproduction of Brain Station 23 Limited \t\t পরিমাণঃ N/A  \t\t মূল্যঃ N/A"
                                            "\n৭।\tরপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত আমদানিকৃত আনুসাংগিক সেবা/উপরকরণাদি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): N/A  \t\tপরিমাণঃ N/A   \t\t\t\t\t  মূল্যঃ  N/A")
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_after = Pt(1)

        # long_text = "৮। \t " + ' ' * 4 + "রপ্তানি সেবা/পণ্যের বিবরণঃ Software Development Services পরিমাণঃ " + all_invoice_qty_str_a + " " + "  মূল্যঃ " + currency+ ' ' + str("{:,.2f}".format(inv_amount))
        # long_text = "৯। \t " + ' ' * 4 + "জাহাজীকরণের তারিখঃ 31/10/22, 31/10/22, 30/11/22, 30/11/22, 30/11/22, 30/11/22, 31/10/22, 30/11/22, 30/11/22, 30/11/22, 30/11/22 \t\tগন্তব্য বন্দরঃ Norway"
        # wrapped_lines = textwrap.wrap(long_text, width=120)
        # formatted_lines = [line + '\t' for line in wrapped_lines]
        # formatted_text = '\n\t'.join(formatted_lines)
        # paragraph2 = document.add_paragraph(formatted_text)
        # paragraph2.style = custom_stylebl
        # paragraph2.paragraph_format.space_before = Pt(1)
        # paragraph2.paragraph_format.space_after = Pt(1)

        if len(self.invoice_line_ids) > 10:
            long_text = "৯। \t " + ' ' * 4 + "জাহাজীকরণের তারিখঃ " + inv_date + " \t\tগন্তব্য বন্দরঃ " + country
            wrapped_lines = textwrap.wrap(long_text, width=120)
            formatted_lines = [line + '\t' for line in wrapped_lines]
            formatted_text = '\n\t'.join(formatted_lines)
            paragraph2 = document.add_paragraph(formatted_text)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
        else:
            if len(self.invoice_line_ids) > 7:
                paragraph2 = document.add_paragraph("৯।\tজাহাজীকরণের তারিখঃ " + inv_date + "\n\tগন্তব্য বন্দরঃ " + country )
                paragraph2.style = custom_stylebl
            else:
                paragraph2 = document.add_paragraph("৯।\tজাহাজীকরণের তারিখঃ " + inv_date + " \tগন্তব্য বন্দরঃ " + country )
                paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph("১০।\tইএক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানরি ক্ষেত্রে): N/A  \t\t\t  মূল্যঃ N/A   \t\t তারিখঃ N/A")
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        # long_text1 = "১১। \t " + ' ' * 5 +  "মোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): USD 158,793.64          নীট এফওবি রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): USD 158,793.64"
        long_text1 = "১১। \t " + ' ' * 4 + "মোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str("{:,.2f}".format(swift_amt)) + "\t\tনীট এফওবি রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str("{:,.2f}".format(swift_amt))
        wrapped_lines1 = textwrap.wrap(long_text1, width=150)
        formatted_lines1 = [line + '\t' for line in wrapped_lines1]
        formatted_text1 = '\n\t'.join(formatted_lines1)
        paragraph2 = document.add_paragraph(formatted_text1)
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        prc_date = ''
        if self.prc_date:
            prc_date = datetime.datetime.strptime(str(self.prc_date), '%Y-%m-%d').strftime('%d/%m/%y')
        prc_ref_code = ''
        if self.prc_ref_code:
            prc_ref_code = self.prc_ref_code

        long_text1 ="১২।\t" + ' ' * 4 + "প্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ " + prc_ref_code+ "\t\t\t তারিখঃ " + prc_date
        wrapped_lines1 = textwrap.wrap(long_text1, width=120)
        formatted_lines1 = [line + '\t' for line in wrapped_lines1]
        formatted_text1 = '\n\t'.join(formatted_lines1)
        paragraph2 = document.add_paragraph(formatted_text1)
        paragraph2.style = custom_stylebl
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        # paragraph2 = document.add_paragraph("১২।\tপ্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ " + prc_ref_code+ "\t\t\t\t\t\t তারিখঃ " + prc_date)
        # paragraph2.style = custom_stylebl

        paragraph2 = document.add_paragraph("\n\n\nরপ্তানিকারকের স্বাক্ষর ও তারিখ")
        # paragraph2.style = custom_style3
        # paragraph2.style = my_styler
        paragraph2.style = custom_stylebsmallbold
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(1)

        paragraph4 = document.add_paragraph("এতদ্বারা প্রত্যয়ন করা যাচ্ছে যে, আমাদের নিজস্ব কারখানায়/প্রতিষ্ঠানে তৈরীকৃত/উৎপাদিত সফটওয়্যার/আইটিইএস/হার্ডওয়্যার উপরোক্ত ৬ ও ৭ নং ক্রমিক বর্ণিত সূত্র হতে সেবা/উপকরাণাদি সংগ্রহের মাধ্যমে রপ্তানির বিপরীতে ভর্তুকির জন্য উপরোক্ত অনুচ্ছেদগুলোতে উল্লিখিত বক্তব্য সঠিক ও নির্ভুল। বিদেশী ক্রেতা/ আমদানিকারকের ক্রয়াদেশের যথার্থতা/বিশ্বাসযোগ্যতা সম্পর্কেও নিশ্চিত করা হলো। ")
        paragraph4.style = custom_stylebsmall
        paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # paragraph4 = document.paragraphs[1]
        # Modify the paragraph format to reduce spacing
        paragraph4.paragraph_format.space_before = Pt(1)
        paragraph4.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph("\n\n\nরপ্তানিকারকের স্বাক্ষর ও তারিখ")
        # paragraph2.style = custom_style3
        # paragraph2.style = my_styler
        paragraph2.style = custom_stylebsmallbold
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph("রপ্তানিকারকের উপরোক্ত ঘোষণার যথার্থতা যাচাইয়ান্তে সঠিক পাওয়া গিয়েছে। ৮নং ক্রমিকে উল্লিখিত ঘোষিত রপ্তানিমূল্য যৌক্তিক ও বিদ্যমান আন্তর্জাতিক বাজার মুল্যের সঙ্গে সংগতিপূর্ণ পাওয়া গিয়েছে এবং বিদেশী ক্রেতার যথার্থতা/বিশ্বাসযোগ্যতা সর্ম্পকেও নিশ্চিত হওয়া গিয়েছে। প্রত্যাবাসিত রপ্তানি মূল্যের (নীট এফওবি মূল্য) রপ্তানি ভর্তুকি পরিশোধের সুপারিশ করা হলো।")
        paragraph2.style = custom_stylebsmall
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph("\n\n\n\n----------------------------------- এসোসিয়েশন এর দুইজন উপযুক্ত কর্মকর্তার স্বাক্ষর, তারিখ ও সীল")
        # paragraph2.style = my_styler
        paragraph2.style = custom_stylebsmall
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph("[কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।]")
        # paragraph2.style = my_stylec
        paragraph2.style = custom_stylebsmall
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(0)
        document.add_page_break()
        # file_name = '%s_kha' % (datetime.datetime.now())
        # bank 55555555 --------------------------------------------------

        section.page_width = Cm(22)  # set the page width to 21 centimeters
        section.page_height = Cm(29.9)  # set the page width to 21 centimeters
        section.left_margin = Cm(2.5)  # 1.5 cm
        section.right_margin = Cm(2)

        self_date = ''
        current_date = datetime.datetime.now().date()
        if self.date:
            self_date = datetime.datetime.strptime(str(current_date), '%Y-%m-%d').strftime('%d/%m/%y')

        # Define a custom style
        custom_style = document.styles.add_style('MyStyle33', WD_STYLE_TYPE.PARAGRAPH)
        custom_style.font.bold = True
        custom_style.font.size = Pt(14)
        custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        custom_style.font.name = 'Arial Narrow'

        custom_style2 = document.styles.add_style('MyStyle34', WD_STYLE_TYPE.PARAGRAPH)
        custom_style2.font.size = Pt(10)
        custom_style2.font.name = 'Calibri'

        custom_style3 = document.styles.add_style('MyStyle35', WD_STYLE_TYPE.PARAGRAPH)
        custom_style3.font.size = Pt(10)
        custom_style3.font.bold = True
        custom_style3.font.name = 'Calibri'

        custom_style4 = document.styles.add_style('MyStyle36', WD_STYLE_TYPE.PARAGRAPH)
        custom_style4.font.bold = True
        custom_style4.font.size = Pt(9)
        custom_style4.font.name = 'Arial Narrow'

        table_total = document.styles.add_style('table_total37', WD_STYLE_TYPE.PARAGRAPH)
        table_total.font.bold = True
        table_total.font.size = Pt(10)
        table_total.font.name = 'Arial Narrow'

        custom_style5 = document.styles.add_style('MyStyle38', WD_STYLE_TYPE.PARAGRAPH)
        custom_style5.font.size = Pt(9)
        custom_style5.font.name = 'Arial Narrow'

        paragraph2 = document.add_paragraph('\n\n\n\n\nRef: ' + self.name + '\n\nDate: ' + self_date) # + '\n'
        paragraph2.style = custom_style3
        paragraph2.paragraph_format.left_indent = Pt(-6)

        paragraph2 = document.add_paragraph("The Branch Manager  \nGulshan Branch \n" + self.bank_id.name + "\nHolding No. 75, Gulshan Avenue \nGulshan, Dhaka \n\n\nDear Sir\n")
        paragraph2.style = custom_style2
        paragraph2.paragraph_format.space_before = Pt(0)
        paragraph2.paragraph_format.space_after = Pt(1)
        paragraph2.paragraph_format.left_indent = Pt(-6)

        paragraph2 = document.add_paragraph("For cash incentive claim.\n")
        paragraph2.style = custom_style3
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)
        paragraph2.paragraph_format.left_indent = Pt(-6)

        paragraph2 = document.add_paragraph("We are submitting herewith necessary documents against following Invoices:")
        paragraph2.style = custom_style2
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(0)
        paragraph2.paragraph_format.left_indent = Pt(-6)

        # table -----------------------

        table = document.add_table(rows=2, cols=8)
        table.style = 'TableGrid'
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # table.autofit = False  # Disable auto-fit behavior
        # table.left_indent = Inches(6)

        cell = table.cell(0, 0)
        # table.columns[0].width = Inches(.2)
        # table.columns[0].left_indent = Pt(100)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('SL#')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style4

        cell = table.cell(0, 1)
        table.columns[1].width = Inches(1.8)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('Client Name')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style4

        cell = table.cell(0, 2)
        table.columns[2].width = Inches(.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('Invoice No')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style4

        cell = table.cell(0, 3)
        table.columns[3].width = Inches(.7)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        cur_name = 'Invoice amount (' + self.fc_currency_id.name + ')'
        p3.add_run(cur_name)
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.style = custom_style4

        cell = table.cell(0, 4)
        table.columns[4].width = Inches(.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('Invoice Date')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style4

        cell = table.cell(0, 5)
        table.columns[5].width = Inches(.6)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('Swift/nostro date')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style4

        cell = table.cell(0, 6)
        table.columns[6].width = Inches(.6)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        cur_name_r = 'Realize Amount (' + self.fc_currency_id.name + ')'
        p3.add_run(cur_name_r)
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.style = custom_style4

        cell = table.cell(0, 7)
        table.columns[7].width = Inches(.6)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        cur_name_i = 'Incentive claim (' + self.fc_currency_id.name + ')'
        p3.add_run(cur_name_i)
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.style = custom_style4

        # add a dynamic row to the table
        cus_name = ''
        if self.swift_customer_name:
            cus_name = self.swift_customer_name

        # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
        row_len = len(self.invoice_line_ids)
        inv_amount = 0
        encashment_amt_bdt = 0
        sl = 1
        total_incentive_rate_fc = 0

        from itertools import groupby

        # group the invoice_line_ids by invoice_id.ref
        groups = groupby(sorted(self.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                         key=lambda x: x.invoice_id.ref)
        # iterate over each group and add a new row for each unique invoice_id.ref
        row = 1
        for invoice_ref, group in groups:
            inv_obj = self.env['cash.incentive.invoice'].search([('invoice_id.ref', '=', invoice_ref), ('head_id', '=', self.id)], limit=1)
            # usd_price = sum(inv_obj.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
            usd_price = inv_obj.invoice_id.invoice_total_fc
            # incentive_rate_fc = usd_price / self.incentive_rate_fc if self.incentive_rate_fc else 0

            inv_amount += inv_obj.invoice_amt
            invoice_date = ''
            if inv_obj.invoice_date:
                invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime('%d-%b-%y')
            gr = 0
            for rec in group:
                gr += 1
                encashment_amt_bdt += rec.encashment_amt_bdt
                incentive_rate_fc = (rec.swift_amt * self.incentive_rate_fc) / 100 if self.incentive_rate_fc else 0
                total_incentive_rate_fc += incentive_rate_fc
                swift_message_date = ''
                if rec.swift_message_id.date:
                    swift_message_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime('%d-%b-%y')

                new_row = table.add_row()
                # add data to the cells in the new row
                new_row.cells[0].text = str(sl)
                new_row.cells[5].text = swift_message_date
                new_row.cells[6].text = str("{:,}".format(round(rec.swift_amt, 2)))
                new_row.cells[7].text = str("{:,}".format(float(format(incentive_rate_fc, '.2f'))))

                first_cell = new_row.cells[0]
                first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph24 = first_cell.paragraphs[0]
                paragraph24.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph24.style = custom_style5

                first_cell = new_row.cells[2]
                first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph21.style = custom_style5

                first_cell = new_row.cells[3]
                first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph21 = first_cell.paragraphs[0]
                paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph21.style = custom_style5

                first_cell1 = new_row.cells[6]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph22.style = custom_style5

                first_cell1 = new_row.cells[7]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph22.style = custom_style5

                # first_cell1 = new_row.cells[2]
                # first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # paragraph22 = first_cell1.paragraphs[0]
                # paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # paragraph22.style = custom_style5

                first_cell1 = new_row.cells[1]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.style = custom_style5

                first_cell1 = new_row.cells[4]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style5

                first_cell1 = new_row.cells[5]
                first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph22 = first_cell1.paragraphs[0]
                paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph22.style = custom_style5

                sl += 1

            cell_1_0 = table.cell(row, 2)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run(invoice_ref)
            cell_2_0 = table.cell((row + gr) - 1, 2)
            cell_1_0.merge(cell_2_0)

            cell_4_0 = table.cell(row, 3)
            p3 = cell_4_0.paragraphs[0]
            p3.add_run(str("{:,}".format(round(usd_price, 2))))
            cell_3_0 = table.cell((row + gr) - 1, 3)
            cell_4_0.merge(cell_3_0)

            cell_5_0 = table.cell(row, 4)
            p3 = cell_5_0.paragraphs[0]
            p3.add_run(invoice_date)
            cell_6_0 = table.cell((row + gr) - 1, 4)
            cell_5_0.merge(cell_6_0)
            row = gr + row

        # for rec in self.invoice_line_ids:
        #     usd_price = sum(rec.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
        #     incentive_rate_fc = usd_price / self.incentive_rate_fc if self.incentive_rate_fc else 0
        #     total_incentive_rate_fc += incentive_rate_fc
        #     inv_amount += rec.invoice_amt
        #     encashment_amt_bdt += rec.encashment_amt_bdt
        #
        #     invoice_date = ''
        #     if rec.invoice_date:
        #         invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d-%b-%y')
        #     swift_message_date = ''
        #     if rec.swift_message_id.date:
        #         swift_message_date = datetime.datetime.strptime(str(rec.swift_message_id.date),
        #                                                         '%Y-%m-%d').strftime('%d-%b-%y')
        #     # add a new row to the table
        #     new_row = table.add_row()
        #
        #     # add data to the cells in the new row
        #     new_row.cells[0].text = str(sl)
        #     new_row.cells[2].text = rec.invoice_id.ref
        #     new_row.cells[3].text = str("{:,}".format(round(usd_price, 2)))
        #     new_row.cells[4].text = invoice_date
        #     new_row.cells[5].text = swift_message_date
        #     new_row.cells[6].text = str("{:,}".format(round(rec.swift_amt, 2)))
        #     new_row.cells[7].text = str("{:,}".format(float(format(incentive_rate_fc, '.2f'))))
        #
        #     first_cell = new_row.cells[0]
        #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph24 = first_cell.paragraphs[0]
        #     paragraph24.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     paragraph24.style = custom_style5
        #
        #     first_cell = new_row.cells[3]
        #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph21 = first_cell.paragraphs[0]
        #     paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        #     paragraph21.style = custom_style5
        #
        #     first_cell1 = new_row.cells[6]
        #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        #     paragraph22.style = custom_style5
        #
        #     first_cell1 = new_row.cells[7]
        #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        #     paragraph22.style = custom_style5
        #
        #     first_cell1 = new_row.cells[2]
        #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     paragraph22.style = custom_style5
        #
        #     first_cell1 = new_row.cells[1]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.style = custom_style5
        #
        #     first_cell1 = new_row.cells[4]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     paragraph22.style = custom_style5
        #
        #     first_cell1 = new_row.cells[5]
        #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        #     paragraph22 = first_cell1.paragraphs[0]
        #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #     paragraph22.style = custom_style5
        #
        #     sl += 1

        # cus name merge
        cell_1_0 = table.cell(1, 1)
        p3 = cell_1_0.paragraphs[0]
        p3.add_run(cus_name)
        cell_2_0 = table.cell(row_len, 1)
        cell_1_0.merge(cell_2_0)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style4

        # total
        total_row = table.add_row()
        total_row.cells[2].text = 'Total'
        total_row.cells[3].text = str("{:,}".format(round(inv_amount, 2)))
        total_row.cells[6].text = str("{:,}".format(round(swift_amt, 2)))
        total_row.cells[7].text = str("{:,}".format(float(format(total_incentive_rate_fc, '.2f'))))

        first_cell = total_row.cells[3]
        paragraph21 = first_cell.paragraphs[0]
        paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph21.style = table_total

        first_cell1 = total_row.cells[6]
        paragraph22 = first_cell1.paragraphs[0]
        paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph22.style = table_total

        first_cell2 = total_row.cells[7]
        paragraph23 = first_cell2.paragraphs[0]
        paragraph23.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph23.style = table_total

        first_cell2 = total_row.cells[2]
        paragraph23 = first_cell2.paragraphs[0]
        paragraph23.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph23.style = table_total

        # after table
        paragraph2 = document.add_paragraph("\nThose documents are as follows:")
        paragraph2.style = custom_style2
        paragraph2.paragraph_format.left_indent = Pt(-6)

        # paragraph2 = document.add_paragraph("\t 1. Form Ka\n\t 2. Commercial Invoice\n\t 3. Form-C (ICT)\n\t 4. Agreement\n\t 5. Certificate of Authorized Dealer\n\t 6. Copy of Swift Message\n\t 7. নগদ ভর্তুকি প্রত্যয়নপত্র \n\t 8. Company's undertaking")
        # paragraph2.style = custom_style2

        # Create a table with 1 row and 2 columns
        table = document.add_table(rows=1, cols=2)
        table.border = None
        table.columns[0].width = Inches(3)  # Adjust the width as needed
        table.columns[1].width = Inches(3)  # Adjust the width as needed

        # Get the first row of the table
        row = table.rows[0]
        # Remove cell borders
        for cell in row.cells:
            cell.border = None

        left_column = table.cell(0, 0)
        left_column = left_column.paragraphs[0]
        left_column.add_run("1. Form Ka\n")
        left_column.add_run("2. Commercial Invoice\n")
        left_column.add_run("3. Form-C (ICT)\n")
        left_column.add_run("4. Agreement\n")
        left_column.paragraph_format.left_indent = Inches(.7)
        left_column.style = custom_style2
        # left_column.paragraph_format.left_indent = Pt(-6)

        # Add the last four items to the right column
        right_column = table.cell(0, 1)
        right_column = right_column.paragraphs[0]
        right_column.add_run("5. Certificate of Authorized Dealer\n")
        right_column.add_run("6. Copy of Swift Message\n")
        run2 = right_column.add_run("7. নগদ ভর্তুকি প্রত্যয়নপত্র\n")
        run2.font.size = Pt(11)
        right_column.add_run("8. Company's undertaking\n")
        right_column.style = custom_style2
        # right_column.paragraph_format.left_indent = Pt(-6)

        paragraph2 = document.add_paragraph("Please proceed at your earliest possible time. \n\nThanking You \n\n\nMasudur Rahman \nSenior Manager, Finance, Accounts & Admin \nBrain Station 23 Limited")
        paragraph2.style = custom_style2
        paragraph2.paragraph_format.left_indent = Pt(-6)
        document.add_page_break()
        # file_name = '%s_flfb' % (datetime.datetime.now())

        # ga 7777777 ------------------------------
        from docx.enum.section import WD_ORIENT
        # document = Document()
        # section = document.sections[0]
        # new_width, new_height = section.page_height, section.page_width
        # # section.orientation = WD_ORIENT.LANDSCAPE
        # section.page_width = new_width
        # section.page_height = new_height
        document.add_section()
        # Set landscape page dimensions
        landscape_section = document.sections[-1]
        landscape_section.page_width = Pt(792)
        landscape_section.page_height = Pt(612)

        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

        custom_bangla_style = document.styles.add_style('custom_bangla_style39', WD_STYLE_TYPE.PARAGRAPH)
        custom_bangla_style.font.size = Pt(11)
        custom_bangla_style.font.name = 'SutonnyOMJ'

        custom_bangla_style_ga = document.styles.add_style('custom_bangla_style_ga40', WD_STYLE_TYPE.PARAGRAPH)
        custom_bangla_style_ga.font.size = Pt(11)
        custom_bangla_style_ga.font.name = 'SutonnyOMJ'

        custom_stylebold = document.styles.add_style('custom_stylebold41', WD_STYLE_TYPE.PARAGRAPH)
        custom_stylebold.font.size = Pt(14)
        custom_stylebold.font.name = 'SutonnyOMJ'
        custom_stylebold.font.bold = True
        custom_stylebold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        custom_styleubl = document.styles.add_style('custom_styleubl42', WD_STYLE_TYPE.PARAGRAPH)
        custom_styleubl.font.size = Pt(11)
        custom_styleubl.font.name = 'SutonnyOMJ'
        custom_styleubl.font.underline = True
        custom_styleubl.font.bold = True

        # p1 = document.add_paragraph('(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য) \t\t\t \t \t \t \t  \t \t  \t \t \t \t   ফরম-‘গ’ \n')
        # p1.style = custom_bangla_style

        paragraph = document.add_paragraph()
        # paragraph.add_run("(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
        run1 = paragraph.add_run("\n(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
        run1.font.size = Pt(11)
        run1.font.name = 'SutonnyOMJ'
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t")
        paragraph.add_run("\t ")
        paragraph.add_run("\t ")
        paragraph.add_run("\t ")
        run2 = paragraph.add_run("ফরম-‘গ’")
        run2.font.size = Pt(14)
        run2.font.bold = True
        run2.font.name = 'SutonnyOMJ'
        # paragraph.add_run("ফরম-‘গ’")
        # run2.style = custom_bangla_style_ga

        # paragraph = document.add_paragraph()
        #
        # # Add runs with different font sizes to the paragraph
        # run1 = paragraph.add_run("This is text with font size 12.")
        # run1.font.size = Pt(12)
        #
        # run2 = paragraph.add_run(" This is text with font size 16.")
        # run2.font.size = Pt(16)

        # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

        # Define a custom style
        custom_style = document.styles.add_style('MyStyle43', WD_STYLE_TYPE.PARAGRAPH)
        # custom_style.font.bold = True
        custom_style.font.size = Pt(14)
        custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        custom_style2 = document.styles.add_style('MyStyle244', WD_STYLE_TYPE.PARAGRAPH)
        custom_style2.font.size = Pt(11)
        custom_style2.font.name = 'Arial Narrow'

        custom_style3 = document.styles.add_style('MyStyle45', WD_STYLE_TYPE.PARAGRAPH)
        custom_style3.font.size = Pt(11)
        custom_style3.font.bold = True

        paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক")
        paragraph2.style = custom_bangla_style_ga
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        # paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক \nবাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব \n-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
        paragraph2 = document.add_paragraph("বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব")
        paragraph2.style = custom_stylebold
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph("-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
        paragraph2.style = custom_bangla_style_ga
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph2.paragraph_format.space_before = Pt(1)

        # table -----------------------
        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name

        a = []
        c = []
        for rec in self.invoice_line_ids:
            a.append(rec.od_sight_rate)
            if rec.encashment_rate_bdt:
                c.append(rec.encashment_rate_bdt)
            # if rec.swift_message_id.encashment_rate_bdt:
            #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
        inv_amount = sum(self.invoice_line_ids.mapped('swift_amt'))
        if a:
            b = min(a)
            od_s_r = b
        else:
            od_s_r = 0

        if c:
            d = min(c)
            encashment_rate_bdt = d
        else:
            encashment_rate_bdt = 0

        if self.od_sight_rate:
            od_sight_rate = self.od_sight_rate
        else:
            if od_s_r:
                od_sight_rate = od_s_r
            else:
                od_sight_rate = encashment_rate_bdt

        rate_fc = inv_amount / self.incentive_rate_fc
        encashment_final_amt = round(rate_fc, 2) * od_sight_rate
        # encashment_final_amt_a = "{:,.0f}".format(round(encashment_final_amt))
        encashment_final_amt_a = self.custom_round(encashment_final_amt)
        usd_bdt_amount = currency + ' ' + str("{:,.2f}".format(rate_fc)) + '\n' + '@' + \
                         str("{:,.2f}".format(od_sight_rate)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a

        table = document.add_table(rows=4, cols=9)
        table.style = 'TableGrid'
        cell_1_0 = table.cell(0, 0)
        p3 = cell_1_0.paragraphs[0]
        p3.add_run('ভর্তুকি আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানা')
        p3.style = custom_bangla_style
        cell_2_0 = table.cell(1, 0)
        cell_1_0.merge(cell_2_0)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(0, 1)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('রপ্তানি সংশ্লিষ্ট তথ্য')
        p3.style = custom_bangla_style
        cell_2_1 = table.cell(0, 5)
        cell_1_1.merge(cell_2_1)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 1)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('সেবা/পণ্যের বিবরণ')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 2)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('জাহাজীকরণের তারিখ ও গন্তব্যস্থল')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 3)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('ইক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানির ক্ষেত্রে)')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 4)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('মোট প্রত্যাবাসিত রপ্তানি (মূল্য  বৈদেশিক মুদ্রা)')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 5)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('নীট এফওবি মূল্য (বৈদেশিক মুদ্রা)')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_3 = table.cell(0, 6)
        p3 = cell_1_3.paragraphs[0]
        p3.add_run('শাখা কর্তৃক আবেদনপত্র গ্রহণের তারিখ')
        p3.style = custom_bangla_style
        cell_2_4 = table.cell(1, 6)
        cell_1_3.merge(cell_2_4)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_5 = table.cell(0, 7)
        p3 = cell_1_5.paragraphs[0]
        p3.add_run('আবেদনকৃত অর্থের পরিমাণ')
        p3.style = custom_bangla_style
        cell_2_6 = table.cell(1, 7)
        cell_1_5.merge(cell_2_6)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_6 = table.cell(0, 8)
        p3 = cell_1_6.paragraphs[0]
        p3.add_run('পরিশোধ্য দাবির পরিমাণ')
        p3.style = custom_bangla_style
        cell_2_7 = table.cell(1, 8)
        cell_1_6.merge(cell_2_7)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 0)
        p3 = cell.paragraphs[0]
        p3.add_run('১')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 1)
        p3 = cell.paragraphs[0]
        p3.add_run('২')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 2)
        p3 = cell.paragraphs[0]
        p3.add_run('৩')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 3)
        p3 = cell.paragraphs[0]
        p3.add_run('৪')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 4)
        p3 = cell.paragraphs[0]
        p3.add_run('৫')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 5)
        p3 = cell.paragraphs[0]
        p3.add_run('৬')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 6)
        p3 = cell.paragraphs[0]
        p3.add_run('৭')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 7)
        p3 = cell.paragraphs[0]
        p3.add_run('৮')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 8)
        p3 = cell.paragraphs[0]
        p3.add_run('৯')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 0)

        cell.height = Inches(12)
        p3 = cell.paragraphs[0]
        p3.add_run('\n\nব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী  বাণিজ্যিক এলাকা, ঢাকা- ১২১২।\n\n')
        p3.style = custom_bangla_style
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # p3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cell = table.cell(3, 1)
        cell.height = Inches(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('Software Development')
        p3.style = custom_style2
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 2)
        cell.height = Inches(12)
        p3 = cell.paragraphs[0]
        p3.add_run('')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 3)
        cell.height = Inches(12)
        p3 = cell.paragraphs[0]
        p3.add_run('')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.height = Inches(12)
        p3 = cell.paragraphs[0]
        p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        cell = table.cell(3, 5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.height = Inches(12)
        p3 = cell.paragraphs[0]
        p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_style3

        cell = table.cell(3, 6)
        cell.height = Inches(12)
        p3 = cell.paragraphs[0]
        p3.add_run('')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 7)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.height = Inches(12)
        p3 = cell.paragraphs[0]
        p3.add_run(usd_bdt_amount)
        p3.style = custom_style3
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 8)
        cell.height = Inches(12)
        p3 = cell.paragraphs[0]
        p3.add_run('')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        column_index = 0
        column_cells = [row.cells[column_index] for row in table.rows]

        # Set the desired height for each cell in the column
        cell_height = Inches(12)  # Adjust the height value as needed
        for cell in column_cells:
            cell.height = cell_height

        p3 = document.add_paragraph('\n\n\n\n\n\t\t দাপ্তরিক সীল  \t\t\t\t\t\t \t\t\t\t  ক্ষমতাপ্রাপ্ত ব্যাংক কর্মকর্তার স্বাক্ষর, তারিখ, নাম ও পদবী')
        p3.style = custom_bangla_style
        # file_name = '%s_ga' % (datetime.datetime.now())
        document.add_page_break()

        # gha 77777777  ------------------------------
        from docx.enum.section import WD_ORIENT


        # section = document.sections[0]
        # section.orientation = WD_ORIENT.LANDSCAPE
        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')
        # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

        # Define a custom style
        custom_style = document.styles.add_style('MyStyle47', WD_STYLE_TYPE.PARAGRAPH)
        # custom_style.font.bold = True
        custom_style.font.size = Pt(14)
        custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        custom_style2 = document.styles.add_style('MyStyle48', WD_STYLE_TYPE.PARAGRAPH)
        custom_style2.font.size = Pt(11)
        custom_style2.font.name = 'Arial Narrow'

        custom_style3 = document.styles.add_style('MyStyle49', WD_STYLE_TYPE.PARAGRAPH)
        custom_style3.font.size = Pt(11)
        custom_style3.font.bold = True

        custom_style4 = document.styles.add_style('MyStyle50', WD_STYLE_TYPE.PARAGRAPH)
        custom_style4.font.bold = True
        custom_style4.font.size = Pt(14)
        custom_style4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        custom_style4.font.name = 'SutonnyOMJ'

        custom_bangla_style = document.styles.add_style('custom_bangla_style51', WD_STYLE_TYPE.PARAGRAPH)
        custom_bangla_style.font.size = Pt(11)
        custom_bangla_style.font.name = 'SutonnyOMJ'

        custom_bangla_style_gha = document.styles.add_style('custom_bangla_style_gha52', WD_STYLE_TYPE.PARAGRAPH)
        custom_bangla_style_gha.font.size = Pt(11)
        custom_bangla_style_gha.font.name = 'SutonnyOMJ'

        custom_stylebold = document.styles.add_style('custom_stylebold53', WD_STYLE_TYPE.PARAGRAPH)
        custom_stylebold.font.size = Pt(16)
        custom_stylebold.font.name = 'SutonnyOMJ'
        custom_stylebold.font.bold = True
        custom_stylebold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        custom_styleubl = document.styles.add_style('custom_styleubl54', WD_STYLE_TYPE.PARAGRAPH)
        custom_styleubl.font.size = Pt(13)
        custom_styleubl.font.name = 'SutonnyOMJ'
        custom_styleubl.font.bold = True
        custom_styleubl.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph1 = document.add_paragraph('\nফরম-‘ঘ’')
        paragraph1.style = custom_style4

        p3 = document.add_paragraph('(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য) \n')
        p3.style = custom_bangla_style_gha

        # paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক \nবাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধিত হিসাবের .........সালের ...... মাসের বিবরণী \nঅনুমোদিত ডিলার ব্যাংক শাখার নামঃ...................................")

        paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক")
        paragraph2.style = custom_bangla_style_gha
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph("বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও")
        paragraph2.style = custom_stylebold
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph("হার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধিত হিসাবের .........সালের ...... মাসের বিবরণী")
        paragraph2.style = custom_styleubl
        paragraph2.paragraph_format.space_before = Pt(1)
        paragraph2.paragraph_format.space_after = Pt(1)

        paragraph2 = document.add_paragraph("অনুমোদিত ডিলার ব্যাংক শাখার নামঃ...................................................................................")
        paragraph2.style = custom_bangla_style_gha
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph2.paragraph_format.space_before = Pt(1)

        # table -----------------------
        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name

        a = []
        c = []
        for rec in self.invoice_line_ids:
            a.append(rec.od_sight_rate)
            if rec.encashment_rate_bdt:
                c.append(rec.encashment_rate_bdt)
            # if rec.swift_message_id.encashment_rate_bdt:
            #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
        inv_amount = sum(self.invoice_line_ids.mapped('swift_amt'))

        if a:
            b = min(a)
            od_s_r = b
        else:
            od_s_r = 0

        if c:
            d = min(c)
            encashment_rate_bdt = d
        else:
            encashment_rate_bdt = 0

        if self.od_sight_rate:
            od_sight_rate = self.od_sight_rate
        else:
            if od_s_r:
                od_sight_rate = od_s_r
            else:
                od_sight_rate = encashment_rate_bdt

        rate_fc = inv_amount / self.incentive_rate_fc
        encashment_final_amt = round(rate_fc, 2) * od_sight_rate
        # encashment_final_amt_a = "{:,.0f}".format(round(encashment_final_amt))
        encashment_final_amt_a = self.custom_round(encashment_final_amt)
        usd_bdt_amount = currency + ' ' + str("{:,.2f}".format(rate_fc)) + '\n' + '@' + \
                         str("{:,.2f}".format(od_sight_rate)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a

        table = document.add_table(rows=4, cols=10)
        table.style = 'TableGrid'

        cell_1_0 = table.cell(0, 0)
        p3 = cell_1_0.paragraphs[0]
        p3.add_run('ভর্তুকি আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানা')
        p3.style = custom_bangla_style_gha
        cell_2_0 = table.cell(1, 0)
        cell_1_0.merge(cell_2_0)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(0, 1)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('রপ্তানি সংশ্লিষ্ট তথ্য')
        p3.style = custom_bangla_style_gha
        cell_2_1 = table.cell(0, 5)
        cell_1_1.merge(cell_2_1)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_2_1 = table.cell(0, 8)
        p3 = cell_2_1.paragraphs[0]
        p3.add_run('শাখায় প্রাপ্ত অনিস্পন্ন আবেদন')
        p3.style = custom_bangla_style_gha
        cell_3_1 = table.cell(0, 9)
        cell_2_1.merge(cell_3_1)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 1)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('সেবা/পণ্যের বিবরণ')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 2)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('জাহাজীকরণের তারিখ ও গন্তব্যস্থল')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 3)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('ইক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানির ক্ষেত্রে)')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 4)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('মোট প্রত্যাবাসিত রপ্তানি (মূল্য  বৈদেশিক মুদ্রা)')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_1 = table.cell(1, 5)
        p3 = cell_1_1.paragraphs[0]
        p3.add_run('নীট এফওবি মূল্য (বৈদেশিক মুদ্রা)')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_3 = table.cell(0, 6)
        table.columns[7].width = Inches(.6)
        p3 = cell_1_3.paragraphs[0]
        p3.add_run('পরিশোধিত ভর্তুকি (টাকা)')
        p3.style = custom_bangla_style_gha
        cell_2_4 = table.cell(1, 6)
        cell_1_3.merge(cell_2_4)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_5 = table.cell(0, 7)
        table.columns[7].width = Inches(.6)
        p3 = cell_1_5.paragraphs[0]
        p3.add_run('পরিশোধের তারিখ')
        p3.style = custom_bangla_style_gha
        cell_2_6 = table.cell(1, 7)
        cell_1_5.merge(cell_2_6)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_6 = table.cell(1, 8)
        table.columns[8].width = Inches(.5)
        p3 = cell_1_6.paragraphs[0]
        p3.add_run('সংখ্যা')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_1_6 = table.cell(1, 9)
        table.columns[9].width = Inches(1.7)
        p3 = cell_1_6.paragraphs[0]
        p3.add_run('দাবীকৃত অর্থ (টাকা)')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 0)
        p3 = cell.paragraphs[0]
        p3.add_run('১')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 1)
        p3 = cell.paragraphs[0]
        p3.add_run('২')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 2)
        p3 = cell.paragraphs[0]
        p3.add_run('৩')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 3)
        p3 = cell.paragraphs[0]
        p3.add_run('৪')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 4)
        p3 = cell.paragraphs[0]
        p3.add_run('৫')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 5)
        p3 = cell.paragraphs[0]
        p3.add_run('৬')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 6)
        p3 = cell.paragraphs[0]
        p3.add_run('৭')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 7)
        p3 = cell.paragraphs[0]
        p3.add_run('৮')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 8)
        p3 = cell.paragraphs[0]
        p3.add_run('৯')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(2, 9)
        p3 = cell.paragraphs[0]
        p3.add_run('১০')
        p3.style = custom_bangla_style_gha
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run(
            '\nব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী  বাণিজ্যিক এলাকা, ঢাকা- ১২১২।\n\n')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.style = custom_bangla_style_gha

        cell = table.cell(3, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run('Software Development')
        p3.style = custom_style2
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 2)
        p3 = cell.paragraphs[0]
        p3.add_run('')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 3)
        p3 = cell.paragraphs[0]
        p3.add_run('')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
        p3.style = custom_style3
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
        p3.style = custom_style3
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 6)
        p3 = cell.paragraphs[0]
        p3.add_run('')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 7)
        p3 = cell.paragraphs[0]
        p3.add_run('')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 8)
        p3 = cell.paragraphs[0]
        p3.add_run('')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell = table.cell(3, 9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p3 = cell.paragraphs[0]
        p3.add_run(usd_bdt_amount)
        p3.style = custom_style3
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p3 = document.add_paragraph(
            '\n\n\n\n\t\t দাপ্তরিক সীল  \t\t\t\t\t\t\t\t\t  ক্ষমতাপ্রাপ্ত ব্যাংক কর্মকর্তার স্বাক্ষর, তারিখ, নাম ও পদবী')
        p3.style = custom_bangla_style_gha
        value = self.name
        split_values = value.split("/")
        desired_value = split_values[-1]
        file_name = 'All_form_%s_%s' % (desired_value, datetime.datetime.now())

        # -------------------
        import os
        dir_path = os.path.dirname(os.path.abspath(__file__))
        base_path = str(dir_path).replace('/models', '')
        docxfile = base_path + '/static/docx/' + file_name + '.docx'
        document.save(docxfile)
        # docx.save(docxfile)

        # return document.save('/home/jobaer/Downloads/jh3.docx')

        return {
            'type': 'ir.actions.act_url',
            'url': 'cash_incentive/static/docx/' + file_name + '.docx',
            'target': 'self',
        }

    import math
    def custom_round(self, value):
        decimal_part = value - math.floor(value)
        if decimal_part <= 0.49:
            a = math.floor(value)
            return str("{:,.0f}".format(round(a)))
        else:
            a = math.ceil(value)
            return str("{:,.0f}".format(round(a)))

    def num2words_fun(self,number):
        num = decimal.Decimal(number)
        decimal_part = num - int(num)
        num = int(num)

        # if decimal_part:
        #     return num2words(num) + " point " + (" ".join(num2words(i) for i in str(decimal_part)[2:]))

        under_20 = ['Zero', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine', 'Ten', 'Eleven',
                    'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen', 'Seventeen', 'Eighteen', 'Nineteen']
        tens = ['Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
        above_100 = {100: 'Hundred', 1000: 'Thousand', 100000: 'Lac', 10000000: 'Crores'}

        if num < 20:
            return under_20[num]

        if num < 100:
            return tens[num // 10 - 2] + ('' if num % 10 == 0 else ' ' + under_20[num % 10])

        # find the appropriate pivot - 'Million' in 3,603,550, or 'Thousand' in 603,550
        pivot = max([key for key in above_100.keys() if key <= num])

        return num2words(num // pivot) + ' ' + above_100[pivot] + (
            '' if num % pivot == 0 else ' ' + num2words(num % pivot))

    def set_cell_margins(self, cell, **kwargs):
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        """
        cell:  actual cell instance you want to modify
        usage:
            set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

        provided values are in twentieths of a point (1/1440 of an inch).
        read more here: http://officeopenxml.com/WPtableCellMargins.php
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for m in ["top", "start", "bottom", "end"]:
            if m in kwargs:
                node = OxmlElement("w:{}".format(m))
                node.set(qn('w:w'), str(kwargs.get(m)))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

        tcPr.append(tcMar)

    def action_download_docx_file2(self):
        type = self.env.context.get('type')
        from docx import Document
        from htmldocx import HtmlToDocx
        from docx.shared import Inches
        # create a new document
        document = Document()
        jh = Document()
        new_parser = HtmlToDocx()

        from docx import Document
        from docx.shared import Pt

        # path to doc with altered style:
        base_doc_location = 'base.docx'
        doc = Document(base_doc_location)
        my_style = doc.styles['Normal']

        # define your desired fonts
        user_cs_font_size = 16
        user_cs_font_name = 'SolaimanLipi'
        user_en_font_size = 12
        user_en_font_name = 'FreeMono'

        # get <w:rPr> element of this style
        rpr = my_style.element.rPr

        # ==================================================
        '''This probably isn't necessary if you already
        have a document with altered style, but just to be
        safe I'm going to add this here'''

        if rpr.rFonts is None:
            rpr._add_rFonts()
        if rpr.sz is None:
            rpr._add_sz()
        # ==================================================

        '''Get the nsmap string for rpr. This is that "w:"
        at the start of elements and element values in xml.
        Like these:
            <w:rPr>
            <w:rFonts>
            w:val

        The nsmap is like a url:
        http://schemas.openxmlformats.org/...

        Now w:rPr translates to:
        {nsmap url string}rPr

        So I made the w_nsmap string like this:'''

        w_nsmap = '{' + rpr.nsmap['w'] + '}'
        # ==================================================

        '''Because I didn't find any better ways to get an
        element based on its tag here's a not so great way
        of getting it:
        '''
        szCs = None
        lang = None

        for element in rpr:
            if element.tag == w_nsmap + 'szCs':
                szCs = element
            elif element.tag == w_nsmap + 'lang':
                lang = element

        '''if there is a szCs and lang element in your style
        those variables will be assigned to it, and if not
        we make those elements and add them to rpr'''

        if szCs is None:
            szCs = rpr.makeelement(w_nsmap + 'szCs', nsmap=rpr.nsmap)
        if lang is None:
            lang = rpr.makeelement(w_nsmap + 'lang', nsmap=rpr.nsmap)

        rpr.append(szCs)
        rpr.append(lang)
        # ==================================================

        '''Now to set our desired values to these elements
        we have to get attrib dictionary of these elements
        and set the name of value as key and our value as
        value for that dict'''

        szCs_attrib = szCs.attrib
        lang_attrib = lang.attrib
        rFonts_atr = rpr.rFonts.attrib

        '''sz and szCs values are string values and 2 times
        the font size so if you want font size to be 11 you
        have to set sz (for western fonts) or szCs (for CTL
        fonts) to "22" '''
        szCs_attrib[w_nsmap + 'val'] = str(int(user_cs_font_size * 2))

        '''Now to change cs font and bidi lang values'''
        rFonts_atr[w_nsmap + 'cs'] = user_cs_font_name
        lang_attrib[w_nsmap + 'bidi'] = 'fa-IR'  # For Persian
        # ==================================================

        '''Because we changed default style we don't even
        need to set style every time we add a new paragraph
        And if you change font name or size the normal way
        it won't change these cs values so you can have a
        font for CTL language and a different font for
        western language
        '''
        doc.add_paragraph('কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।')
        en_font = my_style.font
        en_font.name = user_en_font_name
        en_font.size = Pt(user_en_font_size)
        doc.add_paragraph('some text')

        doc.save('ex.docx')
        # Save the document
            # doc.save('my_document.docx')

        file_name = '%s_prc' % (datetime.datetime.now())

        # -------------------
        import os
        dir_path = os.path.dirname(os.path.abspath(__file__))
        base_path = str(dir_path).replace('/models', '')
        docxfile = base_path + '/static/docx/' + file_name + '.docx'
        document.save(docxfile)
        # docx.save(docxfile)

        # return document.save('/home/jobaer/Downloads/jh3.docx')

        return {
            'type': 'ir.actions.act_url',
            'url': 'cash_incentive/static/docx/' + file_name + '.docx',
            'target': 'self',
        }

    def get_word_prc(self):
        from docx import Document
        from htmldocx import HtmlToDocx
        from docx.shared import Inches
        # create a new document
        document = Document()
        #self.prc_letter_description = self.bank_id.prc_letter_description
        currency = ''
        if self.fc_currency_id:
            currency = self.fc_currency_id.name
        self_date = ''
        if self.date:
            self_date = datetime.datetime.strptime(str(self.date), '%Y-%m-%d').strftime('%d/%m/%y')

        html = ''

        inv_number = ''
        inv_date = ''
        inv_amount = 0
        encashment_amt_bdt = 0
        cus_name = ''
        if self.swift_customer_name:
            cus_name = self.swift_customer_name
        for rec in self.invoice_line_ids:
            inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
            invoice_date = ''
            if rec.invoice_date:
                invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
            inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
            inv_amount += rec.invoice_amt
            encashment_amt_bdt += rec.encashment_amt_bdt
            date_credited_beneficiaries = ''
            if rec.date_credited_beneficiaries:
                date_credited_beneficiaries = datetime.datetime.strptime(str(rec.date_credited_beneficiaries), '%Y-%m-%d').strftime('%d-%b-%y')
            swift_date = ''
            if rec.swift_message_id.date:
                swift_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime('%d-%b-%y')
            ref_online_to_bb = rec.ref_online_to_bb if rec.ref_online_to_bb else ''
            reporting_st_to_bb = rec.reporting_st_to_bb if rec.reporting_st_to_bb else ''

        prc_text = self.bank_id.prc_letter_description

        final_text = ''
        if prc_text:
            final_text = prc_text.replace('$ref_name', self.name).replace('$date', str(self_date)).replace('$contract_number', '' if not self.contract_number else str(self.contract_number))\
                .replace('$contract_date', '' if not self.contract_date_str else str(self.contract_date_str)).replace('$invoice_number', inv_number)\
                .replace('$invoice_date', '& date: '+inv_date).replace('$currency', currency).replace('$total_amount', str("{:,}".format(inv_amount)))
        return final_text

        # import docx, io
        # from htmldocx import HtmlToDocx
        # document = docx.Document()
        # new_parser = HtmlToDocx()
        # # html_text = '<h1>jobaer</h1>'
        # # print(html_text)
        # html_text = self.prc_letter_description.replace('<br>', '<p></p>')
        # new_parser.add_html_to_document(html_text, document)
        # # do more stuff to document
        # # print(datetime.datetime.now())
        # file_name = '%s_prc' % (datetime.datetime.now())
        #
        # # -------------------
        # import os
        # dir_path = os.path.dirname(os.path.abspath(__file__))
        # base_path = str(dir_path).replace('/models', '')
        # docxfile = base_path + '/static/cash_incentive_docx/' + file_name + '.docx'
        # document.save(docxfile)
        # # return document.save('/home/jobaer/Downloads/jh3.docx')
        #
        # return {
        #     'type': 'ir.actions.act_url',
        #     'url': 'cash_incentive/static/cash_incentive_docx/' + file_name + '.docx',
        #     'target': 'self',
        # }


class CashIncentiveInvoices(models.Model):
    _name = "cash.incentive.invoice"
    _description = "Cash Incentive Invoice"
    _rec_name = 'invoice_id'
    _order = 'invoice_ref asc'

    head_id = fields.Many2one('cash.incentive.head', ondelete='cascade', string='Cash Incentive', tracking=True)
    date = fields.Date(string='Preparation Date', required=True, default=fields.Date.context_today, tracking=3)
    partner_id = fields.Many2one('res.partner', string='Customer', tracking=True)
    swift_message_id = fields.Many2one('swift.message', ondelete='cascade', string='SWIFT Message', tracking=True)
    bank_id = fields.Many2one('res.bank', string='Bank', related='swift_message_id.bank_id', domain="[('is_cash_incentive_bank', '=', True)]", tracking=True)
    swift_customer_name = fields.Char(related='swift_message_id.swift_customer_name', tracking=True)
    od_sight_rate = fields.Float(string='OD Sight Rate', digits=(16, 4), tracking=True)
    incentive_rate_fc = fields.Float(string='Incentive Rate (FC)(%)', digits=(16, 4), related="swift_message_id.incentive_rate_fc", default=10, tracking=True)

    application_deadline = fields.Date(string='Application Deadline', tracking=5,
                                       help='Application Deadline Will be 179 Days More than SWIFT Date', compute='_compute_application_deadline')
    swift_date = fields.Date(string='SWIFT Date', related='swift_message_id.date')
    remaining_days = fields.Integer(string='Remaining days', compute='_compute_remaining_day', tracking=6, search='_value_search')

    #-----------
    invoice_id = fields.Many2one('account.move', string='Invoice', required=True, ondelete='cascade', tracking=True)
    invoice_ref = fields.Char(related='invoice_id.ref', store=True, tracking=True)
    invoice_partner_id = fields.Many2one(related='invoice_id.partner_id', tracking=True)
    invoice_customer_type = fields.Selection(related='invoice_partner_id.vendor_type', string='Type (Local/Foreign)', tracking=True)
    invoice_no = fields.Char(related='invoice_id.name', tracking=True)
    foreign_currency_type = fields.Many2one(related='invoice_id.foreign_currency_type', tracking=True)
    invoice_total_fc = fields.Float(related='invoice_id.invoice_total_fc', digits=(16, 2))
    invoice_payment_amount_fc = fields.Float(related='invoice_id.invoice_payment_amount_fc', digits=(16, 2), tracking=True)
    invoice_remaining_amount_fc = fields.Float(related='invoice_id.invoice_remaining_amount_fc', digits=(16, 2), tracking=True)
    amount_total_signed = fields.Monetary(related='invoice_id.amount_total_signed', tracking=True)
    invoice_state = fields.Selection(related='invoice_id.state', tracking=True)
    cash_incentive_date = fields.Date(related='head_id.date', tracking=True)

    contract_id = fields.Many2one("client.contract", string="Contract No.", tracking=20)
    contract_ids = fields.Many2many("client.contract", string="Contract No.", tracking=20)

    swift_bank_id = fields.Many2one(related='swift_message_id.bank_id', tracking=True)
    encashment_bank_id = fields.Many2one(related='swift_message_id.encashment_bank_id', tracking=True)

    usd_rate = fields.Float(string='Invoice BDT Rate', digits=(16, 4), related='invoice_id.usd_rate', tracking=True)
    #self.invoice_id.swift_remaining_amount

    invoice_date = fields.Date(string='Invoice Date', tracking=True)
    invoice_qty_str = fields.Char(string='Quantity', tracking=True)
    invoice_amt = fields.Float(string='Invoice Amount (FC)', digits=(16, 2), tracking=True)
    basis_fee_amt = fields.Float(string='Basis Fee Amount', digits=(16, 2), tracking=True)

    swift_amt = fields.Float(string='Encashable SWIFT Amount (FC)', digits=(16, 2), tracking=True)
    swift_amt_bdt = fields.Float(string='Encashable Amount (BDT)', compute="_compute_swift_amt_bdt", digits=(16, 2), tracking=True)
    swift_charge_fc = fields.Float(string='Bank Charge (FC)', digits=(16, 2), tracking=True)

    #swift_charge_rate = fields.Float(string='Bank Charge Rate (BDT)', default=0)
    swift_charge_bdt = fields.Float(string='Bank Charge (BDT)', compute="_compute_swift_amt_bdt", digits=(16, 2), tracking=True)
    total_swift_amt = fields.Float(string='Receivable Credit (FC)', compute="_compute_swift_amt_bdt", digits=(16, 2), tracking=True)
    total_swift_amt_bdt = fields.Float(string='Receivable Credit (BDT)', compute="_compute_swift_amt_bdt", digits=(16, 2), tracking=True)
    due_swift_amt = fields.Float(string='Due Receivable (FC)', compute="_compute_swift_amt_bdt", digits=(16, 2), tracking=True)
    due_swift_amt_bdt = fields.Float(string='Due Receivable (BDT)', compute="_compute_swift_amt_bdt", digits=(16, 2), tracking=True)

    currency_id = fields.Many2one("res.currency", string="Currency", tracking=True)

    incentive_amt_fc = fields.Float(string='Incentive Amount (FC)', digits=(16, 2), compute="_compute_incentive_amt_fc", tracking=True)
    incentive_amt_bdt = fields.Float(string='Incentive Amount (BDT)', compute="_compute_incentive_amt_bdt", digits=(16, 2), tracking=True)

    #---encashment
    encashment_date = fields.Date(string='Encashment Date', related="swift_message_id.encashment_date", store=True, tracking=True)

    encashment_amt_fc = fields.Float(string='Encashment (FC)', digits=(16, 2), default=0, tracking=True)
    encashment_rate_bdt = fields.Float(string='Encashment Rate (BDT)', digits=(16, 4), related="swift_message_id.encashment_rate_bdt", tracking=True)
    encashment_amt_bdt = fields.Float(string='Encashment Amount (BDT)', digits=(16, 2), compute="_compute_encashment_amt_bdt", tracking=True)

    erq_amt_fc = fields.Float(string='ERQ (FC)', digits=(16, 2))
    erq_rate_bdt = fields.Float(string='ERQ Rate (BDT)', digits=(16, 4), related="swift_message_id.inter_bank_rate_bdt", tracking=True)
    erq_amt_bdt = fields.Float(string='ERQ Amount (BDT)', digits=(16, 2), compute="_compute_encashment_amt_bdt", tracking=True)
    difference_amnt_bdt = fields.Float(string='Foreign Exchange Gain/Loss (BDT)', digits=(16, 2), compute="_compute_encashment_amt_bdt", tracking=True)
    #----------

    invoice_amt_bdt = fields.Float(string='Invoice Amount (BDT)', digits=(16, 2), compute="_compute_invoice_amt_bdt", tracking=True)
    contract_number = fields.Char(string='Contract No.', tracking=True)
    contract_date_str = fields.Char(string='Contract Date', tracking=True)
    contract_price_str = fields.Char(string='Contract Price', tracking=True)

    date_credited_beneficiaries = fields.Date(string='Date Credited Beneficiaries', tracking=10,)
    reporting_st_to_bb = fields.Char(string='Reporting statement/schedule to BB with Month', tracking=11,)
    ref_online_to_bb = fields.Char(string='Reference of Online reporting to BB', tracking=12)
    is_show = fields.Boolean(default=False, compute="_compute_is_show", tracking=True)
    swift_msg_state = fields.Selection([
        ('draft', 'Draft'),
        ('confirm', 'Confirmed'),
        ('er', 'Encashment Rate'),
        ('pay', 'Done'),
        ('cancel', 'Cancelled'),
    ], string='Swift Msg. Status', copy=False, related="swift_message_id.state", tracking=True)

    incentive_state = fields.Selection([
        ('draft', 'Draft'),
        ('confirm', 'Confirmed'),
        ('approve', 'Approved'),
        ('reject', 'Rejected'),
        ('done', 'Completed'),
        ('cancel', 'Cancelled'),
    ], string='Cash Incentive Status', copy=False, related="head_id.state", tracking=True)

    # comment-for-upgrade
    # ---------------
    # partner_cr_acc_id = fields.Many2one('account.account', 'Credit Account',
    #                                     domain="[('user_type_id.type', '!=', 'view')]", related="swift_message_id.partner_cr_acc_id")
    # encashment_acc_id = fields.Many2one('account.account', 'Encashment (Dr.)',
    #                                     domain="[('user_type_id.type', '!=', 'view')]", related="swift_message_id.encashment_acc_id")
    # bank_charge_acc_id = fields.Many2one('account.account', 'Bank Charge (Dr.)',
    #                                      domain="[('user_type_id.type', '!=', 'view')]", related="swift_message_id.bank_charge_acc_id")
    # fc_gain_loss_acc_id = fields.Many2one('account.account', 'Foreign Exchange Gain/Loss',
    #                                       domain="[('user_type_id.type', '!=', 'view')]", related="swift_message_id.fc_gain_loss_acc_id")
    # erq_acc_id = fields.Many2one('account.account', 'ERQ (Dr.)', domain="[('user_type_id.type', '!=', 'view')]", related="swift_message_id.erq_acc_id")

    # -----------------

    partner_cr_acc_id = fields.Many2one('account.account', 'Credit Account',
                                        related="swift_message_id.partner_cr_acc_id")
    encashment_acc_id = fields.Many2one('account.account', 'Encashment (Dr.)',
                                        related="swift_message_id.encashment_acc_id")
    bank_charge_acc_id = fields.Many2one('account.account', 'Bank Charge (Dr.)',
                                         related="swift_message_id.bank_charge_acc_id")
    fc_gain_loss_acc_id = fields.Many2one('account.account', 'Foreign Exchange Gain/Loss',
                                          related="swift_message_id.fc_gain_loss_acc_id")
    erq_acc_id = fields.Many2one('account.account', 'ERQ (Dr.)',
                                 related="swift_message_id.erq_acc_id")

    def _value_search(self, operator, value):
        recs = self.search([]).filtered(lambda x: x.remaining_days <= value)
        if recs:
            return [('id', 'in', [x.id for x in recs] if recs else False)]
        else:
            return [('id', '!=', 0)]

    @api.onchange('date')
    def _onchange_head_id(self):
        for rec in self:
            rec.contract_id = rec.head_id.contract_id.id
            # contract_ids = [x.id for x in rec.contract_id]
            # rec.invoice_id.contract_ids = contract_ids
            # rec.contract_number = self.head_id.contract_id.code
            # rec.contract_date_str = self.head_id.contract_id.date
            # rec.reference_id = self.crm_id.reference_id.id

    @api.depends("invoice_id")
    def _compute_is_show(self):
        is_show = False
        is_show = self.env.context.get('is_show')
        self.is_show = is_show

    @api.depends('swift_date')
    def _compute_application_deadline(self):
        for rec in self:
            if rec.swift_date:
                rec.application_deadline = rec.swift_date + datetime.timedelta(days=179)
            else:
                rec.application_deadline = None

    @api.depends("swift_amt", "swift_charge_fc", "usd_rate")#, "swift_charge_rate"
    def _compute_swift_amt_bdt(self):
        for rec in self:
            rec.swift_amt_bdt = round(rec.swift_amt * rec.usd_rate, 2)
            rec.swift_charge_bdt = round(rec.swift_charge_fc * rec.usd_rate, 2)
            rec.total_swift_amt = round(rec.swift_amt + rec.swift_charge_fc, 2)
            rec.total_swift_amt_bdt = round(rec.total_swift_amt * rec.usd_rate, 2)

            due_swift_amt = round(rec.invoice_amt - rec.total_swift_amt, 2)
            due_swift_amt_bdt = round(rec.due_swift_amt * rec.usd_rate, 2)
            if due_swift_amt == (-0):
                due_swift_amt = 0
            if due_swift_amt_bdt == (-0):
                due_swift_amt_bdt = 0

            rec.due_swift_amt = due_swift_amt
            rec.due_swift_amt_bdt = due_swift_amt_bdt

    # @api.depends("total_swift_amt_bdt", "encashment_amt_bdt")
    # def _compute_difference_amnt_bdt(self):
    #     for rec in self:
    #         rec.difference_amnt_bdt = rec.encashment_amt_bdt - rec.total_swift_amt_bdt

    @api.depends("application_deadline", "swift_date")
    def _compute_remaining_day(self):
        for rec in self:
            today = fields.Date.today()
            application_deadline = rec.application_deadline
            remaining_days=0
            if today and application_deadline:
                try:
                    remaining_days = (application_deadline - today).days
                except:
                    remaining_days = 0
            rec.remaining_days = remaining_days

    def action_confirm(self):
        pass

    def save_data(self):
        return {'type': 'ir.actions.act_window_close'}

    def remove_data(self):
        head_id = self.head_id
        self.head_id = None
        head_id.on_change_invoice_line_ids()
        return {'type': 'ir.actions.act_window_close'}

    # @api.constrains('invoice_id', 'swift_message_id', 'head_id')
    # def _check_unique_invoice_id(self):
    #     envobj = self.env['cash.incentive.invoice']
    #     for rec in self:
    #         msg = '"%s"' % rec.invoice_id.ref
    #         record = []
    #         h_name = ''
    #         if rec.head_id:
    #             record = envobj.sudo().search([('invoice_id', '=', rec.invoice_id.id), ('head_id', '=', rec.head_id.id)])
    #             h_name = rec.head_id.name
    #
    #         elif rec.swift_message_id:
    #             record = envobj.sudo().search([('swift_message_id', '=', rec.swift_message_id.id), ('invoice_id', '=', rec.invoice_id.id)])
    #             h_name = rec.swift_message_id.code
    #
    #         if len(record) > 1:
    #             raise exceptions.ValidationError("'" + msg + "' already exists in  '%s'!" %(h_name))

    def action_open_edit_invoice(self):
        ir_model_data = self.env['ir.model.data']
        try:
            invoice_form_id = ir_model_data.get_object_reference('cash_incentive', 'view_cash_incentive_invoice_form')[1]
        except ValueError:
            invoice_form_id = False
        ctx = dict(self.env.context or {})
        ctx.update({
            'default_model': 'cash.incentive.invoice',
            'active_model': 'cash.incentive.invoice',
            'active_id': self.ids[0],
            'default_res_id': self.ids[0],
        })
        return {
            'name': _('Invoice'),
            'view_type': 'form',
            'view_mode': 'form',
            'res_model': 'cash.incentive.invoice',
            'res_id': self.id,
            'views': [(invoice_form_id, 'form')],
            'view_id': invoice_form_id,
            'type': 'ir.actions.act_window',
            'target': 'new',
            'context': {'is_show': True},
        }

    @api.depends("incentive_rate_fc", "swift_amt")
    def _compute_incentive_amt_fc(self):
        for rec in self:
            incentive_amt_fc = 0
            if rec.incentive_rate_fc > 0 and rec.swift_amt > 0:
                try:
                    incentive_amt_fc = round((rec.incentive_rate_fc * rec.swift_amt) / 100, 2)
                except:
                    incentive_amt_fc = 0
            rec.incentive_amt_fc = incentive_amt_fc

    @api.depends("od_sight_rate", "incentive_amt_fc")
    def _compute_incentive_amt_bdt(self):
        for rec in self:
            incentive_amt_bdt = 0
            if rec.od_sight_rate > 0 and rec.incentive_amt_fc > 0:
                try:
                    incentive_amt_bdt = round((rec.od_sight_rate * rec.incentive_amt_fc), 2)
                except:
                    incentive_amt_bdt = 0
            rec.incentive_amt_bdt = incentive_amt_bdt

    @api.depends("invoice_amt", "usd_rate")
    def _compute_invoice_amt_bdt(self):
        for rec in self:
            invoice_amt_bdt = 0
            if rec.usd_rate > 0 and rec.invoice_amt > 0:
                try:
                    invoice_amt_bdt = round((rec.usd_rate * rec.invoice_amt), 2)
                    # invoice_amt_bdt = round((rec.usd_rate * rec.swift_amt), 2)
                except:
                    invoice_amt_bdt = 0
            rec.invoice_amt_bdt = invoice_amt_bdt

    @api.depends("encashment_rate_bdt", "encashment_amt_fc", "erq_rate_bdt", "erq_amt_fc")
    def _compute_encashment_amt_bdt(self):
        for rec in self:
            encashment_amt_bdt = 0
            erq_amt_bdt = 0
            if rec.encashment_rate_bdt > 0 and rec.encashment_amt_fc > 0:
                try:
                    encashment_amt_bdt = round((rec.encashment_rate_bdt * rec.encashment_amt_fc), 2)
                except:
                    encashment_amt_bdt = 0
            #----------
            if rec.erq_rate_bdt > 0 and rec.erq_amt_fc > 0:
                try:
                    erq_amt_bdt = round((rec.erq_rate_bdt * rec.erq_amt_fc), 2)
                except:
                    erq_amt_bdt = 0

            rec.encashment_amt_bdt = encashment_amt_bdt
            rec.erq_amt_bdt = erq_amt_bdt

            rec.difference_amnt_bdt = (encashment_amt_bdt + erq_amt_bdt) - rec.swift_amt_bdt

    @api.onchange('swift_message_id')
    def onchange_swift_message_id(self):
        if self.swift_message_id:
            self.ref_online_to_bb = self.swift_message_id.ref_online_to_bb
            self.date_credited_beneficiaries = self.swift_message_id.date_credited_beneficiaries
            self.reporting_st_to_bb = self.swift_message_id.reporting_st_to_bb

    # @api.onchange('contract_id')
    # def _onchange_contract_id(self):
    #     if self.contract_id:
    #         self.contract_number = self.contract_id.code
    #         # print([x.id for x in self.contract_id])
    #         contract_ids = [x.id for x in self.contract_id]
    #         # if self.invoice_id.contract_ids:
    #         #     # pre_ids = [x.id for x in self.invoice_id.contract_ids]
    #         #     print(self.invoice_id.contract_ids)
    #         #     print(rec for rec in self.invoice_id.contract_ids)
    #         #     for rec in self.invoice_id.contract_ids:
    #         #         print(rec)
    #         #         contract_ids.append(rec.id)
    #         self.invoice_id.contract_ids = contract_ids
    #         self.contract_number = self.contract_id.code
    #         date_str = datetime.datetime.strptime(str(self.contract_id.date), '%Y-%m-%d').strftime('%d/%m/%y')
    #         self.contract_date_str = date_str
    #         self.contract_price_str = self.contract_id.range
    #     else:
    #         self.invoice_id.contract_ids = None
    #         self.contract_number = ''
    #         self.contract_date_str = ''
    #         self.contract_price_str = ''
    @api.onchange('contract_ids')
    def _onchange_contract_ids(self):
        if self.contract_ids:
            for rec in self:
                contract_ids = [x._origin.id for x in rec.contract_ids]
                contract_number = ''
                date_str = ''
                contract_price_str = ''
                for x in rec.contract_ids:
                    contract_number += x.reference if not contract_number else ', ' + x.reference
                    if x.date:
                        dates = datetime.datetime.strptime(str(x.date), '%Y-%m-%d').strftime('%d/%m/%y')
                        date_str += dates if not date_str else ', ' + dates
                    if x.range:
                        contract_price_str += x.range if not contract_price_str else ', ' + x.range
                rec.contract_number = contract_number
                rec.contract_date_str = date_str
                rec.contract_price_str = contract_price_str
                rec.invoice_id.contract_ids = contract_ids

        else:
            self.invoice_id.contract_ids = None
            self.contract_number = ''
            self.contract_date_str = ''
            self.contract_price_str = ''

    @api.onchange('swift_amt','swift_charge_fc')
    def onchange_swift_amt(self):
        if self.swift_amt or self.swift_charge_fc:
            swift_value = round(self.swift_amt+self.swift_charge_fc, 2)
            inv_value = round(self.invoice_amt, 2)
            if swift_value > inv_value:
                raise UserError(_('Receivable Credit Amount can not be greater than Invoice Amount.'))

        if self.invoice_id.swift_remaining_amount > 0:
            if (self.swift_amt+self.swift_charge_fc) > self.invoice_id.swift_remaining_amount:
                raise UserError(_('Receivable Credit Amount can not be greater than Remaining Invoice Amount.'))

    @api.onchange('head_id')
    def onchange_head_id(self):
        if self.swift_message_id:
            self.onchange_swift_message_id()
            self.head_id.fc_currency_id = self.swift_message_id.currency_id.id
            self.head_id.remiter_bank_name = self.swift_message_id.remiter_bank_name
            self.head_id.remiter_bank_address = self.swift_message_id.remiter_bank_address
            self.od_sight_rate = self.swift_message_id.od_sight_rate
            qty = ''
            h_q = 0
            d_q = 0
            for l in self.invoice_id.invoice_line_ids:
                if l.quantity_type == '0':
                    h_q += l.quantity
                else:
                    d_q += l.quantity
            if h_q:
                if not qty:
                    qty += 'ITES/' + str(h_q) + ' hrs'
                else:
                    qty += ', ITES/' + str(h_q) + ' hrs'
            if d_q:
                if not qty:
                    qty += 'ITES/' + str(d_q) + ' Developers' if d_q > 1 else 'ITES/' + str(d_q) + ' Developer'
                else:
                    qty += ', ITES/' + str(d_q) + ' Developers'  if d_q > 1 else ', ITES/' + str(d_q) + ' Developer'

            self.invoice_qty_str = qty
            self.invoice_id.cash_incentive_id = self.head_id.id
            self.head_id.on_change_invoice_line_ids()

    @api.onchange('invoice_id')
    def onchange_invoice_id(self):
        if self.invoice_id:
            # if self.invoice_id.is_done_inv_amount:
            #     raise UserError(_('No Invoice amount left to Encashment.'))
            # if not self.contract_number or not self.contract_date_str:
            l_contract_number = ''
            l_contract_date_str = ''
            l_contract_price_str = ''

            # set contract to invoice when onchange invoice id------------
            if self.contract_id:
                if not self.invoice_id.contract_ids:
                    contract_ids = [x.id for x in self.contract_id]
                    self.invoice_id.contract_ids = contract_ids
                else:
                    self.contract_id = self.invoice_id.contract_ids[0]
            else:
                if self.invoice_id.contract_ids:
                    self.contract_id = self.invoice_id.contract_ids[0]

            for l in self.invoice_id.contract_ids:
                date_str = datetime.datetime.strptime(str(l.date), '%Y-%m-%d').strftime('%d/%m/%y')
                l_contract_date_str += str(date_str) if not l_contract_date_str else ', ' + str(date_str)
                l_contract_price_str += str(l.range) if not l_contract_price_str else ', ' + str(l.range)
                l_contract_number += str(l.reference) if not l_contract_number else ', ' + str(l.reference)

            self.contract_number = l_contract_number
            self.contract_date_str = l_contract_date_str
            self.contract_price_str = l_contract_price_str

            # inv_amount = sum(self.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
            inv_amount = self.invoice_id.invoice_total_fc
            self.invoice_date = self.invoice_id.date
            self.invoice_amt = inv_amount

            #self.swift_charge_rate = self.invoice_id.usd_rate
            self.swift_charge_fc = 0

            if self.invoice_id.swift_remaining_amount > 0:
                self.swift_amt = self.invoice_id.swift_remaining_amount
            else:
                self.swift_amt = inv_amount
            self.currency_id = self.invoice_id.foreign_currency_type.id
            qty = ''
            h_q = 0
            d_q = 0
            for l in self.invoice_id.invoice_line_ids:
                if l.quantity_type == '0':
                    h_q += l.quantity
                else:
                    d_q += l.quantity
            if h_q:
                if not qty:
                    qty += 'ITES/' + str(h_q) + ' hrs'
                else:
                    qty += ', ITES/' + str(h_q) + ' hrs'
            if d_q:
                if not qty:
                    qty += 'ITES/' + str(d_q) + ' Developers' if d_q > 1 else 'ITES/' + str(d_q) + ' Developer'
                else:
                    qty += ', ITES/' + str(d_q) + ' Developers' if d_q > 1 else ', ITES/' + str(d_q) + ' Developer'
            self.invoice_qty_str = qty
            # '\n' +  'Total Amount: ' ' <b>' + str(math.floor(self.amount_total_tmp)) + '0</b>' '\n' + 'Delivery Date: ' ' <b>' + str(self.date_order) + '</b>' + '\n' +  'Supplier Name: ' ' <b>' + str(self.partner_id.name) + '</b>'

            self.invoice_id.swift_id = self.swift_message_id.id
        # comment-for-upgrade
        # else:
        #     self.prc_letter_description = ''


class CashIncentivePrc(models.Model):
    _name = "cash.incentive.prc"
    _description = "Cash Incentive PRC"

    head_id = fields.Many2one('cash.incentive.head', required=True, ondelete='cascade')


class SwiftMessageDownloadLine(models.TransientModel):
    _name = "incentive.swift.message.download.line"

    head_id = fields.Many2one('cash.incentive.head', ondelete='cascade')
    swift_id = fields.Many2one('swift.message', string='SWIFT Message', ondelete='cascade')
    swift_file = fields.Binary(string='SWIFT File (File- Received from Bank)', attachment=True, related='swift_id.swift_file')
    swift_file_name = fields.Char( related='swift_id.swift_file_name')
    ict_file = fields.Binary(string='ICT Form (Signed ICT Form, Received after Encashment)', attachment=True, related='swift_id.ict_file')
    rate_file = fields.Binary(string='Rate Sheet', attachment=True, related='swift_id.rate_file')
    ict_file_name = fields.Char(related='swift_id.ict_file_name')
    rate_file_name = fields.Char(related='swift_id.rate_file_name')
    inv_names = fields.Char(string="Invoices", related='swift_id.inv_names')
    date = fields.Date(string='SWIFT Date', related='swift_id.date')

# class IrAttachmentInherit(models.Model):
#     _inherit = 'ir.attachment'
#
#     @api.model
#     def _file_read(self, fname, bin_size=False):
#         import os
#         import logging
#         _logger = logging.getLogger(__name__)
#
#         print('fname...',fname)
#         full_path = self._full_path(fname)
#         r = ''
#         try:
#             if bin_size:
#                 r = self.human_size(os.path.getsize(full_path))
#             else:
#                 with open(full_path,'rb') as fd:
#                     r = base64.b64encode(fd.read())
#         except (IOError, OSError):
#             _logger.info("_read_file reading %s", full_path, exc_info=True)
#         return r
#
#     def human_size(sz):
#         """
#         Return the size in a human readable format
#         """
#         if not sz:
#             return False
#         units = ('bytes', 'Kb', 'Mb', 'Gb', 'Tb')
#         if isinstance(sz, str):
#             sz = len(sz)
#         s, i = float(sz), 0
#         while s >= 1024 and i < len(units) - 1:
#             s /= 1024
#             i += 1
#         return "%0.2f %s" % (s, units[i])

