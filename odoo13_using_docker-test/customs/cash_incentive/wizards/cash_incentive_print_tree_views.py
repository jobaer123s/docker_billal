from odoo import models, fields, _, api
from odoo.exceptions import AccessError, UserError, ValidationError
import datetime
from num2words import num2words
import math
import decimal


class CashIncentiveWizard(models.TransientModel):

    _name = 'cash.incentive.print.wizards'

    type = fields.Selection([
        ('01', 'PDF'),
        ('02', 'Word')
    ], string='Type', copy=False, default='01')
    report_name = fields.Selection([
        ('01', 'PRC'),
        ('04', 'Forwarding Letter for BASIS'),
        ('02', 'Form KA'),
        ('03', 'Form KHA'),
        ('05', 'Forwarding Letter for Bank'),
        ('06', 'Form Ga for Bank'),
        ('07', 'Form Gha for Bank'),
        ('08', 'Print All Form'),
    ], string='Report Name', copy=False, default='01')

    def action_print_pdf(self):
        active_ids = self.env.context.get('active_ids')
        if self.type == '02':
            from docx import Document
            from htmldocx import HtmlToDocx
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.text import WD_BREAK
            from docx.shared import Pt
            from docx.enum.style import WD_STYLE_TYPE
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.enum.table import WD_ALIGN_VERTICAL

            from docx.shared import Inches
            file_name = ''
            document = Document()
            from docx.shared import Cm
            section = document.sections[0]
            section.left_margin = Cm(1.5)  # 1.5 cm
            section.right_margin = Cm(1.5)  # 1.5 cm
            mystyle = 0
            page_sl = 1

            file_names = ''
            # for data in active_ids:
            #     file_obj = self.env['cash.incentive.head'].browse(data)
            #     ref_name_val = file_obj.name
            #     split_values = ref_name_val.split("/")
            #     desired_value = split_values[-1]
            #     file_names += desired_value if not file_names else ', ' + desired_value

            for data in active_ids:
                file_obj = self.env['cash.incentive.head'].browse(data)
                ref_name_val = file_obj.name
                split_values = ref_name_val.split("/")
                desired_value = split_values[-1]

                country = ''
                if file_obj.partner_id.country_id:
                    country = file_obj.partner_id.country_id.name

                currency = ''
                if file_obj.fc_currency_id:
                    currency = file_obj.fc_currency_id.name
                inv_number = ''
                inv_date = ''
                inv_amount = 0
                swift_amt = 0
                swift_id = None
                inv_ids = []
                for rec in file_obj.invoice_line_ids:
                    swift_amt += rec.swift_amt
                    invoice_date = ''
                    if rec.invoice_date:
                        invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
                    if rec.invoice_id not in inv_ids:
                        inv_ids.append(rec.invoice_id)
                        inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
                        inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
                        inv_amount += rec.invoice_amt
                    if not swift_id:
                        swift_id = rec.swift_message_id
                client_name = ''
                if file_obj.partner_id:
                    client_name = '' if not file_obj.swift_customer_name else file_obj.swift_customer_name

                customer_address = file_obj.customer_address

                remiter_address = ''
                if swift_id:
                    if swift_id.remiter_bank_name:
                        remiter_address += swift_id.remiter_bank_name
                    if swift_id.remiter_bank_address:
                        remiter_address += swift_id.remiter_bank_address if not remiter_address else ', ' + swift_id.remiter_bank_address

                if self.report_name == '01':
                    self_date = ''
                    if file_obj.date:
                        self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')

                    document.add_paragraph('\n\nRef: ' + file_obj.name + '\t\t\t\t\t\t\t\t\t Format-A' + '\nDate: '+ self_date + '\n')
                    # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

                    # Define a custom style
                    custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style.font.bold = True
                    custom_style.font.size = Pt(14)
                    custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    mystyle += 1
                    custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style2.font.size = Pt(11)
                    custom_style2.font.name = 'Arial Narrow'
                    custom_style2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    mystyle += 1
                    custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style3.font.bold = True
                    custom_style3.font.size = Pt(10)
                    custom_style3.font.name = 'Arial Narrow'

                    mystyle += 1
                    custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style4.font.size = Pt(11)
                    custom_style4.font.name = 'Arial Narrow'
                    mystyle += 1

                    # Add a paragraph and apply the custom style
                    paragraph1 = document.add_paragraph('CERTIFICATE OF AUTHORIZED DEALER ')
                    paragraph1.style = custom_style

                    contract_number = file_obj.contract_number
                    if not file_obj.contract_number:
                        contract_number = ''
                    contract_date_str = file_obj.contract_date_str
                    if not file_obj.contract_date_str:
                        contract_date_str = ''
                    paragraph2 = document.add_paragraph(
                        "This is to certify that we have received following  inward  remittance  in "
                        " the  name  of  Brain Station  23 Limited,  Plot  02  (8  Floor),  Bir  Uttam  A.  K."
                        "  Khandakar Road, Mohakhali C/A, Dhaka- 1212 against "
                        + contract_number + ' & date: ' + str(contract_date_str) + " invoice no. "
                        + inv_number + " & date:  " + inv_date + " amount in " + currency + ' ' + str(
                            "{:,}".format(round(inv_amount, 2)))
                        + " for rendering of software development service. Summary of the transaction is as follows: ")
                    paragraph2.style = custom_style2
                    # Set the paragraph alignment to center

                    table = document.add_table(rows=1, cols=9)
                    table.style = 'TableGrid'

                    cell = table.cell(0, 0)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('(a) Remitter')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3
                    # cell.merge(table.cell(0, 1))

                    cell = table.cell(0, 1)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('(b) Address')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    cell = table.cell(0, 2)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('Invoice No.')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    cell = table.cell(0, 3)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    cur_name = '(c) Amount in FC (' + file_obj.fc_currency_id.name + ')'
                    p3.add_run(cur_name)
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    cell = table.cell(0, 4)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('(d) Date of Credit in banks nostro account')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    cell = table.cell(0, 5)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('(e) Equivalent Taka')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    cell = table.cell(0, 6)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('(f) Credited to beneficiary a/c')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    cell = table.cell(0, 7)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('(g) Reporting statement/schedule to BB with Month')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    cell = table.cell(0, 8)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('(h) Reference of Online reporting to BB')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    # add a dynamic row to the table
                    cus_name = ''
                    if file_obj.swift_customer_name:
                        cus_name = file_obj.swift_customer_name

                    # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
                    row_len = len(file_obj.invoice_line_ids)
                    inv_amount = 0
                    encashment_amt_bdt = 0
                    equivalent_taka_total = 0

                    from itertools import groupby

                    # group the invoice_line_ids by invoice_id.ref
                    groups = groupby(sorted(file_obj.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                                     key=lambda x: x.invoice_id.ref)
                    # iterate over each group and add a new row for each unique invoice_id.ref
                    row = 1
                    swift_amt_total = 0
                    for invoice_ref, group in groups:
                        gr = 0
                        for rec in group:
                            gr += 1
                            new_row = table.add_row()
                            inv_amount += rec.invoice_amt
                            swift_amt_total += rec.swift_amt
                            encashment_amt_bdt += rec.encashment_amt_bdt

                            swift_date = ''
                            if rec.swift_message_id.date:
                                swift_date = datetime.datetime.strptime(str(rec.swift_message_id.date),
                                                                        '%Y-%m-%d').strftime(
                                    '%d-%b-%y')

                            date_credited_beneficiaries = ''
                            if rec.date_credited_beneficiaries:
                                date_credited_beneficiaries = datetime.datetime.strptime(
                                    str(rec.date_credited_beneficiaries),
                                    '%Y-%m-%d').strftime('%d-%b-%y')
                            reporting_st_to_bb = rec.reporting_st_to_bb if rec.reporting_st_to_bb else ''
                            ref_online_to_bb = rec.ref_online_to_bb if rec.ref_online_to_bb else ''
                            equivalent_taka = round(rec.swift_amt * rec.encashment_rate_bdt, 2)
                            equivalent_taka_total += equivalent_taka

                            # add data to the cells in the new row
                            # new_row.cells[2].text = rec.invoice_id.ref
                            new_row.cells[3].text = str("{:,}".format(round(rec.swift_amt, 2)))
                            new_row.cells[4].text = str(swift_date)
                            new_row.cells[5].text = str("{:,}".format(round(equivalent_taka, 2)))
                            new_row.cells[6].text = str(date_credited_beneficiaries)
                            new_row.cells[7].text = str(reporting_st_to_bb)
                            new_row.cells[8].text = str(ref_online_to_bb)

                            first_cell = new_row.cells[3]
                            first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph21 = first_cell.paragraphs[0]
                            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph21.style = custom_style4

                            first_cell1 = new_row.cells[5]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph22.style = custom_style4

                            first_cell1 = new_row.cells[2]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # paragraph22.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            paragraph22.style = custom_style4

                            first_cell1 = new_row.cells[0]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.style = custom_style4
                            first_cell1 = new_row.cells[1]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.style = custom_style4
                            first_cell1 = new_row.cells[4]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style4
                            first_cell1 = new_row.cells[6]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style4
                            first_cell1 = new_row.cells[7]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style4
                            first_cell1 = new_row.cells[8]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style4

                        cell_1_0 = table.cell(row, 2)
                        p3 = cell_1_0.paragraphs[0]
                        p3.add_run(invoice_ref)
                        cell_2_0 = table.cell((row + gr) - 1, 2)
                        cell_1_0.merge(cell_2_0)
                        row = gr + row
                    # cus name merge
                    cell_1_0 = table.cell(1, 0)
                    p3 = cell_1_0.paragraphs[0]
                    p3.add_run(cus_name)
                    cell_2_0 = table.cell(row_len, 0)
                    cell_1_0.merge(cell_2_0)
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    # address merge
                    cell_3_0 = table.cell(1, 1)
                    p3 = cell_3_0.paragraphs[0]
                    p3.add_run(file_obj.customer_address)
                    cell_4_0 = table.cell(row_len, 1)
                    cell_3_0.merge(cell_4_0)
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style3

                    # total
                    total_row = table.add_row()
                    total_row.cells[2].text = 'Total'
                    total_row.cells[3].text = str("{:,}".format(round(swift_amt_total, 2)))
                    total_row.cells[5].text = str("{:,}".format(round(equivalent_taka_total, 2)))

                    first_cell = total_row.cells[2]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style3

                    first_cell = total_row.cells[3]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style3

                    first_cell1 = total_row.cells[5]
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = custom_style3

                    paragraph = document.add_paragraph()
                    paragraph.add_run("\n\n\n\n\n Signature of Head of the branch.")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("Signature of the Issuing Officer.")
                    # add an automatic page break
                    
                    if page_sl != len(active_ids):
                        document.add_page_break()
                        page_sl += 1

                    # file_name = '%s_prc' % (datetime.datetime.now())
                    file_name = 'PRC_Multiple_Files_%s' % (datetime.datetime.now())
                
                elif self.report_name == '04':

                    section.page_width = Cm(22)  # set the page width to 21 centimeters
                    section.page_height = Cm(29.9)  # set the page width to 21 centimeters
                    section.left_margin = Cm(3)  # 1.5 cm
                    section.right_margin = Cm(2)

                    # Define a custom style
                    custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style.font.bold = True
                    custom_style.font.size = Pt(14)
                    custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    custom_style.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style2.font.size = Pt(11)
                    custom_style2.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style3.font.size = Pt(11)
                    custom_style3.font.bold = True
                    custom_style3.font.name = 'Arial Narrow'
                    mystyle += 1

                    # Add a paragraph and apply the custom style
                    # paragraph1 = document.add_paragraph('CERTIFICATE OF AUTHORIZED DEALER \n ')
                    # paragraph1.style = custom_style

                    # paragraph2 = document.add_paragraph("Ref No.: " + self.name + " \nDate: "+ self_date + '\n\n")
                    current_date = datetime.datetime.now().date()
                    self_date = ''
                    if current_date:
                        self_date = datetime.datetime.strptime(str(current_date), '%Y-%m-%d').strftime('%d/%m/%y')
                    paragraph2 = document.add_paragraph('\n\n\n\nRef: ' + file_obj.name + '\nDate: ' + str(self_date) + '\n')
                    paragraph2.style = custom_style2

                    paragraph4 = document.add_paragraph(
                        "The President \nBASIS \nBDBL Bhabon (5th Floor-West), \n12 Kawran Bazar, Dhaka-1215 \n")
                    paragraph4.style = custom_style2

                    paragraph5 = document.add_paragraph(
                        "Subject: Request to Issue BASIS Certificate for Cash Assistance.\n")
                    paragraph5.style = custom_style3

                    paragraph6 = document.add_paragraph("Dear Sir,")
                    paragraph6.style = custom_style2
                    paragraph6.paragraph_format.space_after = Pt(0)

                    contract_number = ''
                    if file_obj.contract_number:
                        contract_number = " and " + file_obj.contract_number

                    paragraph6 = document.add_paragraph(
                        "With reference to the above-mentioned subject, we would like to draw your kind attention to the fact that we are going to draw cash subsidy against our following Invoice No: " + inv_number + contract_number + ".\n")
                    paragraph6.style = custom_style2

                    paragraph6 = document.add_paragraph("There is a checklist as following:  ")
                    paragraph6.style = custom_style2

                    # Create a list of items
                    items = ["Forwarding Letter Addressing BASIS president", "Export Agreement",
                             "Commercial Invoice Related to Export",
                             "Format -A Certificate of Authorized Dealer Issued by Bank",
                             "ICT form-c", "Company's Undertaking", "Form-Ka (Bangladesh Bank)",
                             "Form-Kha (Bangladesh Bank)", "Pay order"]

                    # Create a paragraph object for each list item and set the paragraph style to "List Bullet"
                    left_indent = Inches(.8)  # Adjust the left indentation as needed

                    for item in items:
                        paragraph = document.add_paragraph(style="List Bullet")
                        paragraph.paragraph_format.left_indent = left_indent

                        run = paragraph.add_run(item)
                        run.font.name = 'Arial Narrow'
                        paragraph.paragraph_format.line_spacing = Pt(12)

                    paragraph8 = document.add_paragraph(
                        "\nYour kind co-operation in this regard will be highly appreciated.\n\nThanking You.\n\nYours Truly\n\n\n\n\nMasudur Rahman \nSenior Manager, Finance, Accounts & Admin \nBrain Station 23 Ltd. \nPlot-2, (8th Floor), Bir Uttam Ak Khandaker Road \nMohakhali C/A, Dhaka-1212")
                    paragraph8.style = custom_style2
                    paragraph8.paragraph_format.space_before = Pt(0)
                    paragraph8.paragraph_format.space_after = Pt(0)

                    if page_sl != len(active_ids):
                        document.add_page_break()
                        page_sl += 1
                    # file_name = '%s_flbs' % (datetime.datetime.now())
                    file_name = 'Forwarding_Letter_BASIS_Multiple_Files_%s' % (datetime.datetime.now())
                
                elif self.report_name == '05':
                    section.page_width = Cm(22)  # set the page width to 21 centimeters
                    section.page_height = Cm(29.9)  # set the page width to 21 centimeters
                    section.left_margin = Cm(2.5)  # 1.5 cm
                    section.right_margin = Cm(2)
                    self_date = ''
                    current_date = datetime.datetime.now().date()

                    if file_obj.date:
                        self_date = datetime.datetime.strptime(str(current_date), '%Y-%m-%d').strftime('%d/%m/%y')

                    # Define a custom style
                    custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style.font.bold = True
                    custom_style.font.size = Pt(14)
                    custom_style.font.name = 'Arial Narrow'
                    custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    mystyle += 1

                    custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style2.font.size = Pt(10)
                    custom_style2.font.name = 'Calibri'
                    mystyle += 1

                    custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style3.font.size = Pt(10)
                    custom_style3.font.bold = True
                    custom_style3.font.name = 'Calibri'
                    mystyle += 1

                    custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style4.font.bold = True
                    custom_style4.font.size = Pt(9)
                    custom_style4.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_style5 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style5.font.size = Pt(9)
                    custom_style5.font.name = 'Arial Narrow'
                    mystyle += 1

                    table_total = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    table_total.font.bold = True
                    table_total.font.size = Pt(10)
                    table_total.font.name = 'Arial Narrow'
                    mystyle += 1

                    paragraph2 = document.add_paragraph('\n\n\n\n\nRef: ' + file_obj.name + '\nDate: ' + self_date)
                    paragraph2.style = custom_style3
                    paragraph2.paragraph_format.left_indent = Pt(-6)

                    # paragraph2 = document.add_paragraph("The Branch Manager  \nGulshan Branch \n" + file_obj.bank_id.name + "\nHolding No. 75, Gulshan Avenue \nGulshan, Dhaka \n\n\nDear Sir\n")
                    paragraph2 = document.add_paragraph("The Manager \n" + file_obj.bank_id.name + "\nIslamic Banking Window. \n4 Bir Uttam AK Khandakar Rd, Mohakhali C/A \nDhaka 1212 \n\n\nDear Sir\n")

                    paragraph2.style = custom_style2
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(1)
                    paragraph2.paragraph_format.left_indent = Pt(-6)

                    paragraph2 = document.add_paragraph("For cash incentive claim.\n")
                    paragraph2.style = custom_style3
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)
                    paragraph2.paragraph_format.left_indent = Pt(-6)

                    paragraph2 = document.add_paragraph(
                        "We are submitting herewith necessary documents against following Invoices:")
                    paragraph2.style = custom_style2
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.left_indent = Pt(-6)

                    # table -----------------------
                    # left_indent = Inches(3)
                    table = document.add_table(rows=1, cols=8)
                    table.style = 'TableGrid'
                    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # table.autofit = False  # Disable auto-fit behavior
                    # table.left_indent = Inches(6)

                    cell = table.cell(0, 0)
                    # table.columns[0].width = Inches(.2)
                    # table.columns[0].left_indent = Pt(100)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('SL#')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style4

                    cell = table.cell(0, 1)
                    table.columns[1].width = Inches(1.8)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('Client Name')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style4

                    cell = table.cell(0, 2)
                    table.columns[2].width = Inches(.5)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('Invoice No')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style4

                    cell = table.cell(0, 3)
                    table.columns[3].width = Inches(.7)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    cur_name = 'Invoice amount (' + file_obj.fc_currency_id.name + ')'
                    p3.add_run(cur_name)
                    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p3.style = custom_style4

                    cell = table.cell(0, 4)
                    table.columns[4].width = Inches(.5)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('Invoice Date')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style4

                    cell = table.cell(0, 5)
                    table.columns[5].width = Inches(.6)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('Swift/nostro date')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style4

                    cell = table.cell(0, 6)
                    table.columns[6].width = Inches(.6)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    cur_name_r = 'Realize Amount (' + file_obj.fc_currency_id.name + ')'
                    p3.add_run(cur_name_r)
                    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p3.style = custom_style4

                    cell = table.cell(0, 7)
                    table.columns[7].width = Inches(.6)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    cur_name_i = 'Incentive claim (' + file_obj.fc_currency_id.name + ')'
                    p3.add_run(cur_name_i)
                    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    p3.style = custom_style4

                    # add a dynamic row to the table
                    cus_name = ''
                    if file_obj.swift_customer_name:
                        cus_name = file_obj.swift_customer_name

                    # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
                    row_len = len(file_obj.invoice_line_ids)
                    inv_amount = 0
                    encashment_amt_bdt = 0
                    sl = 1
                    total_incentive_rate_fc = 0

                    from itertools import groupby

                    # group the invoice_line_ids by invoice_id.ref
                    groups = groupby(sorted(file_obj.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                                     key=lambda x: x.invoice_id.ref)
                    # iterate over each group and add a new row for each unique invoice_id.ref
                    row = 1
                    for invoice_ref, group in groups:
                        inv_obj = file_obj.env['cash.incentive.invoice'].search(
                            [('invoice_id.ref', '=', invoice_ref), ('head_id', '=', file_obj.id)], limit=1)
                        # usd_price = sum(inv_obj.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
                        usd_price = inv_obj.invoice_id.invoice_total_fc
                        # incentive_rate_fc = usd_price / file_obj.incentive_rate_fc if file_obj.incentive_rate_fc else 0
                        # total_incentive_rate_fc += incentive_rate_fc
                        inv_amount += inv_obj.invoice_amt
                        invoice_date = ''
                        if inv_obj.invoice_date:
                            invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime(
                                '%d-%b-%y')
                        gr = 0
                        for rec in group:
                            gr += 1

                            incentive_rate_fc = (rec.swift_amt * inv_obj.incentive_rate_fc) / 100 if inv_obj.incentive_rate_fc else 0
                            total_incentive_rate_fc += incentive_rate_fc

                            encashment_amt_bdt += rec.encashment_amt_bdt

                            swift_message_date = ''
                            if rec.swift_message_id.date:
                                swift_message_date = datetime.datetime.strptime(str(rec.swift_message_id.date),
                                                                                '%Y-%m-%d').strftime('%d-%b-%y')

                            new_row = table.add_row()
                            # add data to the cells in the new row
                            new_row.cells[0].text = str(sl)
                            new_row.cells[5].text = swift_message_date
                            new_row.cells[6].text = str("{:,}".format(round(rec.swift_amt, 2)))
                            new_row.cells[7].text = str("{:,}".format(float(format(incentive_rate_fc, '.2f'))))

                            first_cell = new_row.cells[0]
                            first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph24 = first_cell.paragraphs[0]
                            paragraph24.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph24.style = custom_style5

                            first_cell = new_row.cells[2]
                            first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph21 = first_cell.paragraphs[0]
                            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph21.style = custom_style5

                            first_cell = new_row.cells[3]
                            first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph21 = first_cell.paragraphs[0]
                            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph21.style = custom_style5

                            first_cell1 = new_row.cells[6]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph22.style = custom_style5

                            first_cell1 = new_row.cells[7]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph22.style = custom_style5

                            # first_cell1 = new_row.cells[2]
                            # first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            # paragraph22 = first_cell1.paragraphs[0]
                            # paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # paragraph22.style = custom_style5

                            first_cell1 = new_row.cells[1]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.style = custom_style5

                            first_cell1 = new_row.cells[4]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style5

                            first_cell1 = new_row.cells[5]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style5

                            sl += 1

                        cell_1_0 = table.cell(row, 2)
                        p3 = cell_1_0.paragraphs[0]
                        p3.add_run(invoice_ref)
                        cell_2_0 = table.cell((row + gr) - 1, 2)
                        cell_1_0.merge(cell_2_0)

                        cell_4_0 = table.cell(row, 3)
                        p3 = cell_4_0.paragraphs[0]
                        p3.add_run(str("{:,}".format(round(usd_price, 2))))
                        cell_3_0 = table.cell((row + gr) - 1, 3)
                        cell_4_0.merge(cell_3_0)

                        cell_5_0 = table.cell(row, 4)
                        p3 = cell_5_0.paragraphs[0]
                        p3.add_run(invoice_date)
                        cell_6_0 = table.cell((row + gr) - 1, 4)
                        cell_5_0.merge(cell_6_0)
                        row = gr + row

                    # cus name merge
                    cell_1_0 = table.cell(1, 1)
                    p3 = cell_1_0.paragraphs[0]
                    p3.add_run(cus_name)
                    cell_2_0 = table.cell(row_len, 1)
                    cell_1_0.merge(cell_2_0)
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_style4

                    # total
                    total_row = table.add_row()
                    total_row.cells[2].text = 'Total'
                    total_row.cells[3].text = str("{:,}".format(round(inv_amount, 2)))
                    total_row.cells[6].text = str("{:,}".format(round(swift_amt, 2)))
                    total_row.cells[7].text = str("{:,}".format(float(format(total_incentive_rate_fc, '.2f'))))

                    first_cell = total_row.cells[3]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = table_total

                    first_cell1 = total_row.cells[6]
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = table_total

                    first_cell2 = total_row.cells[7]
                    paragraph23 = first_cell2.paragraphs[0]
                    paragraph23.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph23.style = table_total

                    first_cell2 = total_row.cells[2]
                    paragraph23 = first_cell2.paragraphs[0]
                    paragraph23.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph23.style = table_total

                    # after table
                    paragraph2 = document.add_paragraph("\nThose documents are as follows:")
                    paragraph2.style = custom_style2
                    paragraph2.paragraph_format.left_indent = Pt(-6)

                    # paragraph2 = document.add_paragraph(
                    #     "\t 1. Form Ka\n\t 2. Commercial Invoice\n\t 3. Form-C (ICT)\n\t 4. Agreement\n\t 5. Certificate of Authorized Dealer\n\t 6. Copy of Swift Message\n\t 7. নগদ ভর্তুকি প্রত্যয়নপত্র \n\t 8. Company's undertaking")
                    # paragraph2.style = custom_style2

                    table = document.add_table(rows=1, cols=2)
                    table.border = None
                    table.columns[0].width = Inches(3)
                    table.columns[1].width = Inches(3)

                    row = table.rows[0]
                    for cell in row.cells:
                        cell.border = None

                    left_column = table.cell(0, 0)
                    left_column = left_column.paragraphs[0]
                    left_column.add_run("1. Form Ka\n")
                    left_column.add_run("2. Commercial Invoice\n")
                    left_column.add_run("3. Form-C (ICT)\n")
                    left_column.add_run("4. Agreement\n")
                    left_column.paragraph_format.left_indent = Inches(.7)
                    left_column.style = custom_style2

                    # Add the last four items to the right column
                    right_column = table.cell(0, 1)
                    right_column = right_column.paragraphs[0]
                    right_column.add_run("5. Certificate of Authorized Dealer\n")
                    right_column.add_run("6. Copy of Swift Message\n")
                    run2 = right_column.add_run("7. নগদ ভর্তুকি প্রত্যয়নপত্র\n")
                    run2.font.size = Pt(11)
                    right_column.add_run("8. Company's undertaking\n")
                    right_column.style = custom_style2

                    paragraph2 = document.add_paragraph(
                        "Please proceed at your earliest possible time. \n\nThanking You \n\n\nMasudur Rahman \nSenior Manager, Finance, Accounts & Admin \nBrain Station 23 Limited")
                    paragraph2.style = custom_style2
                    paragraph2.paragraph_format.left_indent = Pt(-6)
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)

                    if page_sl != len(active_ids):
                        document.add_page_break()
                        page_sl += 1

                    # file_name = '%s_flfb' % (datetime.datetime.now())
                    file_name = 'Forwarding_Letter_Bank_Multiple_Files_%s' % (datetime.datetime.now())
                
                elif self.report_name == '02':
                    section.page_width = Cm(22)  # set the page width to 21 centimeters
                    section.page_height = Cm(29.9)  # set the page width to 21 centimeters
                    section.left_margin = Cm(2)  # 1.5 cm
                    section.right_margin = Cm(2)
                    self_date = ''
                    if file_obj.date:
                        self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')
                    # document.add_paragraph('Ref :' + self.name + '\nDate: '+ self_date + '\n' )
                    # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)
                    # Define a custom style
                    custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    # custom_style.font.bold = True
                    custom_style.font.size = Pt(14)
                    custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    mystyle += 1

                    custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style2.font.size = Pt(11)
                    mystyle += 1

                    custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style3.font.size = Pt(11)
                    custom_style3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    mystyle += 1

                    custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style4.font.bold = True
                    custom_style4.font.size = Pt(9)
                    custom_style4.font.name = 'Arial Narrow'
                    mystyle += 1

                    tab_total = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    tab_total.font.bold = True
                    tab_total.font.size = Pt(9)
                    tab_total.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_styleextra = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_styleextra.font.size = Pt(9)
                    custom_styleextra.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_style5 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style5.font.size = Pt(9)
                    custom_style5.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_stylebl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_stylebl.font.size = Pt(9)
                    custom_stylebl.font.name = 'SutonnyOMJ'
                    mystyle += 1

                    custom_na = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_na.font.size = Pt(9)
                    custom_na.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_stylebold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_stylebold.font.size = Pt(11)
                    custom_stylebold.font.name = 'SutonnyOMJ'
                    custom_stylebold.font.bold = True
                    mystyle += 1

                    custom_styleubl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_styleubl.font.size = Pt(11)
                    custom_styleubl.font.name = 'SutonnyOMJ'
                    custom_styleubl.font.underline = True
                    custom_styleubl.font.bold = True
                    mystyle += 1

                    document.styles['Normal'].font.size = Pt(8)
                    paragraph = document.add_paragraph()
                    paragraph.add_run("অনুচ্ছেদ ০৫(খ) এফই সার্কুলার নং-০৩/২০১৮ দ্রষ্টব্য)")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run(' ' * 6 + "ফরম-‘ক’")
                    paragraph.style = custom_stylebl

                    # Add a paragraph and apply the custom style
                    paragraph1 = document.add_paragraph(
                        'বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও')
                    paragraph1.style = custom_stylebold
                    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph1.paragraph_format.space_before = Pt(0)
                    paragraph1.paragraph_format.space_after = Pt(0)
                    paragraph1 = document.add_paragraph('হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকির জন্য আবেদনপত্র ')
                    paragraph1.style = custom_styleubl
                    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph1.paragraph_format.space_before = Pt(0)

                    paragraph2 = document.add_paragraph(
                        "(ক) আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানাঃ ব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী বাণিজ্যিক এলাকা, ঢাকা - ১২১২। রপ্তানি নিবন্ধন সনদপত্র (ইআরসি) নম্বরঃ ২৬০৩২৬২১০৬৬৬৪২০")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(10)
                    
                    contract_price_str = file_obj.contract_price_str
                    if not file_obj.contract_price_str:
                        contract_price_str = ''
                    contract_date_str = file_obj.contract_date_str
                    if not file_obj.contract_date_str:
                        contract_date_str = ''
                        
                    con_no = ''
                    if file_obj.contract_number:
                        con_no = file_obj.contract_number.replace('Contract No.', '').replace('Purchase Order No.', '')
                    paragraph2 = document.add_paragraph("(খ) রপ্তানি চুক্তিপত্র নম্বরঃ" + con_no + " \nতারিখঃ " + str(
                        contract_date_str) + "\nমূল্যঃ " + contract_price_str + "\n(পাঠ্যযোগ সত্যায়িত কপি দাখিল করতে হবে)")
                    
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(10)
                    paragraph2 = document.add_paragraph(
                        "(গ) রপ্তানিকৃত সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের স্থানীয় সংগ্রহসূত্র, পরিমাণ ও মূল্যঃ")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(10)
                    # Set the paragraph alignment to center

                    # table1 ------------------------------
                    table5 = document.add_table(rows=3, cols=3)
                    table5.alignment = 2
                    table5.left_indent = Inches(100)
                    table5.style = 'TableGrid'
                    table5.autofit = False  # Disable automatic column width adjustment
                    # Set the width of the table to 6 inches (adjust the value according to your desired width)

                    cell = table5.cell(0, 0)
                    p3 = cell.paragraphs[0]
                    p3.add_run('সরবরাহকারীর নাম ও ঠিকানা')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_stylebl

                    cell = table5.cell(0, 1)
                    p1 = cell.paragraphs[0]
                    p1.add_run('পরিমাণ')
                    p1.style = custom_stylebl
                    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table5.cell(0, 2)
                    p2 = cell.paragraphs[0]
                    p2.add_run('মূল্য')
                    p2.style = custom_stylebl
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table5.cell(1, 0)
                    p4 = cell.paragraphs[0]
                    p4.add_run('১')
                    p4.style = custom_stylebl
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table5.cell(1, 1)
                    p5 = cell.paragraphs[0]
                    p5.add_run('২')
                    p5.style = custom_stylebl
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table5.cell(1, 2)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৩')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table5.cell(2, 0)
                    p7 = cell.paragraphs[0]
                    p7.add_run('N/A')
                    p7.style = custom_na
                    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table5.cell(2, 1)
                    p8 = cell.paragraphs[0]
                    p8.add_run('N/A')
                    p8.style = custom_na
                    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table5.cell(2, 2)
                    p9 = cell.paragraphs[0]
                    p9.add_run('N/A')
                    p9.style = custom_na
                    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # table 2 ----------------------------
                    paragraph2 = document.add_paragraph(
                        "(রপ্তানিকৃত সেবা/পণ্যের বর্ণনা, মূল্য ও সংগ্রহসূত্রের বিষয়ে সেবা/পণ্য সংশ্লিষ্ট এসোসিয়েশন এর প্রত্যয়নপত্র দাখিল করতে হবে) \n(ঘ) রপ্তানিকৃত সেবা/পণ্য উৎপাদনে ব্যবহৃত আমাদানিকৃত সেবা/উপকরণাদিঃ")
                    paragraph2.style = custom_stylebl
                    # paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(12)

                    # Set the paragraph alignment to center
                    table1 = document.add_table(rows=3, cols=4)
                    table1.autofit = False
                    table1.width = Inches(6)
                    table1.alignment = 2
                    table1.style = 'TableGrid'

                    cell = table1.cell(0, 0)
                    p3 = cell.paragraphs[0]
                    p3.add_run('সরবরাহকারীর নাম ও ঠিকানা')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # table1.columns[0].width = Inches(.6)
                    p3.style = custom_stylebl

                    cell = table1.cell(0, 1)
                    p3 = cell.paragraphs[0]
                    p3.add_run('সেবা/পণ্যের নাম ও পরিমাণ ')
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # table1.columns[1].width = Inches(.6)
                    p3.style = custom_stylebl

                    cell = table1.cell(0, 2)
                    p2 = cell.paragraphs[0]
                    p2.add_run('ঋণপত্র/ব্যাংক টু ব্যাংক ঋণপত্র/ডকুমেন্টরী কালেকশন/টিটি রেমিটেন্স নম্বর, তারিখ')
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # table1.columns[2].width = Inches(5.2)
                    p2.style = custom_stylebl

                    cell = table1.cell(0, 3)
                    p2 = cell.paragraphs[0]
                    p2.add_run('মূল্য')
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # table1.columns[3].width = Inches(1)
                    p2.style = custom_stylebl

                    cell = table1.cell(1, 0)
                    p4 = cell.paragraphs[0]
                    p4.add_run('১')
                    p4.style = custom_stylebl
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table1.cell(1, 1)
                    p5 = cell.paragraphs[0]
                    p5.add_run('২')
                    p5.style = custom_stylebl
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table1.cell(1, 2)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৩')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table1.cell(1, 3)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৪')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table1.cell(2, 0)
                    p7 = cell.paragraphs[0]
                    p7.add_run('N/A')
                    p7.style = custom_na
                    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table1.cell(2, 1)
                    p8 = cell.paragraphs[0]
                    p8.add_run('N/A')
                    p8.style = custom_na
                    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table1.cell(2, 2)
                    p9 = cell.paragraphs[0]
                    p9.add_run('N/A')
                    p9.style = custom_na
                    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table1.cell(2, 3)
                    p9 = cell.paragraphs[0]
                    p9.add_run('N/A')
                    p9.style = custom_na
                    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # table 3 ----------------------
                    paragraph2 = document.add_paragraph(
                        "(৩) নং কলামের ঋণপত্রে পাঠযোগ্য সত্যায়িত কপি দাখিল করতে হবে। সেবা আমাদানির ক্ষেত্রে যথাযথ পদ্ধতি অনুসরণ করা হয়েছে মর্মে অনুমোদিত ডিলার শাখাকে নিশ্চিত হতে হবে। উৎপাদন প্রক্রিয়ায় ব্যবহৃত উপকরণাদির জন্য শুল্ক বন্ড সুবিধা ভোগ করা হয়নি/ডিউটি ড্র-ব্যাংক সুবিধা গ্রহণ করা হয়নি ও ভবিষ্যতে আবেদনও করা হবে না মর্মে রপ্তানিকরাকের ঘোষণাপত্র দাখিল করতে হবে।) \n(ঙ) রপ্তানি চালানের বিবরণঃ")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(12)

                    table = document.add_table(rows=2, cols=8)
                    table.alignment = 2  # Left alignment
                    # table.left_indent = Inches(5)
                    table.style = 'TableGrid'
                    # table.autofit = False
                    # table.width = Inches(6) it is worked...............
                    # table.autofit = False
                    # table.width = Inches(4)
                    # table.alignment = 1

                    column_width = Inches(.2)
                    table.columns[0].width = column_width
                    cell = table.cell(0, 0)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p10 = cell.paragraphs[0]
                    p10.add_run('   পণ্যের বর্ণনা   ')
                    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p10.style = custom_stylebl
                    # table.columns[0].width = Inches(.3)
                    # cell.merge(table.cell(0, 1))

                    column_width = Inches(1.8)
                    table.columns[1].width = column_width
                    cell = table.cell(0, 1)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p10 = cell.paragraphs[0]
                    p10.add_run('পরিমাণ')
                    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p10.style = custom_stylebl
                    # table.columns[1].width = Inches(1.69)

                    column_width = Inches(.1)
                    table.columns[2].width = column_width
                    cell = table.cell(0, 2)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p10 = cell.paragraphs[0]
                    p10.add_run('আমদানিকারকের দেশের নাম')
                    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p10.style = custom_stylebl
                    # table.columns[2].width = Inches(.8)

                    table.columns[3].width = Inches(.6)
                    cell = table.cell(0, 3)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p10 = cell.paragraphs[0]
                    p10.add_run('ইনভয়েস মূল্য (বৈদেশিক মুদ্রায়)')
                    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p10.style = custom_stylebl

                    table.columns[4].width = Inches(.6)
                    cell = table.cell(0, 4)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p10 = cell.paragraphs[0]
                    p10.add_run('জাহাজীকরণ/রপ্তানির তারিখ')
                    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p10.style = custom_stylebl

                    column_width = Inches(.2)
                    table.columns[5].width = column_width
                    cell = table.cell(0, 5)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p10 = cell.paragraphs[0]
                    p10.add_run('   ইএক্সপি নম্বর*   ')
                    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p10.style = custom_stylebl

                    table.columns[6].width = Inches(.8)
                    table.columns[7].width = Inches(.3)
                    cell_1_0 = table.cell(0, 6)
                    cell_1_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p10 = cell_1_0.paragraphs[0]
                    p10.add_run('বৈদেশিক মুদ্রায় প্রত্যাবাসিত রপ্তানিমূল্য ও প্রত্যাবাসনের তারিখ')
                    cell_2_0 = table.cell(0, 7)
                    cell_1_0.merge(cell_2_0)
                    p10.style = custom_stylebl
                    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    left_indent = Inches(2)  # Adjust the indent as needed

                    for row in table.rows:
                        for cell in row.cells:
                            cell.left_indent = left_indent

                    # row 2 ----------------------------
                    cell = table.cell(1, 0)
                    p4 = cell.paragraphs[0]
                    p4.add_run('১')
                    p4.style = custom_stylebl
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 1)
                    p5 = cell.paragraphs[0]
                    p5.add_run('২')
                    p5.style = custom_stylebl
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 2)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৩')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 3)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৪')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 4)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৫')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 5)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৬')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell_1_1 = table.cell(1, 6)
                    p10 = cell_1_1.paragraphs[0]
                    p10.add_run('৭')
                    p6.style = custom_stylebl
                    cell_2_1 = table.cell(1, 7)
                    cell_1_1.merge(cell_2_1)
                    p10.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # cell_1_1 = table.cell(1, 6)
                    # p10 = cell_1_1.paragraphs[0]
                    # p10.add_run('৭')
                    # cell_2_1 = table.cell(1, 7)
                    # cell_1_1.merge(cell_2_1)

                    # # add a dynamic row to the table
                    # inv_date = ''
                    inv_amount = 0
                    incentive_rate_fc = 0
                    swift_amt = 0
                    incentive_amt_fc = 0
                    encashment_amt_bdt = 0
                    total_incentive_rate_fc = 0
                    sl = 0

                    inv_ids = []
                    dev_total1 = 0
                    hour_total1 = 0
                    a = []
                    c = []
                    for line in file_obj.invoice_line_ids:
                        a.append(line.od_sight_rate)
                        if line.encashment_rate_bdt:
                            c.append(line.encashment_rate_bdt)
                        if line.invoice_id.id not in inv_ids:
                            dev_total1 += sum(
                                r.quantity if r.quantity_type == '1' else 0 for r in line.invoice_id.invoice_line_ids)
                            hour_total1 += sum(
                                r.quantity if r.quantity_type == '0' else 0 for r in line.invoice_id.invoice_line_ids)
                        inv_ids.append(line.invoice_id.id)
                    if a:
                        b = min(a)
                        od_s_r = b
                    else:
                        od_s_r = 0

                    if c:
                        d = min(c)
                        encashment_rate_bdt = d
                    else:
                        encashment_rate_bdt = 0

                    # dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in
                    #                  file_obj.invoice_line_ids.invoice_id.invoice_line_ids)
                    # hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in
                    #                   file_obj.invoice_line_ids.invoice_id.invoice_line_ids)
                    hour_total = float(format(hour_total1, '.2f'))
                    dev_total = float(format(dev_total1, '.2f'))

                    if hour_total == int(hour_total):
                        hour_total_num = "{:.0f}".format(hour_total)
                    else:
                        hour_total_num = "{:.2f}".format(hour_total)

                    if dev_total == int(dev_total):
                        dev_total_num = "{:.0f}".format(dev_total)
                    else:
                        dev_total_num = "{:.2f}".format(dev_total)

                    all_invoice_qty_str = ''
                    if dev_total > 0:
                        all_invoice_qty_str += str('ITES/ ') + str(
                            dev_total_num) if not all_invoice_qty_str else ' & ' + str(
                            'ITES/ ') + str(
                            dev_total_num)
                    if hour_total > 0:
                        all_invoice_qty_str += str(hour_total_num) + str(
                            ' HRS') if not all_invoice_qty_str else ' & ' + str(
                            hour_total_num) + str(' HRS')

                    currency_symbol = ''
                    if file_obj.fc_currency_id.symbol:
                        currency_symbol = file_obj.fc_currency_id.symbol
                    country_name = ''
                    if file_obj.partner_id.country_id:
                        country_name = file_obj.partner_id.country_id.name
                    row_len = len(file_obj.invoice_line_ids)

                    from itertools import groupby
                    # group the invoice_line_ids by invoice_id.ref
                    groups = groupby(sorted(file_obj.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                                     key=lambda x: x.invoice_id.ref)
                    # iterate over each group and add a new row for each unique invoice_id.ref
                    row = 2

                    for invoice_ref, group in groups:
                        inv_obj = file_obj.env['cash.incentive.invoice'].search(
                            [('invoice_id.ref', '=', invoice_ref), ('head_id', '=', file_obj.id)], limit=1)
                        # usd_price = sum(inv_obj.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
                        usd_price = inv_obj.invoice_id.invoice_total_fc
                        incentive_rate_fc = usd_price / file_obj.incentive_rate_fc if file_obj.incentive_rate_fc else 0
                        total_incentive_rate_fc += incentive_rate_fc
                        inv_amount += inv_obj.invoice_amt
                        encashment_amt_bdt += inv_obj.encashment_amt_bdt
                        invoice_date = ''
                        swift_message_date = ''
                        if inv_obj.invoice_date:
                            invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime(
                                '%d-%b-%y')
                        if inv_obj.swift_message_id.date:
                            swift_message_date = datetime.datetime.strptime(str(inv_obj.swift_message_id.date),
                                                                            '%Y-%m-%d').strftime('%d-%b-%y')

                        # getting hours and developers ===========
                        l_dev_total = sum(
                            r.quantity if r.quantity_type == '1' else 0 for r in inv_obj.invoice_id.invoice_line_ids)
                        l_hour_total = sum(
                            r.quantity if r.quantity_type == '0' else 0 for r in inv_obj.invoice_id.invoice_line_ids)
                        l_all_invoice_qty_str = ''
                        if l_dev_total > 0:
                            if l_dev_total == int(l_dev_total):
                                l_dev_total = "{:.0f}".format(l_dev_total)
                            else:
                                l_dev_total = "{:.2f}".format(l_dev_total)
                            l_all_invoice_qty_str += str('ITES/ ') + str(
                                l_dev_total) if not l_all_invoice_qty_str else ' & ' + str(
                                'ITES/ ') + str(l_dev_total)
                        if l_hour_total > 0:
                            if l_hour_total == int(l_hour_total):
                                l_hour_total = "{:.0f}".format(l_hour_total)
                            else:
                                l_hour_total = "{:.2f}".format(l_hour_total)
                            l_all_invoice_qty_str += str(l_hour_total) + str(
                                ' HRS') if not l_all_invoice_qty_str else ' & ' + str(
                                l_hour_total) + str(' HRS')
                        sl += 1
                        gr = 0
                        for rec in group:
                            swift_amt += rec.swift_amt
                            incentive_amt_fc += rec.incentive_amt_fc
                            # if rec.swift_message_id.encashment_rate_bdt:
                            #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
                            gr += 1
                            new_row = table.add_row()
                            new_row.cells[6].text = str(currency_symbol) + ' ' + str("{:,.2f}".format(rec.swift_amt))
                            new_row.cells[7].text = str(swift_message_date)

                            first_cell = new_row.cells[3]
                            first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph21 = first_cell.paragraphs[0]
                            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph21.style = custom_style5

                            first_cell1 = new_row.cells[6]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph22.style = custom_style5

                            first_cell1 = new_row.cells[0]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.style = custom_style5
                            first_cell1 = new_row.cells[1]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style5
                            first_cell1 = new_row.cells[2]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style5
                            first_cell1 = new_row.cells[4]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style5
                            first_cell1 = new_row.cells[5]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style5
                            first_cell1 = new_row.cells[7]
                            first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph22 = first_cell1.paragraphs[0]
                            paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph22.style = custom_style5

                        cell1_1_0 = table.cell(row, 1)
                        p3 = cell1_1_0.paragraphs[0]
                        p3.add_run(l_all_invoice_qty_str)
                        cell1_2_0 = table.cell((row + gr) - 1, 1)
                        cell1_1_0.merge(cell1_2_0)

                        cell_4_0 = table.cell(row, 3)
                        p3 = cell_4_0.paragraphs[0]
                        p3.add_run(str(currency_symbol) + ' ' + str("{:,.2f}".format(inv_obj.invoice_amt)))
                        cell_3_0 = table.cell((row + gr) - 1, 3)
                        cell_4_0.merge(cell_3_0)

                        cell_5_0 = table.cell(row, 4)
                        p3 = cell_5_0.paragraphs[0]
                        p3.add_run(invoice_date)
                        cell_6_0 = table.cell((row + gr) - 1, 4)
                        cell_5_0.merge(cell_6_0)
                        row = gr + row

                    # #column merge -----------------------
                    cell_1_2 = table.cell(2, 0)
                    p3 = cell_1_2.paragraphs[0]
                    p3.add_run('Software Development')
                    cell_2_2 = table.cell(1 + row_len, 0)
                    cell_1_2.merge(cell_2_2)
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_styleextra

                    cell_1_2 = table.cell(2, 2)
                    p3 = cell_1_2.paragraphs[0]
                    p3.add_run(country_name)
                    cell_2_2 = table.cell(1 + row_len, 2)
                    cell_1_2.merge(cell_2_2)
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.style = custom_styleextra

                    cell_1_2 = table.cell(2, 5)
                    p3 = cell_1_2.paragraphs[0]
                    p3.add_run('N/A')
                    cell_2_2 = table.cell(1 + row_len, 5)
                    cell_1_2.merge(cell_2_2)
                    p3.style = custom_styleextra

                    # total ---------------------------
                    total_row = table.add_row()
                    total_row.cells[0].text = 'Total'
                    total_row.cells[1].text = all_invoice_qty_str

                    # inv_amount_a = "{:.2f}".format(inv_amount).rstrip('0').rstrip('.') + ('0' if inv_amount % 1 else '')
                    inv_amount_a = str("{:,.2f}".format(inv_amount))
                    total_row.cells[3].text = str(currency_symbol) + ' ' + inv_amount_a

                    # total_row.cells[3].text = str(currency_symbol)+ ' ' + str("{:,}".format(round(inv_amount, 2)))
                    # formatted_number = "{:.2f}".format(inv_amount).rstrip('0').rstrip('.') + ('0' if inv_amount % 1 else '')

                    # swift_amt_a = "{:.2f}".format(swift_amt).rstrip('0').rstrip('.') + ('0' if swift_amt % 1 else '')
                    swift_amt_a = str("{:,.2f}".format(swift_amt))
                    total_row.cells[6].text = str(currency_symbol) + ' ' + swift_amt_a

                    first_cell = total_row.cells[0]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = tab_total

                    first_cell = total_row.cells[1]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = tab_total

                    first_cell = total_row.cells[3]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = tab_total

                    first_cell1 = total_row.cells[6]
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = tab_total

                    # table 4 ----------------------------------
                    paragraph2 = document.add_paragraph(
                        "*দৃশ্যমান আকারে পণ্য রপ্তানির ক্ষেত্রে প্রযোজ্য \n (কমার্শিয়াল ইনভয়েস, প্যাকিং লিষ্ট এবং জাহাজীকরণের প্রমাণ স্বরূপ পরিবহন কর্তৃপক্ষ ইস্যুকৃত এবং প্রত্যয়নকৃত বিল অব লোডিং/এয়ারওয়ে বিল, বিল অব এক্সপোর্ট (শুল্ক কর্তৃপক্ষ কর্তৃক ইস্যুকৃত ও পীক্ষিত এবং on-hand হওয়ার স্বপক্ষে পরিবহন কর্তৃপক্ষ প্রত্যয়নকৃত) এর পূর্ণাঙ্গ সেট ইত্যাদির সত্যায়িত পাঠযোগ্য কপি এবং রপ্তানিমূল্য প্রত্যাবাসন সনদপত্র দাখিল করতে হবে। তবে অদৃশ্যকারে সেবা রপ্তানির ক্ষেত্রে জাহাজীকরণের দলিল ও বিল অব এক্সপোর্ট  আবশ্যকতা থাকবে না।) \n(চ) ভর্তুকির আবেদনকৃত অংকঃ")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(12)

                    # Set the paragraph alignment to center
                    table = document.add_table(rows=3, cols=4)
                    table.style = 'TableGrid'
                    table.autofit = False
                    table.width = Inches(6)
                    table.alignment = 2

                    cell = table.cell(0, 0)
                    p3 = cell.paragraphs[0]
                    p3.add_run('প্রত্যাবাসিত রপ্তানিমূল্য')
                    p3.style = custom_stylebl
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[0].width = Inches(1.2)

                    cell = table.cell(0, 1)
                    p3 = cell.paragraphs[0]
                    p3.add_run('প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার পরিমাণ')
                    p3.style = custom_stylebl
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[1].width = Inches(1.7)

                    cell = table.cell(0, 2)
                    p2 = cell.paragraphs[0]
                    p2.add_run('বৈদেশিক মুদ্রায় পরিশোধ্য কমিশন ইন্সুরেন্স ইত্যাদি (যদি থাকে)')
                    p2.style = custom_stylebl
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[2].width = Inches(2.58)

                    cell = table.cell(0, 3)
                    p2 = cell.paragraphs[0]
                    p2.add_run('নীট এফওবি রপ্তানিমূল্য ১-(২+৩)')
                    p2.style = custom_stylebl
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 0)
                    p4 = cell.paragraphs[0]
                    p4.add_run('১')
                    p4.style = custom_stylebl
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 1)
                    p5 = cell.paragraphs[0]
                    p5.add_run('২')
                    p5.style = custom_stylebl
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 2)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৩')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 3)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৪')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 0)
                    p7 = cell.paragraphs[0]
                    p7.add_run(str(currency_symbol) + ' ' + swift_amt_a)
                    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p7.style = custom_style4

                    cell = table.cell(2, 1)
                    p8 = cell.paragraphs[0]
                    p8.add_run('N/A')
                    p8.style = custom_na
                    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 2)
                    p9 = cell.paragraphs[0]
                    p9.add_run('N/A')
                    p9.style = custom_na
                    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 3)
                    p9 = cell.paragraphs[0]
                    p9.add_run(str(currency_symbol) + ' ' + swift_amt_a)
                    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p9.style = custom_style4

                    # table 5 ----------------------------------
                    paragraph2 = document.add_paragraph(
                        "(প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার উল্লেখ সম্বলিত ফ্রেইট সার্টিফিকেটের সত্যায়িত কপি দাখিল করতে হবে)")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(12)

                    vortuki_swift_amt = 0
                    if file_obj.incentive_rate_fc > 0:
                        vortuki_swift_amt = swift_amt / file_obj.incentive_rate_fc

                    # Set the paragraph alignment to center
                    table = document.add_table(rows=4, cols=4)
                    table.style = 'TableGrid'
                    table.autofit = False
                    table.width = Inches(6)
                    table.alignment = 2

                    cell = table.cell(0, 0)
                    p3 = cell.paragraphs[0]
                    p3.add_run('রপ্তানি সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের মূল্য')
                    p3.style = custom_stylebl
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell2 = table.cell(0, 1)
                    cell.merge(cell2)

                    cell = table.cell(0, 2)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p2 = cell.paragraphs[0]
                    p2.add_run('স্থানীয় মূল্য সংযোজনের হার')
                    p2.style = custom_stylebl
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell2 = table.cell(1, 2)
                    cell.merge(cell2)

                    cell = table.cell(0, 3)
                    p2 = cell.paragraphs[0]
                    p2.add_run('প্রাপ্য ভর্তুকি* ৪x১০%')
                    p2.style = custom_stylebl
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 0)
                    p4 = cell.paragraphs[0]
                    p4.add_run('দেশীয় পণ্য/সেবা')
                    p4.style = custom_stylebl
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 1)
                    p5 = cell.paragraphs[0]
                    p5.add_run('আমদানিকৃত সেবা/পণ্য')
                    p5.style = custom_stylebl
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # cell = table.cell(1, 2)
                    # p6 = cell.paragraphs[0]
                    # p6.add_run('')
                    # p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 3)
                    p6 = cell.paragraphs[0]
                    p6.add_run('')
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 0)
                    p4 = cell.paragraphs[0]
                    p4.add_run('৫')
                    p4.style = custom_stylebl
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 1)
                    p5 = cell.paragraphs[0]
                    p5.add_run('৬')
                    p5.style = custom_stylebl
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 2)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৭')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 3)
                    p6 = cell.paragraphs[0]
                    p6.add_run(' ৮')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(3, 0)
                    p7 = cell.paragraphs[0]
                    p7.add_run(str(currency_symbol) + ' ' + swift_amt_a)
                    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p7.style = custom_style4

                    cell = table.cell(3, 1)
                    p8 = cell.paragraphs[0]
                    p8.add_run('N/A')
                    p8.style = custom_na
                    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(3, 2)
                    p9 = cell.paragraphs[0]
                    p9.add_run('100%')
                    p9.style = custom_na
                    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    incentive_amt_fc_a = str("{:,.2f}".format(incentive_amt_fc))
                    cell = table.cell(3, 3)
                    p9 = cell.paragraphs[0]
                    p9.add_run(str(currency_symbol) + ' ' + incentive_amt_fc_a)
                    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p9.style = custom_style4

                    # signature ------------------------------------------
                    paragraph2 = document.add_paragraph(
                        "(*৭ নম্বর কলামের হার আলোচ্য সার্কুলারের ৪ নম্বর অনুচ্ছেদের সাথে সামঞ্জস্যতার ক্ষেত্রে ভর্তুকি প্রাপ্য হবে।) \nএ মর্মে অঙ্গীকার করা হচ্ছে যে, আমাদের নিজস্ব কারখানায়/প্রতিষ্ঠানে তৈরী/উৎপাদিত সফটওয়্যার/আইটিইএস/হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকির জন্য আবেদন করা হলো। এ আবেদনপত্রে প্রদত্ত সকল তথ্য/ঘোষণা সম্পূর্ণ ও সঠিক। যদি পরবর্তীতে কোন ভুল/অসত্য তথ্য/প্রতারণা/জালিয়াতি উদঘাটিত হয় তবে গৃহীত ভর্তুকির সমুদয় অর্থ বা এর অংশবিশেষ আমার/আমাদের নিকট হইতে এবং/অথবা আমার/আমাদের ব্যাংক হিসাব থেকে আদায়/ফেরত নেয়া যাবে। \n\nতারিখঃ..................................... ")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(10)

                    paragraph2 = document.add_paragraph(
                        "........................................ \nআবেদনকারী প্রতিষ্ঠানের স্বত্তাধিকারী/\n ক্ষমতাপ্রাপ্ত  কর্মকর্তার স্বাক্ষর ও পদবী")
                    paragraph2.style = custom_stylebl
                    paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(10)

                    # table 6 ----------------------------------
                    document.add_page_break()
                    paragraph2 = document.add_paragraph("(ছ) ভর্তুকি প্রদানকারী ব্যাংক শাখা কর্তৃক পূরণীয়ঃ \t\t\t\t\t\t\t\t\t" + ' ' * 11 + "(বৈদেশিক মুদ্রায়)")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)
                    paragraph2.paragraph_format.line_spacing = Pt(10)

                    # Set the paragraph alignment to center
                    table = document.add_table(rows=3, cols=4)
                    table.style = 'TableGrid'
                    table.autofit = False
                    table.width = Inches(6)
                    table.alignment = 2

                    cell = table.cell(0, 0)
                    p3 = cell.paragraphs[0]
                    p3.add_run('প্রত্যাবাসিত রপ্তানিমূল্য')
                    p3.style = custom_stylebl
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[0].width = Inches(1.21)

                    cell = table.cell(0, 1)
                    p3 = cell.paragraphs[0]
                    p3.add_run('প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার পরিমাণ')
                    p3.style = custom_stylebl
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[1].width = Inches(1.7)

                    cell = table.cell(0, 2)
                    p2 = cell.paragraphs[0]
                    p2.add_run('বৈদেশিক মুদ্রায় পরিশোধ কমিশন, ইন্সুরেন্স ইত্যাদি (যদি থাকে)')
                    p2.style = custom_stylebl
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[2].width = Inches(2.6)

                    cell = table.cell(0, 3)
                    p2 = cell.paragraphs[0]
                    p2.add_run('নীট এফওবি রপ্তানিমূল্য ১-(২+৩)')
                    p2.style = custom_stylebl
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[3].width = Inches(1.8)

                    cell = table.cell(1, 0)
                    p4 = cell.paragraphs[0]
                    p4.add_run('১')
                    p4.style = custom_stylebl
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 1)
                    p5 = cell.paragraphs[0]
                    p5.add_run('২')
                    p5.style = custom_stylebl
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 2)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৩')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 3)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৪')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 0)
                    p4 = cell.paragraphs[0]
                    p4.add_run(str(currency_symbol) + ' ' + swift_amt_a)
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p4.style = custom_style4

                    cell = table.cell(2, 1)
                    p5 = cell.paragraphs[0]
                    p5.add_run('N/A')
                    p5.style = custom_na
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 2)
                    p6 = cell.paragraphs[0]
                    p6.add_run('N/A')
                    p6.style = custom_na
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 3)
                    p6 = cell.paragraphs[0]
                    p6.add_run(str(currency_symbol) + ' ' + swift_amt_a)
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p6.style = custom_style4

                    # table 7 ----------------------------------

                    rate_fc = swift_amt / file_obj.incentive_rate_fc
                    # encashment_final_amt = encashment_amt_bdt / file_obj.incentive_rate_fc
                    # encashment_final_amt = rate_fc / file_obj.incentive_rate_fc
                    if file_obj.od_sight_rate:
                        od_sight_rate = file_obj.od_sight_rate
                    else:
                        if od_s_r:
                            od_sight_rate = od_s_r
                        else:
                            od_sight_rate = encashment_rate_bdt
                    rate_fc_a = "{:,.2f}".format(rate_fc)
                    encashment_rate_bdt_a = "{:.2f}".format(file_obj.od_sight_rate)
                    encashment_final_amt_a = round(rate_fc, 2) * od_sight_rate
                    # encashment_final_amt = "{:,.0f}".format(round(encashment_final_amt_a))
                    encashment_final_amt = self.custom_round(encashment_final_amt_a)
                    usd_bdt_amount = currency + ' ' + str(
                        rate_fc_a) + '\n' + '@' + str("{:,}".format(round(od_sight_rate, 2))) + '=' + '\n' + 'BDT ' + str(
                        encashment_final_amt)

                    # amount in word
                    amount_in_word_str = ''
                    if encashment_final_amt_a:
                        # amount_in_word = num2words(round(encashment_final_amt_a))
                        amount_in_word = self.num2words_fun(round(encashment_final_amt_a))
                        # upper case function call
                        amount_in_word_str = file_obj.upper_case(amount_in_word)

                    paragraph2 = document.add_paragraph(
                        "(প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়া উল্লেখ সম্বলিত ফ্রেইট সার্টিফিকেটের সত্যায়িত কপি দাখিল করতে হবে)")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.line_spacing = Pt(12)

                    # Set the paragraph alignment to center
                    table = document.add_table(rows=4, cols=4)
                    table.style = 'TableGrid'
                    table.autofit = False
                    table.width = Inches(6)
                    table.alignment = 2

                    cell = table.cell(0, 0)
                    p3 = cell.paragraphs[0]
                    p3.add_run('রপ্তানি সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের মূল্য')
                    p3.style = custom_stylebl
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell2 = table.cell(0, 1)
                    cell.merge(cell2)

                    cell = table.cell(0, 2)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p2 = cell.paragraphs[0]
                    p2.add_run('স্থানীয় মূল্য সংযোজনের হার [(৪-৬)/৪]x ১০০')
                    p2.style = custom_stylebl
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell2 = table.cell(1, 2)
                    cell.merge(cell2)

                    cell = table.cell(0, 3)
                    p2 = cell.paragraphs[0]
                    p2.add_run(
                        'পরিশোধ ভর্তুকির পরিমাণ (টাকায়)*(রপ্তানিমূল্য প্রত্যাবাসনের তারিখে সংশ্লিষ্ট ৪x ১০% বৈদেশিক মুদ্রার ওডি সাইট)')
                    p2.style = custom_stylebl
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell2 = table.cell(1, 3)
                    cell.merge(cell2)

                    cell = table.cell(1, 0)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p4 = cell.paragraphs[0]
                    p4.add_run('দেশীয় পণ্য/সেবা')
                    p4.style = custom_stylebl
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(1, 1)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p5 = cell.paragraphs[0]
                    p5.add_run('আমদানিকৃত সেবা/পণ্য বৈদেশিক মুদ্রায়')
                    p5.style = custom_stylebl
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # cell = table.cell(1, 3)
                    # p6 = cell.paragraphs[0]
                    # p6.add_run('')
                    # p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 0)
                    p4 = cell.paragraphs[0]
                    p4.add_run('৫')
                    p4.style = custom_stylebl
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 1)
                    p5 = cell.paragraphs[0]
                    p5.add_run('৬')
                    p5.style = custom_stylebl
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 2)
                    p6 = cell.paragraphs[0]
                    p6.add_run('৭')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(2, 3)
                    p6 = cell.paragraphs[0]
                    p6.add_run(' ৮')
                    p6.style = custom_stylebl
                    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(3, 0)
                    p7 = cell.paragraphs[0]
                    p7.add_run(str(currency_symbol) + ' ' + swift_amt_a)
                    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p7.style = custom_style4

                    cell = table.cell(3, 1)
                    p8 = cell.paragraphs[0]
                    p8.add_run('N/A')
                    p8.style = custom_na
                    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(3, 2)
                    p9 = cell.paragraphs[0]
                    p9.add_run('100%')
                    p9.style = custom_na
                    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cell = table.cell(3, 3)
                    p9 = cell.paragraphs[0]
                    p9.add_run(usd_bdt_amount)
                    p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p9.style = custom_style4

                    # conclusion ----------------
                    paragraph2 = document.add_paragraph(
                        "(*৭ নম্বর কলামের হার আলোচ্য সার্কুলারের ৪ নম্বর অনুচ্ছেদের সাথে সামঞ্জস্যতার ক্ষেত্রে ভর্তুকি প্রাপ্য হবে) \nভর্তুকি পরিমাণ: " + amount_in_word_str + ' Only')
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.line_spacing = Pt(12)

                    paragraph2 = document.add_paragraph("\n\n\n\n\n\n\nপরিশোধের তারিখঃ --------------------------")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.line_spacing = Pt(10)

                    paragraph2 = document.add_paragraph(
                        "----------------------------------------- \nভর্তুকির অনুমোদনের ক্ষমতাপ্রাপ্ত ব্যাংক  \nকর্মকর্তার স্বাক্ষর, নাম ও পদবী")
                    paragraph2.style = custom_stylebl
                    paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph2.paragraph_format.line_spacing = Pt(10)
                    
                    if page_sl != len(active_ids):
                        document.add_page_break()
                        page_sl += 1
                        
                    # file_name = '%s_ka' % (datetime.datetime.now())
                    file_name = 'Form_KA_Multiple_Files_%s' % (datetime.datetime.now())

                elif self.report_name == '03':
                    # modify the page setup
                    section.page_width = Cm(22)  # set the page width to 21 centimeters
                    section.page_height = Cm(29.9)  # set the page width to 21 centimeters
                    section.left_margin = Cm(2.5)  # 1.5 cm
                    section.right_margin = Cm(2)  # 1.5 cm
                    self_date = ''
                    if file_obj.date:
                        self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')

                    # document.add_paragraph('কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।')
                    # en_font = my_stylew.font
                    # document.add_paragraph('some text')

                    custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    # custom_style.font.bold = True
                    custom_style.font.size = Pt(8)
                    custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    custom_style.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_style_table = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style_table.font.size = Pt(10)
                    custom_style_table.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_stylebsmallbold = document.styles.add_style(str(mystyle),
                                                                       WD_STYLE_TYPE.PARAGRAPH)
                    custom_stylebsmallbold.font.size = Pt(9)
                    custom_stylebsmallbold.font.name = 'SutonnyOMJ'
                    custom_stylebsmallbold.font.bold = True
                    mystyle += 1

                    custom_stylebsmall = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_stylebsmall.font.size = Pt(9)
                    custom_stylebsmall.font.name = 'SutonnyOMJ'
                    mystyle += 1

                    custom_stylebl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_stylebl.font.size = Pt(9)
                    custom_stylebl.font.name = 'SutonnyOMJ'
                    mystyle += 1

                    custom_sonar_bangla = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_sonar_bangla.font.size = Pt(9)
                    custom_sonar_bangla.font.name = 'Shonar Bangla'
                    mystyle += 1

                    custom_tableh = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_tableh.font.size = Pt(9)
                    custom_tableh.font.name = 'SutonnyOMJ'
                    mystyle += 1

                    custom_table = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_table.font.size = Pt(9)
                    custom_table.font.name = 'SutonnyOMJ'
                    mystyle += 1

                    custom_stylebold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_stylebold.font.size = Pt(11)
                    custom_stylebold.font.name = 'SutonnyOMJ'
                    custom_stylebold.font.bold = True
                    mystyle += 1

                    custom_styleubl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_styleubl.font.size = Pt(11)
                    custom_styleubl.font.name = 'SutonnyOMJ'
                    custom_styleubl.font.underline = True
                    custom_styleubl.font.bold = True
                    mystyle += 1

                    document.styles['Normal'].font.size = Pt(8)
                    paragraph = document.add_paragraph()
                    paragraph.add_run("\n(অনুচ্ছেদ ০৬ (ক), এফই সার্কুলার নং-০৩/২০১৮ দ্রষ্টব্য)")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.style = custom_stylebl
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.add_run("ফরম- 'খ'")
                    paragraph = document.add_paragraph()
                    paragraph.add_run("Ref: " + str(file_obj.form_kha_ref_code))
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run("\t")
                    paragraph.add_run(" Date: " + self_date)
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.style = custom_stylebl
                    # paragraph.style = custom_style

                    # Define a custom style
                    custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style2.font.size = Pt(8)
                    mystyle += 1

                    custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style3.font.size = Pt(8)
                    custom_style3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    mystyle += 1

                    custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style4.font.bold = True
                    custom_style4.font.size = Pt(10)
                    custom_style4.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_total = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_total.font.bold = True
                    custom_total.font.size = Pt(9)
                    custom_total.font.name = 'Arial Narrow'
                    mystyle += 1

                    custom_style5 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                    custom_style5.font.size = Pt(8)
                    custom_style5.font.name = 'Arial Narrow'
                    mystyle += 1

                    # Add a paragraph and apply the custom style
                    paragraph = document.add_paragraph(
                        'বেসিস প্রদেয় প্রত্যয়ন সনদপত্র \nবাংলাদেশ হতে সফটওয়্যার, আইটিইএস (Information Technology Enabled Services)')
                    paragraph.style = custom_stylebold
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph = document.add_paragraph(
                        'ও হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকি প্রাপ্তির প্রত্যয়ন সনদপত্র।')
                    paragraph.style = custom_styleubl
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)

                    customer = '' if not file_obj.swift_customer_name else str(file_obj.swift_customer_name) + ', ' + str(
                        file_obj.customer_address)
                    con_no = ''
                    if file_obj.contract_number:
                        con_no = file_obj.contract_number.replace('Contract No.', '').replace('Purchase Order No.', '')

                    # new_parser = HtmlToDocx()
                    # html_text = """
                    #             <div style='position: relative'>
                    #                 <p style='float:left'> ১। </p>
                    #                 <p  style='margin-left: 18px; float:right'> আবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road, Mohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ RA- 260326210666420 </p>
                    #             </div>
                    #
                    #             """
                    # new_parser.add_html_to_document(html_text, document)

                    # paragraph = document.add_paragraph()
                    # paragraph.add_run('Word1')
                    # paragraph.add_run(' ' * 3)
                    # paragraph.add_run('Word2')
                    contract_price_str = file_obj.contract_price_str
                    if not file_obj.contract_price_str:
                        contract_price_str = ''
                    contract_date_str = file_obj.contract_date_str
                    if not file_obj.contract_date_str:
                        contract_date_str = ''

                    import textwrap
                    width = 140

                    # paragraph2 = document.add_paragraph(
                    #     "১।\tআবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road,"
                    #     " \tMohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ 260326210666420 \n২।\tরপ্তানি ঋণপত্র/চুক্তিপত্র নম্বরঃ" + con_no + '\t তারিখঃ ' +
                    #     contract_date_str + '\t মূল্যঃ ' + contract_price_str + '\n৩।\tবিদেশি ক্রেতার নাম ও ঠিকানাঃ ' + customer +
                    #     "\n৪।\tবিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ " + remiter_address + '\n' + '৫।')
                    # paragraph2.style = custom_stylebl
                    # paragraph2.add_run("\t")
                    # paragraph2.paragraph_format.space_before = Pt(0)
                    # paragraph2.paragraph_format.space_after = Pt(1)

                    paragraph2 = document.add_paragraph(
                        "১।\tআবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road,"
                        " \tMohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ 260326210666420")
                    paragraph2.style = custom_stylebl
                    # paragraph2.add_run("\t")
                    paragraph2.paragraph_format.space_before = Pt(0)
                    paragraph2.paragraph_format.space_after = Pt(0)

                    long_text1 = "২।\t" + ' ' * 5 + "রপ্তানি ঋণপত্র/চুক্তিপত্র নম্বরঃ" + con_no + '\t তারিখঃ ' + contract_date_str + '\t মূল্যঃ ' + contract_price_str
                    wrapped_lines1 = textwrap.wrap(long_text1, width=120)
                    formatted_lines1 = [line + '\t' for line in wrapped_lines1]
                    formatted_text1 = '\n\t'.join(formatted_lines1)
                    paragraph2 = document.add_paragraph(formatted_text1)
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    long_text2 = "৩।\t" + ' ' * 5 + "বিদেশি ক্রেতার নাম ও ঠিকানাঃ " + customer
                    wrapped_lines2 = textwrap.wrap(long_text2, width=123)
                    formatted_lines2 = [line + '\n\t' for line in wrapped_lines2]
                    formatted_text2 = ''.join(formatted_lines2).rstrip('\n\t')
                    paragraph2 = document.add_paragraph(formatted_text2)
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    long_text1 = "৪।\t" + ' ' * 5 + "বিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ " + remiter_address
                    wrapped_lines1 = textwrap.wrap(long_text1, width=120)
                    formatted_lines1 = [line + '\t' for line in wrapped_lines1]
                    formatted_text1 = '\n\t'.join(formatted_lines1)
                    paragraph2 = document.add_paragraph(formatted_text1)
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    paragraph2 = document.add_paragraph('৫।')
                    paragraph2.style = custom_stylebl


                    # table = document.add_table(rows=3, cols=4)
                    #
                    # # Set the width of the first column to 1 inch
                    # column_width = Inches(1)
                    # table.columns[0].width = column_width
                    #
                    # # Set the width of the second column to 2 inches
                    # column_width = Inches(1)
                    # table.columns[1].width = column_width
                    #
                    # # Set the width of the third column to 0.5 inches
                    # column_width = Inches(4)
                    # table.columns[2].width = column_width
                    #
                    # # Set the width of the third column to 0.5 inches
                    # column_width = Inches(1)
                    # table.columns[3].width = column_width

                    table = document.add_table(rows=1, cols=4)
                    table.style = 'TableGrid'
                    table.alignment = 1
                    # Set the width of the first column to 1 inch
                    column_width = Inches(1.2)
                    table.columns[0].width = column_width
                    cell = table.cell(0, 0)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('ক) ইনভয়েস নম্বর')
                    p3.style = custom_tableh
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Set the width of the second column to 2 inches
                    column_width = Inches(1)
                    table.columns[1].width = column_width
                    cell = table.cell(0, 1)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    p3.add_run('তারিখ')
                    p3.style = custom_tableh
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Set the width of the third column to 0.5 inches
                    column_width = Inches(4)
                    table.columns[2].width = column_width
                    cell1 = table.cell(0, 2)
                    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell1.paragraphs[0]
                    p3.add_run('খ) ইনভয়েসে উল্লেখিত সেবা/পণ্যের পরিমাণ (Qty & Hours)')
                    p3.style = custom_tableh
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.columns[2].width = Inches(3.8)

                    # Set the width of the third column to 0.5 inches
                    column_width = Inches(1)
                    table.columns[3].width = column_width
                    cell = table.cell(0, 3)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p3 = cell.paragraphs[0]
                    abc = 'মূল্য (' + str(file_obj.fc_currency_id.name) + ')'
                    p3.add_run(abc)
                    p3.style = custom_tableh
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Set the paragraph alignment to center
                    # table = document.add_table(rows=1, cols=4)
                    # table.style = 'TableGrid'
                    # table.autofit = False
                    # # table = document.tables[0]
                    # # Adjust the left indentation of the table
                    # table.alignment = 1  # Center alignment
                    # # table.left_indent = Pt(144)
                    #
                    # cell = table.cell(0, 0)
                    # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # p3 = cell.paragraphs[0]
                    # p3.add_run('ক) ইনভয়েস নম্বর'...............)
                    # p3.style = custom_tableh
                    # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # table.columns[0].width = Inches(1.2)
                    #
                    # cell = table.cell(0, 1)
                    # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # p3 = cell.paragraphs[0]
                    # p3.add_run('তারিখ')
                    # p3.style = custom_tableh
                    # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # table.columns[1].width = Inches(1)
                    #
                    # cell1 = table.cell(0, 2)
                    # cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # p3 = cell1.paragraphs[0]
                    # p3.add_run('খ) ইনভয়েসে উল্লেখিত সেবা/পণ্যের পরিমাণ (Qty & Hours)')
                    # p3.style = custom_tableh
                    # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # table.columns[2].width = Inches(3.8)
                    #
                    # cell = table.cell(0, 3)
                    # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # p3 = cell.paragraphs[0]
                    # p3.add_run('মূল্য (USD)')
                    # p3.style = custom_tableh
                    # p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # table.columns[3].width = Inches(1)
                    #
                    # # add a dynamic row to the table
                    cus_name = ''
                    if file_obj.swift_customer_name:
                        cus_name = file_obj.swift_customer_name

                    # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
                    total_qty = 0
                    dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in
                                     file_obj.invoice_line_ids.invoice_id.invoice_line_ids)
                    hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in
                                      file_obj.invoice_line_ids.invoice_id.invoice_line_ids)
                    dev_total = dev_total1
                    hour_total = hour_total1

                    if hour_total == int(hour_total):
                        hour_total_num = "{:.0f}".format(hour_total)
                    else:
                        hour_total_num = "{:.2f}".format(hour_total)

                    if dev_total == int(dev_total):
                        dev_total_num = "{:.0f}".format(dev_total)
                    else:
                        dev_total_num = "{:.2f}".format(dev_total)

                    qty_str = ''
                    if dev_total > 0:
                        # qty_str += str(dev_total_num) + str(' Developers')
                        qty_str += str(dev_total_num) + (str(' Developers') if dev_total > 1 else str(' Developer'))
                    if hour_total > 0:
                        qty_str += str(hour_total_num) + str(' HRS') if not qty_str else ' & ' + str(
                            hour_total_num) + str(' HRS')
                    all_invoice_qty_str = 'Software Development /' + qty_str
                    all_invoice_qty_str_a = qty_str
                    row_len = len(file_obj.invoice_line_ids)
                    swift_amt = 0
                    inv_amount = 0

                    from itertools import groupby

                    # group the invoice_line_ids by invoice_id.ref
                    groups = groupby(sorted(file_obj.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                                     key=lambda x: x.invoice_id.ref)
                    # iterate over each group and add a new row for each unique invoice_id.ref
                    row = 1
                    for invoice_ref, group in groups:
                        inv_obj = file_obj.env['cash.incentive.invoice'].search(
                            [('invoice_id.ref', '=', invoice_ref), ('head_id', '=', file_obj.id)], limit=1)
                        inv_amount += inv_obj.invoice_amt
                        invoice_date = ''
                        if inv_obj.invoice_date:
                            invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime(
                                '%d-%b-%y')

                        qty = 'Software Development /'
                        h_q = 0
                        d_q = 0
                        for l in inv_obj.invoice_id.invoice_line_ids:
                            if l.quantity_type == '0':
                                h_q += l.quantity
                            else:
                                d_q += l.quantity
                        if d_q:
                            if d_q == int(d_q):
                                d_q_num = "{:.0f}".format(d_q)
                            else:
                                d_q_num = "{:.2f}".format(d_q)
                            qty += str(d_q_num) + (str(' Developers') if d_q > 1 else str(' Developer'))
                        if h_q:
                            if h_q == int(h_q):
                                h_q_new = "{:.0f}".format(h_q)
                            else:
                                h_q_new = "{:.2f}".format(h_q)
                            if not d_q:
                                qty += str(h_q_new) + ' HRS'
                            else:
                                qty += ' & ' + str(h_q_new) + ' HRS'
                        invoice_qty_str = qty

                        gr = 0
                        for rec in group:
                            gr += 1
                            swift_amt += rec.swift_amt
                            # add a new row to the table
                            new_row = table.add_row()
                            # add data to the cells in the new row
                            # new_row.cells[0].text = rec.invoice_id.ref
                            # new_row.cells[1].text = invoice_date
                            # new_row.cells[2].text = invoice_qty_str
                            # new_row.cells[3].text = currency + ' ' + str("{:,}".format(round(rec.invoice_amt, 2)))

                            first_cell = new_row.cells[0]
                            paragraph21 = first_cell.paragraphs[0]
                            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph21.style = custom_style5
                            first_cell = new_row.cells[1]
                            paragraph21 = first_cell.paragraphs[0]
                            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph21.style = custom_style5
                            first_cell = new_row.cells[2]
                            paragraph21 = first_cell.paragraphs[0]
                            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph21.style = custom_style5
                            first_cell = new_row.cells[3]
                            paragraph21 = first_cell.paragraphs[0]
                            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            paragraph21.style = custom_style5

                        cell_1_0 = table.cell(row, 0)
                        p3 = cell_1_0.paragraphs[0]
                        p3.add_run(invoice_ref)
                        p3.style = custom_style_table
                        cell_2_0 = table.cell((row + gr) - 1, 0)
                        cell_1_0.merge(cell_2_0)

                        cell_4_0 = table.cell(row, 1)
                        p3 = cell_4_0.paragraphs[0]
                        p3.add_run(invoice_date)
                        p3.style = custom_style_table
                        cell_3_0 = table.cell((row + gr) - 1, 1)
                        cell_4_0.merge(cell_3_0)

                        cell_5_0 = table.cell(row, 2)
                        p3 = cell_5_0.paragraphs[0]
                        p3.add_run(invoice_qty_str)
                        p3.style = custom_style_table
                        cell_6_0 = table.cell((row + gr) - 1, 2)
                        cell_5_0.merge(cell_6_0)

                        cell_5_0 = table.cell(row, 3)
                        p3 = cell_5_0.paragraphs[0]
                        p3.style = custom_style_table
                        # p3.add_run(str("{:,}".format(round(inv_obj.invoice_amt, 2))))
                        p3.add_run(str("{:,.2f}".format(inv_obj.invoice_amt)))
                        cell_6_0 = table.cell((row + gr) - 1, 3)
                        cell_5_0.merge(cell_6_0)
                        row = gr + row

                    # total
                    total_row = table.add_row()
                    total_row.cells[0].text = 'Total'
                    total_row.cells[2].text = all_invoice_qty_str
                    total_row.cells[3].text = str("{:,.2f}".format(inv_amount))

                    first_cell = total_row.cells[0]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = custom_total
                    first_cell = total_row.cells[2]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = custom_total
                    first_cell = total_row.cells[3]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_total

                    self_date = ''
                    if file_obj.date:
                        self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')

                    # footer -------------------
                    # paragraph2 = document.add_paragraph("৬।\tরপ্তানিকৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত স্থানীয় সেবা/উপকরণাদির সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): 100% in house production \tof Brain Station 23 Limited \t\t\t\t পরিমাণঃ N/A  \t\t মূল্যঃ N/A"
                    #                                     "\n৭।\tরপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত আমদানিকৃত আনুসাংগিক সেবা/উপরকরণাদি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): N/A  \t\t\tপরিমাণঃ N/A   \t\t\t\t\t  মূল্যঃ  N/A"
                    #                                     "\n৮।\tরপ্তানি সেবা/পণ্যের বিবরণঃ Software Development Services  \tপরিমাণঃ " + all_invoice_qty_str_a + " \t  মূল্যঃ " + currency+ ' ' + str("{:,.2f}".format(inv_amount)) +
                    #                                     "\n৯।\tজাহাজীকরণের তারিখঃ " + inv_date + "\tগন্তব্য বন্দরঃ " + country + "\n১০।\tইএক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানরি ক্ষেত্রে): N/A  \t\t\t  মূল্যঃ N/A   \t\t তারিখঃ N/A"
                    #                                     "\n১১।\tমোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency+ ' ' + str("{:,.2f}".format(swift_amt)) + "\t\t নীট এফওবি রপ্তানি মূল্য ( বৈদেশিক মুদ্রায়): "+ currency+ ' ' + str("{:,.2f}".format(swift_amt)) +
                    #                                     "\n১২।\tপ্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ \t\t\t\t\t\t তারিখঃ")
                    # paragraph2.style = custom_stylebl

                    # paragraph2 = document.add_paragraph(
                    #     "৬।\tরপ্তানিকৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত স্থানীয় সেবা/উপকরণাদির সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): 100% in house \tproduction of Brain Station 23 Limited \t\t\t\t পরিমাণঃ N/A  \t\t মূল্যঃ N/A"
                    #     "\n৭।\tরপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত আমদানিকৃত আনুসাংগিক সেবা/উপরকরণাদি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): N/A  \t\tপরিমাণঃ N/A   \t\t\t\t\t  মূল্যঃ  N/A"
                    #     "\n৮।\tরপ্তানি সেবা/পণ্যের বিবরণঃ Software Development Services  \tপরিমাণঃ " + all_invoice_qty_str_a + " " + "  মূল্যঃ " + currency + ' ' + str(
                    #         "{:,.2f}".format(inv_amount))
                    #     )
                    # paragraph2.style = custom_stylebl
                    # paragraph2.paragraph_format.space_after = Pt(1)
                    # if len(file_obj.invoice_line_ids) > 7:
                    #     paragraph2 = document.add_paragraph(
                    #         "৯।\tজাহাজীকরণের তারিখঃ " + inv_date + "\n\tগন্তব্য বন্দরঃ " + country)
                    #     paragraph2.style = custom_stylebl
                    # else:
                    #     paragraph2 = document.add_paragraph(
                    #         "৯।\tজাহাজীকরণের তারিখঃ " + inv_date + " \tগন্তব্য বন্দরঃ " + country)
                    #     paragraph2.style = custom_stylebl
                    # paragraph2.paragraph_format.space_before = Pt(1)
                    # paragraph2.paragraph_format.space_after = Pt(1)
                    # paragraph2 = document.add_paragraph(
                    #     "১০।\tইএক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানরি ক্ষেত্রে): N/A  \t\t\t  মূল্যঃ N/A   \t\t তারিখঃ N/A"
                    #     "\n১১।\tমোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str(
                    #         "{:,.2f}".format(swift_amt)) +
                    #     "\t\t নীট এফওবি রপ্তানি মূল্য ( বৈদেশিক মুদ্রায়): " + currency + ' ' + str(
                    #         "{:,.2f}".format(swift_amt)) +
                    #     "\n১২।\tপ্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ \t\t\t\t\t\t তারিখঃ")
                    # paragraph2.style = custom_stylebl

                    paragraph2 = document.add_paragraph(
                        "৬।\tরপ্তানিকৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত স্থানীয় সেবা/উপকরণাদির সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): 100% in house \tproduction of Brain Station 23 Limited \t\t পরিমাণঃ N/A  \t\t মূল্যঃ N/A"
                        "\n৭।\tরপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত আমদানিকৃত আনুসাংগিক সেবা/উপরকরণাদি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): N/A  \t\tপরিমাণঃ N/A   \t\t\t\t\t  মূল্যঃ  N/A")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_after = Pt(1)

                    long_text = "৮। \t " + ' ' * 4 + "রপ্তানি সেবা/পণ্যের বিবরণঃ Software Development Services পরিমাণঃ " + all_invoice_qty_str_a + " " + "  মূল্যঃ " + currency + ' ' + str(
                        "{:,.2f}".format(inv_amount))
                    wrapped_lines = textwrap.wrap(long_text, width=width)
                    formatted_lines = [line + '\t' for line in wrapped_lines]
                    formatted_text = '\n\t'.join(formatted_lines)
                    paragraph2 = document.add_paragraph(formatted_text)
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    if len(file_obj.invoice_line_ids) > 10:
                        long_text = "৯। \t " + ' ' * 4 + "জাহাজীকরণের তারিখঃ " + inv_date + " \t\tগন্তব্য বন্দরঃ " + country
                        wrapped_lines = textwrap.wrap(long_text, width=120)
                        formatted_lines = [line + '\t' for line in wrapped_lines]
                        formatted_text = '\n\t'.join(formatted_lines)
                        paragraph2 = document.add_paragraph(formatted_text)
                        paragraph2.style = custom_stylebl
                        paragraph2.paragraph_format.space_before = Pt(1)
                        paragraph2.paragraph_format.space_after = Pt(1)
                    else:
                        if len(file_obj.invoice_line_ids) > 7:
                            paragraph2 = document.add_paragraph(
                                "৯।\tজাহাজীকরণের তারিখঃ " + inv_date + "\n\tগন্তব্য বন্দরঃ " + country)
                            paragraph2.style = custom_stylebl
                        else:
                            paragraph2 = document.add_paragraph(
                                "৯।\tজাহাজীকরণের তারিখঃ " + inv_date + " \tগন্তব্য বন্দরঃ " + country)
                            paragraph2.style = custom_stylebl
                        paragraph2.paragraph_format.space_before = Pt(1)
                        paragraph2.paragraph_format.space_after = Pt(1)

                    # if len(file_obj.invoice_line_ids) > 7:
                    #     paragraph2 = document.add_paragraph(
                    #         "৯।\tজাহাজীকরণের তারিখঃ " + inv_date + "\n\tগন্তব্য বন্দরঃ " + country)
                    #     paragraph2.style = custom_stylebl
                    # else:
                    #     paragraph2 = document.add_paragraph(
                    #         "৯।\tজাহাজীকরণের তারিখঃ " + inv_date + " \tগন্তব্য বন্দরঃ " + country)
                    #     paragraph2.style = custom_stylebl
                    # paragraph2.paragraph_format.space_before = Pt(1)
                    # paragraph2.paragraph_format.space_after = Pt(1)

                    paragraph2 = document.add_paragraph(
                        "১০।\tইএক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানরি ক্ষেত্রে): N/A  \t\t\t  মূল্যঃ N/A   \t\t তারিখঃ N/A")
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    # long_text1 = "১১। \t " + ' ' * 5 +  "মোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): USD 158,793.64          নীট এফওবি রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): USD 158,793.64"
                    long_text1 = "১১। \t " + ' ' * 4 + "মোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str(
                        "{:,.2f}".format(
                            swift_amt)) + "\t\tনীট এফওবি রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str(
                        "{:,.2f}".format(swift_amt))
                    wrapped_lines1 = textwrap.wrap(long_text1, width=150)
                    formatted_lines1 = [line + '\t' for line in wrapped_lines1]
                    formatted_text1 = '\n\t'.join(formatted_lines1)
                    paragraph2 = document.add_paragraph(formatted_text1)
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    prc_date = ''
                    if file_obj.prc_date:
                        prc_date = datetime.datetime.strptime(str(file_obj.prc_date), '%Y-%m-%d').strftime('%d/%m/%y')
                    prc_ref_code = ''
                    if file_obj.prc_ref_code:
                        prc_ref_code = file_obj.prc_ref_code

                    long_text1 = "১২।\t" + ' ' * 4 + "প্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ " + prc_ref_code + "\t\t\t তারিখঃ " + prc_date
                    wrapped_lines1 = textwrap.wrap(long_text1, width=120)
                    formatted_lines1 = [line + '\t' for line in wrapped_lines1]
                    formatted_text1 = '\n\t'.join(formatted_lines1)
                    paragraph2 = document.add_paragraph(formatted_text1)
                    paragraph2.style = custom_stylebl
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    # paragraph2 = document.add_paragraph(
                    #     "১২।\tপ্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ \t\t\t\t\t\t তারিখঃ")
                    # paragraph2.style = custom_stylebl

                    paragraph2 = document.add_paragraph("\n\n\nরপ্তানিকারকের স্বাক্ষর ও তারিখ")
                    # paragraph2.style = custom_style3
                    # paragraph2.style = my_styler
                    paragraph2.style = custom_stylebsmallbold
                    paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    paragraph4 = document.add_paragraph(
                        "এতদ্বারা প্রত্যয়ন করা যাচ্ছে যে, আমাদের নিজস্ব কারখানায়/প্রতিষ্ঠানে তৈরীকৃত/উৎপাদিত সফটওয়্যার/আইটিইএস/হার্ডওয়্যার উপরোক্ত ৬ ও ৭ নং ক্রমিক বর্ণিত সূত্র হতে সেবা/উপকরাণাদি সংগ্রহের মাধ্যমে রপ্তানির বিপরীতে ভর্তুকির জন্য উপরোক্ত অনুচ্ছেদগুলোতে উল্লিখিত বক্তব্য সঠিক ও নির্ভুল। বিদেশী ক্রেতা/ আমদানিকারকের ক্রয়াদেশের যথার্থতা/বিশ্বাসযোগ্যতা সম্পর্কেও নিশ্চিত করা হলো। ")
                    paragraph4.style = custom_stylebsmall
                    paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # paragraph4 = document.paragraphs[1]
                    # Modify the paragraph format to reduce spacing
                    paragraph4.paragraph_format.space_before = Pt(1)
                    paragraph4.paragraph_format.space_after = Pt(1)

                    paragraph2 = document.add_paragraph("\n\n\nরপ্তানিকারকের স্বাক্ষর ও তারিখ")
                    # paragraph2.style = custom_style3
                    # paragraph2.style = my_styler
                    paragraph2.style = custom_stylebsmallbold
                    paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    paragraph2 = document.add_paragraph(
                        "রপ্তানিকারকের উপরোক্ত ঘোষণার যথার্থতা যাচাইয়ান্তে সঠিক পাওয়া গিয়েছে। ৮নং ক্রমিকে উল্লিখিত ঘোষিত রপ্তানিমূল্য যৌক্তিক ও বিদ্যমান আন্তর্জাতিক বাজার মুল্যের সঙ্গে সংগতিপূর্ণ পাওয়া গিয়েছে এবং বিদেশী ক্রেতার যথার্থতা/বিশ্বাসযোগ্যতা সর্ম্পকেও নিশ্চিত হওয়া গিয়েছে। প্রত্যাবাসিত রপ্তানি মূল্যের (নীট এফওবি মূল্য) রপ্তানি ভর্তুকি পরিশোধের সুপারিশ করা হলো।")
                    paragraph2.style = custom_stylebsmall
                    paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    paragraph2 = document.add_paragraph(
                        "\n\n\n\n----------------------------------- এসোসিয়েশন এর দুইজন উপযুক্ত কর্মকর্তার স্বাক্ষর, তারিখ ও সীল")
                    # paragraph2.style = my_styler
                    paragraph2.style = custom_stylebsmall
                    paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)

                    paragraph2 = document.add_paragraph(
                        "[কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।]")
                    # paragraph2.style = my_stylec
                    paragraph2.style = custom_stylebsmall
                    paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph2.paragraph_format.space_before = Pt(1)
                    paragraph2.paragraph_format.space_after = Pt(1)
                    if page_sl != len(active_ids):
                        document.add_page_break()
                        page_sl += 1

                    # file_name = '%s_kha' % (datetime.datetime.now())
                    file_name = 'Form_KHA_Multiple_Files_%s' % (datetime.datetime.now())

                # elif self.report_name == '06':
                #     from docx.enum.section import WD_ORIENT
                #     section = document.sections[0]
                #     new_width, new_height = section.page_height, section.page_width
                #     section.orientation = WD_ORIENT.LANDSCAPE
                #     section.page_width = new_width
                #     section.page_height = new_height
                #     self_date = ''
                #     if file_obj.date:
                #         self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')
                #
                #     custom_bangla_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                #     custom_bangla_style.font.size = Pt(11)
                #     custom_bangla_style.font.name = 'SutonnyOMJ'
                #     mystyle += 1
                #
                #     custom_bangla_style_ga = document.styles.add_style(str(mystyle),
                #                                                        WD_STYLE_TYPE.PARAGRAPH)
                #     custom_bangla_style_ga.font.size = Pt(11)
                #     custom_bangla_style_ga.font.name = 'SutonnyOMJ'
                #     mystyle += 1
                #
                #     custom_stylebold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                #     custom_stylebold.font.size = Pt(14)
                #     custom_stylebold.font.name = 'SutonnyOMJ'
                #     custom_stylebold.font.bold = True
                #     custom_stylebold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #     mystyle += 1
                #
                #     custom_styleubl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                #     custom_styleubl.font.size = Pt(11)
                #     custom_styleubl.font.name = 'SutonnyOMJ'
                #     custom_styleubl.font.underline = True
                #     custom_styleubl.font.bold = True
                #     mystyle += 1
                #
                #     # p1 = document.add_paragraph('(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য) \t\t\t \t \t \t \t  \t \t  \t \t \t \t   ফরম-‘গ’ \n')
                #     # p1.style = custom_bangla_style
                #
                #     paragraph = document.add_paragraph()
                #     # paragraph.add_run("(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
                #     run1 = paragraph.add_run("(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
                #     run1.font.size = Pt(11)
                #     run1.font.name = 'SutonnyOMJ'
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t")
                #     paragraph.add_run("\t ")
                #     paragraph.add_run("\t ")
                #     paragraph.add_run("\t ")
                #     paragraph.add_run("\t ")
                #     run2 = paragraph.add_run("ফরম-‘গ’")
                #     run2.font.size = Pt(14)
                #     run2.font.bold = True
                #     run2.font.name = 'SutonnyOMJ'
                #     # paragraph.add_run("ফরম-‘গ’")
                #     # run2.style = custom_bangla_style_ga
                #
                #     # paragraph = document.add_paragraph()
                #     #
                #     # # Add runs with different font sizes to the paragraph
                #     # run1 = paragraph.add_run("This is text with font size 12.")
                #     # run1.font.size = Pt(12)
                #     #
                #     # run2 = paragraph.add_run(" This is text with font size 16.")
                #     # run2.font.size = Pt(16)
                #
                #     # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)
                #
                #     # Define a custom style
                #     custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                #     # custom_style.font.bold = True
                #     custom_style.font.size = Pt(14)
                #     custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #     mystyle += 1
                #
                #     custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                #     custom_style2.font.size = Pt(11)
                #     custom_style2.font.name = 'Arial Narrow'
                #     mystyle += 1
                #
                #     custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                #     custom_style3.font.size = Pt(11)
                #     custom_style3.font.bold = True
                #     mystyle += 1
                #
                #     paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক")
                #     paragraph2.style = custom_bangla_style_ga
                #     paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #     paragraph2.paragraph_format.space_before = Pt(1)
                #     paragraph2.paragraph_format.space_after = Pt(1)
                #
                #     # paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক \nবাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব \n-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
                #     paragraph2 = document.add_paragraph(
                #         "বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব")
                #     paragraph2.style = custom_stylebold
                #     paragraph2.paragraph_format.space_before = Pt(1)
                #     paragraph2.paragraph_format.space_after = Pt(1)
                #
                #     paragraph2 = document.add_paragraph(
                #         "-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
                #     paragraph2.style = custom_bangla_style_ga
                #     paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #     paragraph2.paragraph_format.space_before = Pt(1)
                #
                #     # table -----------------------
                #     currency = ''
                #     if file_obj.fc_currency_id:
                #         currency = file_obj.fc_currency_id.name
                #
                #     a = []
                #     c = []
                #     for rec in file_obj.invoice_line_ids:
                #         a.append(rec.od_sight_rate)
                #         if rec.encashment_rate_bdt:
                #             c.append(rec.encashment_rate_bdt)
                #         # if rec.swift_message_id.encashment_rate_bdt:
                #         #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
                #     inv_amount = sum(file_obj.invoice_line_ids.mapped('swift_amt'))
                #     if a:
                #         b = min(a)
                #         od_s_r = b
                #     else:
                #         od_s_r = 0
                #
                #     if c:
                #         d = min(c)
                #         encashment_rate_bdt = d
                #     else:
                #         encashment_rate_bdt = 0
                #
                #     if file_obj.od_sight_rate:
                #         od_sight_rate = file_obj.od_sight_rate
                #     else:
                #         if od_s_r:
                #             od_sight_rate = od_s_r
                #         else:
                #             od_sight_rate = encashment_rate_bdt
                #
                #     rate_fc = inv_amount / file_obj.incentive_rate_fc
                #     encashment_final_amt = round(rate_fc, 2) * od_sight_rate
                #     # encashment_final_amt_a = "{:,.0f}".format(round(encashment_final_amt))
                #     encashment_final_amt_a = file_obj.custom_round(encashment_final_amt)
                #     usd_bdt_amount = currency + ' ' + str("{:,.2f}".format(rate_fc)) + '\n' + '@' + \
                #                      str("{:,.2f}".format(od_sight_rate)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a
                #
                #     table = document.add_table(rows=4, cols=9)
                #     table.style = 'TableGrid'
                #     cell_1_0 = table.cell(0, 0)
                #     p3 = cell_1_0.paragraphs[0]
                #     p3.add_run('ভর্তুকি আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানা')
                #     p3.style = custom_bangla_style
                #     cell_2_0 = table.cell(1, 0)
                #     cell_1_0.merge(cell_2_0)
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell_1_1 = table.cell(0, 1)
                #     p3 = cell_1_1.paragraphs[0]
                #     p3.add_run('রপ্তানি সংশ্লিষ্ট তথ্য')
                #     p3.style = custom_bangla_style
                #     cell_2_1 = table.cell(0, 5)
                #     cell_1_1.merge(cell_2_1)
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell_1_1 = table.cell(1, 1)
                #     p3 = cell_1_1.paragraphs[0]
                #     p3.add_run('সেবা/পণ্যের বিবরণ')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell_1_1 = table.cell(1, 2)
                #     p3 = cell_1_1.paragraphs[0]
                #     p3.add_run('জাহাজীকরণের তারিখ ও গন্তব্যস্থল')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell_1_1 = table.cell(1, 3)
                #     p3 = cell_1_1.paragraphs[0]
                #     p3.add_run('ইক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানির ক্ষেত্রে)')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell_1_1 = table.cell(1, 4)
                #     p3 = cell_1_1.paragraphs[0]
                #     p3.add_run('মোট প্রত্যাবাসিত রপ্তানি (মূল্য  বৈদেশিক মুদ্রা)')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell_1_1 = table.cell(1, 5)
                #     p3 = cell_1_1.paragraphs[0]
                #     p3.add_run('নীট এফওবি মূল্য (বৈদেশিক মুদ্রা)')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell_1_3 = table.cell(0, 6)
                #     p3 = cell_1_3.paragraphs[0]
                #     p3.add_run('শাখা কর্তৃক আবেদনপত্র গ্রহণের তারিখ')
                #     p3.style = custom_bangla_style
                #     cell_2_4 = table.cell(1, 6)
                #     cell_1_3.merge(cell_2_4)
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell_1_5 = table.cell(0, 7)
                #     p3 = cell_1_5.paragraphs[0]
                #     p3.add_run('আবেদনকৃত অর্থের পরিমাণ')
                #     p3.style = custom_bangla_style
                #     cell_2_6 = table.cell(1, 7)
                #     cell_1_5.merge(cell_2_6)
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell_1_6 = table.cell(0, 8)
                #     p3 = cell_1_6.paragraphs[0]
                #     p3.add_run('পরিশোধ্য দাবির পরিমাণ')
                #     p3.style = custom_bangla_style
                #     cell_2_7 = table.cell(1, 8)
                #     cell_1_6.merge(cell_2_7)
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(2, 0)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('১')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(2, 1)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('২')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(2, 2)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('৩')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(2, 3)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('৪')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(2, 4)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('৫')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(2, 5)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('৬')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(2, 6)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('৭')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(2, 7)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('৮')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(2, 8)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('৯')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(3, 0)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run(
                #         'ব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী  বাণিজ্যিক এলাকা, ঢাকা- ১২১২।')
                #     p3.style = custom_bangla_style
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(3, 1)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('Software Development')
                #     p3.style = custom_style2
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(3, 2)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('')
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(3, 3)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('')
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(3, 4)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #     p3.style = custom_style3
                #
                #     cell = table.cell(3, 5)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #     p3.style = custom_style3
                #
                #     cell = table.cell(3, 6)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('')
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(3, 7)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run(usd_bdt_amount)
                #     p3.style = custom_style3
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     cell = table.cell(3, 8)
                #     p3 = cell.paragraphs[0]
                #     p3.add_run('')
                #     p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #
                #     p3 = document.add_paragraph(
                #         '\n\n\n\n\n\n\n\n\t\t দাপ্তরিক সীল  \t\t\t\t\t\t \t\t  ক্ষমতাপ্রাপ্ত ব্যাংক কর্মকর্তার স্বাক্ষর, তারিখ, নাম ও পদবী')
                #     p3.style = custom_bangla_style
                #
                #     if page_sl != len(active_ids):
                #         document.add_page_break()
                #         page_sl += 1
                #
                #     file_name = '%s_ga' % (datetime.datetime.now())
                    
                # if self.report_name == '07':
                #     self.action_print_pdf_gha()
                else:
                    raise UserError(_('Failed! API Authentication Error!'))

            # -------------------
            import os
            dir_path = os.path.dirname(os.path.abspath(__file__))
            base_path = str(dir_path).replace('/models', '').replace('/wizards', '')
            docxfile = base_path + '/static/docx/' + file_name + '.docx'
            document.save(docxfile)
            # docx.save(docxfile)

            # return document.save('/home/jobaer/Downloads/jh3.docx')

            return {
                'type': 'ir.actions.act_url',
                'url': 'cash_incentive/static/docx/' + file_name + '.docx',
                'target': 'self',
            }

        else:
            with_head = self.env.context.get('with_head')
            data = {}
            result = []
            for data in active_ids:
                file_obj = self.env['cash.incentive.head'].browse(data)
                if self.report_name == '01':
                    file_obj.action_refresh_prc3()
                    # vals = {
                    #     'details': file_obj.prc_letter_description,
                    # }
                    # result.append(vals)
                if self.report_name == '02':
                    file_obj.action_refresh_form_ka()
                    # vals = {
                    #     'details': file_obj.form_ka_basis_description,
                    # }
                    # result.append(vals)
                if self.report_name == '03':
                    file_obj.action_refresh_basis_kha()
                    # vals = {
                    #     'details': file_obj.form_kha_basis_description,
                    # }
                    # result.append(vals)

            data = {
                'model': "cash.incentive.print.wizards",
                'form': self.read()[0],
                # 'csr': result,
                'with_head': with_head,
            }
            data['incentive_ids'] = active_ids
            data['rpt_name'] = self.report_name
            # get data from sql
            return self.env.ref('cash_incentive.report_incentive_letter_action_tree_print_ids').with_context(
                landscape=False).report_action(self, data=data)
    
    def action_print_pdf_gha(self):
        active_ids = self.env.context.get('active_ids')
        from docx import Document
        from htmldocx import HtmlToDocx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.text import WD_BREAK
        from docx.shared import Pt
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.section import WD_ORIENT

        from docx.shared import Inches
        file_name = ''
        document = Document()
        from docx.shared import Cm
        section = document.sections[0]
        section.left_margin = Cm(1.5)  # 1.5 cm
        section.right_margin = Cm(1.5)  # 1.5 cm

        from docx.enum.section import WD_ORIENT
        section = document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        mystyle = 0
        page_sl = 1
        file_names = ''
        # for data in active_ids:
        #     file_obj = self.env['cash.incentive.head'].browse(data)
        #     ref_name_val = file_obj.name
        #     split_values = ref_name_val.split("/")
        #     desired_value = split_values[-1]
        #     file_names += desired_value if not file_names else ', ' + desired_value

        for data in active_ids:
            if self.report_name == '07':
                file_obj = self.env['cash.incentive.head'].browse(data)
                ref_name_val = file_obj.name
                split_values = ref_name_val.split("/")
                desired_value = split_values[-1]

                country = ''
                if file_obj.partner_id.country_id:
                    country = file_obj.partner_id.country_id.name

                currency = ''
                if file_obj.fc_currency_id:
                    currency = file_obj.fc_currency_id.name
                inv_number = ''
                inv_date = ''
                inv_amount = 0
                swift_amt = 0
                swift_id = None
                for rec in file_obj.invoice_line_ids:
                    swift_amt += rec.swift_amt
                    invoice_date = ''
                    if rec.invoice_date:
                        invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
                    inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
                    inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
                    inv_amount += rec.invoice_amt
                    if not swift_id:
                        swift_id = rec.swift_message_id
                client_name = ''
                if file_obj.partner_id:
                    client_name = '' if not file_obj.swift_customer_name else file_obj.swift_customer_name

                customer_address = file_obj.customer_address

                remiter_address = ''
                if swift_id:
                    if swift_id.remiter_bank_name:
                        remiter_address += swift_id.remiter_bank_name
                    if swift_id.remiter_bank_address:
                        remiter_address += swift_id.remiter_bank_address if not remiter_address else ', ' + swift_id.remiter_bank_address

                self_date = ''
                if file_obj.date:
                    self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')
                # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

                # Define a custom style
                custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                # custom_style.font.bold = True
                custom_style.font.size = Pt(14)
                custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mystyle += 1

                custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_style2.font.size = Pt(11)
                custom_style2.font.name = 'Arial Narrow'
                mystyle += 1

                custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_style3.font.size = Pt(11)
                custom_style3.font.bold = True
                mystyle += 1

                custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_style4.font.bold = True
                custom_style4.font.size = Pt(14)
                custom_style4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                custom_style4.font.name = 'SutonnyOMJ'
                mystyle += 1

                custom_bangla_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_bangla_style.font.size = Pt(11)
                custom_bangla_style.font.name = 'SutonnyOMJ'
                mystyle += 1

                custom_bangla_style_gha = document.styles.add_style(str(mystyle),
                                                                    WD_STYLE_TYPE.PARAGRAPH)
                custom_bangla_style_gha.font.size = Pt(11)
                custom_bangla_style_gha.font.name = 'SutonnyOMJ'
                mystyle += 1

                custom_stylebold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_stylebold.font.size = Pt(16)
                custom_stylebold.font.name = 'SutonnyOMJ'
                custom_stylebold.font.bold = True
                custom_stylebold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mystyle += 1

                custom_styleubl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_styleubl.font.size = Pt(13)
                custom_styleubl.font.name = 'SutonnyOMJ'
                custom_styleubl.font.bold = True
                custom_styleubl.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mystyle += 1

                paragraph1 = document.add_paragraph('\nফরম-‘ঘ’')
                paragraph1.style = custom_style4

                p3 = document.add_paragraph('(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য) \n')
                p3.style = custom_bangla_style_gha

                # paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক \nবাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধিত হিসাবের .........সালের ...... মাসের বিবরণী \nঅনুমোদিত ডিলার ব্যাংক শাখার নামঃ...................................")

                paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক")
                paragraph2.style = custom_bangla_style_gha
                paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph2.paragraph_format.space_before = Pt(1)
                paragraph2.paragraph_format.space_after = Pt(1)

                paragraph2 = document.add_paragraph(
                    "বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও")
                paragraph2.style = custom_stylebold
                paragraph2.paragraph_format.space_before = Pt(1)
                paragraph2.paragraph_format.space_after = Pt(1)

                paragraph2 = document.add_paragraph(
                    "হার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধিত হিসাবের .........সালের ...... মাসের বিবরণী")
                paragraph2.style = custom_styleubl
                paragraph2.paragraph_format.space_before = Pt(1)
                paragraph2.paragraph_format.space_after = Pt(1)

                paragraph2 = document.add_paragraph(
                    "অনুমোদিত ডিলার ব্যাংক শাখার নামঃ...................................................................................")
                paragraph2.style = custom_bangla_style_gha
                paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph2.paragraph_format.space_before = Pt(1)

                # table -----------------------
                currency = ''
                if file_obj.fc_currency_id:
                    currency = file_obj.fc_currency_id.name

                a = []
                c = []
                for rec in file_obj.invoice_line_ids:
                    a.append(rec.od_sight_rate)
                    if rec.encashment_rate_bdt:
                        c.append(rec.encashment_rate_bdt)
                    # if rec.swift_message_id.encashment_rate_bdt:
                    #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
                inv_amount = sum(file_obj.invoice_line_ids.mapped('swift_amt'))

                if a:
                    b = min(a)
                    od_s_r = b
                else:
                    od_s_r = 0

                if c:
                    d = min(c)
                    encashment_rate_bdt = d
                else:
                    encashment_rate_bdt = 0

                if file_obj.od_sight_rate:
                    od_sight_rate = file_obj.od_sight_rate
                else:
                    if od_s_r:
                        od_sight_rate = od_s_r
                    else:
                        od_sight_rate = encashment_rate_bdt

                rate_fc = inv_amount / file_obj.incentive_rate_fc
                encashment_final_amt = round(rate_fc, 2) * od_sight_rate
                # encashment_final_amt_a = "{:,.0f}".format(round(encashment_final_amt))
                encashment_final_amt_a = file_obj.custom_round(encashment_final_amt)
                usd_bdt_amount = currency + ' ' + str("{:,.2f}".format(rate_fc)) + '\n' + '@' + \
                                 str("{:,.2f}".format(od_sight_rate)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a

                table = document.add_table(rows=4, cols=10)
                table.style = 'TableGrid'

                cell_1_0 = table.cell(0, 0)
                p3 = cell_1_0.paragraphs[0]
                p3.add_run('ভর্তুকি আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানা')
                p3.style = custom_bangla_style_gha
                cell_2_0 = table.cell(1, 0)
                cell_1_0.merge(cell_2_0)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(0, 1)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('রপ্তানি সংশ্লিষ্ট তথ্য')
                p3.style = custom_bangla_style_gha
                cell_2_1 = table.cell(0, 5)
                cell_1_1.merge(cell_2_1)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_2_1 = table.cell(0, 8)
                p3 = cell_2_1.paragraphs[0]
                p3.add_run('শাখায় প্রাপ্ত অনিস্পন্ন আবেদন')
                p3.style = custom_bangla_style_gha
                cell_3_1 = table.cell(0, 9)
                cell_2_1.merge(cell_3_1)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 1)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('সেবা/পণ্যের বিবরণ')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 2)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('জাহাজীকরণের তারিখ ও গন্তব্যস্থল')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 3)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('ইক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানির ক্ষেত্রে)')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 4)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('মোট প্রত্যাবাসিত রপ্তানি (মূল্য  বৈদেশিক মুদ্রা)')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 5)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('নীট এফওবি মূল্য (বৈদেশিক মুদ্রা)')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_3 = table.cell(0, 6)
                table.columns[7].width = Inches(.6)
                p3 = cell_1_3.paragraphs[0]
                p3.add_run('পরিশোধিত ভর্তুকি (টাকা)')
                p3.style = custom_bangla_style_gha
                cell_2_4 = table.cell(1, 6)
                cell_1_3.merge(cell_2_4)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_5 = table.cell(0, 7)
                table.columns[7].width = Inches(.6)
                p3 = cell_1_5.paragraphs[0]
                p3.add_run('পরিশোধের তারিখ')
                p3.style = custom_bangla_style_gha
                cell_2_6 = table.cell(1, 7)
                cell_1_5.merge(cell_2_6)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_6 = table.cell(1, 8)
                table.columns[8].width = Inches(.5)
                p3 = cell_1_6.paragraphs[0]
                p3.add_run('সংখ্যা')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_6 = table.cell(1, 9)
                table.columns[9].width = Inches(1.5)
                p3 = cell_1_6.paragraphs[0]
                p3.add_run('দাবীকৃত অর্থ (টাকা)')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 0)
                p3 = cell.paragraphs[0]
                p3.add_run('১')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 1)
                p3 = cell.paragraphs[0]
                p3.add_run('২')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 2)
                p3 = cell.paragraphs[0]
                p3.add_run('৩')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 3)
                p3 = cell.paragraphs[0]
                p3.add_run('৪')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 4)
                p3 = cell.paragraphs[0]
                p3.add_run('৫')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 5)
                p3 = cell.paragraphs[0]
                p3.add_run('৬')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 6)
                p3 = cell.paragraphs[0]
                p3.add_run('৭')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 7)
                p3 = cell.paragraphs[0]
                p3.add_run('৮')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 8)
                p3 = cell.paragraphs[0]
                p3.add_run('৯')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 9)
                p3 = cell.paragraphs[0]
                p3.add_run('১০')
                p3.style = custom_bangla_style_gha
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 0)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run(
                    '\nব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী  বাণিজ্যিক এলাকা, ঢাকা- ১২১২।\n\n')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p3.style = custom_bangla_style_gha

                cell = table.cell(3, 1)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run('Software Development')
                p3.style = custom_style2
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 2)
                p3 = cell.paragraphs[0]
                p3.add_run('')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 3)
                p3 = cell.paragraphs[0]
                p3.add_run('')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 4)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
                p3.style = custom_style3
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 5)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
                p3.style = custom_style3
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 6)
                p3 = cell.paragraphs[0]
                p3.add_run('')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 7)
                p3 = cell.paragraphs[0]
                p3.add_run('')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 8)
                p3 = cell.paragraphs[0]
                p3.add_run('')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 9)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run(usd_bdt_amount)
                p3.style = custom_style3
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p3 = document.add_paragraph('\n\n\n\n\t\t দাপ্তরিক সীল  \t\t\t\t\t\t\t\t\t  ক্ষমতাপ্রাপ্ত ব্যাংক কর্মকর্তার স্বাক্ষর, তারিখ, নাম ও পদবী')
                p3.style = custom_bangla_style_gha

                if page_sl != len(active_ids):
                    document.add_page_break()
                    page_sl += 1

                file_name = 'Form_GHA_Multiple_Files_%s' % (datetime.datetime.now())
                # file_name = '%s_gha' % (datetime.datetime.now())

            else:
                file_obj = self.env['cash.incentive.head'].browse(data)
                self_date = ''
                if file_obj.date:
                    self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')

                custom_bangla_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_bangla_style.font.size = Pt(11)
                custom_bangla_style.font.name = 'SutonnyOMJ'
                mystyle += 1

                custom_bangla_style_ga = document.styles.add_style(str(mystyle),
                                                                   WD_STYLE_TYPE.PARAGRAPH)
                custom_bangla_style_ga.font.size = Pt(11)
                custom_bangla_style_ga.font.name = 'SutonnyOMJ'
                mystyle += 1

                custom_stylebold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_stylebold.font.size = Pt(14)
                custom_stylebold.font.name = 'SutonnyOMJ'
                custom_stylebold.font.bold = True
                custom_stylebold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mystyle += 1

                custom_styleubl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_styleubl.font.size = Pt(11)
                custom_styleubl.font.name = 'SutonnyOMJ'
                custom_styleubl.font.underline = True
                custom_styleubl.font.bold = True
                mystyle += 1

                # p1 = document.add_paragraph('(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য) \t\t\t \t \t \t \t  \t \t  \t \t \t \t   ফরম-‘গ’ \n')
                # p1.style = custom_bangla_style

                paragraph = document.add_paragraph()
                # paragraph.add_run("(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
                run1 = paragraph.add_run("\n(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
                run1.font.size = Pt(11)
                run1.font.name = 'SutonnyOMJ'
                paragraph.add_run("\t")
                paragraph.add_run("\t")
                paragraph.add_run("\t")
                paragraph.add_run("\t")
                paragraph.add_run("\t")
                paragraph.add_run("\t")
                paragraph.add_run("\t")
                paragraph.add_run("\t")
                paragraph.add_run("\t")
                paragraph.add_run("\t")
                paragraph.add_run("\t ")
                paragraph.add_run("\t ")
                paragraph.add_run("\t ")
                run2 = paragraph.add_run("ফরম-‘গ’")
                run2.font.size = Pt(14)
                run2.font.bold = True
                run2.font.name = 'SutonnyOMJ'
                # paragraph.add_run("ফরম-‘গ’")
                # run2.style = custom_bangla_style_ga

                # paragraph = document.add_paragraph()
                #
                # # Add runs with different font sizes to the paragraph
                # run1 = paragraph.add_run("This is text with font size 12.")
                # run1.font.size = Pt(12)
                #
                # run2 = paragraph.add_run(" This is text with font size 16.")
                # run2.font.size = Pt(16)

                # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)

                # Define a custom style
                custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                # custom_style.font.bold = True
                custom_style.font.size = Pt(14)
                custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                mystyle += 1

                custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_style2.font.size = Pt(11)
                custom_style2.font.name = 'Arial Narrow'
                mystyle += 1

                custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
                custom_style3.font.size = Pt(11)
                custom_style3.font.bold = True
                mystyle += 1

                paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক")
                paragraph2.style = custom_bangla_style_ga
                paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph2.paragraph_format.space_before = Pt(1)
                paragraph2.paragraph_format.space_after = Pt(1)

                # paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক \nবাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব \n-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
                paragraph2 = document.add_paragraph(
                    "বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব")
                paragraph2.style = custom_stylebold
                paragraph2.paragraph_format.space_before = Pt(1)
                paragraph2.paragraph_format.space_after = Pt(1)

                paragraph2 = document.add_paragraph(
                    "-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
                paragraph2.style = custom_bangla_style_ga
                paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph2.paragraph_format.space_before = Pt(1)

                # table -----------------------
                currency = ''
                if file_obj.fc_currency_id:
                    currency = file_obj.fc_currency_id.name

                a = []
                c = []
                for rec in file_obj.invoice_line_ids:
                    a.append(rec.od_sight_rate)
                    if rec.encashment_rate_bdt:
                        c.append(rec.encashment_rate_bdt)
                    # if rec.swift_message_id.encashment_rate_bdt:
                    #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
                inv_amount = sum(file_obj.invoice_line_ids.mapped('swift_amt'))
                if a:
                    b = min(a)
                    od_s_r = b
                else:
                    od_s_r = 0

                if c:
                    d = min(c)
                    encashment_rate_bdt = d
                else:
                    encashment_rate_bdt = 0

                if file_obj.od_sight_rate:
                    od_sight_rate = file_obj.od_sight_rate
                else:
                    if od_s_r:
                        od_sight_rate = od_s_r
                    else:
                        od_sight_rate = encashment_rate_bdt

                rate_fc = inv_amount / file_obj.incentive_rate_fc
                encashment_final_amt = round(rate_fc, 2) * od_sight_rate
                # encashment_final_amt_a = "{:,.0f}".format(round(encashment_final_amt))
                encashment_final_amt_a = file_obj.custom_round(encashment_final_amt)
                usd_bdt_amount = currency + ' ' + str("{:,.2f}".format(rate_fc)) + '\n' + '@' + \
                                 str("{:,.2f}".format(od_sight_rate)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a

                table = document.add_table(rows=4, cols=9)
                table.style = 'TableGrid'
                cell_1_0 = table.cell(0, 0)
                p3 = cell_1_0.paragraphs[0]
                p3.add_run('ভর্তুকি আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানা')
                p3.style = custom_bangla_style
                cell_2_0 = table.cell(1, 0)
                cell_1_0.merge(cell_2_0)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(0, 1)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('রপ্তানি সংশ্লিষ্ট তথ্য')
                p3.style = custom_bangla_style
                cell_2_1 = table.cell(0, 5)
                cell_1_1.merge(cell_2_1)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 1)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('সেবা/পণ্যের বিবরণ')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 2)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('জাহাজীকরণের তারিখ ও গন্তব্যস্থল')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 3)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('ইক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানির ক্ষেত্রে)')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 4)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('মোট প্রত্যাবাসিত রপ্তানি (মূল্য  বৈদেশিক মুদ্রা)')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_1 = table.cell(1, 5)
                p3 = cell_1_1.paragraphs[0]
                p3.add_run('নীট এফওবি মূল্য (বৈদেশিক মুদ্রা)')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_3 = table.cell(0, 6)
                p3 = cell_1_3.paragraphs[0]
                p3.add_run('শাখা কর্তৃক আবেদনপত্র গ্রহণের তারিখ')
                p3.style = custom_bangla_style
                cell_2_4 = table.cell(1, 6)
                cell_1_3.merge(cell_2_4)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_5 = table.cell(0, 7)
                p3 = cell_1_5.paragraphs[0]
                p3.add_run('আবেদনকৃত অর্থের পরিমাণ')
                p3.style = custom_bangla_style
                cell_2_6 = table.cell(1, 7)
                cell_1_5.merge(cell_2_6)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_1_6 = table.cell(0, 8)
                p3 = cell_1_6.paragraphs[0]
                p3.add_run('পরিশোধ্য দাবির পরিমাণ')
                p3.style = custom_bangla_style
                cell_2_7 = table.cell(1, 8)
                cell_1_6.merge(cell_2_7)
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 0)
                p3 = cell.paragraphs[0]
                p3.add_run('১')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 1)
                p3 = cell.paragraphs[0]
                p3.add_run('২')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 2)
                p3 = cell.paragraphs[0]
                p3.add_run('৩')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 3)
                p3 = cell.paragraphs[0]
                p3.add_run('৪')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 4)
                p3 = cell.paragraphs[0]
                p3.add_run('৫')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 5)
                p3 = cell.paragraphs[0]
                p3.add_run('৬')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 6)
                p3 = cell.paragraphs[0]
                p3.add_run('৭')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 7)
                p3 = cell.paragraphs[0]
                p3.add_run('৮')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(2, 8)
                p3 = cell.paragraphs[0]
                p3.add_run('৯')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 0)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run(
                    '\n\nব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী  বাণিজ্যিক এলাকা, ঢাকা- ১২১২।\n\n')
                p3.style = custom_bangla_style
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 1)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run('Software Development')
                p3.style = custom_style2
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 2)
                p3 = cell.paragraphs[0]
                p3.add_run('')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 3)
                p3 = cell.paragraphs[0]
                p3.add_run('')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 4)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p3.style = custom_style3

                cell = table.cell(3, 5)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p3.style = custom_style3

                cell = table.cell(3, 6)
                p3 = cell.paragraphs[0]
                p3.add_run('')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 7)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p3 = cell.paragraphs[0]
                p3.add_run(usd_bdt_amount)
                p3.style = custom_style3
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell = table.cell(3, 8)
                p3 = cell.paragraphs[0]
                p3.add_run('')
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p3 = document.add_paragraph('\n\n\n\n\n\t\t দাপ্তরিক সীল  \t\t\t\t\t\t \t\t  ক্ষমতাপ্রাপ্ত ব্যাংক কর্মকর্তার স্বাক্ষর, তারিখ, নাম ও পদবী')
                p3.style = custom_bangla_style

                if page_sl != len(active_ids):
                    document.add_page_break()
                    page_sl += 1

                # file_name = '%s_ga' % (datetime.datetime.now())
                file_name = 'Form_GA_Multiple_Files_%s' % (datetime.datetime.now())

        # -------------------
        import os
        dir_path = os.path.dirname(os.path.abspath(__file__))
        base_path = str(dir_path).replace('/models', '').replace('/wizards', '')
        docxfile = base_path + '/static/docx/' + file_name + '.docx'
        document.save(docxfile)
        # docx.save(docxfile)

        # return document.save('/home/jobaer/Downloads/jh3.docx')

        return {
            'type': 'ir.actions.act_url',
            'url': 'cash_incentive/static/docx/' + file_name + '.docx',
            'target': 'self',
        }
    
    # all_form_word -----------------------------------
    def action_download_all_docx_file(self):
        active_ids = self.env.context.get('active_ids')
        type = self.env.context.get('type')
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        from docx.oxml.shared import OxmlElement, qn

        from docx.shared import Inches
        # create a new document

        # get the section object and adjust the left and right margins
        from docx.shared import Cm
        # section.left_margin = Cm(1.5)  # 1.5 cm
        # section.right_margin = Cm(1.5)  # 1.5 cm
        # jh = Document()
        # new_parser = HtmlToDocx()
        document = Document()
        mystyle = 0
        page_sl = 1
        file_names = ''

        for data in active_ids:
            if mystyle == 0:
                section = document.sections[0]
            else:
                section = document.add_section()
            section.orientation = 1
            # section1.page_width = Pt(612)
            # section1.page_height = Pt(792)

            file_obj = self.env['cash.incentive.head'].browse(data)

            country = ''
            if file_obj.partner_id.country_id:
                country = file_obj.partner_id.country_id.name

            currency = ''
            if file_obj.fc_currency_id:
                currency = file_obj.fc_currency_id.name
            inv_number = ''
            inv_date = ''
            inv_amount = 0
            swift_amt = 0
            swift_id = None
            inv_ids = []
            for rec in file_obj.invoice_line_ids:
                swift_amt += rec.swift_amt
                invoice_date = ''
                if rec.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d/%m/%y')
                if rec.invoice_id not in inv_ids:
                    inv_ids.append(rec.invoice_id)
                    inv_number += str(rec.invoice_id.ref) if not inv_number else ', ' + str(rec.invoice_id.ref)
                    inv_date += str(invoice_date) if not inv_date else ', ' + str(invoice_date)
                    inv_amount += rec.invoice_amt
                if not swift_id:
                    swift_id = rec.swift_message_id
            client_name = ''
            if file_obj.partner_id:
                client_name = '' if not file_obj.swift_customer_name else file_obj.swift_customer_name

            customer_address = file_obj.customer_address

            remiter_address = ''
            if swift_id:
                if swift_id.remiter_bank_name:
                    remiter_address += swift_id.remiter_bank_name
                if swift_id.remiter_bank_address:
                    remiter_address += swift_id.remiter_bank_address if not remiter_address else ', ' + swift_id.remiter_bank_address

            file_name = ''
            # prc 11111111111-------------------------------------------------------------------------------------------------------
            self_date = ''
            if file_obj.date:
                self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')
    
            document.add_paragraph('\n\nRef: ' + file_obj.name + '\t\t\t\t\t\t\t\t\t Format-A' + '\nDate: '+ self_date + '\n' )
            # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)
    
            # Define a custom style
            custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mystyle += 1
    
            custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(11)
            custom_style2.font.name = 'Arial Narrow'
            custom_style2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            mystyle += 1
    
            custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.bold = True
            custom_style3.font.size = Pt(10)
            custom_style3.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.size = Pt(11)
            custom_style4.font.name = 'Arial Narrow'
            mystyle += 1
    
            # Add a paragraph and apply the custom style
            paragraph1 = document.add_paragraph('CERTIFICATE OF AUTHORIZED DEALER')
            paragraph1.style = custom_style
            contract_number = file_obj.contract_number
            if not file_obj.contract_number:
                contract_number = ''
            contract_date_str = file_obj.contract_date_str
            if not file_obj.contract_date_str:
                contract_date_str = ''
            paragraph2 = document.add_paragraph("This is to certify that we have received following inward remittance in "
                                                "the name of Brain Station 23 Limited, Plot 02  (8th Floor), Bir Uttam A. K."
                                                " Khandakar Road, Mohakhali C/A, Dhaka-1212 against "
                                                + contract_number + ' & date: ' + str(contract_date_str) + " invoice no. "
                                                + inv_number + " & date:  " + inv_date +" amount in " +currency + ' ' +  str("{:,}".format(round(inv_amount, 2)))
                                                + " for rendering of software development service. Summary of the transaction is as follows: ")
            paragraph2.style = custom_style2
            # Set the paragraph alignment to center
    
            table = document.add_table(rows=1, cols=9)
            table.style = 'TableGrid'
    
            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(a) Remitter')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
            # cell.merge(table.cell(0, 1))
    
            cell = table.cell(0, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(b) Address')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            cell = table.cell(0, 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Invoice No.')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            cell = table.cell(0, 3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name = '(c) Amount in FC (' + file_obj.fc_currency_id.name + ')'
            p3.add_run(cur_name)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            cell = table.cell(0, 4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(d) Date of Credit in banks nostro account')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            cell = table.cell(0, 5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(e) Equivalent Taka')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            cell = table.cell(0, 6)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(f) Credited to beneficiary a/c')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            cell = table.cell(0, 7)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(g) Reporting statement/schedule to BB with Month')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            cell = table.cell(0, 8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('(h) Reference of Online reporting to BB')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            # add a dynamic row to the table
            cus_name = ''
            if file_obj.swift_customer_name:
                cus_name = file_obj.swift_customer_name
    
            # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
            row_len = len(file_obj.invoice_line_ids)
            inv_amount = 0
            encashment_amt_bdt = 0
            equivalent_taka_total = 0
    
            from itertools import groupby
    
            # group the invoice_line_ids by invoice_id.ref
            groups = groupby(sorted(file_obj.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                             key=lambda x: x.invoice_id.ref)
            # iterate over each group and add a new row for each unique invoice_id.ref
            row = 1
            swift_amt_total = 0
            for invoice_ref, group in groups:
                gr = 0
                for rec in group:
                    gr += 1
                    new_row = table.add_row()
                    inv_amount += rec.invoice_amt
                    swift_amt_total += rec.swift_amt
                    encashment_amt_bdt += rec.encashment_amt_bdt
    
                    swift_date = ''
                    if rec.swift_message_id.date:
                        swift_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime(
                            '%d-%b-%y')
    
                    date_credited_beneficiaries = ''
                    if rec.date_credited_beneficiaries:
                        date_credited_beneficiaries = datetime.datetime.strptime(str(rec.date_credited_beneficiaries),
                                                                                 '%Y-%m-%d').strftime('%d-%b-%y')
                    reporting_st_to_bb = rec.reporting_st_to_bb if rec.reporting_st_to_bb else ''
                    ref_online_to_bb = rec.ref_online_to_bb if rec.ref_online_to_bb else ''
                    equivalent_taka = round(rec.swift_amt * rec.encashment_rate_bdt, 2)
                    equivalent_taka_total += equivalent_taka
    
                    equivalent_taka_decimal = "{:,.2f}".format(equivalent_taka)
    
                    # add data to the cells in the new row
                    # new_row.cells[2].text = rec.invoice_id.ref
                    new_row.cells[3].text = str("{:,.2f}".format(rec.swift_amt))
                    new_row.cells[4].text = str(swift_date)
                    new_row.cells[5].text = str(equivalent_taka_decimal)
                    new_row.cells[6].text = str(date_credited_beneficiaries)
                    new_row.cells[7].text = str(reporting_st_to_bb)
                    new_row.cells[8].text = str(ref_online_to_bb)
    
                    first_cell = new_row.cells[3]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style4
    
                    first_cell1 = new_row.cells[5]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = custom_style4
    
                    first_cell1 = new_row.cells[2]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # paragraph22.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph22.style = custom_style4
    
                    first_cell1 = new_row.cells[0]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[1]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[4]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[6]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[7]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style4
                    first_cell1 = new_row.cells[8]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style4
    
                cell_1_0 = table.cell(row, 2)
                p3 = cell_1_0.paragraphs[0]
                p3.add_run(invoice_ref)
                cell_2_0 = table.cell((row + gr) - 1, 2)
                cell_1_0.merge(cell_2_0)
                row = gr + row

            # cus name merge
            cell_1_0 = table.cell(1, 0)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run(cus_name)
            cell_2_0 = table.cell(row_len, 0)
            cell_1_0.merge(cell_2_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            # address merge
            cell_3_0 = table.cell(1, 1)
            p3 = cell_3_0.paragraphs[0]
            p3.add_run(file_obj.customer_address)
            cell_4_0 = table.cell(row_len, 1)
            cell_3_0.merge(cell_4_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            # total
            equivalent_taka_total_decimal = "{:,.0f}".format(equivalent_taka_total)
            total_row = table.add_row()
            total_row.cells[2].text = 'Total'
            total_row.cells[3].text = "{:,.2f}".format(swift_amt_total)
            total_row.cells[5].text = str(equivalent_taka_total_decimal)
    
            first_cell = total_row.cells[2]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = custom_style3
    
            first_cell = total_row.cells[3]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = custom_style3
    
            first_cell1 = total_row.cells[5]
            paragraph22 = first_cell1.paragraphs[0]
            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph22.style = custom_style3
    
            paragraph = document.add_paragraph()
            paragraph.add_run("\n\n\n\n\n Signature of Head of the branch.")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("Signature of the Issuing Officer.")
    
            document.add_page_break()


            #forwarding 2222222222-------------------------------------------------
            section.page_width = Cm(22)  # set the page width to 21 centimeters
            section.page_height = Cm(29.9)  # set the page width to 21 centimeters
            section.left_margin = Cm(3)  # 1.5 cm
            section.right_margin = Cm(2)
            # Define a custom style
            custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            custom_style.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.name = 'Arial Narrow'
            custom_style2.font.size = Pt(11)
            mystyle += 1
    
            custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(11)
            custom_style3.font.name = 'Arial Narrow'
            custom_style3.font.bold = True
            mystyle += 1

            current_date = datetime.datetime.now().date()
            self_date = ''
            if current_date:
                self_date = datetime.datetime.strptime(str(current_date), '%Y-%m-%d').strftime('%d/%m/%y')
            paragraph2 = document.add_paragraph('\n\n\n\nRef: ' + file_obj.name + '\nDate: ' + str(self_date) + '\n')
            paragraph2.style = custom_style2
    
            paragraph4 = document.add_paragraph("The President \nBASIS \nBDBL Bhabon (5th Floor-West), \n12 Kawran Bazar, Dhaka-1215 \n")
            paragraph4.style = custom_style2
    
            paragraph5 = document.add_paragraph("Subject: Request to Issue BASIS Certificate for Cash Assistance.\n")
            paragraph5.style = custom_style3
    
            paragraph6 = document.add_paragraph("Dear Sir,")
            paragraph6.style = custom_style2
            paragraph6.paragraph_format.space_after = Pt(0)
    
            contract_number = ''
            if file_obj.contract_number:
                contract_number = " and "+ file_obj.contract_number
    
            paragraph6 = document.add_paragraph("With reference to the above-mentioned subject, we would like to draw your kind attention to the fact that we are going to draw cash subsidy against our following Invoice No: " + inv_number + contract_number + ".\n")
            paragraph6.style = custom_style2
            paragraph6.style = custom_style2
    
            paragraph6 = document.add_paragraph("There is a checklist as following:  ")
            paragraph6.style = custom_style2
    
            # Create a list of items
            items = ["Forwarding Letter Addressing BASIS president", "Export Agreement", "Commercial Invoice Related to Export", "Format -A Certificate of Authorized Dealer Issued by Bank",
                     "ICT form-c", "Company's Undertaking", "Form-Ka (Bangladesh Bank)", "Form-Kha (Bangladesh Bank)", "Pay order"]
    
            left_indent = Inches(.8)  # Adjust the left indentation as needed
    
            for item in items:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.left_indent = left_indent
    
                run = paragraph.add_run(item)
                run.font.name = 'Arial Narrow'
                paragraph.paragraph_format.line_spacing = Pt(12)
    
            paragraph8 = document.add_paragraph("\nYour kind co-operation in this regard will be highly appreciated.\n\nThanking You.\n\nYours Truly\n\n\n\n\nMasudur Rahman \nSenior Manager, Finance, Accounts & Admin \nBrain Station 23 Ltd. \nPlot-2, (5th Floor), Bir Uttam Ak Khandaker Road \nMohakhali C/A, Dhaka-1212")
            paragraph8.style = custom_style2
            paragraph8.paragraph_format.space_before = Pt(0)
            paragraph8.paragraph_format.space_after = Pt(0)
            document.add_page_break()
            # file_name = '%s_flbs' % (datetime.datetime.now())
    
            # ka 333333333333333 -------------------------------------------
            section.page_width = Cm(22)  # set the page width to 21 centimeters
            section.page_height = Cm(29.9)  # set the page width to 21 centimeters
            section.left_margin = Cm(2)  # 1.5 cm
            section.right_margin = Cm(2)
    
            self_date = ''
            if file_obj.date:
                self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')
    
            # end bangla font maker ====================================
            # Define a custom style
            custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            # custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mystyle += 1
    
            custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(11)
            mystyle += 1
    
            custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(11)
            custom_style3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            mystyle += 1
    
            custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.bold = True
            custom_style4.font.size = Pt(9)
            custom_style4.font.name = 'Arial Narrow'
            mystyle += 1
    
            tab_total = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            tab_total.font.bold = True
            tab_total.font.size = Pt(9)
            tab_total.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_styleextra = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_styleextra.font.size = Pt(9)
            custom_styleextra.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_style5 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style5.font.size = Pt(9)
            custom_style5.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_stylebl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebl.font.size = Pt(9)
            custom_stylebl.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_na = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_na.font.size = Pt(9)
            custom_na.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_stylebold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebold.font.size = Pt(11)
            custom_stylebold.font.name = 'SutonnyOMJ'
            custom_stylebold.font.bold = True
            mystyle += 1
    
            custom_styleubl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_styleubl.font.size = Pt(11)
            custom_styleubl.font.name = 'SutonnyOMJ'
            custom_styleubl.font.underline = True
            custom_styleubl.font.bold = True
            mystyle += 1
    
            document.styles['Normal'].font.size = Pt(8)
            paragraph = document.add_paragraph()
            paragraph.add_run("অনুচ্ছেদ ০৫(খ) এফই সার্কুলার নং-০৩/২০১৮ দ্রষ্টব্য)")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run(' ' * 7 + "ফরম-‘ক’")
            paragraph.style = custom_stylebl
    
            # Add a paragraph and apply the custom style
            paragraph1 = document.add_paragraph('বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও')
            paragraph1.style = custom_stylebold
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph1.paragraph_format.space_before = Pt(0)
            paragraph1.paragraph_format.space_after = Pt(0)
            paragraph1 = document.add_paragraph('হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকির জন্য আবেদনপত্র ')
            paragraph1.style = custom_styleubl
            paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph1.paragraph_format.space_before = Pt(0)
    
            paragraph2 = document.add_paragraph("(ক) আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানাঃ ব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী বাণিজ্যিক এলাকা, ঢাকা - ১২১২। রপ্তানি নিবন্ধন সনদপত্র (ইআরসি) নম্বরঃ ২৬০৩২৬২১০৬৬৬৪২০")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)
    
            con_no = ''
            if file_obj.contract_number:
                con_no = file_obj.contract_number.replace('Contract No.', '').replace('Purchase Order No.', '')
            paragraph2 = document.add_paragraph("(খ) রপ্তানি চুক্তিপত্র নম্বরঃ"+ con_no + " \nতারিখঃ " + str(file_obj.contract_date_str) + "\nমূল্যঃ " +file_obj.contract_price_str + "\n(পাঠ্যযোগ সত্যায়িত কপি দাখিল করতে হবে)" )
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)
            paragraph2 = document.add_paragraph("(গ) রপ্তানিকৃত সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের স্থানীয় সংগ্রহসূত্র, পরিমাণ ও মূল্যঃ" )
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)
            # Set the paragraph alignment to center
    
            # table1 ------------------------------
            table5 = document.add_table(rows=3, cols=3)
            table5.alignment = 2
            table5.left_indent = Inches(100)
            table5.style = 'TableGrid'
            table5.autofit = False  # Disable automatic column width adjustment
              # Set the width of the table to 6 inches (adjust the value according to your desired width)
    
            cell = table5.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('সরবরাহকারীর নাম ও ঠিকানা')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_stylebl
    
            cell = table5.cell(0, 1)
            p1 = cell.paragraphs[0]
            p1.add_run('পরিমাণ')
            p1.style = custom_stylebl
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table5.cell(0, 2)
            p2 = cell.paragraphs[0]
            p2.add_run('মূল্য')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table5.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table5.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table5.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table5.cell(2, 0)
            p7 = cell.paragraphs[0]
            p7.add_run('N/A')
            p7.style = custom_na
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table5.cell(2, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table5.cell(2, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('N/A')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # table 2 ----------------------------
            paragraph2 = document.add_paragraph("(রপ্তানিকৃত সেবা/পণ্যের বর্ণনা, মূল্য ও সংগ্রহসূত্রের বিষয়ে সেবা/পণ্য সংশ্লিষ্ট এসোসিয়েশন এর প্রত্যয়নপত্র দাখিল করতে হবে) \n(ঘ) রপ্তানিকৃত সেবা/পণ্য উৎপাদনে ব্যবহৃত আমাদানিকৃত সেবা/উপকরণাদিঃ" )
            paragraph2.style = custom_stylebl
            # paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(12)
    
            # Set the paragraph alignment to center
            table1 = document.add_table(rows=3, cols=4)
            table1.autofit = False
            table1.width = Inches(6)
            table1.alignment = 2
            table1.style = 'TableGrid'
    
            cell = table1.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('সরবরাহকারীর নাম ও ঠিকানা')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table1.columns[0].width = Inches(.6)
            p3.style = custom_stylebl
    
            cell = table1.cell(0, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('সেবা/পণ্যের নাম ও পরিমাণ ')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table1.columns[1].width = Inches(.6)
            p3.style = custom_stylebl
    
            cell = table1.cell(0, 2)
            p2 = cell.paragraphs[0]
            p2.add_run('ঋণপত্র/ব্যাংক টু ব্যাংক ঋণপত্র/ডকুমেন্টরী কালেকশন/টিটি রেমিটেন্স নম্বর, তারিখ')
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table1.columns[2].width = Inches(5.2)
            p2.style = custom_stylebl
    
            cell = table1.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('মূল্য')
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table1.columns[3].width = Inches(1)
            p2.style = custom_stylebl
    
            cell = table1.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table1.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table1.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table1.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('৪')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table1.cell(2, 0)
            p7 = cell.paragraphs[0]
            p7.add_run('N/A')
            p7.style = custom_na
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table1.cell(2, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table1.cell(2, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('N/A')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table1.cell(2, 3)
            p9 = cell.paragraphs[0]
            p9.add_run('N/A')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # table 3 ----------------------
            paragraph2 = document.add_paragraph("(৩) নং কলামের ঋণপত্রে পাঠযোগ্য সত্যায়িত কপি দাখিল করতে হবে। সেবা আমাদানির ক্ষেত্রে যথাযথ পদ্ধতি অনুসরণ করা হয়েছে মর্মে অনুমোদিত ডিলার শাখাকে নিশ্চিত হতে হবে। উৎপাদন প্রক্রিয়ায় ব্যবহৃত উপকরণাদির জন্য শুল্ক বন্ড সুবিধা ভোগ করা হয়নি/ডিউটি ড্র-ব্যাংক সুবিধা গ্রহণ করা হয়নি ও ভবিষ্যতে আবেদনও করা হবে না মর্মে রপ্তানিকরাকের ঘোষণাপত্র দাখিল করতে হবে।) \n(ঙ) রপ্তানি চালানের বিবরণঃ" )
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(12)
    
            table = document.add_table(rows=2, cols=8)
            table.alignment = 2  # Left alignment
            # table.left_indent = Inches(5)
            table.style = 'TableGrid'
    
            column_width = Inches(.2)
            table.columns[0].width = column_width
            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('   পণ্যের বর্ণনা   ')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl
            # table.columns[0].width = Inches(.3)
            # cell.merge(table.cell(0, 1))
    
            column_width = Inches(1.8)
            table.columns[1].width = column_width
            cell = table.cell(0, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('পরিমাণ')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl
            # table.columns[1].width = Inches(1.69)
    
            column_width = Inches(.1)
            table.columns[2].width = column_width
            cell = table.cell(0, 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('আমদানিকারকের দেশের নাম')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl
            # table.columns[2].width = Inches(.8)
    
            table.columns[3].width = Inches(.6)
            cell = table.cell(0, 3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('ইনভয়েস মূল্য (বৈদেশিক মুদ্রায়)')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl
    
            table.columns[4].width = Inches(.6)
            cell = table.cell(0, 4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('জাহাজীকরণ/রপ্তানির তারিখ')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl
    
            column_width = Inches(.2)
            table.columns[5].width = column_width
            cell = table.cell(0, 5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell.paragraphs[0]
            p10.add_run('   ইএক্সপি নম্বর*   ')
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10.style = custom_stylebl
    
            table.columns[6].width = Inches(.8)
            table.columns[7].width = Inches(.3)
            cell_1_0 = table.cell(0, 6)
            cell_1_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p10 = cell_1_0.paragraphs[0]
            p10.add_run('বৈদেশিক মুদ্রায় প্রত্যাবাসিত রপ্তানিমূল্য ও প্রত্যাবাসনের তারিখ')
            cell_2_0 = table.cell(0, 7)
            cell_1_0.merge(cell_2_0)
            p10.style = custom_stylebl
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            left_indent = Inches(2)  # Adjust the indent as needed
    
            for row in table.rows:
                for cell in row.cells:
                    cell.left_indent = left_indent
    
            # row 2 ----------------------------
            cell = table.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('৪')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 4)
            p6 = cell.paragraphs[0]
            p6.add_run('৫')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 5)
            p6 = cell.paragraphs[0]
            p6.add_run('৬')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 6)
            p10 = cell_1_1.paragraphs[0]
            p10.add_run('৭')
            p6.style = custom_stylebl
            cell_2_1 = table.cell(1, 7)
            cell_1_1.merge(cell_2_1)
            p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # # add a dynamic row to the table
            # inv_date = ''
            inv_amount = 0
            incentive_rate_fc = 0
            swift_amt = 0
            incentive_amt_fc = 0
            encashment_amt_bdt = 0
            total_incentive_rate_fc = 0
            sl = 0
    
            inv_ids = []
            dev_total1 = 0
            hour_total1 = 0
            a = []
            c = []
            for line in file_obj.invoice_line_ids:
                a.append(line.od_sight_rate)
                c.append(line.encashment_rate_bdt)
                if line.invoice_id.id not in inv_ids:
                    dev_total1 += sum(
                        r.quantity if r.quantity_type == '1' else 0 for r in line.invoice_id.invoice_line_ids)
                    hour_total1 += sum(
                        r.quantity if r.quantity_type == '0' else 0 for r in line.invoice_id.invoice_line_ids)
                inv_ids.append(line.invoice_id.id)
    
            if a:
                b = min(a)
                od_s_r = b
            else:
                od_s_r = 0
    
            if c:
                d = min(c)
                encashment_rate_bdt = d
            else:
                encashment_rate_bdt = 0
    
            hour_total = float(format(hour_total1, '.2f'))
            dev_total = float(format(dev_total1, '.2f'))
    
            if hour_total == int(hour_total):
                hour_total_num = "{:.0f}".format(hour_total)
            else:
                hour_total_num = "{:.2f}".format(hour_total)
    
            if dev_total == int(dev_total):
                dev_total_num = "{:.0f}".format(dev_total)
            else:
                dev_total_num = "{:.2f}".format(dev_total)
    
            all_invoice_qty_str = ''
            if dev_total > 0:
                all_invoice_qty_str += str('ITES/ ') + str(dev_total_num) if not all_invoice_qty_str else ' & ' + str(
                    'ITES/ ') + str(
                    dev_total_num)
            if hour_total > 0:
                all_invoice_qty_str += str(hour_total_num) + str(' HRS') if not all_invoice_qty_str else ' & ' + str(
                    hour_total_num) + str(' HRS')
    
            currency_symbol = ''
            if file_obj.fc_currency_id.symbol:
                currency_symbol = file_obj.fc_currency_id.symbol
            country_name = ''
            if file_obj.partner_id.country_id:
                country_name = file_obj.partner_id.country_id.name
            row_len = len(file_obj.invoice_line_ids)
    
            from itertools import groupby
            # group the invoice_line_ids by invoice_id.ref
            groups = groupby(sorted(file_obj.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                             key=lambda x: x.invoice_id.ref)
            # iterate over each group and add a new row for each unique invoice_id.ref
            row = 2
    
            for invoice_ref, group in groups:
                inv_obj = file_obj.env['cash.incentive.invoice'].search(
                    [('invoice_id.ref', '=', invoice_ref), ('head_id', '=', file_obj.id)], limit=1)
                # usd_price = sum(inv_obj.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
                usd_price = inv_obj.invoice_id.invoice_total_fc
                incentive_rate_fc = usd_price / file_obj.incentive_rate_fc if file_obj.incentive_rate_fc else 0
                total_incentive_rate_fc += incentive_rate_fc
                inv_amount += inv_obj.invoice_amt
                encashment_amt_bdt += inv_obj.encashment_amt_bdt
                invoice_date = ''
                swift_message_date = ''
                if inv_obj.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime(
                        '%d-%b-%y')
                if inv_obj.swift_message_id.date:
                    swift_message_date = datetime.datetime.strptime(str(inv_obj.swift_message_id.date),
                                                                    '%Y-%m-%d').strftime('%d-%b-%y')
    
                # getting hours and developers ===========
                l_dev_total = sum(
                    r.quantity if r.quantity_type == '1' else 0 for r in inv_obj.invoice_id.invoice_line_ids)
                l_hour_total = sum(
                    r.quantity if r.quantity_type == '0' else 0 for r in inv_obj.invoice_id.invoice_line_ids)
                l_all_invoice_qty_str = ''
                if l_dev_total > 0:
                    if l_dev_total == int(l_dev_total):
                        l_dev_total = "{:.0f}".format(l_dev_total)
                    else:
                        l_dev_total = "{:.2f}".format(l_dev_total)
                    l_all_invoice_qty_str += str('ITES/ ') + str(
                        l_dev_total) if not l_all_invoice_qty_str else ' & ' + str(
                        'ITES/ ') + str(l_dev_total)
                if l_hour_total > 0:
                    if l_hour_total == int(l_hour_total):
                        l_hour_total = "{:.0f}".format(l_hour_total)
                    else:
                        l_hour_total = "{:.2f}".format(l_hour_total)
                    l_all_invoice_qty_str += str(l_hour_total) + str(
                        ' HRS') if not l_all_invoice_qty_str else ' & ' + str(
                        l_hour_total) + str(' HRS')
                sl += 1
                gr = 0
                for rec in group:
                    swift_amt += rec.swift_amt
                    incentive_amt_fc += rec.incentive_amt_fc
                    # if rec.swift_message_id.encashment_rate_bdt:
                    #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
                    gr += 1
                    new_row = table.add_row()
                    new_row.cells[6].text = str(currency_symbol) + ' ' + str("{:,.2f}".format(rec.swift_amt))
                    new_row.cells[7].text = str(swift_message_date)
    
                    first_cell = new_row.cells[3]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style5
    
                    first_cell1 = new_row.cells[6]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = custom_style5
    
                    first_cell1 = new_row.cells[0]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[1]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[2]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[4]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[5]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
                    first_cell1 = new_row.cells[7]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
    
                cell1_1_0 = table.cell(row, 1)
                p3 = cell1_1_0.paragraphs[0]
                p3.add_run(l_all_invoice_qty_str)
                cell1_2_0 = table.cell((row + gr) - 1, 1)
                cell1_1_0.merge(cell1_2_0)
    
                cell_4_0 = table.cell(row, 3)
                p3 = cell_4_0.paragraphs[0]
                p3.add_run(str(currency_symbol)+ ' ' + str("{:,.2f}".format(inv_obj.invoice_amt)))
                cell_3_0 = table.cell((row + gr) - 1, 3)
                cell_4_0.merge(cell_3_0)
    
                cell_5_0 = table.cell(row, 4)
                p3 = cell_5_0.paragraphs[0]
                p3.add_run(invoice_date)
                cell_6_0 = table.cell((row + gr) - 1, 4)
                cell_5_0.merge(cell_6_0)
                row = gr + row
    
            # #column merge -----------------------
            column_width = Inches(.3)
            # table.columns[0].width = column_width
            cell_1_2 = table.cell(2, 0)
            cell_1_2.width = column_width
            p3 = cell_1_2.paragraphs[0]
            p3.add_run('Software Development')
            cell_2_2 = table.cell(1+row_len, 0)
            cell_1_2.merge(cell_2_2)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_styleextra
    
            cell_1_2 = table.cell(2, 2)
            p3 = cell_1_2.paragraphs[0]
            p3.add_run(country_name)
            cell_2_2 = table.cell(1+row_len, 2)
            cell_1_2.merge(cell_2_2)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_styleextra
    
            cell_1_2 = table.cell(2, 5)
            p3 = cell_1_2.paragraphs[0]
            p3.add_run('N/A')
            cell_2_2 = table.cell(1+row_len, 5)
            cell_1_2.merge(cell_2_2)
            p3.style = custom_styleextra
    
            # total ---------------------------
            total_row = table.add_row()
            total_row.cells[0].text = 'Total'
            total_row.cells[1].text = all_invoice_qty_str
    
            # inv_amount_a = "{:.2f}".format(inv_amount).rstrip('0').rstrip('.') + ('0' if inv_amount % 1 else '')
            inv_amount_a = str("{:,.2f}".format(inv_amount))
            total_row.cells[3].text = str(currency_symbol)+ ' ' + inv_amount_a
    
            # total_row.cells[3].text = str(currency_symbol)+ ' ' + str("{:,}".format(round(inv_amount, 2)))
            # formatted_number = "{:.2f}".format(inv_amount).rstrip('0').rstrip('.') + ('0' if inv_amount % 1 else '')
    
            # swift_amt_a = "{:.2f}".format(swift_amt).rstrip('0').rstrip('.') + ('0' if swift_amt % 1 else '')
            swift_amt_a = str("{:,.2f}".format(swift_amt))
            total_row.cells[6].text = str(currency_symbol)+ ' ' + swift_amt_a
    
            first_cell = total_row.cells[0]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph21.style = tab_total
    
            first_cell = total_row.cells[1]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph21.style = tab_total
    
            first_cell = total_row.cells[3]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = tab_total
    
            first_cell1 = total_row.cells[6]
            paragraph22 = first_cell1.paragraphs[0]
            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph22.style = tab_total
    
            # table 4 ----------------------------------
            paragraph2 = document.add_paragraph("*দৃশ্যমান আকারে পণ্য রপ্তানির ক্ষেত্রে প্রযোজ্য \n (কমার্শিয়াল ইনভয়েস, প্যাকিং লিষ্ট এবং জাহাজীকরণের প্রমাণ স্বরূপ পরিবহন কর্তৃপক্ষ ইস্যুকৃত এবং প্রত্যয়নকৃত বিল অব লোডিং/এয়ারওয়ে বিল, বিল অব এক্সপোর্ট (শুল্ক কর্তৃপক্ষ কর্তৃক ইস্যুকৃত ও পীক্ষিত এবং on-hand হওয়ার স্বপক্ষে পরিবহন কর্তৃপক্ষ প্রত্যয়নকৃত) এর পূর্ণাঙ্গ সেট ইত্যাদির সত্যায়িত পাঠযোগ্য কপি এবং রপ্তানিমূল্য প্রত্যাবাসন সনদপত্র দাখিল করতে হবে। তবে অদৃশ্যকারে সেবা রপ্তানির ক্ষেত্রে জাহাজীকরণের দলিল ও বিল অব এক্সপোর্ট  আবশ্যকতা থাকবে না।) \n(চ) ভর্তুকির আবেদনকৃত অংকঃ")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(12)
    
            # Set the paragraph alignment to center
            table = document.add_table(rows=3, cols=4)
            table.style = 'TableGrid'
            table.autofit = False
            table.width = Inches(6)
            table.alignment = 2
    
            cell = table.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('প্রত্যাবাসিত রপ্তানিমূল্য')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[0].width = Inches(1.2)
    
            cell = table.cell(0, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার পরিমাণ')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[1].width = Inches(1.7)
    
            cell = table.cell(0, 2)
            p2 = cell.paragraphs[0]
            p2.add_run('বৈদেশিক মুদ্রায় পরিশোধ্য কমিশন ইন্সুরেন্স ইত্যাদি (যদি থাকে)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[2].width = Inches(2.58)
    
            cell = table.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('নীট এফওবি রপ্তানিমূল্য ১-(২+৩)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('৪')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 0)
            p7 = cell.paragraphs[0]
            p7.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p7.style = custom_style4
    
            cell = table.cell(2, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('N/A')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 3)
            p9 = cell.paragraphs[0]
            p9.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p9.style = custom_style4
    
            # table 5 ----------------------------------
            paragraph2 = document.add_paragraph("(প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার উল্লেখ সম্বলিত ফ্রেইট সার্টিফিকেটের সত্যায়িত কপি দাখিল করতে হবে)")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(12)
    
            vortuki_swift_amt = 0
            if file_obj.incentive_rate_fc > 0:
                vortuki_swift_amt = swift_amt / file_obj.incentive_rate_fc
    
            # Set the paragraph alignment to center
            table = document.add_table(rows=4, cols=4)
            table.style = 'TableGrid'
            table.autofit = False
            table.width = Inches(6)
            table.alignment = 2
    
            cell = table.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('রপ্তানি সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের মূল্য')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(0, 1)
            cell.merge(cell2)
    
            cell = table.cell(0, 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p2 = cell.paragraphs[0]
            p2.add_run('স্থানীয় মূল্য সংযোজনের হার')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(1, 2)
            cell.merge(cell2)
    
            cell = table.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('প্রাপ্য ভর্তুকি* ৪x১০%')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('দেশীয় পণ্য/সেবা')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('আমদানিকৃত সেবা/পণ্য')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # cell = table.cell(1, 2)
            # p6 = cell.paragraphs[0]
            # p6.add_run('')
            # p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('')
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('৫')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('৬')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৭')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 3)
            p6 = cell.paragraphs[0]
            p6.add_run(' ৮')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 0)
            p7 = cell.paragraphs[0]
            p7.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p7.style = custom_style4
    
            cell = table.cell(3, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('100%')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            incentive_amt_fc_a = str("{:,.2f}".format(incentive_amt_fc))
            cell = table.cell(3, 3)
            p9 = cell.paragraphs[0]
            p9.add_run(str(currency_symbol)+ ' ' + incentive_amt_fc_a)
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p9.style = custom_style4
    
            # signature ------------------------------------------
            paragraph2 = document.add_paragraph("(*৭ নম্বর কলামের হার আলোচ্য সার্কুলারের ৪ নম্বর অনুচ্ছেদের সাথে সামঞ্জস্যতার ক্ষেত্রে ভর্তুকি প্রাপ্য হবে।) \nএ মর্মে অঙ্গীকার করা হচ্ছে যে, আমাদের নিজস্ব কারখানায়/প্রতিষ্ঠানে তৈরী/উৎপাদিত সফটওয়্যার/আইটিইএস/হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকির জন্য আবেদন করা হলো। এ আবেদনপত্রে প্রদত্ত সকল তথ্য/ঘোষণা সম্পূর্ণ ও সঠিক। যদি পরবর্তীতে কোন ভুল/অসত্য তথ্য/প্রতারণা/জালিয়াতি উদঘাটিত হয় তবে গৃহীত ভর্তুকির সমুদয় অর্থ বা এর অংশবিশেষ আমার/আমাদের নিকট হইতে এবং/অথবা আমার/আমাদের ব্যাংক হিসাব থেকে আদায়/ফেরত নেয়া যাবে। \n\nতারিখঃ..................................... ")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)
    
            paragraph2 = document.add_paragraph("........................................ \nআবেদনকারী প্রতিষ্ঠানের স্বত্তাধিকারী/\n ক্ষমতাপ্রাপ্ত  কর্মকর্তার স্বাক্ষর ও পদবী")
            paragraph2.style = custom_stylebl
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)
    
            # table 6 ----------------------------------
            document.add_page_break()
            paragraph2 = document.add_paragraph("(ছ) ভর্তুকি প্রদানকারী ব্যাংক শাখা কর্তৃক পূরণীয়ঃ \t\t\t\t\t\t\t\t\t" + ' ' * 11 + "(বৈদেশিক মুদ্রায়)")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.line_spacing = Pt(10)
    
            # Set the paragraph alignment to center
            table = document.add_table(rows=3, cols=4)
            table.style = 'TableGrid'
            table.autofit = False
            table.width = Inches(6)
            table.alignment = 2
    
            cell = table.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('প্রত্যাবাসিত রপ্তানিমূল্য')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[0].width = Inches(1.21)
    
            cell = table.cell(0, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়ার পরিমাণ')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[1].width = Inches(1.7)
    
            cell = table.cell(0, 2)
            p2 = cell.paragraphs[0]
            p2.add_run('বৈদেশিক মুদ্রায় পরিশোধ কমিশন, ইন্সুরেন্স ইত্যাদি (যদি থাকে)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[2].width = Inches(2.6)
    
            cell = table.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('নীট এফওবি রপ্তানিমূল্য ১-(২+৩)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.columns[3].width = Inches(1.8)
    
            cell = table.cell(1, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('১')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('২')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৩')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 3)
            p6 = cell.paragraphs[0]
            p6.add_run('৪')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 0)
            p4 = cell.paragraphs[0]
            p4.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.style = custom_style4
    
            cell = table.cell(2, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('N/A')
            p5.style = custom_na
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('N/A')
            p6.style = custom_na
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 3)
            p6 = cell.paragraphs[0]
            p6.add_run(str(currency_symbol)+ ' '+ swift_amt_a)
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p6.style = custom_style4
    
            # table 7 ----------------------------------
            if file_obj.od_sight_rate:
                od_sight_rate = file_obj.od_sight_rate
            else:
                if od_s_r:
                    od_sight_rate = od_s_r
                else:
                    od_sight_rate = encashment_rate_bdt
    
            rate_fc = swift_amt / file_obj.incentive_rate_fc
            # encashment_final_amt = encashment_amt_bdt / file_obj.incentive_rate_fc
            # encashment_final_amt = rate_fc / file_obj.incentive_rate_fc
            rate_fc_a = "{:,.2f}".format(rate_fc)
            # od_sight_rate = "{:.2f}".format(file_obj.od_sight_rate)
            encashment_final_amt_a = round(rate_fc, 2) * od_sight_rate
            # encashment_final_amt = "{:,.0f}".format(round(encashment_final_amt_a))
            encashment_final_amt = file_obj.custom_round(encashment_final_amt_a)
            usd_bdt_amount = currency + ' ' + str(rate_fc_a) + '\n' + '@' + str("{:,}".format(round(od_sight_rate, 2))) + '=' + '\n' + 'BDT ' + encashment_final_amt
            # amount in word
            amount_in_word_str = ''
            if encashment_final_amt_a:
                # amount_in_word = num2words(round(encashment_final_amt_a))
                amount_in_word = file_obj.num2words_fun(round(encashment_final_amt_a))
                # upper case function call
                amount_in_word_str = file_obj.upper_case(amount_in_word)
    
            paragraph2 = document.add_paragraph("(প্রযোজ্য ক্ষেত্রে জাহাজ ভাড়া উল্লেখ সম্বলিত ফ্রেইট সার্টিফিকেটের সত্যায়িত কপি দাখিল করতে হবে)")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.line_spacing = Pt(12)
    
            # Set the paragraph alignment to center
            table = document.add_table(rows=4, cols=4)
            table.style = 'TableGrid'
            table.autofit = False
            table.width = Inches(6)
            table.alignment = 2
    
            cell = table.cell(0, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('রপ্তানি সেবা/পণ্য উৎপাদনে ব্যবহৃত সেবা/পণ্যের মূল্য')
            p3.style = custom_stylebl
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(0, 1)
            cell.merge(cell2)
    
            cell = table.cell(0, 2)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p2 = cell.paragraphs[0]
            p2.add_run('স্থানীয় মূল্য সংযোজনের হার [(৪-৬)/৪]x ১০০')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(1, 2)
            cell.merge(cell2)
    
            cell = table.cell(0, 3)
            p2 = cell.paragraphs[0]
            p2.add_run('পরিশোধ ভর্তুকির পরিমাণ (টাকায়)*(রপ্তানিমূল্য প্রত্যাবাসনের তারিখে সংশ্লিষ্ট ৪x ১০% বৈদেশিক মুদ্রার ওডি সাইট)')
            p2.style = custom_stylebl
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell2 = table.cell(1, 3)
            cell.merge(cell2)
    
            cell = table.cell(1, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p4 = cell.paragraphs[0]
            p4.add_run('দেশীয় পণ্য/সেবা')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(1, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p5 = cell.paragraphs[0]
            p5.add_run('আমদানিকৃত সেবা/পণ্য বৈদেশিক মুদ্রায়')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # cell = table.cell(1, 3)
            # p6 = cell.paragraphs[0]
            # p6.add_run('')
            # p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 0)
            p4 = cell.paragraphs[0]
            p4.add_run('৫')
            p4.style = custom_stylebl
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 1)
            p5 = cell.paragraphs[0]
            p5.add_run('৬')
            p5.style = custom_stylebl
            p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 2)
            p6 = cell.paragraphs[0]
            p6.add_run('৭')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 3)
            p6 = cell.paragraphs[0]
            p6.add_run(' ৮')
            p6.style = custom_stylebl
            p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 0)
            p7 = cell.paragraphs[0]
            p7.add_run(str(currency_symbol) + ' ' + swift_amt_a)
            p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p7.style = custom_style4
    
            cell = table.cell(3, 1)
            p8 = cell.paragraphs[0]
            p8.add_run('N/A')
            p8.style = custom_na
            p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 2)
            p9 = cell.paragraphs[0]
            p9.add_run('100%')
            p9.style = custom_na
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 3)
            p9 = cell.paragraphs[0]
            p9.add_run(usd_bdt_amount)
            p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p9.style = custom_style4
    
            # conclusion ----------------
            paragraph2 = document.add_paragraph("(*৭ নম্বর কলামের হার আলোচ্য সার্কুলারের ৪ নম্বর অনুচ্ছেদের সাথে সামঞ্জস্যতার ক্ষেত্রে ভর্তুকি প্রাপ্য হবে) \nভর্তুকি পরিমাণ: " + amount_in_word_str + ' Only')
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.line_spacing = Pt(12)
    
            paragraph2 = document.add_paragraph("\n\n\n\n\n\n\nপরিশোধের তারিখঃ --------------------------")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.line_spacing = Pt(10)
    
            paragraph2 = document.add_paragraph("----------------------------------------- \nভর্তুকির অনুমোদনের ক্ষমতাপ্রাপ্ত ব্যাংক  \nকর্মকর্তার স্বাক্ষর, নাম ও পদবী")
            paragraph2.style = custom_stylebl
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.line_spacing = Pt(10)
            document.add_page_break()

            # file_name = '%s_ka' % (datetime.datetime.now())
            # KHA 444444444444444 -------------------------------------------
            # modify the page setup
            section.page_width = Cm(22)  # set the page width to 21 centimeters
            section.page_height = Cm(29.9)  # set the page width to 21 centimeters
            section.left_margin = Cm(2.5)  # 1.5 cm
            section.right_margin = Cm(2)  # 1.5 cm
            self_date = ''
            if file_obj.date:
                self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')
    
            # document.add_paragraph('কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।')
            # en_font = my_stylew.font
            # document.add_paragraph('some text')
    
            custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            # custom_style.font.bold = True
            custom_style.font.size = Pt(8)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            custom_style.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_style_table = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style_table.font.size = Pt(10)
            custom_style_table.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_stylebsmallbold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebsmallbold.font.size = Pt(9)
            custom_stylebsmallbold.font.name = 'SutonnyOMJ'
            custom_stylebsmallbold.font.bold = True
            mystyle += 1
    
            custom_stylebsmall = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebsmall.font.size = Pt(9)
            custom_stylebsmall.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_stylebl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebl.font.size = Pt(9)
            custom_stylebl.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_sonar_bangla = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_sonar_bangla.font.size = Pt(9)
            custom_sonar_bangla.font.name = 'Shonar Bangla'
            mystyle += 1
    
            custom_tableh = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_tableh.font.size = Pt(9)
            custom_tableh.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_table = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_table.font.size = Pt(9)
            custom_table.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_stylebold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebold.font.size = Pt(11)
            custom_stylebold.font.name = 'SutonnyOMJ'
            custom_stylebold.font.bold = True
            mystyle += 1
    
            custom_styleubl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_styleubl.font.size = Pt(11)
            custom_styleubl.font.name = 'SutonnyOMJ'
            custom_styleubl.font.underline = True
            custom_styleubl.font.bold = True
            mystyle += 1
    
            document.styles['Normal'].font.size = Pt(8)
            paragraph = document.add_paragraph()
            paragraph.add_run("\n(অনুচ্ছেদ ০৬ (ক), এফই সার্কুলার নং-০৩/২০১৮ দ্রষ্টব্য)")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.style = custom_stylebl
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.add_run("ফরম- 'খ'")
            paragraph = document.add_paragraph()
            paragraph.add_run("Ref: " + str(file_obj.form_kha_ref_code))
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run(" Date: " + self_date)
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.style = custom_stylebl
    
            # paragraph.style = custom_style
    
            # Define a custom style
            custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(8)
            mystyle += 1
    
            custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(8)
            custom_style3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            mystyle += 1
    
            custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.bold = True
            custom_style4.font.size = Pt(10)
            custom_style4.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_total = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_total.font.bold = True
            custom_total.font.size = Pt(9)
            custom_total.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_style5 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style5.font.size = Pt(8)
            custom_style5.font.name = 'Arial Narrow'
            mystyle += 1
    
            # Add a paragraph and apply the custom style
            paragraph = document.add_paragraph('বেসিস প্রদেয় প্রত্যয়ন সনদপত্র \nবাংলাদেশ হতে সফটওয়্যার, আইটিইএস (Information Technology Enabled Services)')
            paragraph.style = custom_stylebold
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph = document.add_paragraph('ও হার্ডওয়্যার রপ্তানির বিপরীতে ভর্তুকি প্রাপ্তির প্রত্যয়ন সনদপত্র।')
            paragraph.style = custom_styleubl
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
    
            customer = '' if not file_obj.swift_customer_name else str(file_obj.swift_customer_name) + ', ' + str(
                file_obj.customer_address)
            con_no = ''
            if file_obj.contract_number:
                con_no = file_obj.contract_number.replace('Contract No.', '').replace('Purchase Order No.', '')

            import textwrap
            width = 140
    
            paragraph2 = document.add_paragraph("১।\tআবদেনকারীর নাম, ইআরসি নম্বর ও ঠিকানাঃ Brain Station 23 Limited, Plot No.02 (8th Floor), Bir Uttam A.K. Khandakar Road,"
                                                " \tMohakhali C/A, Dhaka-1212.  ইআরসি নম্বরঃ 260326210666420")
            paragraph2.style = custom_stylebl
            # paragraph2.add_run("\t")
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(0)
    
            long_text1 = "২।\t" + ' ' * 5 + "রপ্তানি ঋণপত্র/চুক্তিপত্র নম্বরঃ" + con_no + '\t তারিখঃ ' + file_obj.contract_date_str + '\t মূল্যঃ ' +file_obj.contract_price_str
            wrapped_lines1 = textwrap.wrap(long_text1, width=120)
            formatted_lines1 = [line + '\t' for line in wrapped_lines1]
            formatted_text1 = '\n\t'.join(formatted_lines1)
            paragraph2 = document.add_paragraph(formatted_text1)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            long_text2 = "৩।\t" + ' ' * 5 + "বিদেশি ক্রেতার নাম ও ঠিকানাঃ " + customer
            # long_text2 = "৩।\t" + ' ' * 5 + "বিদেশি ক্রেতার নাম ও ঠিকানাঃ ABBVIE INC AP, DEPT V312 AP34-2 1 N WAUKEGAN RD, NORTH CHICAGO"
            wrapped_lines2 = textwrap.wrap(long_text2, width=123)
            formatted_lines2 = [line + '\n\t' for line in wrapped_lines2]
            # formatted_text2 = ''.join(formatted_lines2)
            formatted_text2 = ''.join(formatted_lines2).rstrip('\n\t')
            paragraph2 = document.add_paragraph(formatted_text2)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            long_text1 = "৪।\t" + ' ' * 5 + "বিদেশি ক্রেতার ব্যাংকের নাম ও ঠিকানাঃ " + remiter_address
            wrapped_lines1 = textwrap.wrap(long_text1, width=120)
            formatted_lines1 = [line + '\t' for line in wrapped_lines1]
            formatted_text1 = '\n\t'.join(formatted_lines1)
            paragraph2 = document.add_paragraph(formatted_text1)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph('৫।')
            paragraph2.style = custom_stylebl
    
            table = document.add_table(rows=1, cols=4)
            table.style = 'TableGrid'
            table.alignment = 1
            # Set the width of the first column to 1 inch
            column_width = Inches(1.2)
            table.columns[0].width = column_width
            cell = table.cell(0, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('ক) ইনভয়েস নম্বর')
            p3.style = custom_tableh
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # Set the width of the second column to 2 inches
            column_width = Inches(1)
            table.columns[1].width = column_width
            cell = table.cell(0, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('তারিখ')
            p3.style = custom_tableh
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            # Set the width of the third column to 0.5 inches
            column_width = Inches(4.3)
            table.columns[2].width = column_width
            cell1 = table.cell(0, 2)
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell1.paragraphs[0]
            p3.add_run('খ) ইনভয়েসে উল্লেখিত সেবা/পণ্যের পরিমাণ (Qty & Hours)')
            p3.style = custom_tableh
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # table.columns[2].width = Inches(3.8)
    
            # Set the width of the third column to 0.5 inches
            column_width = Inches(1)
            table.columns[3].width = column_width
            cell = table.cell(0, 3)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name = 'মূল্য (' + file_obj.fc_currency_id.name + ')'
            # p3.add_run('মূল্য (USD)')
            p3.add_run(cur_name)
            p3.style = custom_tableh
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            #
            # # add a dynamic row to the table
            cus_name = ''
            if file_obj.swift_customer_name:
                cus_name = file_obj.swift_customer_name
    
            # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
            total_qty = 0
            dev_total1 = sum(line.quantity if line.quantity_type == '1' else 0 for line in
                             file_obj.invoice_line_ids.invoice_id.invoice_line_ids)
            hour_total1 = sum(line.quantity if line.quantity_type == '0' else 0 for line in
                              file_obj.invoice_line_ids.invoice_id.invoice_line_ids)
            dev_total = dev_total1
            hour_total = hour_total1
    
            if hour_total == int(hour_total):
                hour_total_num = "{:.0f}".format(hour_total)
            else:
                hour_total_num = "{:.2f}".format(hour_total)
    
            if dev_total == int(dev_total):
                dev_total_num = "{:.0f}".format(dev_total)
            else:
                dev_total_num = "{:.2f}".format(dev_total)
    
            qty_str = ''
            if dev_total > 0:
                qty_str += str(dev_total_num) + (str(' Developers') if dev_total > 1 else str(' Developer'))
            if hour_total > 0:
                qty_str += str(hour_total_num) + str(' HRS') if not qty_str else ' & ' + str(
                    hour_total_num) + str(' HRS')
            all_invoice_qty_str = 'Software Development /' + qty_str
            all_invoice_qty_str_a = qty_str
            row_len = len(file_obj.invoice_line_ids)
            swift_amt = 0
            inv_amount = 0
    
            from itertools import groupby
    
            # group the invoice_line_ids by invoice_id.ref
            groups = groupby(sorted(file_obj.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                             key=lambda x: x.invoice_id.ref)
            # iterate over each group and add a new row for each unique invoice_id.ref
            row = 1
            for invoice_ref, group in groups:
                inv_obj = file_obj.env['cash.incentive.invoice'].search(
                    [('invoice_id.ref', '=', invoice_ref), ('head_id', '=', file_obj.id)], limit=1)
                inv_amount += inv_obj.invoice_amt
                invoice_date = ''
                if inv_obj.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime(
                        '%d-%b-%y')
    
                qty = 'Software Development /'
                h_q = 0
                d_q = 0
                for l in inv_obj.invoice_id.invoice_line_ids:
                    if l.quantity_type == '0':
                        h_q += l.quantity
                    else:
                        d_q += l.quantity
                if d_q:
                    if d_q == int(d_q):
                        d_q_num = "{:.0f}".format(d_q)
                    else:
                        d_q_num = "{:.2f}".format(d_q)
                    qty += str(d_q_num) + (str(' Developers') if d_q > 1 else str(' Developer'))
                if h_q:
                    if h_q == int(h_q):
                        h_q_new = "{:.0f}".format(h_q)
                    else:
                        h_q_new = "{:.2f}".format(h_q)
                    if not d_q:
                        qty += str(h_q_new) + ' HRS'
                    else:
                        qty += ' & ' + str(h_q_new) + ' HRS'
                invoice_qty_str = qty
    
                gr = 0
                for rec in group:
                    gr += 1
                    swift_amt += rec.swift_amt
                    # add a new row to the table
                    new_row = table.add_row()
                    first_cell = new_row.cells[0]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = custom_style5
                    first_cell = new_row.cells[1]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = custom_style5
                    first_cell = new_row.cells[2]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph21.style = custom_style5
                    first_cell = new_row.cells[3]
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style5
    
                cell_1_0 = table.cell(row, 0)
                p3 = cell_1_0.paragraphs[0]
                p3.add_run(invoice_ref)
                p3.style = custom_style_table
                cell_2_0 = table.cell((row + gr) - 1, 0)
                cell_1_0.merge(cell_2_0)
    
                cell_4_0 = table.cell(row, 1)
                p3 = cell_4_0.paragraphs[0]
                p3.add_run(invoice_date)
                p3.style = custom_style_table
                cell_3_0 = table.cell((row + gr) - 1, 1)
                cell_4_0.merge(cell_3_0)
    
                cell_5_0 = table.cell(row, 2)
                p3 = cell_5_0.paragraphs[0]
                p3.add_run(invoice_qty_str)
                p3.style = custom_style_table
                cell_6_0 = table.cell((row + gr) - 1, 2)
                cell_5_0.merge(cell_6_0)
    
                cell_5_0 = table.cell(row, 3)
                p3 = cell_5_0.paragraphs[0]
                p3.style = custom_style_table
                # p3.add_run(str("{:,}".format(round(inv_obj.invoice_amt, 2))))
                p3.add_run(str("{:,.2f}".format(inv_obj.invoice_amt)))
                cell_6_0 = table.cell((row + gr) - 1, 3)
                cell_5_0.merge(cell_6_0)
                row = gr + row
    
            # total
            total_row = table.add_row()
            total_row.cells[0].text = 'Total'
            total_row.cells[2].text = all_invoice_qty_str
            total_row.cells[3].text = str("{:,.2f}".format(inv_amount))
    
            first_cell = total_row.cells[0]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph21.style = custom_total
            first_cell = total_row.cells[2]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph21.style = custom_total
            first_cell = total_row.cells[3]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = custom_total
    
            self_date = ''
            if file_obj.date:
                self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')

    
            paragraph2 = document.add_paragraph("৬।\tরপ্তানিকৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত স্থানীয় সেবা/উপকরণাদির সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): 100% in house \tproduction of Brain Station 23 Limited \t\t পরিমাণঃ N/A  \t\t মূল্যঃ N/A"
                                                "\n৭।\tরপ্তানকিৃত সফটওয়্যার/ আইটিএস/ হার্ডওয়্যার তৈরীতে ব্যবহৃত আমদানিকৃত আনুসাংগিক সেবা/উপরকরণাদি সংগ্রহ সূত্র (সরবরাহকারীর নাম ও ঠিকানা): N/A  \t\tপরিমাণঃ N/A   \t\t\t\t\t  মূল্যঃ  N/A")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_after = Pt(1)
    
    
            if len(file_obj.invoice_line_ids) > 10:
                long_text = "৯। \t " + ' ' * 4 + "জাহাজীকরণের তারিখঃ " + inv_date + " \t\tগন্তব্য বন্দরঃ " + country
                wrapped_lines = textwrap.wrap(long_text, width=120)
                formatted_lines = [line + '\t' for line in wrapped_lines]
                formatted_text = '\n\t'.join(formatted_lines)
                paragraph2 = document.add_paragraph(formatted_text)
                paragraph2.style = custom_stylebl
                paragraph2.paragraph_format.space_before = Pt(1)
                paragraph2.paragraph_format.space_after = Pt(1)
            else:
                if len(file_obj.invoice_line_ids) > 7:
                    paragraph2 = document.add_paragraph("৯।\tজাহাজীকরণের তারিখঃ " + inv_date + "\n\tগন্তব্য বন্দরঃ " + country )
                    paragraph2.style = custom_stylebl
                else:
                    paragraph2 = document.add_paragraph("৯।\tজাহাজীকরণের তারিখঃ " + inv_date + " \tগন্তব্য বন্দরঃ " + country )
                    paragraph2.style = custom_stylebl
                paragraph2.paragraph_format.space_before = Pt(1)
                paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph("১০।\tইএক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানরি ক্ষেত্রে): N/A  \t\t\t  মূল্যঃ N/A   \t\t তারিখঃ N/A")
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            long_text1 = "১১। \t " + ' ' * 4 + "মোট প্রত্যাবাসিত রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str("{:,.2f}".format(swift_amt)) + "\t\tনীট এফওবি রপ্তানি মূল্য (বৈদেশিক মুদ্রায়): " + currency + ' ' + str("{:,.2f}".format(swift_amt))
            wrapped_lines1 = textwrap.wrap(long_text1, width=150)
            formatted_lines1 = [line + '\t' for line in wrapped_lines1]
            formatted_text1 = '\n\t'.join(formatted_lines1)
            paragraph2 = document.add_paragraph(formatted_text1)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            prc_date = ''
            if file_obj.prc_date:
                prc_date = datetime.datetime.strptime(str(file_obj.prc_date), '%Y-%m-%d').strftime('%d/%m/%y')
            prc_ref_code = ''
            if file_obj.prc_ref_code:
                prc_ref_code = file_obj.prc_ref_code
    
            long_text1 ="১২।\t" + ' ' * 4 + "প্রত্যাবাসিত রপ্তানি মূল্যের সনদপত্র নম্বরঃ " + prc_ref_code+ "\t\t\t তারিখঃ " + prc_date
            wrapped_lines1 = textwrap.wrap(long_text1, width=120)
            formatted_lines1 = [line + '\t' for line in wrapped_lines1]
            formatted_text1 = '\n\t'.join(formatted_lines1)
            paragraph2 = document.add_paragraph(formatted_text1)
            paragraph2.style = custom_stylebl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)

            paragraph2 = document.add_paragraph("\n\n\nরপ্তানিকারকের স্বাক্ষর ও তারিখ")
            paragraph2.style = custom_stylebsmallbold
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph4 = document.add_paragraph("এতদ্বারা প্রত্যয়ন করা যাচ্ছে যে, আমাদের নিজস্ব কারখানায়/প্রতিষ্ঠানে তৈরীকৃত/উৎপাদিত সফটওয়্যার/আইটিইএস/হার্ডওয়্যার উপরোক্ত ৬ ও ৭ নং ক্রমিক বর্ণিত সূত্র হতে সেবা/উপকরাণাদি সংগ্রহের মাধ্যমে রপ্তানির বিপরীতে ভর্তুকির জন্য উপরোক্ত অনুচ্ছেদগুলোতে উল্লিখিত বক্তব্য সঠিক ও নির্ভুল। বিদেশী ক্রেতা/ আমদানিকারকের ক্রয়াদেশের যথার্থতা/বিশ্বাসযোগ্যতা সম্পর্কেও নিশ্চিত করা হলো। ")
            paragraph4.style = custom_stylebsmall
            paragraph4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph4.paragraph_format.space_before = Pt(1)
            paragraph4.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph("\n\n\nরপ্তানিকারকের স্বাক্ষর ও তারিখ")
            paragraph2.style = custom_stylebsmallbold
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph("রপ্তানিকারকের উপরোক্ত ঘোষণার যথার্থতা যাচাইয়ান্তে সঠিক পাওয়া গিয়েছে। ৮নং ক্রমিকে উল্লিখিত ঘোষিত রপ্তানিমূল্য যৌক্তিক ও বিদ্যমান আন্তর্জাতিক বাজার মুল্যের সঙ্গে সংগতিপূর্ণ পাওয়া গিয়েছে এবং বিদেশী ক্রেতার যথার্থতা/বিশ্বাসযোগ্যতা সর্ম্পকেও নিশ্চিত হওয়া গিয়েছে। প্রত্যাবাসিত রপ্তানি মূল্যের (নীট এফওবি মূল্য) রপ্তানি ভর্তুকি পরিশোধের সুপারিশ করা হলো।")
            paragraph2.style = custom_stylebsmall
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph("\n\n\n\n----------------------------------- এসোসিয়েশন এর দুইজন উপযুক্ত কর্মকর্তার স্বাক্ষর, তারিখ ও সীল")
            # paragraph2.style = my_styler
            paragraph2.style = custom_stylebsmall
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph("[কোন প্রকার ঘষামাজা, কাটাছেড়া বা সংশোধন করা হলে এ প্রত্যয়নপত্র বাতিল বলে গণ্য হবে।]")
            # paragraph2.style = my_stylec
            paragraph2.style = custom_stylebsmall
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(0)
            document.add_page_break()
            # file_name = '%s_kha' % (datetime.datetime.now())
            # bank 55555555 --------------------------------------------------
    
            section.page_width = Cm(22)  # set the page width to 21 centimeters
            section.page_height = Cm(29.9)  # set the page width to 21 centimeters
            section.left_margin = Cm(2.5)  # 1.5 cm
            section.right_margin = Cm(2)
    
            self_date = ''
            current_date = datetime.datetime.now().date()
            if file_obj.date:
                self_date = datetime.datetime.strptime(str(current_date), '%Y-%m-%d').strftime('%d/%m/%y')
    
            # Define a custom style
            custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            custom_style.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(10)
            custom_style2.font.name = 'Calibri'
            mystyle += 1
    
            custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(10)
            custom_style3.font.bold = True
            custom_style3.font.name = 'Calibri'
            mystyle += 1
    
            custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.bold = True
            custom_style4.font.size = Pt(9)
            custom_style4.font.name = 'Arial Narrow'
            mystyle += 1
    
            table_total = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            table_total.font.bold = True
            table_total.font.size = Pt(10)
            table_total.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_style5 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style5.font.size = Pt(9)
            custom_style5.font.name = 'Arial Narrow'
            mystyle += 1
    
            paragraph2 = document.add_paragraph('\n\n\n\n\nRef: ' + file_obj.name + '\n\nDate: ' + self_date) # + '\n'
            paragraph2.style = custom_style3
            paragraph2.paragraph_format.left_indent = Pt(-6)
    
            paragraph2 = document.add_paragraph("The Branch Manager  \nGulshan Branch \n" + file_obj.bank_id.name + "\nHolding No. 75, Gulshan Avenue \nGulshan, Dhaka \n\n\nDear Sir\n")
            paragraph2.style = custom_style2
            paragraph2.paragraph_format.space_before = Pt(0)
            paragraph2.paragraph_format.space_after = Pt(1)
            paragraph2.paragraph_format.left_indent = Pt(-6)
    
            paragraph2 = document.add_paragraph("For cash incentive claim.\n")
            paragraph2.style = custom_style3
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
            paragraph2.paragraph_format.left_indent = Pt(-6)
    
            paragraph2 = document.add_paragraph("We are submitting herewith necessary documents against following Invoices:")
            paragraph2.style = custom_style2
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(0)
            paragraph2.paragraph_format.left_indent = Pt(-6)
    
            # table -----------------------
    
            table = document.add_table(rows=1, cols=8)
            table.style = 'TableGrid'
            table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
            # table.autofit = False  # Disable auto-fit behavior
            # table.left_indent = Inches(6)
    
            cell = table.cell(0, 0)
            # table.columns[0].width = Inches(.2)
            # table.columns[0].left_indent = Pt(100)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('SL#')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4
    
            cell = table.cell(0, 1)
            table.columns[1].width = Inches(1.8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Client Name')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4
    
            cell = table.cell(0, 2)
            table.columns[2].width = Inches(.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Invoice No')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4
    
            cell = table.cell(0, 3)
            table.columns[3].width = Inches(.7)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name = 'Invoice amount (' + file_obj.fc_currency_id.name + ')'
            p3.add_run(cur_name)
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p3.style = custom_style4
    
            cell = table.cell(0, 4)
            table.columns[4].width = Inches(.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Invoice Date')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4
    
            cell = table.cell(0, 5)
            table.columns[5].width = Inches(.6)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Swift/nostro date')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4
    
            cell = table.cell(0, 6)
            table.columns[6].width = Inches(.6)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name_r = 'Realize Amount (' + file_obj.fc_currency_id.name + ')'
            p3.add_run(cur_name_r)
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p3.style = custom_style4
    
            cell = table.cell(0, 7)
            table.columns[7].width = Inches(.6)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            cur_name_i = 'Incentive claim (' + file_obj.fc_currency_id.name + ')'
            p3.add_run(cur_name_i)
            p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p3.style = custom_style4
    
            # add a dynamic row to the table
            cus_name = ''
            if file_obj.swift_customer_name:
                cus_name = file_obj.swift_customer_name
    
            # row_data = [('John', '35', 'Male'), ('Mary', '27', 'Female'), ('Bob', '42', 'Male')]
            row_len = len(file_obj.invoice_line_ids)
            inv_amount = 0
            encashment_amt_bdt = 0
            sl = 1
            total_incentive_rate_fc = 0
    
            from itertools import groupby
    
            # group the invoice_line_ids by invoice_id.ref
            groups = groupby(sorted(file_obj.invoice_line_ids, key=lambda x: x.invoice_id.ref),
                             key=lambda x: x.invoice_id.ref)
            # iterate over each group and add a new row for each unique invoice_id.ref
            row = 1
            for invoice_ref, group in groups:
                inv_obj = file_obj.env['cash.incentive.invoice'].search([('invoice_id.ref', '=', invoice_ref), ('head_id', '=', file_obj.id)], limit=1)
                # usd_price = sum(inv_obj.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
                usd_price = inv_obj.invoice_id.invoice_total_fc
                # incentive_rate_fc = usd_price / file_obj.incentive_rate_fc if file_obj.incentive_rate_fc else 0
    
                inv_amount += inv_obj.invoice_amt
                invoice_date = ''
                if inv_obj.invoice_date:
                    invoice_date = datetime.datetime.strptime(str(inv_obj.invoice_date), '%Y-%m-%d').strftime('%d-%b-%y')
                gr = 0
                for rec in group:
                    gr += 1
                    encashment_amt_bdt += rec.encashment_amt_bdt
                    incentive_rate_fc = (rec.swift_amt * file_obj.incentive_rate_fc) / 100 if file_obj.incentive_rate_fc else 0
                    total_incentive_rate_fc += incentive_rate_fc
                    swift_message_date = ''
                    if rec.swift_message_id.date:
                        swift_message_date = datetime.datetime.strptime(str(rec.swift_message_id.date), '%Y-%m-%d').strftime('%d-%b-%y')
    
                    new_row = table.add_row()
                    # add data to the cells in the new row
                    new_row.cells[0].text = str(sl)
                    new_row.cells[5].text = swift_message_date
                    new_row.cells[6].text = str("{:,}".format(round(rec.swift_amt, 2)))
                    new_row.cells[7].text = str("{:,}".format(float(format(incentive_rate_fc, '.2f'))))
    
                    first_cell = new_row.cells[0]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph24 = first_cell.paragraphs[0]
                    paragraph24.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph24.style = custom_style5
    
                    first_cell = new_row.cells[2]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style5
    
                    first_cell = new_row.cells[3]
                    first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph21 = first_cell.paragraphs[0]
                    paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph21.style = custom_style5
    
                    first_cell1 = new_row.cells[6]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = custom_style5
    
                    first_cell1 = new_row.cells[7]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    paragraph22.style = custom_style5
    
                    # first_cell1 = new_row.cells[2]
                    # first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # paragraph22 = first_cell1.paragraphs[0]
                    # paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # paragraph22.style = custom_style5
    
                    first_cell1 = new_row.cells[1]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.style = custom_style5
    
                    first_cell1 = new_row.cells[4]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
    
                    first_cell1 = new_row.cells[5]
                    first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph22 = first_cell1.paragraphs[0]
                    paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph22.style = custom_style5
    
                    sl += 1
    
                cell_1_0 = table.cell(row, 2)
                p3 = cell_1_0.paragraphs[0]
                p3.add_run(invoice_ref)
                cell_2_0 = table.cell((row + gr) - 1, 2)
                cell_1_0.merge(cell_2_0)
    
                cell_4_0 = table.cell(row, 3)
                p3 = cell_4_0.paragraphs[0]
                p3.add_run(str("{:,}".format(round(usd_price, 2))))
                cell_3_0 = table.cell((row + gr) - 1, 3)
                cell_4_0.merge(cell_3_0)
    
                cell_5_0 = table.cell(row, 4)
                p3 = cell_5_0.paragraphs[0]
                p3.add_run(invoice_date)
                cell_6_0 = table.cell((row + gr) - 1, 4)
                cell_5_0.merge(cell_6_0)
                row = gr + row
    
            # for rec in file_obj.invoice_line_ids:
            #     usd_price = sum(rec.invoice_id.invoice_line_ids.mapped('qty_usd_price'))
            #     incentive_rate_fc = usd_price / file_obj.incentive_rate_fc if file_obj.incentive_rate_fc else 0
            #     total_incentive_rate_fc += incentive_rate_fc
            #     inv_amount += rec.invoice_amt
            #     encashment_amt_bdt += rec.encashment_amt_bdt
            #
            #     invoice_date = ''
            #     if rec.invoice_date:
            #         invoice_date = datetime.datetime.strptime(str(rec.invoice_date), '%Y-%m-%d').strftime('%d-%b-%y')
            #     swift_message_date = ''
            #     if rec.swift_message_id.date:
            #         swift_message_date = datetime.datetime.strptime(str(rec.swift_message_id.date),
            #                                                         '%Y-%m-%d').strftime('%d-%b-%y')
            #     # add a new row to the table
            #     new_row = table.add_row()
            #
            #     # add data to the cells in the new row
            #     new_row.cells[0].text = str(sl)
            #     new_row.cells[2].text = rec.invoice_id.ref
            #     new_row.cells[3].text = str("{:,}".format(round(usd_price, 2)))
            #     new_row.cells[4].text = invoice_date
            #     new_row.cells[5].text = swift_message_date
            #     new_row.cells[6].text = str("{:,}".format(round(rec.swift_amt, 2)))
            #     new_row.cells[7].text = str("{:,}".format(float(format(incentive_rate_fc, '.2f'))))
            #
            #     first_cell = new_row.cells[0]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph24 = first_cell.paragraphs[0]
            #     paragraph24.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph24.style = custom_style5
            #
            #     first_cell = new_row.cells[3]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph21 = first_cell.paragraphs[0]
            #     paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #     paragraph21.style = custom_style5
            #
            #     first_cell1 = new_row.cells[6]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[7]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[2]
            #     first_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[1]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[4]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style5
            #
            #     first_cell1 = new_row.cells[5]
            #     first_cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            #     paragraph22 = first_cell1.paragraphs[0]
            #     paragraph22.alignment = WD_ALIGN_PARAGRAPH.CENTER
            #     paragraph22.style = custom_style5
            #
            #     sl += 1
    
            # cus name merge
            cell_1_0 = table.cell(1, 1)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run(cus_name)
            cell_2_0 = table.cell(row_len, 1)
            cell_1_0.merge(cell_2_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style4
    
            # total
            total_row = table.add_row()
            total_row.cells[2].text = 'Total'
            total_row.cells[3].text = str("{:,}".format(round(inv_amount, 2)))
            total_row.cells[6].text = str("{:,}".format(round(swift_amt, 2)))
            total_row.cells[7].text = str("{:,}".format(float(format(total_incentive_rate_fc, '.2f'))))
    
            first_cell = total_row.cells[3]
            paragraph21 = first_cell.paragraphs[0]
            paragraph21.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph21.style = table_total
    
            first_cell1 = total_row.cells[6]
            paragraph22 = first_cell1.paragraphs[0]
            paragraph22.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph22.style = table_total
    
            first_cell2 = total_row.cells[7]
            paragraph23 = first_cell2.paragraphs[0]
            paragraph23.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph23.style = table_total
    
            first_cell2 = total_row.cells[2]
            paragraph23 = first_cell2.paragraphs[0]
            paragraph23.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph23.style = table_total
    
            # after table
            paragraph2 = document.add_paragraph("\nThose documents are as follows:")
            paragraph2.style = custom_style2
            paragraph2.paragraph_format.left_indent = Pt(-6)
    
            # paragraph2 = document.add_paragraph("\t 1. Form Ka\n\t 2. Commercial Invoice\n\t 3. Form-C (ICT)\n\t 4. Agreement\n\t 5. Certificate of Authorized Dealer\n\t 6. Copy of Swift Message\n\t 7. নগদ ভর্তুকি প্রত্যয়নপত্র \n\t 8. Company's undertaking")
            # paragraph2.style = custom_style2
    
            # Create a table with 1 row and 2 columns
            table = document.add_table(rows=1, cols=2)
            table.border = None
            table.columns[0].width = Inches(3)  # Adjust the width as needed
            table.columns[1].width = Inches(3)  # Adjust the width as needed
    
            # Get the first row of the table
            row = table.rows[0]
            # Remove cell borders
            for cell in row.cells:
                cell.border = None
    
            left_column = table.cell(0, 0)
            left_column = left_column.paragraphs[0]
            left_column.add_run("1. Form Ka\n")
            left_column.add_run("2. Commercial Invoice\n")
            left_column.add_run("3. Form-C (ICT)\n")
            left_column.add_run("4. Agreement\n")
            left_column.paragraph_format.left_indent = Inches(.7)
            left_column.style = custom_style2
            # left_column.paragraph_format.left_indent = Pt(-6)
    
            # Add the last four items to the right column
            right_column = table.cell(0, 1)
            right_column = right_column.paragraphs[0]
            right_column.add_run("5. Certificate of Authorized Dealer\n")
            right_column.add_run("6. Copy of Swift Message\n")
            run2 = right_column.add_run("7. নগদ ভর্তুকি প্রত্যয়নপত্র\n")
            run2.font.size = Pt(11)
            right_column.add_run("8. Company's undertaking\n")
            right_column.style = custom_style2
            # right_column.paragraph_format.left_indent = Pt(-6)
    
            paragraph2 = document.add_paragraph("Please proceed at your earliest possible time. \n\nThanking You \n\n\nMasudur Rahman \nSenior Manager, Finance, Accounts & Admin \nBrain Station 23 Limited")
            paragraph2.style = custom_style2
            paragraph2.paragraph_format.left_indent = Pt(-6)
            document.add_page_break()
            # file_name = '%s_flfb' % (datetime.datetime.now())
    
            # ga 7777777 ------------------------------

            section = document.add_section()
            section.orientation = 0
            section.page_width = Pt(792)
            section.page_height = Pt(612)
    
            self_date = ''
            if file_obj.date:
                self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')
    
            custom_bangla_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_bangla_style.font.size = Pt(11)
            custom_bangla_style.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_bangla_style_ga = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_bangla_style_ga.font.size = Pt(11)
            custom_bangla_style_ga.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_stylebold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebold.font.size = Pt(14)
            custom_stylebold.font.name = 'SutonnyOMJ'
            custom_stylebold.font.bold = True
            custom_stylebold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mystyle += 1
    
            custom_styleubl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_styleubl.font.size = Pt(11)
            custom_styleubl.font.name = 'SutonnyOMJ'
            custom_styleubl.font.underline = True
            custom_styleubl.font.bold = True
            mystyle += 1
    
            # p1 = document.add_paragraph('(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য) \t\t\t \t \t \t \t  \t \t  \t \t \t \t   ফরম-‘গ’ \n')
            # p1.style = custom_bangla_style
    
            paragraph = document.add_paragraph()
            # paragraph.add_run("(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
            run1 = paragraph.add_run("\n(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য)")
            run1.font.size = Pt(11)
            run1.font.name = 'SutonnyOMJ'
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t")
            paragraph.add_run("\t ")
            paragraph.add_run("\t ")
            paragraph.add_run("\t ")
            run2 = paragraph.add_run("ফরম-‘গ’")
            run2.font.size = Pt(14)
            run2.font.bold = True
            run2.font.name = 'SutonnyOMJ'

            # Define a custom style
            custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            # custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mystyle += 1
    
            custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(11)
            custom_style2.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(11)
            custom_style3.font.bold = True
            mystyle += 1
    
            paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক")
            paragraph2.style = custom_bangla_style_ga
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            # paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক \nবাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব \n-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
            paragraph2 = document.add_paragraph("বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধ্য অর্থের দাবী প্রস্তাব")
            paragraph2.style = custom_stylebold
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph("-------------------------------------- ব্যাংক, প্রধান কার্যালয়, ঢাকা।")
            paragraph2.style = custom_bangla_style_ga
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)
    
            # table -----------------------
            currency = ''
            if file_obj.fc_currency_id:
                currency = file_obj.fc_currency_id.name
    
            a = []
            c = []
            for rec in file_obj.invoice_line_ids:
                a.append(rec.od_sight_rate)
                if rec.encashment_rate_bdt:
                    c.append(rec.encashment_rate_bdt)
                # if rec.swift_message_id.encashment_rate_bdt:
                #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
            inv_amount = sum(file_obj.invoice_line_ids.mapped('swift_amt'))
            if a:
                b = min(a)
                od_s_r = b
            else:
                od_s_r = 0
    
            if c:
                d = min(c)
                encashment_rate_bdt = d
            else:
                encashment_rate_bdt = 0
    
            if file_obj.od_sight_rate:
                od_sight_rate = file_obj.od_sight_rate
            else:
                if od_s_r:
                    od_sight_rate = od_s_r
                else:
                    od_sight_rate = encashment_rate_bdt
    
            rate_fc = inv_amount / file_obj.incentive_rate_fc
            encashment_final_amt = round(rate_fc, 2) * od_sight_rate
            # encashment_final_amt_a = "{:,.0f}".format(round(encashment_final_amt))
            encashment_final_amt_a = file_obj.custom_round(encashment_final_amt)
            usd_bdt_amount = currency + ' ' + str("{:,.2f}".format(rate_fc)) + '\n' + '@' + \
                             str("{:,.2f}".format(od_sight_rate)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a
    
            table = document.add_table(rows=4, cols=9)
            table.style = 'TableGrid'
            cell_1_0 = table.cell(0, 0)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run('ভর্তুকি আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানা')
            p3.style = custom_bangla_style
            cell_2_0 = table.cell(1, 0)
            cell_1_0.merge(cell_2_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(0, 1)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('রপ্তানি সংশ্লিষ্ট তথ্য')
            p3.style = custom_bangla_style
            cell_2_1 = table.cell(0, 5)
            cell_1_1.merge(cell_2_1)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 1)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('সেবা/পণ্যের বিবরণ')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 2)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('জাহাজীকরণের তারিখ ও গন্তব্যস্থল')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 3)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('ইক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানির ক্ষেত্রে)')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 4)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('মোট প্রত্যাবাসিত রপ্তানি (মূল্য  বৈদেশিক মুদ্রা)')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 5)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('নীট এফওবি মূল্য (বৈদেশিক মুদ্রা)')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_3 = table.cell(0, 6)
            p3 = cell_1_3.paragraphs[0]
            p3.add_run('শাখা কর্তৃক আবেদনপত্র গ্রহণের তারিখ')
            p3.style = custom_bangla_style
            cell_2_4 = table.cell(1, 6)
            cell_1_3.merge(cell_2_4)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_5 = table.cell(0, 7)
            p3 = cell_1_5.paragraphs[0]
            p3.add_run('আবেদনকৃত অর্থের পরিমাণ')
            p3.style = custom_bangla_style
            cell_2_6 = table.cell(1, 7)
            cell_1_5.merge(cell_2_6)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_6 = table.cell(0, 8)
            p3 = cell_1_6.paragraphs[0]
            p3.add_run('পরিশোধ্য দাবির পরিমাণ')
            p3.style = custom_bangla_style
            cell_2_7 = table.cell(1, 8)
            cell_1_6.merge(cell_2_7)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('১')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('২')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 2)
            p3 = cell.paragraphs[0]
            p3.add_run('৩')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 3)
            p3 = cell.paragraphs[0]
            p3.add_run('৪')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 4)
            p3 = cell.paragraphs[0]
            p3.add_run('৫')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 5)
            p3 = cell.paragraphs[0]
            p3.add_run('৬')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 6)
            p3 = cell.paragraphs[0]
            p3.add_run('৭')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 7)
            p3 = cell.paragraphs[0]
            p3.add_run('৮')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 8)
            p3 = cell.paragraphs[0]
            p3.add_run('৯')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 0)
    
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('\n\nব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী  বাণিজ্যিক এলাকা, ঢাকা- ১২১২।\n\n')
            p3.style = custom_bangla_style
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # p3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
            cell = table.cell(3, 1)
            cell.height = Inches(12)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Software Development')
            p3.style = custom_style2
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 2)
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 3)
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            cell = table.cell(3, 5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_style3
    
            cell = table.cell(3, 6)
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 7)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run(usd_bdt_amount)
            p3.style = custom_style3
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 8)
            cell.height = Inches(12)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            column_index = 0
            column_cells = [row.cells[column_index] for row in table.rows]
    
            # Set the desired height for each cell in the column
            cell_height = Inches(12)  # Adjust the height value as needed
            for cell in column_cells:
                cell.height = cell_height
    
            p3 = document.add_paragraph('\n\n\n\n\n\t\t দাপ্তরিক সীল  \t\t\t\t\t\t \t\t\t\t  ক্ষমতাপ্রাপ্ত ব্যাংক কর্মকর্তার স্বাক্ষর, তারিখ, নাম ও পদবী')
            p3.style = custom_bangla_style
            # file_name = '%s_ga' % (datetime.datetime.now())
            document.add_page_break()
    
            # gha 77777777  ------------------------------
    
            self_date = ''
            if file_obj.date:
                self_date = datetime.datetime.strptime(str(file_obj.date), '%Y-%m-%d').strftime('%d/%m/%y')
            # document.add_heading('CERTIFICATE OF AUTHORIZED DEALER', level=1)
    
            # Define a custom style
            custom_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            # custom_style.font.bold = True
            custom_style.font.size = Pt(14)
            custom_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mystyle += 1
    
            custom_style2 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style2.font.size = Pt(11)
            custom_style2.font.name = 'Arial Narrow'
            mystyle += 1
    
            custom_style3 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style3.font.size = Pt(11)
            custom_style3.font.bold = True
            mystyle += 1
    
            custom_style4 = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_style4.font.bold = True
            custom_style4.font.size = Pt(14)
            custom_style4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            custom_style4.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_bangla_style = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_bangla_style.font.size = Pt(11)
            custom_bangla_style.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_bangla_style_gha = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_bangla_style_gha.font.size = Pt(11)
            custom_bangla_style_gha.font.name = 'SutonnyOMJ'
            mystyle += 1
    
            custom_stylebold = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_stylebold.font.size = Pt(16)
            custom_stylebold.font.name = 'SutonnyOMJ'
            custom_stylebold.font.bold = True
            custom_stylebold.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mystyle += 1
    
            custom_styleubl = document.styles.add_style(str(mystyle), WD_STYLE_TYPE.PARAGRAPH)
            custom_styleubl.font.size = Pt(13)
            custom_styleubl.font.name = 'SutonnyOMJ'
            custom_styleubl.font.bold = True
            custom_styleubl.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            mystyle += 1
    
            paragraph1 = document.add_paragraph('\nফরম-‘ঘ’')
            paragraph1.style = custom_style4
    
            p3 = document.add_paragraph('(অনুচ্ছেদ ৬(গ), এফই সার্কুলার নম্বর-০৩/২০১৮ দ্রষ্টব্য) \n')
            p3.style = custom_bangla_style_gha
    
            # paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক \nবাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও \nহার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধিত হিসাবের .........সালের ...... মাসের বিবরণী \nঅনুমোদিত ডিলার ব্যাংক শাখার নামঃ...................................")
    
            paragraph2 = document.add_paragraph("এফই সার্কুলার নং-০৩/২০১৮ মোতাবেক")
            paragraph2.style = custom_bangla_style_gha
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph("বাংলাদেশ হতে সফটওয়্যার, আইটিএস (Information Technology Enabled Services) ও")
            paragraph2.style = custom_stylebold
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph("হার্ডওয়্যার রপ্তানি খাতে রপ্তানি ভর্তুকি বাবদ পরিশোধিত হিসাবের .........সালের ...... মাসের বিবরণী")
            paragraph2.style = custom_styleubl
            paragraph2.paragraph_format.space_before = Pt(1)
            paragraph2.paragraph_format.space_after = Pt(1)
    
            paragraph2 = document.add_paragraph("অনুমোদিত ডিলার ব্যাংক শাখার নামঃ...................................................................................")
            paragraph2.style = custom_bangla_style_gha
            paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph2.paragraph_format.space_before = Pt(1)
    
            # table -----------------------
            currency = ''
            if file_obj.fc_currency_id:
                currency = file_obj.fc_currency_id.name
    
            a = []
            c = []
            for rec in file_obj.invoice_line_ids:
                a.append(rec.od_sight_rate)
                if rec.encashment_rate_bdt:
                    c.append(rec.encashment_rate_bdt)
                # if rec.swift_message_id.encashment_rate_bdt:
                #     encashment_rate_bdt = rec.swift_message_id.encashment_rate_bdt
            inv_amount = sum(file_obj.invoice_line_ids.mapped('swift_amt'))
    
            if a:
                b = min(a)
                od_s_r = b
            else:
                od_s_r = 0
    
            if c:
                d = min(c)
                encashment_rate_bdt = d
            else:
                encashment_rate_bdt = 0
    
            if file_obj.od_sight_rate:
                od_sight_rate = file_obj.od_sight_rate
            else:
                if od_s_r:
                    od_sight_rate = od_s_r
                else:
                    od_sight_rate = encashment_rate_bdt
    
            rate_fc = inv_amount / file_obj.incentive_rate_fc
            encashment_final_amt = round(rate_fc, 2) * od_sight_rate
            # encashment_final_amt_a = "{:,.0f}".format(round(encashment_final_amt))
            encashment_final_amt_a = file_obj.custom_round(encashment_final_amt)
            usd_bdt_amount = currency + ' ' + str("{:,.2f}".format(rate_fc)) + '\n' + '@' + \
                             str("{:,.2f}".format(od_sight_rate)) + '=' + '\n' + 'BDT ' + encashment_final_amt_a
    
            table = document.add_table(rows=4, cols=10)
            table.style = 'TableGrid'
    
            cell_1_0 = table.cell(0, 0)
            p3 = cell_1_0.paragraphs[0]
            p3.add_run('ভর্তুকি আবেদনকারী প্রতিষ্ঠানের নাম ও ঠিকানা')
            p3.style = custom_bangla_style_gha
            cell_2_0 = table.cell(1, 0)
            cell_1_0.merge(cell_2_0)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(0, 1)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('রপ্তানি সংশ্লিষ্ট তথ্য')
            p3.style = custom_bangla_style_gha
            cell_2_1 = table.cell(0, 5)
            cell_1_1.merge(cell_2_1)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_2_1 = table.cell(0, 8)
            p3 = cell_2_1.paragraphs[0]
            p3.add_run('শাখায় প্রাপ্ত অনিস্পন্ন আবেদন')
            p3.style = custom_bangla_style_gha
            cell_3_1 = table.cell(0, 9)
            cell_2_1.merge(cell_3_1)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 1)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('সেবা/পণ্যের বিবরণ')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 2)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('জাহাজীকরণের তারিখ ও গন্তব্যস্থল')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 3)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('ইক্সপি নম্বর (দৃশ্যমান পণ্য রপ্তানির ক্ষেত্রে)')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 4)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('মোট প্রত্যাবাসিত রপ্তানি (মূল্য  বৈদেশিক মুদ্রা)')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_1 = table.cell(1, 5)
            p3 = cell_1_1.paragraphs[0]
            p3.add_run('নীট এফওবি মূল্য (বৈদেশিক মুদ্রা)')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_3 = table.cell(0, 6)
            table.columns[7].width = Inches(.6)
            p3 = cell_1_3.paragraphs[0]
            p3.add_run('পরিশোধিত ভর্তুকি (টাকা)')
            p3.style = custom_bangla_style_gha
            cell_2_4 = table.cell(1, 6)
            cell_1_3.merge(cell_2_4)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_5 = table.cell(0, 7)
            table.columns[7].width = Inches(.6)
            p3 = cell_1_5.paragraphs[0]
            p3.add_run('পরিশোধের তারিখ')
            p3.style = custom_bangla_style_gha
            cell_2_6 = table.cell(1, 7)
            cell_1_5.merge(cell_2_6)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_6 = table.cell(1, 8)
            table.columns[8].width = Inches(.5)
            p3 = cell_1_6.paragraphs[0]
            p3.add_run('সংখ্যা')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell_1_6 = table.cell(1, 9)
            table.columns[9].width = Inches(1.7)
            p3 = cell_1_6.paragraphs[0]
            p3.add_run('দাবীকৃত অর্থ (টাকা)')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 0)
            p3 = cell.paragraphs[0]
            p3.add_run('১')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 1)
            p3 = cell.paragraphs[0]
            p3.add_run('২')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 2)
            p3 = cell.paragraphs[0]
            p3.add_run('৩')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 3)
            p3 = cell.paragraphs[0]
            p3.add_run('৪')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 4)
            p3 = cell.paragraphs[0]
            p3.add_run('৫')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 5)
            p3 = cell.paragraphs[0]
            p3.add_run('৬')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 6)
            p3 = cell.paragraphs[0]
            p3.add_run('৭')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 7)
            p3 = cell.paragraphs[0]
            p3.add_run('৮')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 8)
            p3 = cell.paragraphs[0]
            p3.add_run('৯')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(2, 9)
            p3 = cell.paragraphs[0]
            p3.add_run('১০')
            p3.style = custom_bangla_style_gha
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run(
                '\nব্রেইন স্টেশন ২৩ লিমিটেড, প্লট ০২ (৮ম এবং ৯ম তলা), বীর উত্তম এ কে খন্দকার রোড, মহাখালী  বাণিজ্যিক এলাকা, ঢাকা- ১২১২।\n\n')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.style = custom_bangla_style_gha
    
            cell = table.cell(3, 1)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run('Software Development')
            p3.style = custom_style2
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 2)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 3)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 4)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
            p3.style = custom_style3
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run(currency + ' ' + str("{:,.2f}".format(inv_amount)))
            p3.style = custom_style3
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 6)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 7)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 8)
            p3 = cell.paragraphs[0]
            p3.add_run('')
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            cell = table.cell(3, 9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p3 = cell.paragraphs[0]
            p3.add_run(usd_bdt_amount)
            p3.style = custom_style3
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
            p3 = document.add_paragraph(
                '\n\n\n\n\t\t দাপ্তরিক সীল  \t\t\t\t\t\t\t\t\t  ক্ষমতাপ্রাপ্ত ব্যাংক কর্মকর্তার স্বাক্ষর, তারিখ, নাম ও পদবী')
            p3.style = custom_bangla_style_gha
            # value = file_obj.name
            # split_values = value.split("/")
            # desired_value = split_values[-1]
            # file_names += desired_value if not file_names else ', ' + desired_value
            document.add_page_break()

            from docx import Document
            from docx.shared import Pt

            # Create a new Word document

        file_name = 'All_form_%s' % (datetime.datetime.now())
        # -------------------
        import os
        dir_path = os.path.dirname(os.path.abspath(__file__))
        base_path = str(dir_path).replace('/models', '').replace('/wizards', '')
        docxfile = base_path + '/static/docx/' + file_name + '.docx'
        document.save(docxfile)
    
            # return document.save('/home/jobaer/Downloads/jh3.docx')

        return {
            'type': 'ir.actions.act_url',
            'url': 'cash_incentive/static/docx/' + file_name + '.docx',
            'target': 'self',
        }
    
    def custom_round(self, value):
        decimal_part = value - math.floor(value)
        if decimal_part <= 0.49:
            a = math.floor(value)
            return str("{:,.0f}".format(round(a)))
        else:
            a = math.ceil(value)
            return str("{:,.0f}".format(round(a)))

    def num2words_fun(self,number):
        num = decimal.Decimal(number)
        decimal_part = num - int(num)
        num = int(num)

        # if decimal_part:
        #     return num2words(num) + " point " + (" ".join(num2words(i) for i in str(decimal_part)[2:]))

        under_20 = ['Zero', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine', 'Ten', 'Eleven',
                    'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen', 'Seventeen', 'Eighteen', 'Nineteen']
        tens = ['Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
        above_100 = {100: 'Hundred', 1000: 'Thousand', 100000: 'Lac', 10000000: 'Crores'}

        if num < 20:
            return under_20[num]

        if num < 100:
            return tens[num // 10 - 2] + ('' if num % 10 == 0 else ' ' + under_20[num % 10])

        # find the appropriate pivot - 'Million' in 3,603,550, or 'Thousand' in 603,550
        pivot = max([key for key in above_100.keys() if key <= num])

        return num2words(num // pivot) + ' ' + above_100[pivot] + (
            '' if num % pivot == 0 else ' ' + num2words(num % pivot))